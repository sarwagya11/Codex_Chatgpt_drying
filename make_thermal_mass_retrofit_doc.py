from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.top_margin    = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin   = Cm(2.3)
section.right_margin  = Cm(2.3)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)


def add_heading(text, size=11, bold=True, space_before=7, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'


def add_body(text, size=10.5, space_after=5, line_spacing=12.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing  = Pt(line_spacing)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'


def add_ref(number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(1.5)
    p.paragraph_format.left_indent       = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    p.paragraph_format.line_spacing      = Pt(10)
    run = p.add_run(f'[{number}] {text}')
    run.font.size = Pt(8.5)
    run.font.name = 'Times New Roman'


# ── TITLE ────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(0)
t.paragraph_format.space_after  = Pt(8)
r = t.add_run('Retrofitting Existing Buildings with Thermal Mass and Insulation in Cold Climates')
r.bold = True
r.font.size = Pt(13)
r.font.name = 'Times New Roman'

# ── PARA 1: Why retrofit / existing stock ────────────────────
add_body(
    'In cold-climate countries such as Sweden and Norway, the most significant near-term '
    'opportunity for reducing building energy use lies in the renovation of the existing stock '
    'rather than in new construction, which already meets stringent standards. A large share of '
    'Nordic housing was built between the 1950s and 1980s using lightweight or prefabricated '
    'techniques with low inherent thermal mass and poor envelope insulation. The Swedish Million '
    'Programme (miljonprogrammet) \u2014 approximately one million dwellings built between 1965 '
    'and 1975 as prefabricated concrete panel blocks \u2014 is the principal renovation target. '
    'Without sufficient thermal storage, these buildings lose heat rapidly during interruptions, '
    'overheat under solar gain, and offer little flexibility for demand shifting [1, 3, 13].'
)

# ── PARA 2: PCM retrofit ──────────────────────────────────────
add_body(
    'The most practically accessible method of adding thermal mass to an existing lightweight '
    'building is through phase change materials (PCMs), which can be incorporated into thin '
    'plasterboard or wallboard panels fixed directly to interior surfaces without structural '
    'alteration [4]. For cold climates, an optimal PCM melting point of 20\u201322\u00b0C is '
    'recommended, aligned with the indoor comfort set point [5]. Studies report reductions in '
    'indoor temperature fluctuations of up to 46%, heat transfer reductions of up to 47.6%, '
    'and heating demand reductions of 5\u201320% depending on thickness and placement [6, 7, 8]. '
    'Where existing concrete slabs are present, thermally activated building systems (TABS) '
    'offer a complementary pathway, embedding pipes in the slab to turn passive mass into a '
    'controllable thermal battery pre-charged with off-peak energy; TABS has been shown to '
    'reduce whole-life building energy costs by 16\u201327% [10, 11, 12].'
)

# ── PARA 3: Insulation retrofit — ETICS vs interior ──────────
add_heading('3.4  Interior Versus Exterior Insulation of Existing Walls',
            size=10.5, space_before=6, space_after=2)

add_body(
    'Deep insulation retrofits of the Million Programme stock have demonstrated purchased energy '
    'reductions of approximately 52% through new prefabricated insulated facade panels, '
    'triple-glazed windows, and heat-recovery ventilation, at a cost of around 200 \u20ac/m\u00b2 '
    'of treated floor area [14]. More modest measures, such as attic insulation with cellulose '
    'fibre, have halved heating demand alone [13]. External thermal insulation composite systems '
    '(ETICS) are the preferred approach: they keep the structural wall warm, reduce condensation '
    'risk, and preserve the thermal mass of the original wall on the interior side. Interior '
    'insulation is used where facade appearance must be preserved, but it drives the outer wall '
    'to lower temperatures, increasing risk of interstitial condensation [16]. A 2024 study of '
    'three Swedish heritage buildings confirmed that capillary-active materials \u2014 calcium '
    'silicate board and wood fibre board \u2014 substantially reduce this hygrothermal failure '
    'risk compared to vapour-tight solutions [17, 18].'
)

# ── PARA 4: Prefabricated facades ────────────────────────────
add_heading('3.5  Prefabricated Facade Systems',
            size=10.5, space_before=6, space_after=2)

add_body(
    'Factory-manufactured facade modules integrating insulation, cladding, windows, and '
    'ventilation terminals into a single prefabricated unit represent a growing approach to '
    'large-scale cold-climate renovation. Craned into place in one operation, they minimise '
    'on-site time and occupant disruption. EU-funded demonstration projects on Swedish Million '
    'Programme housing have delivered measured energy savings of up to 40%, bringing energy '
    'use to levels comparable with current new-build standards [15]. The approach is now '
    'recognised in Swedish renovation policy and continues to develop towards lower cost and '
    'greater integration with building energy management systems.',
    space_after=8
)

# ── REFERENCES ───────────────────────────────────────────────
add_heading('References', size=10, space_before=4, space_after=2)

refs = [
    ('1',  'Hamdy M., Hasan A. Enhancing Building Resilience During Power Outages in Cold Climates. Building Simulation, Springer, 2025. doi:10.1007/s12273-025-1371-2'),
    ('2',  'Overheating Risk and Energy Demand of Nordic Old and New Apartment Buildings. Applied Sciences, MDPI, 2021. doi:10.3390/app11093972'),
    ('3',  'Retrofitting Buildings into Thermal Batteries for Demand-Side Flexibility and Thermal Safety during Power Outages in Winter. Energies, MDPI, 2022. doi:10.3390/en15124405'),
    ('4',  'Soares N. et al. Retrofitting Buildings with PCM: Effects of Location and Climatic Condition. Building and Environment, Elsevier, 2023. doi:10.1016/j.buildenv.2023.110221'),
    ('5',  'SINTEF Blog. What are Phase Change Materials? blog.sintef.com/energy/phase-change-materials-pcm/'),
    ('6',  'Enhancing Building Thermal Performance: A Review of PCM Integration. Energies, MDPI, 2025. doi:10.3390/en18123200'),
    ('7',  'Jayalath A. et al. PCM Integration in Concrete for Thermal Energy Storage. Sustainable Energy Research, Springer, 2024. doi:10.1186/s40807-024-00138-8'),
    ('8',  'Kishore R.A. et al. PCM-Incorporated Building Envelope Under Cold Climates. Springer LNCE, 2024. doi:10.1007/978-981-97-8305-2_53'),
    ('9',  'California Energy Commission. PCM-Enhanced Insulation for Residential Wall Retrofits. CEC, 2024. energy.ca.gov/publications/2024'),
    ('10', 'REHVA Journal. Operation and Control of TABS. rehva.eu/rehva-journal/chapter/operation-and-control-of-thermally-activated-building-systems-tabs'),
    ('11', 'Roman\u00ed J. et al. A Review of TABS as an Alternative for Improving Indoor Environment. Energies, MDPI, 2022. doi:10.3390/en15176179'),
    ('12', 'Kensby J. et al. Energy Flexibility Using Thermal Mass for Swedish Single-Family Houses. E3S Web of Conf., 2024. doi:10.1051/e3sconf/202446504003'),
    ('13', 'H\u00f6gberg L. et al. Incentives for Energy Efficiency When Renovating the Swedish Million Homes Programme. Sustainability, MDPI, 2009. doi:10.3390/su1041349'),
    ('14', 'Passipedia. Step-by-Step Deep Retrofit on a Million Program House. passipedia.org/planning/refurbishment_with_passive_house_components/'),
    ('15', 'IVL Swedish Environmental Research Institute. The Million Programme is Being Updated. Press release, 2019. ivl.se/english/ivl/press/press-releases/2019-02-26'),
    ('16', 'Lstiburek J. BSD-163: Controlling Cold-Weather Condensation Using Insulation. Building Science Corp. buildingscience.com/documents/digests/bsd-controlling-cold-weather-condensation-using-insulation'),
    ('17', 'Nusser B., Teibinger M. Assessment of Interior Insulation on Exterior Walls in Three Swedish Buildings. Science and Technology for the Built Environment, 2024. doi:10.1080/23744731.2024.2366135'),
    ('18', 'Gomes A.P. et al. Dataset of Criteria on Thermal Insulation Solutions in Building Facades in Norway, Portugal and Italy. Data in Brief, PMC, 2023. pmc.ncbi.nlm.nih.gov/articles/PMC10558703/'),
]

for num, text in refs:
    add_ref(num, text)

# ── SAVE ─────────────────────────────────────────────────────
out_path = 'D:/Masters/RQ5/Codex_Chatgpt_drying/Thermal_Mass_Retrofit_Background.docx'
doc.save(out_path)
print('Saved:', out_path)
