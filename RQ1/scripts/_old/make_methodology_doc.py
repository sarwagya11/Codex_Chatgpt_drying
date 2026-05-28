"""Generate Word document: Methodology Chapter for Masters Thesis.

Creates a comprehensive methodology section covering:
1. System description and configurations
2. Mathematical models (psychrometrics, solar, HP, kinetics, chamber)
3. Drying kinetics model
4. Simulation procedure
5. Weather data and locations
6. Performance metrics

Run:
    python scripts/make_methodology_doc.py
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).resolve().parents[1]


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def set_doc_margins(doc, top=2.54, bottom=2.54, left=3.0, right=2.54):
    """Set document margins in cm."""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, bold=False, italic=False, size=11, alignment=None,
             space_after=6, first_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    run.italic = italic
    if alignment:
        p.paragraph_format.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    return p


def add_equation(doc, text, label=None):
    """Add an equation (centered, monospace). Optional label right-aligned."""
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)
    if label:
        run2 = p.add_run(f'    ({label})')
        run2.font.size = Pt(10)
        run2.font.name = 'Times New Roman'
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.name = 'Times New Roman'
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        r2.font.name = 'Times New Roman'
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
    return p


def add_table(doc, headers, rows, caption=None, col_widths=None):
    """Add a formatted table with optional caption."""
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Times New Roman'

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacer
    return table


# ===========================================================================
# MAIN DOCUMENT
# ===========================================================================

def main():
    doc = Document()
    set_doc_margins(doc)

    # --- Title ---
    title = doc.add_heading('CHAPTER 3: METHODOLOGY', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    # ======================================================================
    # 3.1 SYSTEM DESCRIPTION
    # ======================================================================
    add_heading(doc, '3.1 System Description', level=1)

    add_para(doc,
        'This chapter presents the mathematical models and simulation methodology '
        'developed for the Solar-Assisted Heat Pump Dryer (SAHPD) system. The SAHPD '
        'is designed for drying apple slices in the climatic conditions of Nepal, with '
        'the objective of minimizing electrical energy consumption while maintaining '
        'consistent product quality. The system integrates three thermal energy sources: '
        'a vapor compression heat pump (R134a refrigerant), flat-plate solar air collectors, '
        'and a sensible heat recovery exchanger (HRX). Ten distinct system configurations '
        'are evaluated, each representing a different integration strategy for these '
        'components.')

    add_para(doc,
        'The dryer operates as a through-flow system with 10 trays arranged in series. '
        'Conditioned air enters at the top tray (Tray 1) and exits from the bottom tray '
        '(Tray 10), progressively picking up moisture from the product. The drying chamber '
        'geometry is determined by the product loading and target air velocity. The simulation '
        'advances in discrete timesteps of 60 seconds, updating air conditions, product moisture, '
        'and energy flows at each step until all trays reach the target moisture content.')

    # --- 3.1.1 Operating parameters ---
    add_heading(doc, '3.1.1 Design Parameters', level=2)

    add_table(doc,
        ['Parameter', 'Symbol', 'Value', 'Unit'],
        [
            ['Target drying temperature', 'T_set', '45', '\u00b0C'],
            ['Air velocity (through trays)', 'v', '1.1', 'm/s'],
            ['Dry mass of product', 'm_p,dry', '3.0', 'kg'],
            ['Initial moisture content (d.b.)', 'X_0', '6.5', 'kg/kg'],
            ['Final moisture content (d.b.)', 'X_f', '0.10', 'kg/kg'],
            ['Product (apple slices)', 'thickness', '6', 'mm'],
            ['Number of trays', 'N', '10', '-'],
            ['Loading density', '\u03c1_load', '3.6', 'kg_fresh/m\u00b2'],
            ['Timestep', '\u0394t', '60', 's'],
            ['Maximum simulation time', 't_max', '72', 'h'],
            ['Refrigerant', '-', 'R134a', '-'],
        ],
        caption='Table 3.1: Design parameters for the SAHPD system')

    # --- 3.1.2 Chamber geometry ---
    add_heading(doc, '3.1.2 Drying Chamber Geometry', level=2)

    add_para(doc,
        'The drying chamber dimensions and air mass flow rate are derived from the product '
        'specifications rather than being specified independently. This ensures physical '
        'consistency between the product loading, air velocity, and mass flow rate.')

    add_para(doc, 'The fresh mass of product loaded per tray is:')
    add_equation(doc, 'm_fresh = m_p,dry \u00d7 (1 + X_0) / N', '3.1')

    add_para(doc, 'The tray area is determined from the loading density:')
    add_equation(doc, 'A_tray = m_fresh,tray / \u03c1_load', '3.2')

    add_para(doc,
        'For the through-flow configuration, air flows horizontally through the gap '
        'between trays. The cross-sectional area for air flow is:')
    add_equation(doc, 'A_cross = \u221aA_tray \u00d7 \u03b4_gap', '3.3')

    add_para(doc, 'where \u03b4_gap = 0.12 m is the vertical gap between trays. '
             'The dry air mass flow rate is then:')
    add_equation(doc, '\u1e41_da = \u03c1_air \u00d7 v \u00d7 A_cross', '3.4')

    add_para(doc,
        'The air density \u03c1_air is computed from the ideal gas law at the drying temperature '
        'and local atmospheric pressure:')
    add_equation(doc, '\u03c1_air = P_atm / (R_da \u00d7 T_set,K)', '3.5')

    add_para(doc,
        'where P_atm is obtained from the standard atmosphere model based on site elevation:')
    add_equation(doc, 'P_atm = 101325 \u00d7 (1 \u2212 0.0065 \u00d7 z / 288.15)^5.2561', '3.6')

    # ======================================================================
    # 3.2 SYSTEM CONFIGURATIONS
    # ======================================================================
    add_heading(doc, '3.2 System Configurations', level=1)

    add_para(doc,
        'Ten configurations are investigated, grouped into four families based on their '
        'thermal energy integration strategy. All configurations share the same drying '
        'chamber, product loading, and kinetics model. They differ only in how the inlet '
        'air is conditioned before entering the chamber.')

    # --- Config A ---
    add_heading(doc, '3.2.1 Configuration A: Heat Pump Convective Dryer (Baseline)', level=2)

    add_para(doc,
        'Configuration A serves as the baseline heat pump convective dryer (HPCD). '
        'In the open-loop mode (r = 0), ambient air is drawn through the heat pump '
        'condenser, heated to the target temperature, passed through the drying chamber, '
        'and exhausted to the atmosphere. The heat pump evaporator draws heat from a '
        'separate ambient air stream.')

    add_para(doc, 'Air flow path (open-loop):')
    add_equation(doc, 'Ambient \u2192 HP Condenser (\u219245\u00b0C) \u2192 Chamber \u2192 Exhaust')

    add_para(doc,
        'In this mode, the evaporator and condenser air streams are independent. '
        'The heat pump is sized to deliver the required condenser heat duty:')
    add_equation(doc, 'Q_cond = \u1e41_da \u00d7 (h_out \u2212 h_in)', '3.7')

    add_para(doc,
        'where h_out and h_in are the moist air enthalpies at the condenser outlet '
        '(T_set) and inlet (T_amb), respectively.')

    # --- Config B ---
    add_heading(doc, '3.2.2 Configuration B: Solar + Heat Pump Series', level=2)

    add_para(doc,
        'Configuration B introduces a flat-plate solar air collector in series before the '
        'heat pump condenser. The solar collector preheats ambient air, reducing the thermal '
        'lift required from the heat pump. During periods of high solar irradiance, the '
        'collector may heat the air to or above the target temperature, allowing the '
        'compressor to operate at reduced capacity or shut off entirely.')

    add_para(doc, 'Air flow path:')
    add_equation(doc, 'Ambient \u2192 Solar Collector \u2192 HP Condenser (\u219245\u00b0C) \u2192 Chamber \u2192 Exhaust')

    add_para(doc, 'The condenser heat duty is reduced to:')
    add_equation(doc, 'Q_cond = \u1e41_da \u00d7 c_p \u00d7 max(T_set \u2212 T_solar,out, 0)', '3.8')

    # --- Config C ---
    add_heading(doc, '3.2.3 Configurations C1 and C2: Solar-Assisted Evaporator', level=2)

    add_para(doc,
        'Configurations C1 and C2 direct solar-heated air to the heat pump evaporator '
        'rather than the condenser. By raising the evaporator source temperature, these '
        'configurations increase the heat pump COP, reducing compressor work for the same '
        'condenser heat delivery. C1 and C2 differ in the mixing point of recirculated '
        'and ambient air relative to the solar collector:')

    add_bullet(doc, ' C1: ambient air is mixed with recirculated air before the solar collector.',
               bold_prefix='Config')
    add_bullet(doc, ' C2: ambient air passes through the solar collector first, then mixes with '
               'recirculated air before the evaporator.', bold_prefix='Config')

    # --- Config D ---
    add_heading(doc, '3.2.4 Configurations D1, D2, and D3: Heat Recovery Exchanger + HP', level=2)

    add_para(doc,
        'The D-family configurations introduce a sensible heat recovery exchanger (HRX) '
        'between the warm, humid chamber exhaust and the cold, dry incoming ambient air. '
        'The HRX pre-warms the inlet air using waste heat from the exhaust, reducing the '
        'thermal load on the heat pump condenser. All D configurations operate in open-loop '
        'mode (r = 0) with no solar collector.')

    add_bullet(doc,
        ' D1: Ambient air is heated by HRX, then by the HP condenser. Evaporator uses '
        'separate ambient air.', bold_prefix='Config')
    add_bullet(doc,
        ' D2: Same as D1, but the evaporator air stream includes a dynamic ambient '
        'supplement derived from the exhaust composition.', bold_prefix='Config')
    add_bullet(doc,
        ' D3: Air streams are swapped \u2014 exhaust passes through the HRX hot side then '
        'to the condenser; ambient passes through HRX cold side then to the evaporator.',
        bold_prefix='Config')

    add_para(doc, 'Air flow path (D1):')
    add_equation(doc, 'Ambient \u2192 HRX(cold) \u2192 HP Condenser \u2192 Chamber \u2192 HRX(hot) \u2192 Exhaust')

    # --- Config E ---
    add_heading(doc, '3.2.5 Configurations E1, E2, and E3: HRX + Solar + HP', level=2)

    add_para(doc,
        'The E-family configurations combine all three thermal recovery strategies: '
        'heat recovery, solar preheating, and heat pump conditioning. These represent '
        'the most thermally integrated designs.')

    add_bullet(doc,
        ' E1: Ambient \u2192 HRX \u2192 Solar \u2192 HP Condenser \u2192 Chamber. '
        'Evaporator uses separate ambient air.', bold_prefix='Config')
    add_bullet(doc,
        ' E2: Same condenser path as E1, but the evaporator air stream is supplemented '
        'with ambient air via an iterative sizing procedure.', bold_prefix='Config')
    add_bullet(doc,
        ' E3: Solar collector heats the evaporator stream instead of the condenser stream. '
        'Ambient \u2192 HRX \u2192 HP Condenser \u2192 Chamber; Exhaust+Ambient \u2192 Solar \u2192 Evaporator.',
        bold_prefix='Config')

    # Configuration summary table
    add_table(doc,
        ['Config', 'Solar', 'HRX', 'HP', 'Key Mechanism'],
        [
            ['A', 'No', 'No', 'Yes', 'HP-only baseline (open-loop)'],
            ['B', 'Yes', 'No', 'Yes', 'Solar preheats condenser inlet air'],
            ['C1', 'Yes', 'No', 'Yes', 'Solar preheats evaporator air (mix before solar)'],
            ['C2', 'Yes', 'No', 'Yes', 'Solar preheats evaporator air (mix after solar)'],
            ['D1', 'No', 'Yes', 'Yes', 'HRX pre-warms condenser inlet; evap: ambient'],
            ['D2', 'No', 'Yes', 'Yes', 'HRX pre-warms condenser inlet; evap: dynamic mix'],
            ['D3', 'No', 'Yes', 'Yes', 'HRX streams swapped (exhaust to condenser)'],
            ['E1', 'Yes', 'Yes', 'Yes', 'HRX + Solar on condenser stream; evap: ambient'],
            ['E2', 'Yes', 'Yes', 'Yes', 'HRX + Solar on condenser stream; evap: dynamic mix'],
            ['E3', 'Yes', 'Yes', 'Yes', 'HRX + Solar on evaporator stream'],
        ],
        caption='Table 3.2: Summary of SAHPD configurations')

    # ======================================================================
    # 3.3 MATHEMATICAL MODELS
    # ======================================================================
    add_heading(doc, '3.3 Mathematical Models', level=1)

    # --- 3.3.1 Psychrometrics ---
    add_heading(doc, '3.3.1 Psychrometric Model', level=2)

    add_para(doc,
        'Air properties at every point in the system are described by three interdependent '
        'state variables: dry-bulb temperature T (\u00b0C), humidity ratio \u03c9 (kg water per '
        'kg dry air), and relative humidity RH. All psychrometric calculations follow '
        'ASHRAE standards (ASHRAE, 2021).')

    add_para(doc, 'Saturation vapor pressure (Tetens correlation, valid \u221240 to +50\u00b0C):')
    add_equation(doc, 'p_sat(T) = 610.94 \u00d7 exp(17.625 \u00d7 T / (T + 243.04))  [Pa]', '3.9')

    add_para(doc, 'Humidity ratio from temperature and relative humidity:')
    add_equation(doc, '\u03c9 = 0.62198 \u00d7 RH \u00d7 p_sat(T) / (P_atm \u2212 RH \u00d7 p_sat(T))', '3.10')

    add_para(doc, 'Moist air specific enthalpy (per kg dry air):')
    add_equation(doc, 'h = c_p,da \u00d7 T + \u03c9 \u00d7 (h_fg,0 + c_p,v \u00d7 T)  [kJ/kg_da]', '3.11')

    add_para(doc,
        'where c_p,da = 1.006 kJ/(kg\u00b7K) is the specific heat of dry air, '
        'c_p,v = 1.86 kJ/(kg\u00b7K) is the specific heat of water vapor, and '
        'h_fg,0 = 2501 kJ/kg is the latent heat of vaporization at 0\u00b0C.')

    add_para(doc, 'Relative humidity from temperature and humidity ratio (inverse calculation):')
    add_equation(doc, 'RH = (\u03c9 \u00d7 P_atm) / ((0.62198 + \u03c9) \u00d7 p_sat(T))', '3.12')

    # --- 3.3.2 Solar collector ---
    add_heading(doc, '3.3.2 Solar Collector Model', level=2)

    add_para(doc,
        'The flat-plate solar air collector is modeled using the Hottel\u2013Whillier\u2013Bliss '
        'equation (Duffie and Beckman, 2013). The useful heat gain is:')
    add_equation(doc,
        'Q_u = A_c \u00d7 F_R \u00d7 [\u03b7_opt \u00d7 G \u2212 U_L \u00d7 (T_in \u2212 T_amb)]', '3.13')

    add_para(doc, 'where the heat removal factor F_R is:')
    add_equation(doc,
        "F_R = (\u1e41_da \u00d7 c_p) / (A_c \u00d7 U_L) \u00d7 [1 \u2212 exp(\u2212F' \u00d7 U_L \u00d7 A_c / (\u1e41_da \u00d7 c_p))]",
        '3.14')

    add_para(doc, 'The collector outlet temperature is:')
    add_equation(doc, 'T_out = T_in + Q_u / (\u1e41_da \u00d7 c_p)', '3.15')

    add_table(doc,
        ['Parameter', 'Symbol', 'Value', 'Unit'],
        [
            ['Optical efficiency', '\u03b7_opt', '0.75', '-'],
            ['Overall loss coefficient', 'U_L', '5.0', 'W/(m\u00b2\u00b7K)'],
            ['Collector efficiency factor', "F'", '0.90', '-'],
            ['Collector area', 'A_c', '0\u201320', 'm\u00b2'],
            ['Incidence angle modifier', 'K_\u03b8', '1.0', '-'],
        ],
        caption='Table 3.3: Solar collector parameters')

    # --- 3.3.3 Heat pump ---
    add_heading(doc, '3.3.3 Heat Pump Thermodynamic Cycle', level=2)

    add_para(doc,
        'The heat pump operates on a standard vapor compression cycle with R134a refrigerant. '
        'Thermodynamic properties at each state point are computed using the CoolProp library '
        '(Bell et al., 2014), which provides NIST-quality equations of state. The cycle '
        'consists of four processes:')

    add_para(doc,
        '(1) Evaporation with superheat: the refrigerant absorbs heat from the air stream '
        '(or ambient environment) at low pressure, exiting as superheated vapor at '
        'T_1 = T_evap + \u0394T_sh.')
    add_para(doc,
        '(2) Compression: the superheated vapor is compressed to condenser pressure. '
        'The actual enthalpy rise is related to the isentropic enthalpy rise by the '
        'compressor isentropic efficiency:')
    add_equation(doc, 'h_2 = h_1 + (h_2s \u2212 h_1) / \u03b7_is', '3.16')

    add_para(doc,
        '(3) Condensation with subcooling: the high-pressure refrigerant rejects heat to '
        'the drying air, exiting as subcooled liquid at T_3 = T_cond \u2212 \u0394T_sc.')
    add_para(doc,
        '(4) Expansion: isenthalpic throttling through the expansion valve (h_4 = h_3).')

    add_para(doc, 'The refrigerant mass flow rate is determined from the condenser heat duty:')
    add_equation(doc, '\u1e41_ref = Q_cond / (h_2 \u2212 h_3)', '3.17')

    add_para(doc, 'Energy flows and performance:')
    add_equation(doc, 'Q_evap = \u1e41_ref \u00d7 (h_1 \u2212 h_4)', '3.18')
    add_equation(doc, 'W_comp = \u1e41_ref \u00d7 (h_2 \u2212 h_1) / \u03b7_mech', '3.19')
    add_equation(doc, 'COP = Q_cond / W_comp', '3.20')

    add_para(doc, 'First law of thermodynamics (energy conservation):')
    add_equation(doc, 'Q_cond = Q_evap + W_comp \u00d7 \u03b7_mech', '3.21')

    add_table(doc,
        ['Parameter', 'Symbol', 'Value', 'Unit'],
        [
            ['Refrigerant', '-', 'R134a', '-'],
            ['Compressor isentropic efficiency', '\u03b7_is', '0.75', '-'],
            ['Mechanical efficiency', '\u03b7_mech', '0.95', '-'],
            ['Superheat', '\u0394T_sh', '5', 'K'],
            ['Subcooling', '\u0394T_sc', '5', 'K'],
            ['Evaporator effectiveness', '\u03b5_evap', '0.85', '-'],
            ['Condenser effectiveness', '\u03b5_cond', '0.85', '-'],
            ['Evaporator coil approach', '\u0394T_approach', '3', 'K'],
            ['Min. evaporator temperature', 'T_evap,min', '\u22125', '\u00b0C'],
            ['Max. condenser temperature', 'T_cond,max', '70', '\u00b0C'],
            ['Max. pressure ratio', 'PR_max', '10', '-'],
        ],
        caption='Table 3.4: Heat pump parameters')

    # --- HP sizing methods ---
    add_heading(doc, '3.3.3.1 Heat Pump Sizing: Open-Loop vs. Closed-Loop', level=3)

    add_para(doc,
        'Two distinct heat pump sizing approaches are used depending on whether the '
        'evaporator and condenser process the same or different air streams.')

    add_para(doc, 'Open-loop sizing (Configs A(r=0), B, C1, C2, D, E):', bold=True)
    add_para(doc,
        'In open-loop configurations, the condenser and evaporator air streams are independent. '
        'The heat pump is sized to deliver the required air heating duty:')
    add_equation(doc, 'Q_cond = \u1e41_da \u00d7 (h(T_set, \u03c9_in) \u2212 h(T_in, \u03c9_in))', '3.22')
    add_para(doc,
        'The evaporator temperature is set based on the ambient (or source) air temperature '
        'with a fixed approach:')
    add_equation(doc, 'T_evap = T_source \u2212 \u0394T_approach,evap', '3.23')

    add_para(doc, 'Closed-loop sizing (Config A with r > 0):', bold=True)
    add_para(doc,
        'When exhaust air is recirculated, the same air stream passes through both the '
        'evaporator and condenser. The first law of thermodynamics (Eq. 3.21) creates a '
        'direct coupling: the condenser heat output is constrained by the evaporator heat '
        'absorption plus compressor work. The evaporator extracts heat from the warm, humid '
        'mixed air (recirculated exhaust + fresh ambient):')
    add_equation(doc, 'Q_evap,air = \u1e41_da \u00d7 (h_before,evap \u2212 h_after,evap)', '3.24')
    add_equation(doc, 'Q_cond,1st law = Q_evap,air \u00d7 COP / (COP \u2212 \u03b7_mech)', '3.25')
    add_para(doc,
        'The actual condenser output is the minimum of the first-law limit and the heating '
        'requirement. When Q_cond,1st law is less than the heating requirement, the chamber '
        'inlet temperature floats below T_set.')

    # --- Evaporator model ---
    add_heading(doc, '3.3.3.2 Evaporator Air-Side Model', level=3)

    add_para(doc,
        'The evaporator uses an effectiveness-based model equivalent to the contact factor '
        'method of Chou and Chua (1994). The air outlet temperature is:')
    add_equation(doc, 'T_air,out = T_air,in \u2212 \u03b5_evap \u00d7 (T_air,in \u2212 T_coil)', '3.26')

    add_para(doc, 'where the coil surface temperature is:')
    add_equation(doc, 'T_coil = T_evap,sat + \u0394T_approach', '3.27')

    add_para(doc,
        'Dehumidification occurs when the air is cooled below its dewpoint. In this case, '
        'the outlet air exits at saturation:')
    add_equation(doc, 'If T_air,out < T_dp:  \u03c9_out = \u03c9_sat(T_air,out)  (saturated)', '3.28')
    add_equation(doc, 'If T_air,out \u2265 T_dp:  \u03c9_out = \u03c9_in  (no condensation)', '3.29')

    # --- Condenser model ---
    add_heading(doc, '3.3.3.3 Condenser Air-Side Model', level=3)

    add_para(doc,
        'The condenser performs sensible heating only (no moisture change). The air outlet '
        'temperature is limited by the condenser effectiveness:')
    add_equation(doc, 'T_air,out = T_air,in + \u03b5_cond \u00d7 (T_cond,sat \u2212 T_air,in)', '3.30')

    add_para(doc,
        'where T_cond,sat = T_set + 10 K is the condenser saturation temperature. '
        'The air outlet temperature is clamped to T_set to prevent overheating. '
        'The humidity ratio is unchanged through the condenser (\u03c9_out = \u03c9_in).')

    # --- 3.3.4 Heat Recovery Exchanger ---
    add_heading(doc, '3.3.4 Heat Recovery Exchanger (HRX)', level=2)

    add_para(doc,
        'The HRX is modeled as a counter-flow sensible heat exchanger that transfers '
        'heat from the warm chamber exhaust (hot side) to the cold incoming ambient air '
        '(cold side). Only sensible heat is exchanged through the solid wall; no moisture '
        'transfer occurs on the cold side. This is preferred over an enthalpy recovery '
        'ventilator (ERV) because moisture transfer from the humid exhaust back to the '
        'inlet air would be counterproductive for drying (Fadhel et al., 2011).')

    add_para(doc, 'Outlet temperatures:')
    add_equation(doc, 'T_amb,heated = T_amb + \u03b5_HRX \u00d7 (T_exhaust \u2212 T_amb)', '3.31')
    add_equation(doc, 'T_exh,cooled = T_exhaust \u2212 \u03b5_HRX \u00d7 (T_exhaust \u2212 T_amb)', '3.32')

    add_para(doc,
        'If the exhaust stream is cooled below its dewpoint on the hot side, condensation '
        'occurs and the exhaust exits at saturation conditions. The HRX effectiveness is '
        'set to \u03b5_HRX = 0.70, consistent with values reported for plate-type heat '
        'exchangers in drying applications (Jouhara et al., 2017).')

    # --- 3.3.5 VPD bypass ---
    add_heading(doc, '3.3.5 VPD-Based Exhaust Bypass Strategy', level=2)

    add_para(doc,
        'For configurations D and E, a vapor pressure deficit (VPD) exhaust bypass strategy '
        'is implemented to improve energy efficiency during the late drying phase. When the '
        'product approaches equilibrium moisture, the exhaust air leaves the chamber '
        'relatively dry and warm. In this condition, the HRX transfers heat from exhaust '
        'air that has substantial drying potential remaining. The VPD bypass routes exhaust '
        'air directly to exhaust (bypassing the HRX) when drying potential is low, '
        'preventing unnecessary heat recovery from air that should be expelled.')

    add_para(doc, 'The VPD utilization fraction is defined as:')
    add_equation(doc, 'VPD_util = (\u03c9_exhaust \u2212 \u03c9_inlet) / \u03c9_inlet', '3.33')

    add_para(doc,
        'The bypass engages when VPD_util falls below a threshold (typically 0.05), with '
        'a 3\u00d7 hysteresis band and a 600-second dwell time to prevent oscillation.')

    # ======================================================================
    # 3.4 DRYING KINETICS
    # ======================================================================
    add_heading(doc, '3.4 Drying Kinetics Model', level=1)

    add_para(doc,
        'The drying kinetics are modeled using a parametric effective drying coefficient '
        'K_eff fitted from experimental data. The moisture ratio (MR) is defined as:')
    add_equation(doc, 'MR = (X \u2212 X_eq) / (X_0 \u2212 X_eq)', '3.34')

    add_para(doc,
        'where X is the current moisture content (kg water/kg dry solid), X_0 = 6.5 is '
        'the initial moisture content, and X_eq is the dynamic equilibrium moisture content '
        'from the GAB isotherm.')

    # GAB isotherm
    add_heading(doc, '3.4.1 GAB Sorption Isotherm', level=2)

    add_para(doc,
        'The equilibrium moisture content is computed using the temperature-dependent GAB '
        '(Guggenheim\u2013Anderson\u2013de Boer) model (Kaymak-Ertekin and Gedik, 2004):')
    add_equation(doc, 'X_eq = (X_m \u00d7 C \u00d7 K \u00d7 a_w) / ((1 \u2212 K\u00b7a_w) \u00d7 (1 \u2212 K\u00b7a_w + C\u00b7K\u00b7a_w))', '3.35')

    add_para(doc,
        'where a_w = RH (water activity equals relative humidity at equilibrium), and '
        'X_m, C, K are temperature-dependent parameters with Arrhenius-type dependence. '
        'Typical values at drying conditions (45\u00b0C, 10\u201340% RH) yield '
        'X_eq \u2248 0.04\u20130.08 kg/kg d.b.')

    # Parametric K_eff
    add_heading(doc, '3.4.2 Parametric Effective Drying Coefficient', level=2)

    add_para(doc,
        'The effective drying coefficient K_eff is modeled as a log-linear function of '
        'temperature, relative humidity, and air velocity, fitted by ordinary least squares '
        'regression to 13 experimental drying curves from Phase 2:')
    add_equation(doc, 'ln(K_eff) = b_0 + b_1\u00b7T + b_2\u00b7RH + b_3\u00b7v + b_4\u00b7(T\u00d7RH)', '3.36')

    add_table(doc,
        ['Coefficient', 'Value', 'Description'],
        [
            ['b_0', '\u221211.58', 'Intercept'],
            ['b_1', '0.0524', 'Temperature effect [\u00b0C\u207b\u00b9]'],
            ['b_2', '\u22121.876', 'Relative humidity effect'],
            ['b_3', '0.611', 'Air velocity effect [s/m]'],
            ['b_4', '0.0132', 'T\u00d7RH interaction'],
        ],
        caption='Table 3.5: Parametric K_eff model coefficients (R\u00b2 = 0.90)')

    add_para(doc,
        'This parametric model provides smooth, continuous K_eff values across the '
        'operating envelope, naturally handling interpolation between experimental points.')

    # Moisture update
    add_heading(doc, '3.4.3 Moisture Content Update', level=2)

    add_para(doc,
        'The moisture content of each tray is updated at each timestep using a first-order '
        'finite difference scheme:')
    add_equation(doc, 'X_{k+1} = X_k \u2212 K_eff(T, RH, v) \u00d7 (X_k \u2212 X_eq) \u00d7 \u0394t', '3.37')

    add_para(doc,
        'subject to the constraint X_{k+1} \u2265 X_eq. The actual water removal rate from '
        'each tray is the minimum of three physical limits:')

    add_bullet(doc, 'Kinetic limit: d\u1e41_w,kinetic = K_eff \u00d7 (X \u2212 X_eq) \u00d7 m_dry / \u0394t',
               bold_prefix='(1) ')
    add_bullet(doc, 'Air capacity limit: d\u1e41_w,air = \u1e41_da \u00d7 \u0394\u03c9_max  '
               '(prevents RH_out > 95%)', bold_prefix='(2) ')
    add_bullet(doc, 'Removable moisture: m_removable = (X \u2212 X_eq) \u00d7 m_dry',
               bold_prefix='(3) ')

    # ======================================================================
    # 3.5 MULTI-TRAY CHAMBER MODEL
    # ======================================================================
    add_heading(doc, '3.5 Multi-Tray Chamber Model', level=1)

    add_para(doc,
        'The drying chamber consists of 10 trays arranged in a series cascade. Air enters '
        'at Tray 1 (hottest, driest air) and exits from Tray 10 (coolest, most humid). '
        'At each timestep, air conditions are updated sequentially through the tray stack:')

    add_equation(doc,
        'Tray j: (T_j,in, \u03c9_j,in) \u2192 drying \u2192 (T_j,out, \u03c9_j,out) = (T_{j+1},in, \u03c9_{j+1},in)')

    add_para(doc, 'For each tray, the enthalpy balance accounts for the latent heat of evaporation:')
    add_equation(doc, 'h_out = h_in \u2212 (d\u1e41_w / \u1e41_da) \u00d7 h_fg', '3.38')
    add_equation(doc, '\u03c9_out = \u03c9_in + d\u1e41_w / (\u1e41_da \u00d7 \u0394t)', '3.39')

    add_para(doc,
        'where h_fg = 2400 kJ/kg is the latent heat of vaporization at drying conditions. '
        'The outlet temperature is recovered from the enthalpy and humidity ratio using '
        'the inverse psychrometric relation (Eq. 3.11). The simulation terminates when all '
        '10 trays satisfy X \u2264 X_f = 0.10 kg/kg d.b.')

    add_para(doc,
        'Due to the cascade arrangement, Tray 1 dries fastest (~4\u20135 hours) while '
        'Tray 10 dries slowest (~15\u201317 hours). The total drying time is determined by '
        'the slowest tray.')

    # ======================================================================
    # 3.6 FAN POWER
    # ======================================================================
    add_heading(doc, '3.6 Fan Power Model', level=1)

    add_para(doc,
        'Fan electrical consumption is computed from the air-side pressure drop through '
        'the drying chamber using the Darcy\u2013Weisbach equation for each tray channel:')
    add_equation(doc, '\u0394P = f \u00d7 (L/D_h) \u00d7 (\u03c1 \u00d7 v\u00b2 / 2) + K_minor \u00d7 (\u03c1 \u00d7 v\u00b2 / 2)', '3.40')

    add_para(doc, 'Fan electrical power:')
    add_equation(doc, 'W_fan = (\u1e41_da \u00d7 \u0394P_total) / (\u03c1_air \u00d7 \u03b7_fan)', '3.41')

    add_para(doc,
        'where \u03b7_fan = 0.60 is the overall fan efficiency (motor + impeller). The fan power '
        'is typically 2\u20134% of the total electrical consumption.')

    # ======================================================================
    # 3.7 WEATHER DATA & LOCATIONS
    # ======================================================================
    add_heading(doc, '3.7 Weather Data and Study Locations', level=1)

    add_para(doc,
        'Hourly weather data for each location is obtained from the PVGIS Typical '
        'Meteorological Year (TMY) database (European Commission Joint Research Centre), '
        'which provides representative hourly profiles based on ERA5 reanalysis data. '
        'Each TMY file contains 8,760 hourly records of ambient temperature, relative '
        'humidity, global horizontal irradiance (GHI), wind speed, and atmospheric pressure.')

    add_table(doc,
        ['Location', 'Elevation (m)', 'P_atm (Pa)', 'Climate Zone', 'Mean T (\u00b0C)', 'Mean RH (%)'],
        [
            ['Kathmandu', '1,350', '86,120', 'Subtropical highland', '18.4', '65'],
            ['Biratnagar', '72', '100,460', 'Subtropical lowland', '24.5', '78'],
            ['Taplejung', '1,820', '81,600', 'Temperate highland', '14.5', '72'],
        ],
        caption='Table 3.6: Study locations and climatic characteristics')

    add_heading(doc, '3.7.1 Seasonal Analysis', level=2)

    add_para(doc,
        'To assess system performance under varying climatic conditions, simulations are '
        'conducted for three representative seasons aligned with the Nepalese agricultural '
        'calendar:')

    add_bullet(doc, ' (October\u2013November): Post-monsoon apple harvest and primary drying '
               'window. Moderate temperatures, declining humidity.', bold_prefix='Sharad/Autumn')
    add_bullet(doc, ' (December\u2013January): Cold, dry conditions with low solar irradiance. '
               'Stress test for heat pump performance.', bold_prefix='Hemanta/Winter')
    add_bullet(doc, ' (March\u2013April): Warm, dry, high solar irradiance. Best-case scenario '
               'for solar-assisted configurations.', bold_prefix='Basanta/Spring')

    add_para(doc,
        'Seasonal weather files are extracted from the annual TMY data and re-indexed '
        'from hour 0 for simulation compatibility.')

    # ======================================================================
    # 3.8 PERFORMANCE METRICS
    # ======================================================================
    add_heading(doc, '3.8 Performance Metrics', level=1)

    add_para(doc, 'The following metrics are used to compare configurations:')

    add_para(doc, 'Specific energy consumption (electrical):', bold=True)
    add_equation(doc, 'SEC = (W_comp + W_fan) / m_w,removed  [kWh/kg]', '3.42')

    add_para(doc, 'Coefficient of performance:', bold=True)
    add_equation(doc, 'COP = Q_cond / W_comp', '3.43')

    add_para(doc, 'Solar fraction:', bold=True)
    add_equation(doc, 'SF = Q_solar / (Q_solar + Q_cond)  [\u2212]', '3.44')

    add_para(doc, 'Energy savings relative to baseline:', bold=True)
    add_equation(doc, 'ES = (SEC_A \u2212 SEC_x) / SEC_A \u00d7 100  [%]', '3.45')

    add_para(doc, 'Second-law (Carnot) efficiency:', bold=True)
    add_equation(doc, '\u03b7_II = COP / COP_Carnot', '3.46')
    add_equation(doc, 'COP_Carnot = T_cond,K / (T_cond,K \u2212 T_evap,K)', '3.47')

    # ======================================================================
    # 3.9 MODEL VALIDATION
    # ======================================================================
    add_heading(doc, '3.9 Model Validation', level=1)

    add_para(doc,
        'The simulation model is validated through four independent consistency checks '
        'applied to all configurations and locations:')

    add_bullet(doc, ' The first law of thermodynamics is verified at every timestep: '
               'Q_cond = Q_evap + W_comp \u00d7 \u03b7_mech. The maximum observed error across all '
               'configurations is less than 10\u207b\u2076.', bold_prefix='First law: ')
    add_bullet(doc, ' The cumulative water removed from all trays equals the sum of '
               'individual tray moisture decrements, with zero discrepancy.',
               bold_prefix='Water mass balance: ')
    add_bullet(doc, ' Temperature, humidity ratio, and relative humidity are verified '
               'for mutual consistency at every air state point (error < 4 \u00d7 10\u207b\u2076).',
               bold_prefix='Psychrometric consistency: ')
    add_bullet(doc, ' COP values range from 3.5 to 4.8, with Carnot efficiency '
               '\u03b7_II = 0.61\u20130.62, consistent with published values for R134a heat pump '
               'dryers (Chua et al., 2002).', bold_prefix='COP validation: ')

    # ======================================================================
    # 3.10 SIMULATION PROCEDURE
    # ======================================================================
    add_heading(doc, '3.10 Simulation Procedure', level=1)

    add_para(doc,
        'The simulation follows a quasi-steady-state approach where each 60-second timestep '
        'is solved as a sequence of steady-state component calculations. The procedure for '
        'each timestep is:')

    add_para(doc, '1. Read ambient weather conditions (T_amb, RH_amb, GHI) from the TMY file.')
    add_para(doc, '2. Compute inlet air conditions through the thermal conditioning chain '
             '(HRX, solar collector, HP components) according to the configuration.')
    add_para(doc, '3. Pass conditioned air through the 10-tray cascade, updating moisture '
             'content and air state at each tray.')
    add_para(doc, '4. Compute energy flows (Q_cond, Q_evap, W_comp, Q_solar, W_fan) and '
             'cumulative totals.')
    add_para(doc, '5. Check termination criterion (all trays at X \u2264 0.10).')
    add_para(doc, '6. Advance to next timestep or terminate.')

    add_para(doc,
        'The simulation is implemented in Python 3.11 using NumPy for numerical operations, '
        'CoolProp for refrigerant thermodynamic properties, and Pandas for data management. '
        'Hourly weather data is interpolated linearly between timesteps. All simulations are '
        'deterministic and reproducible.')

    # --- Save ---
    output_path = PROJECT_ROOT / 'Methodology_Chapter.docx'
    doc.save(str(output_path))
    print(f"Document saved: {output_path}")
    print(f"  Size: {output_path.stat().st_size / 1024:.0f} KB")


if __name__ == '__main__':
    main()
