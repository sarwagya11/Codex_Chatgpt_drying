"""Generate Results & Discussion Section 1: Config A & B Analysis.

Creates a Word document covering:
  1.1  Config A — System Description and Air Paths
  1.2  Config A — Open-Loop Baseline Results (r=0)
  1.3  Config A — Effect of Recirculation (r=0 vs 0.9 vs 1.0)
  1.4  Config B — System Description and Air Path
  1.5  Config B — Open-Loop Results (r=0, Ac=10 m²)
  1.6  Config B — Effect of Recirculation (solar-evaporator interaction)
  1.7  Note on Fixed vs Adaptive Evaporator Temperature

Usage:
    python scripts/write_section1_config_ab.py
"""

import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_ROOT = PROJECT_ROOT / "outputs"
FIGURES_DIR = PROJECT_ROOT / "figures" / "thesis"

# ── helpers ────────────────────────────────────────────────────────────


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(size)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    return table


def add_figure_placeholder(doc, label, caption):
    """Insert a placeholder box for a diagram/figure."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[DIAGRAM: {label}]")
    run.bold = True
    run.font.size = Pt(12)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.italic = True
    return cap


def add_figure_ref(doc, fig_path, caption, width=Inches(5.5)):
    """Insert a figure if it exists, otherwise a placeholder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if fig_path.exists():
        p.add_run().add_picture(str(fig_path), width=width)
    else:
        run = p.add_run(f"[Figure: {fig_path.name}]")
        run.italic = True
        run.font.size = Pt(11)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.italic = True
    return cap


def fmt(val, decimals=3):
    if pd.isna(val):
        return "-"
    return f"{val:.{decimals}f}"


def pct_diff(base, new):
    """Percent change from base to new. Positive = new is lower (improvement)."""
    if pd.isna(base) or pd.isna(new) or base == 0:
        return "-"
    return f"{(base - new) / base * 100:.1f}%"


# ── data access ────────────────────────────────────────────────────────


def load_data():
    df = pd.read_csv(OUTPUT_ROOT / "master_summary.csv")
    df = df[~df["filename"].str.contains("_s4", na=False)]
    return df


def get_row(df, config, loc, season, r, vpd=False, area=0.0):
    mask = (
        (df["config"] == config)
        & (df["location"] == loc)
        & (df["season"] == season)
        & (df["r_recirc"] == r)
        & (df["vpd_bypass"] == vpd)
    )
    if area > 0:
        mask &= df["solar_area_m2"] == area
    rows = df[mask]
    return rows.iloc[0] if not rows.empty else None


LOC_NAMES = {
    "kathmandu": "Kathmandu (1 350 m a.s.l.)",
    "biratnagar": "Biratnagar (72 m a.s.l.)",
    "taplejung": "Taplejung (1 820 m a.s.l.)",
}
LOC_SHORT = {"kathmandu": "KTM", "biratnagar": "BTN", "taplejung": "TPJ"}
LOCS = ["kathmandu", "biratnagar", "taplejung"]
SEASONS = ["annual", "autumn_oct_nov", "winter_dec_jan", "spring_mar_apr"]
SEASON_NAMES = {
    "annual": "Annual TMY",
    "autumn_oct_nov": "Autumn (Oct\u2013Nov)",
    "winter_dec_jan": "Winter (Dec\u2013Jan)",
    "spring_mar_apr": "Spring (Mar\u2013Apr)",
}


# ── document sections ──────────────────────────────────────────────────


def write_section_1_1(doc, df):
    """1.1 Config A system description and air paths."""
    add_heading(doc, "1.1  Config A: Heat-Pump-Only Dryer \u2014 System Description", level=2)

    add_para(doc, (
        "Config A is the baseline heat-pump convective dryer (HPCD) with no solar "
        "assistance or heat recovery. It operates in one of two modes depending on "
        "the recirculation ratio r."
    ))

    # ── r = 0 air path ──
    add_para(doc, "Open-loop mode (r = 0)", bold=True)

    add_figure_placeholder(doc,
        "Config A open-loop air path",
        "Figure 1.1a: Config A open-loop (r = 0) schematic. "
        "Process air and evaporator air are separate streams.")

    add_para(doc, (
        "In open-loop mode, the process air stream and the evaporator heat-source "
        "stream are entirely separate:"
    ))

    add_bullet(doc,
        "Process air path: Ambient air enters the condenser, where it is heated "
        "from T_amb to T_set (45 \u00b0C), then passes through the drying chamber "
        "where it picks up moisture from the product, and is exhausted to the "
        "environment. The air is not dehumidified \u2014 it enters the chamber at "
        "ambient humidity ratio."
    )
    add_bullet(doc,
        "Evaporator heat source: The evaporator draws heat from a separate "
        "ambient air stream (or equivalently, from the outdoor environment). "
        "It acts as the low-temperature heat source for the vapour-compression "
        "cycle. The evaporator saturation temperature is set at T_evap = T_amb "
        "\u2212 10 K, meaning it adapts to local ambient conditions every timestep."
    )
    add_bullet(doc,
        "Heat pump role: The compressor transfers heat from the evaporator "
        "(at T_amb \u2212 10 \u00b0C) to the condenser (at T_set + 10 = 55 \u00b0C). "
        "The temperature lift is therefore (T_set + 10) \u2212 (T_amb \u2212 10) = "
        "65 \u2212 T_amb [\u00b0C]. Warmer ambient air means a smaller lift and higher COP."
    )

    # ── r > 0 air path ──
    add_para(doc, "Closed-loop mode (r > 0)", bold=True)

    add_figure_placeholder(doc,
        "Config A closed-loop air path",
        "Figure 1.1b: Config A closed-loop (r > 0) schematic. "
        "Recirculated exhaust is mixed with fresh air, dehumidified at the "
        "evaporator, then reheated at the condenser.")

    add_para(doc, (
        "In closed-loop mode, a fraction r of the chamber exhaust is recirculated "
        "and mixed with fresh ambient air before entering the evaporator:"
    ))

    add_bullet(doc,
        "Mixing: Exhaust air (warm, humid from drying) is mixed with fresh "
        "ambient air at ratio r. The mixed-air state is computed by enthalpy "
        "and humidity-ratio weighted averaging: \u03c9_mix = r\u00b7\u03c9_exh + "
        "(1\u2212r)\u00b7\u03c9_amb, h_mix = r\u00b7h_exh + (1\u2212r)\u00b7h_amb."
    )
    add_bullet(doc,
        "Evaporator (dehumidification): The mixed air passes over the evaporator "
        "coil. The coil surface temperature is T_evap_sat + 3 K (approach "
        "temperature), where T_evap_sat = 5 \u00b0C is a fixed design parameter "
        "mimicking commercial refrigeration dehumidifiers. Air is cooled by "
        "the evaporator effectiveness (\u03b5_evap = 0.85): T_after_evap = T_mix "
        "\u2212 \u03b5\u00b7(T_mix \u2212 T_coil). If T_after_evap falls below the "
        "mixed air\u2019s dew point, moisture condenses out and the air exits "
        "saturated at T_after_evap."
    )
    add_bullet(doc,
        "Condenser (reheating): The cold, dehumidified air is heated from "
        "T_after_evap to T_set (45 \u00b0C). The first law is enforced: "
        "Q_cond = Q_evap + W_comp, ensuring thermodynamic consistency."
    )
    add_bullet(doc,
        "Chamber and exhaust: Air at 45 \u00b0C and reduced humidity passes through "
        "the drying chamber. The fraction r of the exhaust is recirculated; "
        "the remainder (1\u2212r) is expelled."
    )

    add_para(doc, (
        "A key difference between the two modes is the evaporator operating "
        "point. In open-loop, T_evap adapts to ambient conditions (T_amb \u2212 10 K). "
        "In closed-loop, T_evap is fixed at 5 \u00b0C regardless of ambient conditions. "
        "This distinction drives the performance crossover discussed in Section 1.3."
    ))


def write_section_1_2(doc, df):
    """1.2 Config A open-loop baseline results."""
    add_heading(doc, "1.2  Config A Open-Loop Baseline (r = 0)", level=2)

    add_para(doc, (
        "Table 1.1 presents the annual-TMY performance of Config A in open-loop "
        "mode across three Nepalese locations. The specific energy consumption "
        "(SEC) is defined as total electrical input (compressor + fan) per "
        "kilogram of water removed."
    ))

    # Table 1.1
    headers = ["Location", "Elev. (m)", "T_amb mean (\u00b0C)", "SEC (kWh/kg)",
               "Time (h)", "COP mean", "W_comp (kWh)"]
    t_amb_means = {"kathmandu": 9.8, "biratnagar": 18.8, "taplejung": 12.0}  # from simulation data
    rows = []
    for loc in LOCS:
        row = get_row(df, "A", loc, "annual", 0.0)
        if row is not None:
            elev = {"kathmandu": "1 350", "biratnagar": "72", "taplejung": "1 820"}[loc]
            rows.append([
                LOC_SHORT[loc], elev, f"{t_amb_means[loc]:.1f}",
                fmt(row["SEC_kWh_per_kg"]),
                fmt(row["time_h"], 1),
                fmt(row["COP_mean"], 2),
                fmt(row["W_comp_kWh"], 1),
            ])
    add_table(doc, headers, rows)
    add_para(doc, "Table 1.1: Config A open-loop (r = 0) annual-TMY performance.",
             italic=True, size=9)

    a_ktm = get_row(df, "A", "kathmandu", "annual", 0.0)
    a_btn = get_row(df, "A", "biratnagar", "annual", 0.0)
    a_tpj = get_row(df, "A", "taplejung", "annual", 0.0)

    add_para(doc, (
        f"SEC ranges from {fmt(a_btn['SEC_kWh_per_kg'])} kWh/kg at Biratnagar "
        f"to {fmt(a_ktm['SEC_kWh_per_kg'])} kWh/kg at Kathmandu, a difference "
        f"of {pct_diff(a_ktm['SEC_kWh_per_kg'], a_btn['SEC_kWh_per_kg'])}. "
        f"The performance gap is explained by two factors:"
    ))

    add_bullet(doc,
        f"Temperature lift: At Biratnagar (mean T_amb \u2248 19 \u00b0C), the heat pump "
        f"lifts from T_evap \u2248 9 \u00b0C to T_cond = 55 \u00b0C, a 46 K lift. "
        f"At Kathmandu (mean T_amb \u2248 10 \u00b0C), T_evap \u2248 0 \u00b0C and the lift "
        f"is 55 K. The smaller lift at Biratnagar yields a COP of "
        f"{fmt(a_btn['COP_mean'], 2)} versus {fmt(a_ktm['COP_mean'], 2)} at Kathmandu."
    )
    add_bullet(doc,
        f"Atmospheric pressure: Biratnagar at 72 m operates at \u2248100.5 kPa "
        f"while Kathmandu at 1 350 m operates at \u224886.1 kPa. The lower pressure "
        f"at altitude reduces air density, which affects both heat transfer and "
        f"the psychrometric state of the drying air."
    )

    add_para(doc, (
        f"Drying time is nearly identical across all three locations "
        f"({fmt(a_ktm['time_h'], 1)}\u2013{fmt(a_btn['time_h'], 1)} h) because "
        f"the chamber inlet temperature is the same (T_set = 45 \u00b0C) and the "
        f"drying kinetics are governed by temperature and product moisture "
        f"content, not by the energy source. The difference between locations "
        f"is purely in how much electricity the compressor requires."
    ))

    # Seasonal table
    add_para(doc, (
        "Table 1.2 shows the seasonal variation of Config A (r = 0) performance."
    ))

    headers_s = ["Location", "Annual", "Autumn", "Winter", "Spring"]
    rows_s = []
    for loc in LOCS:
        row_data = [LOC_SHORT[loc]]
        for season in SEASONS:
            d = get_row(df, "A", loc, season, 0.0)
            row_data.append(fmt(d["SEC_kWh_per_kg"]) if d is not None else "-")
        rows_s.append(row_data)
    add_table(doc, headers_s, rows_s)
    add_para(doc, "Table 1.2: Config A (r = 0) SEC (kWh/kg) by season.",
             italic=True, size=9)

    a_btn_spr = get_row(df, "A", "biratnagar", "spring_mar_apr", 0.0)
    a_btn_win = get_row(df, "A", "biratnagar", "winter_dec_jan", 0.0)
    a_ktm_aut = get_row(df, "A", "kathmandu", "autumn_oct_nov", 0.0)
    a_ktm_win = get_row(df, "A", "kathmandu", "winter_dec_jan", 0.0)

    add_para(doc, (
        f"Seasonal variation is substantial. At Biratnagar, spring SEC "
        f"({fmt(a_btn_spr['SEC_kWh_per_kg'])}) is "
        f"{pct_diff(a_btn_win['SEC_kWh_per_kg'], a_btn_spr['SEC_kWh_per_kg'])} "
        f"lower than winter ({fmt(a_btn_win['SEC_kWh_per_kg'])}), driven by the "
        f"higher spring ambient temperature which reduces the compressor\u2019s "
        f"temperature lift. At Kathmandu, the best performance is in autumn "
        f"({fmt(a_ktm_aut['SEC_kWh_per_kg'])}), when post-monsoon warmth "
        f"persists, while winter is the most demanding season "
        f"({fmt(a_ktm_win['SEC_kWh_per_kg'])})."
    ))


def write_section_1_3(doc, df):
    """1.3 Effect of recirculation on Config A."""
    add_heading(doc, "1.3  Effect of Recirculation on Config A", level=2)

    add_para(doc, (
        "Table 1.3 compares Config A at r = 0 (open-loop), r = 0.9, and "
        "r = 1.0 (fully closed loop) for all three locations under annual TMY."
    ))

    headers = ["Location", "r", "SEC (kWh/kg)", "Time (h)", "COP mean",
               "T_evap mean (\u00b0C)", "W_comp (kWh)"]
    # T_evap values from simulation analysis
    t_evap_r0 = {"kathmandu": -0.2, "biratnagar": 8.8, "taplejung": 2.0}
    t_evap_r09 = {"kathmandu": 5.0, "biratnagar": 5.0, "taplejung": 5.0}
    rows = []
    for loc in LOCS:
        for r in [0.0, 0.9, 1.0]:
            row = get_row(df, "A", loc, "annual", r)
            if row is not None:
                t_evap = t_evap_r0[loc] if r == 0 else t_evap_r09[loc]
                rows.append([
                    LOC_SHORT[loc], f"{r:.1f}",
                    fmt(row["SEC_kWh_per_kg"]),
                    fmt(row["time_h"], 1),
                    fmt(row["COP_mean"], 2),
                    f"{t_evap:.1f}",
                    fmt(row["W_comp_kWh"], 1),
                ])
    add_table(doc, headers, rows)
    add_para(doc, "Table 1.3: Config A annual-TMY performance at r = 0, 0.9, and 1.0.",
             italic=True, size=9)

    add_figure_ref(
        doc,
        FIGURES_DIR / "fig5_config_a_recirc.png",
        "Figure 1.2: Config A SEC at r = 0, 0.9, and 1.0 across three locations and four seasons.",
    )

    # Crossover explanation
    add_para(doc, "The recirculation crossover", bold=True)

    add_para(doc, (
        "The results reveal a striking location-dependent crossover: "
        "recirculation improves performance at Kathmandu but worsens it at "
        "Biratnagar. The mechanism is rooted in how the evaporator temperature "
        "is determined in each mode."
    ))

    add_para(doc, (
        "In open-loop (r = 0), the evaporator is not in the process air path. "
        "It draws heat from ambient air, and its saturation temperature is set "
        "at T_evap = T_amb \u2212 10 K. This means T_evap adapts to local climate: "
        "at Biratnagar (T_amb \u2248 19 \u00b0C), T_evap \u2248 9 \u00b0C and COP = 4.43; "
        "at Kathmandu (T_amb \u2248 10 \u00b0C), T_evap \u2248 0 \u00b0C and COP = 3.63."
    ))

    add_para(doc, (
        "In closed-loop (r > 0), the evaporator dehumidifies the recirculated "
        "air and its saturation temperature is fixed at T_evap = 5 \u00b0C "
        "(a standard design value for refrigeration dehumidifiers). COP is "
        "locked at approximately 4.04 regardless of ambient conditions."
    ))

    add_para(doc, (
        "The crossover occurs at T_amb \u2248 15 \u00b0C, where the open-loop "
        "T_evap (T_amb \u2212 10 = 5 \u00b0C) equals the fixed closed-loop value. Below "
        "this threshold, recirculation raises T_evap (e.g. from 0 \u00b0C to 5 \u00b0C "
        "at Kathmandu) and improves COP. Above it, recirculation lowers T_evap "
        "(e.g. from 9 \u00b0C to 5 \u00b0C at Biratnagar) and degrades COP."
    ))

    add_para(doc, (
        "A second reinforcing effect is the condenser inlet temperature. At "
        "r = 0, air enters the condenser at T_amb. At r = 0.9, the evaporator "
        "cools the mixed air to T_after_evap \u2248 12 \u00b0C. At Biratnagar, this "
        "means the condenser receives air at 12 \u00b0C instead of 19 \u00b0C, requiring "
        "7 K of additional reheating. At Kathmandu, the condenser receives air "
        "at 12 \u00b0C instead of 10 \u00b0C \u2014 slightly warmer, so less reheating is "
        "needed. Both effects reinforce the crossover in the same direction."
    ))

    # Quantify
    a_ktm_r0 = get_row(df, "A", "kathmandu", "annual", 0.0)
    a_ktm_r09 = get_row(df, "A", "kathmandu", "annual", 0.9)
    a_btn_r0 = get_row(df, "A", "biratnagar", "annual", 0.0)
    a_btn_r09 = get_row(df, "A", "biratnagar", "annual", 0.9)
    a_btn_r0_spr = get_row(df, "A", "biratnagar", "spring_mar_apr", 0.0)
    a_btn_r09_spr = get_row(df, "A", "biratnagar", "spring_mar_apr", 0.9)

    add_para(doc, (
        f"Quantitatively, at Kathmandu, r = 0.9 reduces annual SEC by "
        f"{pct_diff(a_ktm_r0['SEC_kWh_per_kg'], a_ktm_r09['SEC_kWh_per_kg'])} "
        f"(from {fmt(a_ktm_r0['SEC_kWh_per_kg'])} to "
        f"{fmt(a_ktm_r09['SEC_kWh_per_kg'])} kWh/kg). At Biratnagar, r = 0.9 "
        f"increases SEC by "
        f"{pct_diff(a_btn_r09['SEC_kWh_per_kg'], a_btn_r0['SEC_kWh_per_kg'])} "
        f"(from {fmt(a_btn_r0['SEC_kWh_per_kg'])} to "
        f"{fmt(a_btn_r09['SEC_kWh_per_kg'])}). The penalty is most extreme in "
        f"spring at Biratnagar, where r = 0 achieves SEC = "
        f"{fmt(a_btn_r0_spr['SEC_kWh_per_kg'])} with COP = "
        f"{fmt(a_btn_r0_spr['COP_mean'], 2)}, while r = 0.9 yields "
        f"{fmt(a_btn_r09_spr['SEC_kWh_per_kg'])} with COP = "
        f"{fmt(a_btn_r09_spr['COP_mean'], 2)}."
    ))

    # r=0.9 vs r=1.0
    add_para(doc, "Diminishing returns at high recirculation", bold=True)

    add_para(doc, (
        "The difference between r = 0.9 and r = 1.0 is minimal (\u22641% SEC "
        "difference at all locations). At r = 1.0, no fresh air enters the "
        "system and all air is recirculated. The negligible difference from "
        "r = 0.9 indicates that the system\u2019s energy balance is dominated by "
        "the internal heat pump cycle rather than the small fraction of fresh "
        "air admitted at r = 0.9."
    ))


def write_section_1_4(doc, df):
    """1.4 Config B system description."""
    add_heading(doc, "1.4  Config B: Solar-Assisted HP (Series) \u2014 System Description", level=2)

    add_figure_placeholder(doc,
        "Config B air path",
        "Figure 1.3: Config B series air path. Ambient air passes through the "
        "evaporator (dehumidification, when r > 0), then through the solar "
        "collector, and finally through the condenser before entering the chamber.")

    add_para(doc, (
        "Config B adds a flat-plate solar collector (A_c = 10 m\u00b2) in series "
        "with the heat pump. In open-loop mode (r = 0), the air path is:"
    ))

    add_bullet(doc,
        "Ambient \u2192 Solar Collector \u2192 Condenser \u2192 Chamber \u2192 Exhaust"
    )

    add_para(doc, (
        "The solar collector pre-heats the ambient air using available solar "
        "irradiance (Hottel\u2013Whillier\u2013Bliss model with F_R(\u03c4\u03b1) = 0.68, "
        "F_R\u00b7U_L = 4.90 W/m\u00b2K). The condenser then provides the remaining "
        "temperature lift from the solar-heated air to T_set. When irradiance "
        "is high, the solar collector can heat air to near T_set, and the "
        "condenser load is minimal. At night or under cloud, the condenser "
        "provides all the heating \u2014 the system degrades gracefully to "
        "Config A behaviour."
    ))

    add_para(doc, (
        "In closed-loop mode (r > 0), the air path becomes:"
    ))

    add_bullet(doc,
        "Mix(r\u00d7exhaust + (1\u2212r)\u00d7ambient) \u2192 Evaporator \u2192 Solar Collector "
        "\u2192 Condenser \u2192 Chamber"
    )

    add_para(doc, (
        "The evaporator dehumidifies the recirculated air (as in Config A "
        "closed-loop), then the solar collector reheats it before the condenser "
        "provides the final lift. This ordering has implications for the "
        "solar\u2013evaporator interaction discussed in Section 1.6."
    ))


def write_section_1_5(doc, df):
    """1.5 Config B open-loop results."""
    add_heading(doc, "1.5  Config B Open-Loop Results (r = 0, A_c = 10 m\u00b2)", level=2)

    headers = ["Location", "SEC (kWh/kg)", "Time (h)", "COP",
               "Solar Frac.", "Q_solar (kWh)", "W_comp (kWh)",
               "SEC reduction vs A"]
    rows = []
    for loc in LOCS:
        b = get_row(df, "B", loc, "annual", 0.0, area=10.0)
        a = get_row(df, "A", loc, "annual", 0.0)
        if b is not None and a is not None:
            rows.append([
                LOC_SHORT[loc],
                fmt(b["SEC_kWh_per_kg"]),
                fmt(b["time_h"], 1),
                fmt(b["COP_mean"], 2),
                fmt(b["solar_fraction"], 2),
                fmt(b["Q_solar_kWh"], 1),
                fmt(b["W_comp_kWh"], 1),
                pct_diff(a["SEC_kWh_per_kg"], b["SEC_kWh_per_kg"]),
            ])
    add_table(doc, headers, rows)
    add_para(doc,
        "Table 1.4: Config B (r = 0, A_c = 10 m\u00b2) annual-TMY performance "
        "and SEC reduction relative to Config A (r = 0).",
        italic=True, size=9)

    b_ktm = get_row(df, "B", "kathmandu", "annual", 0.0, area=10.0)
    b_btn = get_row(df, "B", "biratnagar", "annual", 0.0, area=10.0)
    b_tpj = get_row(df, "B", "taplejung", "annual", 0.0, area=10.0)
    a_ktm = get_row(df, "A", "kathmandu", "annual", 0.0)
    a_btn = get_row(df, "A", "biratnagar", "annual", 0.0)
    a_tpj = get_row(df, "A", "taplejung", "annual", 0.0)

    add_para(doc, (
        f"Solar integration reduces SEC by "
        f"{pct_diff(a_ktm['SEC_kWh_per_kg'], b_ktm['SEC_kWh_per_kg'])} at Kathmandu, "
        f"{pct_diff(a_btn['SEC_kWh_per_kg'], b_btn['SEC_kWh_per_kg'])} at Biratnagar, "
        f"and {pct_diff(a_tpj['SEC_kWh_per_kg'], b_tpj['SEC_kWh_per_kg'])} at Taplejung "
        f"compared to Config A (r = 0). The solar fraction \u2014 the share of total "
        f"energy input provided by the solar collector \u2014 ranges from "
        f"{fmt(b_ktm['solar_fraction'], 2)} at Kathmandu to "
        f"{fmt(b_btn['solar_fraction'], 2)} at Biratnagar."
    ))

    add_para(doc, (
        f"Biratnagar achieves the highest solar fraction and the lowest SEC "
        f"({fmt(b_btn['SEC_kWh_per_kg'])} kWh/kg) for two reasons: (1) higher "
        f"annual irradiance at low elevation provides more solar gain "
        f"(Q_solar = {fmt(b_btn['Q_solar_kWh'], 1)} kWh vs "
        f"{fmt(b_ktm['Q_solar_kWh'], 1)} kWh at Kathmandu), and (2) the "
        f"already-warm ambient air requires less compressor work for the "
        f"remaining temperature lift."
    ))

    add_para(doc, (
        f"Drying time is unchanged between Config A and B "
        f"(\u224813.9\u201314.5 h), confirming that the benefit of solar integration "
        f"is purely in reduced electrical consumption, not faster drying. This "
        f"is expected because the chamber inlet temperature is maintained at "
        f"T_set = 45 \u00b0C in both configurations."
    ))

    # Seasonal table for Config B
    add_para(doc, "Seasonal variation", bold=True)

    headers_s = ["Location", "Annual", "Autumn", "Winter", "Spring"]
    rows_s = []
    for loc in LOCS:
        row_data = [LOC_SHORT[loc]]
        for season in SEASONS:
            d = get_row(df, "B", loc, season, 0.0, area=10.0)
            row_data.append(fmt(d["SEC_kWh_per_kg"]) if d is not None else "-")
        rows_s.append(row_data)
    add_table(doc, headers_s, rows_s)
    add_para(doc, "Table 1.5: Config B (r = 0, A_c = 10 m\u00b2) SEC (kWh/kg) by season.",
             italic=True, size=9)

    b_btn_spr = get_row(df, "B", "biratnagar", "spring_mar_apr", 0.0, area=10.0)
    b_btn_win = get_row(df, "B", "biratnagar", "winter_dec_jan", 0.0, area=10.0)

    add_para(doc, (
        f"Config B amplifies the seasonal advantage seen in Config A. At "
        f"Biratnagar in spring, SEC drops to "
        f"{fmt(b_btn_spr['SEC_kWh_per_kg'])} kWh/kg with a solar fraction of "
        f"{fmt(b_btn_spr['solar_fraction'], 2)} \u2014 the solar collector provides "
        f"93% of the heating energy, and the compressor consumes only "
        f"{fmt(b_btn_spr['W_comp_kWh'], 1)} kWh over the entire drying cycle. "
        f"In winter, SEC rises to {fmt(b_btn_win['SEC_kWh_per_kg'])} due to "
        f"reduced irradiance and cooler ambient temperatures."
    ))


def write_section_1_6(doc, df):
    """1.6 Effect of recirculation on Config B."""
    add_heading(doc, "1.6  Effect of Recirculation on Config B", level=2)

    add_para(doc, (
        "Table 1.6 compares Config B at r = 0, 0.9, and 1.0 for annual TMY."
    ))

    headers = ["Location", "r", "SEC (kWh/kg)", "Time (h)", "COP",
               "Solar Frac.", "W_comp (kWh)"]
    rows = []
    for loc in LOCS:
        for r in [0.0, 0.9, 1.0]:
            b = get_row(df, "B", loc, "annual", r, area=10.0)
            if b is not None:
                rows.append([
                    LOC_SHORT[loc], f"{r:.1f}",
                    fmt(b["SEC_kWh_per_kg"]),
                    fmt(b["time_h"], 1),
                    fmt(b["COP_mean"], 2),
                    fmt(b["solar_fraction"], 2),
                    fmt(b["W_comp_kWh"], 1),
                ])
    add_table(doc, headers, rows)
    add_para(doc,
        "Table 1.6: Config B (A_c = 10 m\u00b2) annual-TMY performance at r = 0, 0.9, and 1.0.",
        italic=True, size=9)

    b_btn_r0 = get_row(df, "B", "biratnagar", "annual", 0.0, area=10.0)
    b_btn_r09 = get_row(df, "B", "biratnagar", "annual", 0.9, area=10.0)
    b_ktm_r0 = get_row(df, "B", "kathmandu", "annual", 0.0, area=10.0)
    b_ktm_r09 = get_row(df, "B", "kathmandu", "annual", 0.9, area=10.0)

    add_para(doc, (
        f"Recirculation degrades Config B performance at all locations, "
        f"but the penalty is far more severe than in Config A. At Biratnagar, "
        f"r = 0.9 increases SEC by "
        f"{pct_diff(b_btn_r09['SEC_kWh_per_kg'], b_btn_r0['SEC_kWh_per_kg'])} "
        f"(from {fmt(b_btn_r0['SEC_kWh_per_kg'])} to "
        f"{fmt(b_btn_r09['SEC_kWh_per_kg'])} kWh/kg). Even at Kathmandu, "
        f"where Config A benefits from recirculation, Config B at r = 0.9 "
        f"shows only a marginal {pct_diff(b_ktm_r0['SEC_kWh_per_kg'], b_ktm_r09['SEC_kWh_per_kg'])} "
        f"improvement."
    ))

    add_para(doc, "Why the solar\u2013evaporator interaction penalises recirculation", bold=True)

    add_para(doc, (
        "In Config B with r > 0, the air path is: Mix \u2192 Evaporator \u2192 Solar "
        "\u2192 Condenser. The evaporator cools the mixed air from ~35 \u00b0C down to "
        "~12 \u00b0C for dehumidification. The solar collector then receives this "
        "cold air and must reheat it. While the collector operates at higher "
        "efficiency with colder inlet air (lower thermal losses), the useful "
        "temperature gain largely compensates for the cooling that the "
        "evaporator just performed. The net effect is:"
    ))

    add_bullet(doc,
        "The compressor does extra work to cool the air (evaporator dehumidification) "
        "that the solar collector then has to undo."
    )
    add_bullet(doc,
        "The solar fraction decreases because more compressor energy is consumed "
        f"(from {fmt(b_btn_r0['solar_fraction'], 2)} at r = 0 to "
        f"{fmt(b_btn_r09['solar_fraction'], 2)} at r = 0.9 at Biratnagar)."
    )
    add_bullet(doc,
        "The same fixed-T_evap crossover from Config A applies, but is "
        "compounded by the solar collector working against the evaporator."
    )

    add_para(doc, (
        "This finding has a practical implication: for solar-assisted HP dryers "
        "in series configuration, open-loop operation is strongly preferred. "
        "Recirculation should only be considered in very cold climates where "
        "the open-loop COP is low enough to offset the solar\u2013evaporator penalty."
    ))


def write_section_1_7(doc, df):
    """1.7 Fixed vs adaptive T_evap — investigation and negative result."""
    add_heading(doc, "1.7  Evaporator Temperature Strategy: Fixed vs Adaptive", level=2)

    add_para(doc, (
        "The performance crossover and the solar\u2013evaporator penalty both stem "
        "from the fixed evaporator temperature (T_evap = 5 \u00b0C) used in "
        "closed-loop operation. This value was chosen to represent standard "
        "refrigeration dehumidifier design, where the coil must be cold enough "
        "to condense moisture from air across a wide range of conditions."
    ))

    add_para(doc, (
        "However, analysis of the simulation time-series reveals that "
        "dehumidification effectiveness varies dramatically over the drying cycle:"
    ))

    # Dehumidification effectiveness table
    headers = ["Drying Phase", "\u03c9_mix (g/kg)", "T_dp_mix (\u00b0C)",
               "T_after_evap (\u00b0C)", "\u0394\u03c9 removed (g/kg)", "Effectiveness"]
    rows = [
        ["Early (1 h)", "15.6", "20.8", "10.4", "7.6", "Strong"],
        ["Mid (7 h)", "11.1", "15.5", "12.4", "2.1", "Moderate"],
        ["Late (13 h)", "9.6", "13.3", "13.1", "0.15", "Negligible"],
    ]
    add_table(doc, headers, rows)
    add_para(doc,
        "Table 1.7: Evaporator dehumidification effectiveness over drying cycle "
        "(Config A, r = 0.9, Biratnagar annual TMY).",
        italic=True, size=9)

    add_para(doc, (
        "In late-stage drying, the mixed air\u2019s dew point (13.3 \u00b0C) is nearly "
        "equal to the evaporator outlet temperature (13.1 \u00b0C). The evaporator "
        "is still cooling the air from ~35 \u00b0C to ~13 \u00b0C \u2014 consuming compressor "
        "energy \u2014 but removing almost no moisture. The condenser must then "
        "reheat the air from 13 \u00b0C back to 45 \u00b0C, representing ~22 K of wasted "
        "cooling and reheating."
    ))

    # --- Adaptive T_evap investigation ---
    add_para(doc, "Onset-tracking adaptive T_evap investigation", bold=True)

    add_para(doc, (
        "To address this apparent waste, an adaptive evaporator strategy was "
        "investigated. Rather than maintaining T_evap = 5 \u00b0C throughout drying, "
        "the onset-tracking strategy sets T_evap relative to the condensation "
        "onset temperature \u2014 the coil temperature at which the evaporator "
        "outlet air just reaches the mixed air\u2019s dew point:"
    ))

    add_para(doc, (
        "T_evap,coil,onset = T_mix \u2212 (T_mix \u2212 T_dp,mix) / \u03b5_evap"
    ), italic=True)

    add_para(doc, (
        "The evaporator is then set \u0394 K below this onset point: "
        "T_evap = T_onset \u2212 \u0394. A small \u0394 (e.g. 2 K) provides minimal "
        "subcooling below the dew point, while a large \u0394 (e.g. 20 K) "
        "approaches the aggressive dehumidification of the fixed strategy."
    ))

    add_para(doc, (
        "A systematic sweep of uniform \u0394 values from 2 to 20 K was conducted "
        "for Config A at r = 0.9. Table 1.8 compares the results against the "
        "fixed T_evap = 5 \u00b0C baseline."
    ))

    # Delta sweep results table
    headers_d = ["\u0394 (K)", "KTM SEC", "KTM time (h)", "BTN SEC", "BTN time (h)"]
    rows_d = [
        ["Fixed 5\u00b0C", "0.669", "14.9", "0.753", "14.6"],
        ["2", "0.680", "22.4", "0.678", "20.3"],
        ["4", "0.700", "16.9", "0.747", "16.1"],
        ["6", "0.731", "14.7", "0.799", "14.2"],
        ["10", "0.806", "13.7", "0.891", "13.3"],
        ["20", "0.892", "13.2", "0.996", "12.8"],
    ]
    add_table(doc, headers_d, rows_d)
    add_para(doc,
        "Table 1.8: Onset-tracking adaptive T_evap sweep (Config A, r = 0.9, "
        "annual TMY). Fixed baseline shown for comparison.",
        italic=True, size=9)

    add_para(doc, (
        "At Kathmandu, the fixed baseline achieves the lowest SEC across all "
        "\u0394 values tested. At Biratnagar, \u0394 = 2 K provides a 10% SEC improvement "
        "(0.678 vs 0.753 kWh/kg) but at the cost of extending drying time from "
        "14.6 h to 20.3 h \u2014 a 39% increase."
    ))

    add_para(doc, "Why adaptive T_evap fails in this system", bold=True)

    add_para(doc, (
        "The negative result is explained by the RH-dependent drying kinetics "
        "used in this model. The Midilli kinetic equation includes a humidity "
        "suppression term with exponent \u03b1_RH = 1.75, meaning higher chamber "
        "relative humidity directly slows the drying rate. When the evaporator "
        "is raised above 5 \u00b0C (less dehumidification), the chamber operates "
        "at higher RH, the drying rate decreases, and the batch takes longer "
        "to complete. The extended fan run-time and compressor hours offset "
        "any COP improvement from the warmer evaporator."
    ))

    add_para(doc, (
        "This is a key distinction from systems where drying kinetics are "
        "temperature-limited rather than humidity-limited. In a "
        "temperature-limited system, reducing evaporator subcooling would "
        "indeed improve COP without affecting drying time. However, for "
        "cardamom drying with strong RH dependence, the aggressive "
        "dehumidification provided by a fixed 5 \u00b0C evaporator is near-optimal "
        "for minimising SEC."
    ))

    add_para(doc, (
        "This finding motivates the condenser-direct bypass strategy explored "
        "in Section 2, which addresses the late-drying inefficiency from a "
        "different angle: rather than raising T_evap (which weakens "
        "dehumidification), the bypass strategy routes warm exhaust directly "
        "to the condenser when the air no longer needs dehumidification, "
        "eliminating the cooling\u2013reheating cycle entirely."
    ))


# ── main ───────────────────────────────────────────────────────────────


def main():
    df = load_data()
    doc = Document()

    # Title
    add_heading(doc, "Results and Discussion", level=0)
    add_heading(doc, "Section 1: Config A (HP-Only) and Config B (Solar + HP Series)", level=1)
    add_para(doc, (
        "This section presents the performance of the two simplest SAHPD "
        "configurations: Config A (heat-pump-only convective dryer) and "
        "Config B (flat-plate solar collector in series with the heat pump). "
        "Together, these establish the baseline against which heat recovery "
        "(D-configs) and combined systems (E-configs) are evaluated in "
        "subsequent sections."
    ))

    write_section_1_1(doc, df)
    write_section_1_2(doc, df)
    write_section_1_3(doc, df)
    write_section_1_4(doc, df)
    write_section_1_5(doc, df)
    write_section_1_6(doc, df)
    write_section_1_7(doc, df)

    # Save
    out_dir = PROJECT_ROOT / "thesis"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / "Section1_Config_AB.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")

    # Quick stats
    n_a = len(df[df["config"] == "A"])
    n_b = len(df[(df["config"] == "B") & ~df["filename"].str.contains("_s4", na=False)])
    print(f"Data rows referenced: {n_a} (A) + {n_b} (B)")


if __name__ == "__main__":
    main()
