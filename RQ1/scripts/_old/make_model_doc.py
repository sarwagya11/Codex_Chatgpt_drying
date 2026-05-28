"""Generate Word document: SAHPD Model Description for Professor Discussion.

Creates a concise document explaining:
1. How air properties are calculated (psychrometrics)
2. How HP components are calculated (vapor compression cycle)
3. How air and HP interact (first-law coupling)
4. Air paths for Config A and Config B
5. Why recirculation works/doesn't work
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

PROJECT_ROOT = Path(__file__).resolve().parents[1]


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_equation(doc, text):
    """Add an equation-style paragraph (monospace, indented)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    return table


def main():
    doc = Document()

    # Title
    title = doc.add_heading('SAHPD Model: Air-HP Interaction & Recirculation Analysis', level=0)
    add_para(doc, 'Prepared for: Professor Discussion | Date: 2026-03-24', italic=True, size=10)
    add_para(doc, 'System: Solar-Assisted Heat Pump Dryer for Apple Drying in Nepal', italic=True, size=10)
    add_para(doc, 'Refrigerant: R134a | T_set: 45 C | Product: 30 kg apple (10 trays)', italic=True, size=10)

    # ====================================================================
    add_heading(doc, '1. Air Property Calculations (Psychrometrics)', level=1)
    # ====================================================================

    add_para(doc, 'All air properties are computed from first principles using ASHRAE psychrometric relations. '
             'The key state variables at every point in the system are Temperature (T), '
             'Humidity Ratio (omega, kg water/kg dry air), and Relative Humidity (RH).')

    add_heading(doc, '1.1 Saturation Vapor Pressure', level=2)
    add_para(doc, 'The saturation vapor pressure of water at temperature T (in C) is computed using the Buck equation:')
    add_equation(doc, 'p_sat(T) = 611.21 * exp((18.678 - T/234.5) * T / (257.14 + T))  [Pa]')

    add_heading(doc, '1.2 Humidity Ratio', level=2)
    add_para(doc, 'From temperature and relative humidity:')
    add_equation(doc, 'omega = 0.622 * RH * p_sat(T) / (P_atm - RH * p_sat(T))  [kg/kg_da]')
    add_para(doc, 'P_atm depends on elevation: Kathmandu (1350m) = 86,120 Pa; Biratnagar (72m) = 100,460 Pa.')

    add_heading(doc, '1.3 Moist Air Enthalpy', level=2)
    add_para(doc, 'Enthalpy per kg of dry air (ASHRAE):')
    add_equation(doc, 'h = 1.006*T + omega*(2501 + 1.86*T)  [kJ/kg_da]')
    add_para(doc, 'This captures both sensible heat (1.006*T) and latent heat (omega*2501) of the moisture in air. '
             'The latent component is critical: humid air carries significantly more energy than dry air at the same temperature.')

    add_heading(doc, '1.4 Dewpoint Temperature', level=2)
    add_para(doc, 'The dewpoint is computed by inverting the saturation pressure equation:')
    add_equation(doc, 'p_v = omega * P_atm / (0.622 + omega)')
    add_equation(doc, 'T_dp = f_inverse(p_v)  [C]')
    add_para(doc, 'If air is cooled below T_dp, moisture condenses out. This is how the evaporator dehumidifies.')

    # ====================================================================
    add_heading(doc, '2. Heat Pump Component Calculations', level=1)
    # ====================================================================

    add_heading(doc, '2.1 Vapor Compression Cycle (R134a)', level=2)
    add_para(doc, 'The HP uses a standard 4-point vapor compression cycle computed with CoolProp (NIST-quality thermodynamic properties):')

    add_heading(doc, 'State 1: Evaporator Outlet (Superheated Vapor)', level=3)
    add_equation(doc, 'T_1 = T_evap + DT_superheat  (5 + 5 = 10 C)')
    add_equation(doc, 'P_1 = P_sat(T_evap)  [from CoolProp]')
    add_equation(doc, 'h_1, s_1 = CoolProp(T_1, P_1)  [kJ/kg]')

    add_heading(doc, 'State 2: Compressor Outlet (Superheated Vapor)', level=3)
    add_equation(doc, 'P_2 = P_sat(T_cond)  = P_sat(55 C)  [condenser pressure]')
    add_equation(doc, 'h_2s = CoolProp(P_2, s_1)  [isentropic compression]')
    add_equation(doc, 'h_2 = h_1 + (h_2s - h_1) / eta_isentropic  [actual, eta=0.72]')
    add_equation(doc, 'W_comp_specific = (h_2 - h_1) / eta_mechanical  [kJ/kg_ref, eta=0.90]')

    add_heading(doc, 'State 3: Condenser Outlet (Subcooled Liquid)', level=3)
    add_equation(doc, 'T_3 = T_cond - DT_subcooling  (55 - 5 = 50 C)')
    add_equation(doc, 'h_3 = CoolProp(T_3, P_cond)  [kJ/kg]')
    add_equation(doc, 'Q_cond_specific = h_2 - h_3  [kJ/kg_ref]')

    add_heading(doc, 'State 4: Expansion Valve Outlet (Two-Phase)', level=3)
    add_equation(doc, 'h_4 = h_3  [isenthalpic expansion]')
    add_equation(doc, 'Q_evap_specific = h_1 - h_4  [kJ/kg_ref]')

    add_heading(doc, '2.2 Sizing the Refrigerant Flow', level=2)
    add_para(doc, 'The refrigerant mass flow rate is sized to deliver the required Q_cond:')
    add_equation(doc, 'm_ref = Q_cond_target / (h_2 - h_3)  [kg/s]')
    add_para(doc, 'Then all energy flows scale with m_ref:')
    add_equation(doc, 'Q_evap = m_ref * (h_1 - h_4)  [kW]')
    add_equation(doc, 'W_comp = m_ref * (h_2 - h_1) / eta_mech  [kW]')
    add_equation(doc, 'Q_cond = m_ref * (h_2 - h_3)  [kW]')
    add_equation(doc, 'COP = Q_cond / W_comp')
    add_para(doc, 'The compressor is modeled as infinitely variable (no maximum displacement). '
             'Feasibility flags warn when Q_cond > 4 kW (1-ton AC limit).')

    add_heading(doc, '2.3 First Law Enforcement', level=2)
    add_para(doc, 'The HP first law requires: Q_cond = Q_evap + W_comp. '
             'In closed-loop mode (r > 0), the SAME air stream passes through the evaporator and then the condenser. '
             'This means Q_cond is constrained by how much heat the evaporator can extract from the air:')
    add_equation(doc, 'Q_evap_air = m_da * (h_before_evap - h_after_evap)  [kW]')
    add_equation(doc, 'Q_cond_1st_law = Q_evap_air * COP / (COP - eta_mech)')
    add_equation(doc, 'Q_cond_actual = min(Q_cond_1st_law, Q_cond_air_requirement, Q_cond_effectiveness)')
    add_para(doc, 'When Q_cond_1st_law < Q_cond_air_requirement, the condenser CANNOT heat air to T_set (45 C). '
             'T_to_chamber floats below T_set. This is the physical root of the recirculation problem.')

    add_heading(doc, '2.4 Evaporator Air-Side Model', level=2)
    add_para(doc, 'The evaporator uses an effectiveness-based model (equivalent to the contact factor method of Chou & Chua, 1994):')
    add_equation(doc, 'T_out = T_air_in - epsilon_evap * (T_air_in - T_coil)')
    add_equation(doc, 'T_coil = T_evap_sat + DT_approach  (5 + 3 = 8 C default)')
    add_para(doc, 'Dehumidification: if T_out < T_dewpoint of incoming air, the air exits saturated at T_out:')
    add_equation(doc, 'if T_out < T_dp:  omega_out = omega_sat(T_out)  [saturated]')
    add_equation(doc, 'else:             omega_out = omega_in          [no condensation]')
    add_para(doc, 'epsilon_evap = 0.85 (fixed). T_evap_sat = 5 C (fixed target), modulated down to -5 C '
             'when T_air_in approaches T_coil (minimum 5K temperature difference maintained).')

    add_heading(doc, '2.5 Condenser Air-Side Model', level=2)
    add_para(doc, 'The condenser effectiveness (epsilon_cond = 0.85) limits the maximum air outlet temperature:')
    add_equation(doc, 'T_out_max = T_air_in + epsilon_cond * (T_cond_sat - T_air_in)')
    add_para(doc, 'With T_cond_sat = 55 C and epsilon = 0.85, the condenser can reach 45 C whenever T_air_in < 37 C. '
             'The condenser does not add or remove moisture (sensible heating only, omega unchanged).')

    # ====================================================================
    add_heading(doc, '3. Config A: HP-Only Dryer', level=1)
    # ====================================================================

    add_heading(doc, '3.1 Open-Loop Air Path (r = 0)', level=2)
    add_para(doc, 'Air flows in a single pass with no recirculation:')
    add_equation(doc, 'Ambient (8.8 C) --> CONDENSER --> 45 C --> CHAMBER --> Exhaust (discarded)')
    add_equation(doc, '                       ^')
    add_equation(doc, 'Separate ambient --> EVAPORATOR --> (heat absorbed, air discarded)')
    add_para(doc, 'The evaporator and condenser process DIFFERENT air streams. There is no first-law coupling '
             'between them on the air side. The HP is sized freely to heat ambient air to T_set.')
    add_equation(doc, 'Q_cond = m_da * (h(45 C, omega_amb) - h(T_amb, omega_amb))')
    add_equation(doc, '       = 0.098 * (h(45) - h(8.8)) = 3.63 kW')
    add_equation(doc, 'T_evap = T_amb - 10 = -1.2 C  (external heat source)')
    add_equation(doc, 'T_to_chamber = 45 C  (ALWAYS reaches T_set)')

    add_heading(doc, '3.2 Closed-Loop Air Path (r > 0)', level=2)
    add_para(doc, 'A fraction r of exhaust is recirculated and mixed with (1-r) fresh ambient air. '
             'The SAME air stream passes through evaporator then condenser:')
    add_equation(doc, 'r*Exhaust + (1-r)*Ambient --> MIX --> EVAPORATOR --> CONDENSER --> CHAMBER')
    add_equation(doc, '      ^                                                              |')
    add_equation(doc, '      |______________________________________________________________|')

    add_para(doc, 'Step-by-step calculation (example: r=0.1, Kathmandu, t=1h):')
    add_equation(doc, '1. MIXING:')
    add_equation(doc, '   T_mix = 0.1*35 + 0.9*8.8 = 11.4 C')
    add_equation(doc, '   omega_mix = 0.1*omega_exh + 0.9*omega_amb = 0.0063 kg/kg')
    add_equation(doc, '')
    add_equation(doc, '2. EVAPORATOR:')
    add_equation(doc, '   T_coil = 5 + 3 = 8 C, but T_mix - T_coil = 3.4K < 5K minimum')
    add_equation(doc, '   --> Modulate: T_coil = 11.4 - 5 = 6.4 C, T_evap = 3.4 C')
    add_equation(doc, '   T_out = 11.4 - 0.85*(11.4 - 6.4) = 7.2 C')
    add_equation(doc, '   T_dp(omega=0.0063) = 3 C --> T_out > T_dp --> NO dehumidification')
    add_equation(doc, '   Q_evap = 0.098 * 1.006 * (11.4 - 7.2) = 0.41 kW  (sensible only)')
    add_equation(doc, '')
    add_equation(doc, '3. FIRST LAW CONSTRAINT:')
    add_equation(doc, '   COP = 4.0 (at T_evap=3.4 C, T_cond=55 C)')
    add_equation(doc, '   Q_cond = Q_evap * COP/(COP - eta_m) = 0.41 * 4.0/3.1 = 0.53 kW')
    add_equation(doc, '   vs Q_cond needed for 45 C = 3.73 kW')
    add_equation(doc, '   Q_cond_actual = min(0.53, 3.73) = 0.53 kW  <-- FIRST LAW LIMITS')
    add_equation(doc, '')
    add_equation(doc, '4. CONDENSER:')
    add_equation(doc, '   T_to_chamber = 7.2 + 0.53/(0.098*1.006) = 7.2 + 5.4 = 12.6 C')
    add_equation(doc, '   (vs 45 C at r=0)')

    add_heading(doc, '3.3 The Negative Feedback Loop', level=2)
    add_para(doc, 'At r=0.1 in cold climate (T_amb=8.8 C), a vicious cycle develops:')
    add_equation(doc, 'Low T_to_chamber (13 C)')
    add_equation(doc, '  --> Slow drying')
    add_equation(doc, '  --> Low T_exhaust (10 C)')
    add_equation(doc, '  --> Low T_mix (9 C)')
    add_equation(doc, '  --> Tiny Q_evap (0.4 kW)')
    add_equation(doc, '  --> Tiny Q_cond (0.5 kW)')
    add_equation(doc, '  --> Low T_to_chamber ...')
    add_para(doc, 'The system gets stuck at 10-15 C. Drying takes 72+ hours and never completes. '
             'This is the "valley of death" at intermediate r values (0.1-0.7) in cold climates.')

    add_heading(doc, '3.4 Why r >= 0.9 Works', level=2)
    add_para(doc, 'At r=0.9, the mix is dominated by warm exhaust:')
    add_equation(doc, 'T_mix = 0.9*38 + 0.1*8.8 = 35.1 C')
    add_equation(doc, 'DeltaT = 35.1 - 8.0 = 27.1 K  --> large Q_evap')
    add_equation(doc, 'Also: omega_mix is high --> T_dp > T_coil --> DEHUMIDIFICATION occurs')
    add_equation(doc, '--> Q_evap includes LATENT heat (much larger than sensible alone)')
    add_equation(doc, '--> Q_cond is large --> T_to_chamber reaches 45 C')

    # ====================================================================
    add_heading(doc, '4. Config B: Solar + HP Series', level=1)
    # ====================================================================

    add_heading(doc, '4.1 Open-Loop Air Path (r = 0)', level=2)
    add_equation(doc, 'Ambient --> SOLAR COLLECTOR --> T_after_solar --> CONDENSER --> CHAMBER')
    add_equation(doc, '                                                    ^')
    add_equation(doc, '                        Separate ambient --> EVAPORATOR (heat source)')
    add_para(doc, 'Solar preheats the air before the HP condenser. If T_after_solar >= 45 C, the HP shuts off '
             '(W_comp = 0, free solar drying). Otherwise, the HP only boosts from T_after_solar to 45 C '
             '(smaller delta-T = less compressor work).')

    add_heading(doc, '4.2 Closed-Loop Air Path (r > 0)', level=2)
    add_equation(doc, 'r*Exhaust + (1-r)*Ambient --> MIX --> EVAPORATOR --> SOLAR --> CONDENSER --> CHAMBER')
    add_equation(doc, '      ^                                                                       |')
    add_equation(doc, '      |_______________________________________________________________________|')
    add_para(doc, 'Key difference from Config A: the SOLAR COLLECTOR sits between evaporator and condenser. '
             'After the evaporator cools air to ~6 C, the solar collector reheats it (to 15-25 C during daytime) '
             'before it enters the condenser.')
    add_para(doc, 'The first-law Q_evap is computed from the mix-to-after-evaporator enthalpy drop (same as Config A). '
             'But Q_cond_required is smaller because the condenser inlet is now T_after_solar (warmer). '
             'When Q_cond_required < Q_cond_1st_law, the HP does less work than the first law allows.')

    add_heading(doc, '4.3 Why Config B Solves the Valley of Death', level=2)
    add_para(doc, 'At intermediate r (e.g., r=0.2, Kathmandu):')
    add_equation(doc, 'Config A:  Mix(11 C) --> Evap(6 C) --> Cond --> 13 C  (STUCK)')
    add_equation(doc, 'Config B:  Mix(11 C) --> Evap(6 C) --> Solar(20 C) --> Cond --> 45 C  (WORKS)')
    add_para(doc, 'The solar collector provides the "missing" thermal energy that the HP first law cannot deliver. '
             'During daytime, solar adds 10-30 C to the air temperature. At night, Config B behaves like Config A '
             '(slow drying), but daytime solar progress carries the batch to completion.')
    add_para(doc, 'Result: ALL recirculation ratios converge in Config B (22/22 runs completed).')

    # ====================================================================
    add_heading(doc, '5. Recirculation Ratio Results', level=1)
    # ====================================================================

    add_heading(doc, '5.1 Config A (HP-Only)', level=2)

    headers = ['r', 'Time (h)', 'W_comp (kWh)', 'SEC (kWh/kg)', 'Converged']
    rows_ktm = [
        ['0.0', '13.8', '13.52', '0.717', 'Yes'],
        ['0.1', '72.0', '15.55', '0.974', 'No'],
        ['0.2', '72.0', '15.07', '0.951', 'No'],
        ['0.3', '72.0', '14.67', '0.931', 'No'],
        ['0.4', '72.0', '14.35', '0.915', 'No'],
        ['0.5', '72.0', '14.08', '0.901', 'No'],
        ['0.6', '72.0', '13.89', '0.890', 'No'],
        ['0.7', '72.0', '14.02', '0.892', 'No'],
        ['0.8', '31.4', '15.33', '0.849', 'Yes'],
        ['0.9', '14.9', '12.35', '0.669', 'Yes'],
        ['1.0', '15.2', '12.45', '0.674', 'Yes'],
    ]
    add_para(doc, 'Kathmandu (1350m, T_amb = 8.8 C):', bold=True)
    add_table(doc, headers, rows_ktm)
    doc.add_paragraph()  # spacer

    rows_btn = [
        ['0.0', '14.4', '10.13', '0.543', 'Yes'],
        ['0.1', '72.0', '28.26', '1.602', 'No'],
        ['0.2', '72.0', '28.49', '1.609', 'No'],
        ['0.3', '72.0', '29.33', '1.647', 'No'],
        ['0.4', '72.0', '30.84', '1.716', 'No'],
        ['0.5', '33.5', '17.85', '0.995', 'Yes'],
        ['0.6', '29.8', '17.99', '0.994', 'Yes'],
        ['0.7', '14.8', '14.21', '0.765', 'Yes'],
        ['0.8', '14.6', '14.01', '0.755', 'Yes'],
        ['0.9', '14.6', '13.98', '0.753', 'Yes'],
        ['1.0', '14.7', '13.95', '0.752', 'Yes'],
    ]
    add_para(doc, 'Biratnagar (72m, T_amb = 25 C):', bold=True)
    add_table(doc, headers, rows_btn)
    doc.add_paragraph()

    add_heading(doc, '5.2 Config B (Solar 10m2 + HP)', level=2)

    headers_b = ['r', 'Time (h)', 'W_comp (kWh)', 'SEC (kWh/kg)', 'Converged']
    rows_b_ktm = [
        ['0.0', '13.8', '9.09', '0.488', 'Yes'],
        ['0.1', '31.6', '6.47', '0.398', 'Yes'],
        ['0.2', '31.4', '6.26', '0.386', 'Yes'],
        ['0.3', '31.1', '6.37', '0.392', 'Yes'],
        ['0.5', '30.1', '6.90', '0.416', 'Yes'],
        ['0.7', '29.0', '7.84', '0.460', 'Yes'],
        ['0.9', '14.9', '8.49', '0.469', 'Yes'],
        ['1.0', '15.2', '8.62', '0.477', 'Yes'],
    ]
    add_para(doc, 'Kathmandu (Solar 10m2):', bold=True)
    add_table(doc, headers_b, rows_b_ktm)
    doc.add_paragraph()

    rows_b_btn = [
        ['0.0', '14.4', '5.18', '0.287', 'Yes'],
        ['0.1', '27.8', '9.50', '0.548', 'Yes'],
        ['0.2', '27.7', '9.62', '0.553', 'Yes'],
        ['0.5', '26.9', '10.91', '0.618', 'Yes'],
        ['0.7', '14.8', '7.76', '0.431', 'Yes'],
        ['0.9', '14.6', '7.59', '0.423', 'Yes'],
        ['1.0', '14.7', '7.60', '0.423', 'Yes'],
    ]
    add_para(doc, 'Biratnagar (Solar 10m2):', bold=True)
    add_table(doc, headers_b, rows_b_btn)
    doc.add_paragraph()

    add_para(doc, 'Key finding: Config B achieves 100% convergence (22/22 runs) across all r values. '
             'Solar preheating eliminates the valley of death. Best SEC: KTM r=0.2 (0.386), BTN r=0.0 (0.287).', bold=True)

    # ====================================================================
    add_heading(doc, '6. Fixed Energy / Fixed Time Comparison', level=1)
    # ====================================================================

    add_heading(doc, '6.1 At Fixed Energy (10 kWh electrical)', level=2)
    add_para(doc, 'How much water is removed when the same electrical energy is consumed?')

    headers_fe = ['r', 'Time to 10 kWh', 'Water (kg)', 'MR']
    rows_fe = [
        ['0.0', '10.0 h', '18.80', '0.032'],
        ['0.1', '39.4 h', '17.23', '0.089'],
        ['0.5', '39.2 h', '17.34', '0.075'],
        ['0.8', '11.8 h', '18.91', '0.022'],
        ['0.9', '11.4 h', '18.98', '0.022'],
        ['1.0', '11.6 h', '18.98', '0.022'],
    ]
    add_table(doc, headers_fe, rows_fe)
    doc.add_paragraph()
    add_para(doc, 'At r=0.1-0.7, the HP runs at very low power (~0.3 kW vs ~1.0 kW). '
             'It takes 39 hours to consume 10 kWh because Q_cond is throttled by the first law.')

    add_heading(doc, '6.2 At Fixed Time (14 hours)', level=2)
    add_para(doc, 'How much energy is consumed and water removed after 14 hours?')

    headers_ft = ['r', 'W_total (kWh)', 'Water (kg)', 'MR', 'T_cham (C)']
    rows_ft = [
        ['0.0', '13.86', '19.34', '0.005', '45.0'],
        ['0.1', '3.79', '9.70', '0.483', '13.5'],
        ['0.5', '4.30', '11.08', '0.409', '14.2'],
        ['0.7', '4.51', '12.17', '0.351', '15.0'],
        ['0.9', '12.15', '19.28', '0.007', '45.0'],
        ['1.0', '12.03', '19.26', '0.007', '45.0'],
    ]
    add_table(doc, headers_ft, rows_ft)
    doc.add_paragraph()
    add_para(doc, 'At 14h, r=0.1 has used only 3.79 kWh (27% of r=0.0) and removed only 50% of the water. '
             'The HP is effectively "starved" by the first-law constraint. The chamber temperature is 13.5 C '
             'instead of 45 C.')

    # ====================================================================
    add_heading(doc, '7. Key Physics Summary', level=1)
    # ====================================================================

    add_para(doc, '1. The minimum thermal energy to evaporate 19.3 kg of water is fixed: '
             'Q_latent = m_w * h_fg = 19.3 * 2400/3600 = 12.9 kWh. This does not change with r.', bold=False)
    add_para(doc, '2. What changes with r is HOW MUCH of the condenser heat goes to useful drying '
             'vs sensible losses (exhaust waste at r=0, evaporator re-cooling at r=1).', bold=False)
    add_para(doc, '3. The HP first law (Q_cond = Q_evap + W_comp) creates a fundamental coupling '
             'in closed-loop mode that does not exist in open-loop mode.', bold=False)
    add_para(doc, '4. In cold climates, intermediate r (0.1-0.7) mixes too much cold ambient air, '
             'starving the evaporator and creating a negative feedback loop.', bold=False)
    add_para(doc, '5. Solar preheating (Config B) breaks this feedback by adding thermal energy '
             'between the evaporator and condenser, independent of the HP first law.', bold=False)
    add_para(doc, '6. Literature support: Braun & Bansal (ORNL, 2022) showed vented and unvented HP dryers '
             'have fundamentally different Carnot limits. A 2019 Energies study found 39% COP degradation '
             'in cold-climate HP dryers with the same mechanism we observe.', bold=False)

    # Save
    output_path = PROJECT_ROOT / 'SAHPD_Model_Description.docx'
    doc.save(str(output_path))
    print(f"Document saved: {output_path}")


if __name__ == '__main__':
    main()
