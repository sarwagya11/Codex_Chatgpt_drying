"""Generate E1/E2/E3 findings reference document (system-level only).

A single .docx the user can lift numbers, tables, and figures from straight
into the paper. No kinetics, no model derivation, no literature.

Output: RQ1/E123_FINDINGS.docx
"""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

ROOT = Path(__file__).resolve().parents[1]
AUDIT = ROOT.parent / "outputs" / "audit"
PLOTS = ROOT / "plots" / "_audit"
OUT_DOCX = ROOT / "E123_FINDINGS.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_table_from_df(doc, df: pd.DataFrame, num_fmt: dict | None = None,
                      first_col_bold=True):
    cols = list(df.columns)
    tbl = doc.add_table(rows=1 + len(df), cols=len(cols))
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    for j, c in enumerate(cols):
        hdr[j].text = str(c)
        for p in hdr[j].paragraphs:
            for r in p.runs:
                r.bold = True
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        for j, c in enumerate(cols):
            v = row[c]
            if num_fmt and c in num_fmt and isinstance(v, (int, float)):
                txt = num_fmt[c].format(v)
            elif isinstance(v, float):
                txt = f"{v:.4f}"
            else:
                txt = str(v)
            tbl.rows[i].cells[j].text = txt
            if first_col_bold and j == 0:
                for p in tbl.rows[i].cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True
    doc.add_paragraph()


def add_figure(doc, png_path: Path, caption: str, width_cm=15):
    if not png_path.exists():
        add_para(doc, f"[MISSING FIGURE: {png_path.name}]", italic=True)
        return
    doc.add_picture(str(png_path), width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(10)


# ---------- data prep ----------
def load_pivots():
    nv = pd.read_csv(AUDIT / "step4_pivot_SEC_no_vpd.csv")
    vp = pd.read_csv(AUDIT / "step4_pivot_SEC_vpd.csv")
    return nv, vp


def fmt_pivot(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df["case"] = df["location"].str[:3].str.upper() + " / " + df["season"].str.replace("_", " ")
    out = df[["case", "E1", "E2", "E3", "best",
              "E2_vs_E1_pct", "E3_vs_E1_pct", "E3_vs_E2_pct"]].copy()
    out.columns = ["case", "E1", "E2", "E3", "best",
                   "E2 vs E1 [%]", "E3 vs E1 [%]", "E3 vs E2 [%]"]
    return out


def build_energy_split_table(es: pd.DataFrame) -> pd.DataFrame:
    """Per-config means across all (location, season) cases."""
    grp = es.groupby("config").agg(
        SEC=("SEC", "mean"),
        SMER=("SMER", "mean"),
        COP=("COP", "mean"),
        share_HRX=("share_hrx", "mean"),
        share_solar=("share_solar", "mean"),
        share_HP=("share_cond", "mean"),
        eta_overall=("eta_overall", "mean"),
        clip_frac=("Q_sol_clipped", lambda s: s.sum() / es.loc[s.index, "Q_sol_gross"].sum()),
    ).reset_index()
    return grp


# ---------- main ----------
def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading("E1 / E2 / E3 System-Level Findings", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Reference document for paper writing — numbers, tables, "
                     "and figures only (system-level; kinetic-model audit "
                     "documented separately).")
    sr.italic = True

    add_para(doc, "All results: HRX effectiveness e=0.70, collector area "
             "A_c=10 m^2, R134a heat pump, T_set=45 C, batch m_p_dry=3.0 kg "
             "(10 trays). VPD bypass: threshold=0.05 kPa, 3x hysteresis, 600 s "
             "dwell. Validation: zero first-law violations across 92 canonical "
             "runs (Section 7).")

    # =========================================================================
    # 1. Configuration recap
    # =========================================================================
    add_heading(doc, "1. Configuration recap", level=1)

    add_para(doc, "All three configurations combine an exhaust-air heat-recovery "
             "exchanger (HRX), a flat-plate solar air collector, and an R134a "
             "heat pump. They differ only in (a) where the solar collector sits "
             "in the air path and (b) how the evaporator is fed.")

    add_heading(doc, "E1 - HRX + Solar before condenser, ambient evaporator",
                level=2)
    add_para(doc, "Air path: Ambient -> HRX (preheat) -> Solar (preheat) -> "
             "Condenser (top-up to 45 C) -> Chamber. The evaporator is fed by "
             "ambient air only. This is the simplest topology: solar acts as a "
             "second preheater, the heat pump never sees recovered exhaust on "
             "the source side.")

    add_heading(doc, "E2 - HRX + Solar before condenser, exhaust+ambient evaporator",
                level=2)
    add_para(doc, "Air path: identical to E1 on the supply side (Amb -> HRX -> "
             "Solar -> Cond -> Chamber). The difference is the evaporator: it "
             "is fed by the post-HRX exhaust mixed with an ambient supplement "
             "sized iteratively so the evaporator can absorb Q_evap = Q_cond - "
             "eta_mech*W_comp. Warmer evaporator inlet -> higher COP -> lower "
             "compressor work for the same Q_cond.")

    add_heading(doc, "E3 - HRX + Solar after condenser, exhaust+ambient evaporator",
                level=2)
    add_para(doc, "Air path: Ambient -> HRX -> Condenser -> Solar -> Chamber. "
             "The collector is plumbed downstream of the condenser, so when "
             "irradiance is high the heat pump can run at a lower lift "
             "(partial-lift control law) and let the collector finish heating "
             "to T_set. The evaporator side matches E2.")

    add_para(doc, "VPD bypass (all three configs): when the chamber-inlet "
             "vapour-pressure deficit relative to recovered exhaust falls below "
             "0.05 kPa, the supply air bypasses the HRX. For E1/E2 the bypass "
             "path is Exhaust -> Solar -> Cond -> Chamber. For E3 the bypass "
             "preserves topology: Exhaust -> Cond -> Solar -> Chamber. A 3x "
             "hysteresis band and 600 s dwell prevent rapid switching.")

    # =========================================================================
    # 2. Headline SEC numbers
    # =========================================================================
    add_heading(doc, "2. Headline performance numbers", level=1)

    nv, vp = load_pivots()

    add_heading(doc, "2.1 SEC without VPD bypass [kWh / kg water]", level=2)
    nv_disp = fmt_pivot(nv)
    add_table_from_df(doc, nv_disp,
                      num_fmt={"E1": "{:.4f}", "E2": "{:.4f}", "E3": "{:.4f}",
                               "E2 vs E1 [%]": "{:+.2f}",
                               "E3 vs E1 [%]": "{:+.2f}",
                               "E3 vs E2 [%]": "{:+.2f}"})
    add_para(doc, f"Source: outputs/audit/step4_pivot_SEC_no_vpd.csv "
             f"({len(nv)} cases). E2 wins SEC in {(nv['best']=='E2').sum()}/{len(nv)} "
             "cases. Mean E2-vs-E1 saving = "
             f"{nv['E2_vs_E1_pct'].mean():.2f} %; "
             f"mean E3-vs-E2 penalty = {nv['E3_vs_E2_pct'].mean():.2f} %.",
             italic=True)

    add_heading(doc, "2.2 SEC with VPD bypass [kWh / kg water]", level=2)
    vp_disp = fmt_pivot(vp)
    add_table_from_df(doc, vp_disp,
                      num_fmt={"E1": "{:.4f}", "E2": "{:.4f}", "E3": "{:.4f}",
                               "E2 vs E1 [%]": "{:+.2f}",
                               "E3 vs E1 [%]": "{:+.2f}",
                               "E3 vs E2 [%]": "{:+.2f}"})
    add_para(doc, f"Source: outputs/audit/step4_pivot_SEC_vpd.csv ({len(vp)} cases). "
             f"E2 wins {(vp['best']=='E2').sum()}/{len(vp)}. Mean E2-vs-E1 saving = "
             f"{vp['E2_vs_E1_pct'].mean():.2f} %; mean E3-vs-E2 penalty = "
             f"{vp['E3_vs_E2_pct'].mean():.2f} %. Dhulikhel intentionally "
             "omitted from the VPD set (no VPD experiment run for that location).",
             italic=True)

    # VPD reduction table
    add_heading(doc, "2.3 VPD bypass reduction in SEC vs no-VPD baseline", level=2)
    merged = nv.merge(vp, on=["location", "season"], suffixes=("_nv", "_vp"))
    rows = []
    for cfg in ("E1", "E2", "E3"):
        for _, r in merged.iterrows():
            rows.append(dict(case=f"{r['location'][:3].upper()} / {r['season'].replace('_',' ')}",
                             config=cfg,
                             SEC_no_vpd=r[f"{cfg}_nv"], SEC_vpd=r[f"{cfg}_vp"],
                             reduction_pct=100 * (r[f"{cfg}_vp"] - r[f"{cfg}_nv"]) / r[f"{cfg}_nv"]))
    rd = pd.DataFrame(rows)
    summary = (rd.groupby("config")["reduction_pct"]
               .agg(mean="mean", min="min", max="max")
               .round(2).reset_index())
    summary.columns = ["config", "mean reduction [%]", "min [%]", "max [%]"]
    add_table_from_df(doc, summary)
    add_para(doc, "Negative values indicate SEC improvement under VPD. "
             "Source: derived from step4_pivot_SEC_*.csv.", italic=True)

    # =========================================================================
    # 3. Head-to-head verdict
    # =========================================================================
    add_heading(doc, "3. Head-to-head verdict", level=1)

    wr = pd.read_csv(AUDIT / "step4_winrate.csv")
    wr = wr.rename(columns={wr.columns[0]: "config"})
    add_heading(doc, "3.1 Win count per KPI (best per case)", level=2)
    add_table_from_df(doc, wr)
    add_para(doc, "Read: 'SEC=13' means that config achieves the best SEC in "
             "13 of 13 cases. Source: outputs/audit/step4_winrate.csv.",
             italic=True)

    add_para(doc, "")
    add_para(doc, "Key takeaways:", bold=True)
    bullets = [
        "E2 wins SEC and SMER in every single case (13/13 no-VPD, 12/12 VPD).",
        "E1 always finishes drying fastest (lowest t_h) because its evaporator "
        "is colder, which forces a higher condenser duty and pushes more heat "
        "into the chamber per unit time, even though it is energetically "
        "expensive.",
        "E3 has the highest median COP in 11/12 VPD cases (lower lift -> "
        "higher Carnot ceiling), but the lift it enables is partial, so total "
        "compressor energy and SEC are both worse than E2.",
        "Q_solar_usable wins are split E1/E3 because their collector inlet "
        "temperatures differ; raw solar capture does not predict SEC.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    # =========================================================================
    # 4. Energy split (where each kWh comes from)
    # =========================================================================
    add_heading(doc, "4. Energy split - where each kWh comes from", level=1)

    es = pd.read_csv(AUDIT / "step3_energy_split.csv")
    add_heading(doc, "4.1 Mean energy shares by config (all 13 cases)", level=2)

    grp = build_energy_split_table(es)
    grp_disp = grp.copy()
    grp_disp["share_HRX"] = grp_disp["share_HRX"] * 100
    grp_disp["share_solar"] = grp_disp["share_solar"] * 100
    grp_disp["share_HP"] = grp_disp["share_HP"] * 100
    grp_disp["eta_overall"] = grp_disp["eta_overall"] * 100
    grp_disp["clip_frac"] = grp_disp["clip_frac"] * 100
    grp_disp.columns = ["config", "SEC [kWh/kg]", "SMER [kg/kWh]", "COP",
                        "HRX share [%]", "Solar share [%]", "HP share [%]",
                        "eta_overall [%]", "Solar clipped [%]"]
    add_table_from_df(doc, grp_disp,
                      num_fmt={c: "{:.2f}" for c in grp_disp.columns[1:]})
    add_para(doc, "Shares are fractions of the chamber-inlet heat budget "
             "(Q_HRX + Q_solar_usable + Q_cond). 'eta_overall' is "
             "Q_useful_latent / (W_total + Q_solar_gross). 'Solar clipped' is "
             "the fraction of gross solar capture that exceeded what the "
             "control loop could absorb at T_set. Source: "
             "outputs/audit/step3_energy_split.csv.", italic=True)

    add_para(doc, "Headline: across all configurations, ~50 % of inlet-air "
             "heat comes from the HRX, ~25-30 % from solar, ~22-30 % from the "
             "heat pump. The HRX is the largest single source of heat, "
             "followed by the heat pump and solar in roughly equal proportion.",
             bold=True)

    # Energy split figures
    add_heading(doc, "4.2 Energy-split figures", level=2)
    for cfg in ("E1", "E2", "E3"):
        add_figure(doc, PLOTS / f"step3_energy_split_{cfg}.png",
                   f"Figure 4.{cfg[-1]}. Energy split for {cfg}: chamber-inlet "
                   "heat budget and collector-irradiance fate, by location.")

    # Air-state psychrometric plots
    add_heading(doc, "4.3 Psychrometric trace (Kathmandu annual reference case)",
                level=2)
    for cfg in ("E1", "E2", "E3"):
        add_figure(doc, PLOTS / f"step2e_air_states_{cfg}_kathmandu_annual.png",
                   f"Figure 4.4-{cfg}. Air state at each station (Kathmandu, "
                   f"annual) for {cfg}: T-omega trajectory through HRX, solar, "
                   "condenser, chamber, evaporator.")

    # =========================================================================
    # 5. VPD bypass
    # =========================================================================
    add_heading(doc, "5. VPD bypass strategy and threshold sensitivity", level=1)

    add_para(doc, "The VPD bypass exists because once recovered exhaust "
             "becomes nearly saturated (low VPD vs chamber inlet), routing it "
             "through the HRX no longer transfers useful drying potential. "
             "Bypass diverts the exhaust around the HRX (and through the heat "
             "pump source side / solar collector instead, depending on "
             "config), keeping the supply hot and dry.")

    add_heading(doc, "5.1 Threshold sensitivity sweep (E2)", level=2)
    sw = pd.read_csv(AUDIT / "step4_vpd_sweep.csv")
    sw_disp = sw.copy()
    sw_disp["bypass_frac"] = sw_disp["bypass_frac"] * 100
    sw_disp = sw_disp[["case", "threshold", "t_h", "SEC", "SMER",
                       "bypass_frac"]]
    sw_disp.columns = ["case", "VPD threshold [kPa]", "t [h]",
                       "SEC [kWh/kg]", "SMER [kg/kWh]",
                       "Bypass active [%]"]
    add_table_from_df(doc, sw_disp,
                      num_fmt={"VPD threshold [kPa]": "{:.2f}",
                               "t [h]": "{:.2f}",
                               "SEC [kWh/kg]": "{:.4f}",
                               "SMER [kg/kWh]": "{:.3f}",
                               "Bypass active [%]": "{:.1f}"})
    add_para(doc, "Source: outputs/audit/step4_vpd_sweep.csv. The 0.05 kPa "
             "threshold is the SEC optimum on both representative cases "
             "(KTM annual, BTN autumn). Above 0.10 kPa the bypass becomes too "
             "aggressive: drying time grows steeply and re-inflates SEC by "
             "starving the chamber of dry-air mass flow.", italic=True)

    add_figure(doc, PLOTS / "step4_vpd_sweep.png",
               "Figure 5.1. VPD threshold sweep on E2. Left: SEC vs threshold "
               "(U-shape, optimum ~0.05). Centre: drying time vs threshold "
               "(monotonically increasing). Right: bypass duty cycle vs "
               "threshold.")

    add_heading(doc, "5.2 Bypass effect at the recommended threshold (0.05 kPa)",
                level=2)
    add_para(doc, "Activating the bypass at the recommended threshold yields "
             "an average SEC reduction across configurations of:", italic=False)
    for cfg in ("E1", "E2", "E3"):
        sub = rd[rd["config"] == cfg]
        add_para(doc, f"  - {cfg}: mean reduction "
                 f"{sub['reduction_pct'].mean():+.2f} % "
                 f"(range {sub['reduction_pct'].min():+.2f} % to "
                 f"{sub['reduction_pct'].max():+.2f} %).")
    add_para(doc, "The bypass benefits all three configurations but most "
             "strongly E1 (which has the coldest evaporator and largest "
             "compressor share). Across configs the bypass collapses the SEC "
             "spread: at no-VPD the E2-vs-E1 gap is ~9 %; at VPD-on it shrinks "
             "to ~3-4 %.")

    # =========================================================================
    # 6. Sensitivities
    # =========================================================================
    add_heading(doc, "6. Sensitivity analyses", level=1)

    add_heading(doc, "6.1 Collector area sweep (E2)", level=2)
    aw = pd.read_csv(AUDIT / "step5_area_sweep.csv")
    aw_pivot = aw.pivot_table(index="A_collector_m2", columns="location",
                              values="SEC")
    aw_disp = aw_pivot.reset_index()
    aw_disp.columns = ["A_c [m^2]"] + list(aw_disp.columns[1:])
    add_table_from_df(doc, aw_disp,
                      num_fmt={c: "{:.4f}" for c in aw_disp.columns[1:]})
    add_para(doc, "Source: outputs/audit/step5_area_sweep.csv. SEC drops "
             "steeply 2 -> 8 m^2, knee at ~10 m^2, diminishing returns above. "
             "Going from 10 to 20 m^2 saves only ~10 % SEC at the cost of "
             "doubling collector capital, while clipping fraction passes "
             "30-40 % at the larger sizes.", italic=True)
    add_figure(doc, PLOTS / "step5_area_sweep.png",
               "Figure 6.1. E2 collector-area sweep across BTN/KTM/TPJ "
               "(annual). Left: SEC vs A_c. Centre: drying time vs A_c. "
               "Right: solar clipping fraction vs A_c.")

    add_heading(doc, "6.2 Season x location heatmap (E2)", level=2)
    sl = pd.read_csv(AUDIT / "step5_season_location.csv")
    sl_pv = sl.pivot_table(index="location", columns="season", values="SEC")
    cols_order = [c for c in ["winter_dec_jan", "spring_mar_apr",
                              "autumn_oct_nov", "annual"] if c in sl_pv.columns]
    sl_pv = sl_pv[cols_order].reset_index()
    add_table_from_df(doc, sl_pv,
                      num_fmt={c: "{:.4f}" for c in sl_pv.columns[1:]})
    add_para(doc, "Source: outputs/audit/step5_season_location.csv. Spring is "
             "the cheapest operating window (highest irradiance + moderate "
             "humidity). Annual values lie between extremes. The Kathmandu "
             "annual SEC is the highest among the three primary locations "
             "due to lower irradiance and altitude effects on dry-air "
             "density.", italic=True)
    add_figure(doc, PLOTS / "step5_season_heatmap.png",
               "Figure 6.2. E2 SEC heatmap, season x location.")

    # =========================================================================
    # 7. Numerical audit (anomaly hunt)
    # =========================================================================
    add_heading(doc, "7. Numerical audit and anomaly hunt", level=1)

    add_para(doc, "92 canonical CSVs (E1/E2/E3 x all locations x all seasons "
             "x {no-VPD, VPD}) were screened for seven classes of anomaly: "
             "(i) per-step first-law residual |Q_cond - (Q_evap + 0.95 W_comp)| "
             "/ expected > 2 %; (ii) chamber-inlet temperature deficit > "
             "0.5 K below T_set when not in transient or bypass; (iii) heat-"
             "exchanger oversizing flags; (iv) HP at-capacity flags; (v) VPD "
             "bypass churn (> 6 switches/h vs 600 s dwell); (vi) E1/E2 "
             "cross-config consistency (Q_HRX and Q_solar_usable must be "
             "identical); (vii) 72 h time-limit DNF.")

    al = pd.read_csv(AUDIT / "step6_anomaly_log.csv")
    add_heading(doc, "7.1 All flagged cases", level=2)
    add_table_from_df(doc, al)
    add_para(doc, "Source: outputs/audit/step6_anomaly_log.csv. "
             "Findings: zero first-law violations across 92 runs, zero "
             "capacity flags, zero T_chamber deficits, zero bypass-churn "
             "warnings, zero E1<->E2 energy-bookkeeping discrepancies. "
             "The single warn case (E2 / BTN / autumn / VPD) reaches 99 % "
             "of the moisture target at the 72 h cutoff and is a known "
             "humid-lowland stress case. The two info-level COP > 8 flags "
             "are E3 partial-lift transients at very low T_cond targets, "
             "which are physically valid (high Carnot ceiling).", italic=True)

    # =========================================================================
    # 8. Plot index
    # =========================================================================
    add_heading(doc, "8. Plot and data index", level=1)

    add_para(doc, "Every figure referenced in this document plus a few "
             "additional E2-vs-E3 deep dives.", italic=True)

    plot_table = pd.DataFrame([
        ("step3_energy_split_E1.png", "Energy split, E1, all 13 cases"),
        ("step3_energy_split_E2.png", "Energy split, E2, all 13 cases"),
        ("step3_energy_split_E3.png", "Energy split, E3, all 13 cases"),
        ("step2e_air_states_E1_kathmandu_annual.png",
         "E1 psychrometric trace, KTM annual"),
        ("step2e_air_states_E2_kathmandu_annual.png",
         "E2 psychrometric trace, KTM annual"),
        ("step2e_air_states_E3_kathmandu_annual.png",
         "E3 psychrometric trace, KTM annual"),
        ("step4_vpd_sweep.png",
         "VPD threshold sweep on E2 (KTM annual + BTN autumn)"),
        ("step5_area_sweep.png", "E2 collector area sweep, BTN/KTM/TPJ"),
        ("step5_season_heatmap.png", "E2 SEC heatmap, season x location"),
        ("e2_vs_e3_cumulative_kathmandu.png",
         "E2 vs E3 cumulative drying / energy, KTM annual"),
        ("e2_vs_e3_power_kathmandu.png",
         "E2 vs E3 instantaneous power streams, KTM annual"),
    ], columns=["filename (plots/_audit/)", "what it shows"])
    add_table_from_df(doc, plot_table, first_col_bold=False)

    data_table = pd.DataFrame([
        ("step3_energy_split.csv", "Per-case energy partition (39 rows)"),
        ("step4_pivot_SEC_no_vpd.csv", "SEC pivot, no-VPD (13 cases)"),
        ("step4_pivot_SEC_vpd.csv", "SEC pivot, VPD-on (12 cases)"),
        ("step4_winrate.csv", "Win count per KPI"),
        ("step4_vpd_sweep.csv", "VPD threshold sweep raw data"),
        ("step5_area_sweep.csv", "E2 collector area sweep"),
        ("step5_season_location.csv", "E2 season x location grid"),
        ("step6_anomaly_log.csv", "Anomaly hunt findings"),
    ], columns=["filename (outputs/audit/)", "what it contains"])
    add_para(doc, "")
    add_para(doc, "Backing data files:", bold=True)
    add_table_from_df(doc, data_table, first_col_bold=False)

    # Save
    doc.save(str(OUT_DOCX))
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
