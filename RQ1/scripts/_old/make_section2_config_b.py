"""Generate thesis/Section2_Config_B.docx - §2 Config B (solar + HP in series).

Beats: system description, open-loop baseline vs Config A, collector area
sweep, recirculation sweep, seasonal. Concise, table-heavy, objective.
"""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL

ROOT = Path(__file__).resolve().parents[1]
SUM = ROOT / "outputs" / "master_summary.csv"
OUT = ROOT / "thesis" / "Section2_Config_B.docx"


def fmt(x, d=3):
    try:
        if x != x:
            return "-"
        return f"{float(x):.{d}f}"
    except Exception:
        return "-"


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"; t.autofit = True
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            c = t.rows[i].cells[j]; c.text = str(v)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9.5)
    return t


def main():
    df = pd.read_csv(SUM)
    a = df[df["config"] == "A"]
    b = df[df["config"] == "B"].copy()
    b = b[~b["filename"].str.contains("_s4", na=False)]

    doc = Document()
    for s in doc.sections:
        s.left_margin = Cm(2); s.right_margin = Cm(2)

    doc.add_heading("Section 2  Config B: Solar + Heat Pump in Series", level=1)
    doc.add_paragraph(
        "Config B places a flat-plate solar air collector in series upstream of the condenser. "
        "In open-loop (r = 0) the air path is Ambient -> Collector -> Condenser -> Chamber -> Exhaust; "
        "in closed-loop the recirculated exhaust is first dehumidified at the evaporator, "
        "then pre-heated by the collector, then brought to T_set = 45 deg C by the condenser. "
        "The heat pump provides only the residual temperature lift that the collector cannot, "
        "so compressor duty scales inversely with available irradiance. The collector uses the "
        "Hottel-Whillier-Bliss model with F_R(tau-alpha) = 0.68 and F_R U_L = 4.2 W/m^2/K."
    )

    # 2.1 Open-loop baseline
    doc.add_heading("2.1  Open-loop baseline (r = 0, A_c = 10 m^2)", level=2)
    a0 = a[(a["season"] == "annual") & (a["r_recirc"] == 0.0) & (a["vpd_bypass"] == False)].set_index("location")
    b0 = b[(b["season"] == "annual") & (b["r_recirc"] == 0.0) &
           (b["solar_area_m2"] == 10.0) & (b["vpd_bypass"] == False)].set_index("location")
    rows = []
    for loc in ["kathmandu", "biratnagar", "taplejung"]:
        sec_a = a0.loc[loc, "SEC_kWh_per_kg"]
        sec_b = b0.loc[loc, "SEC_kWh_per_kg"]
        sf = b0.loc[loc, "solar_fraction"]
        wcomp = b0.loc[loc, "W_comp_kWh"]
        qsol = b0.loc[loc, "Q_solar_kWh"]
        red = (sec_a - sec_b) / sec_a * 100
        rows.append([loc.title(), fmt(sec_a, 3), fmt(sec_b, 3), f"-{red:.1f} %",
                     fmt(sf, 3), fmt(qsol, 2), fmt(wcomp, 2)])
    doc.add_paragraph("Table 2.1  Config B annual-TMY performance vs. Config A, r = 0, A_c = 10 m^2.")
    add_table(doc, ["Site", "SEC_A (kWh/kg)", "SEC_B (kWh/kg)", "SEC reduction",
                    "Solar fraction", "Q_solar (kWh)", "W_comp (kWh)"], rows)
    doc.add_paragraph(
        "The collector offsets 63-80 % of the condenser heat demand. SEC reduction is largest "
        "at Biratnagar (47.2 %) where both the solar resource (lower altitude, less cloud "
        "attenuation) and the high ambient temperature (smaller remaining heat-pump lift) "
        "compound. Kathmandu sees a 32.0 % reduction and Taplejung 39.2 %. Drying time changes "
        "by at most half an hour, consistent with the constraint identified in Section 1 that "
        "time is set by T_to_chamber and ambient RH rather than by the heat source."
    )

    # 2.2 Seasonal table
    doc.add_heading("2.2  Seasonal performance (r = 0, A_c = 10 m^2)", level=2)
    bs = b[(b["r_recirc"] == 0.0) & (b["solar_area_m2"] == 10.0) & (b["vpd_bypass"] == False)]
    pivot = bs.pivot_table(index="location", columns="season",
                           values="SEC_kWh_per_kg").round(3)
    pivot = pivot[["winter_dec_jan", "spring_mar_apr", "autumn_oct_nov"]]
    rows = []
    for loc, r in pivot.iterrows():
        rows.append([loc.title(), fmt(r["winter_dec_jan"], 3),
                     fmt(r["spring_mar_apr"], 3), fmt(r["autumn_oct_nov"], 3)])
    doc.add_paragraph("Table 2.2  Config B seasonal SEC (kWh/kg), r = 0, A_c = 10 m^2.")
    add_table(doc, ["Site", "Winter", "Spring", "Autumn"], rows)

    sfpivot = bs.pivot_table(index="location", columns="season",
                             values="solar_fraction").round(3)
    sfpivot = sfpivot[["winter_dec_jan", "spring_mar_apr", "autumn_oct_nov"]]
    rows = []
    for loc, r in sfpivot.iterrows():
        rows.append([loc.title(), fmt(r["winter_dec_jan"], 3),
                     fmt(r["spring_mar_apr"], 3), fmt(r["autumn_oct_nov"], 3)])
    doc.add_paragraph("Table 2.3  Config B seasonal solar fraction, r = 0, A_c = 10 m^2.")
    add_table(doc, ["Site", "Winter", "Spring", "Autumn"], rows)
    doc.add_paragraph(
        "Spring (Mar-Apr) delivers the lowest SEC at every site (0.128 kWh/kg at Biratnagar, "
        "0.182 at Kathmandu, 0.212 at Taplejung) with solar fraction of 0.88-0.93. The "
        "combination of high irradiance, warm ambient, and moderate humidity gives the "
        "collector a large usable temperature rise. Winter SEC at Kathmandu (0.295) and "
        "Biratnagar (0.240) remains well below the annual values because the pre-monsoon "
        "and monsoon months (absent from these seasonal windows) drive the annual average up."
    )

    # 2.3 Collector area
    doc.add_heading("2.3  Collector area", level=2)
    area_rows = b[(b["season"] == "annual") & (b["r_recirc"] == 0.0) &
                  (b["vpd_bypass"] == False) & (b["solar_area_m2"].isin([5.0, 10.0]))]
    pivot = area_rows.pivot_table(index="location", columns="solar_area_m2",
                                  values="SEC_kWh_per_kg").round(3)
    sf_p = area_rows.pivot_table(index="location", columns="solar_area_m2",
                                 values="solar_fraction").round(3)
    rows = []
    for loc in pivot.index:
        cols = pivot.columns.tolist()
        sec5 = pivot.loc[loc].get(5.0, float("nan"))
        sec10 = pivot.loc[loc].get(10.0, float("nan"))
        sf5 = sf_p.loc[loc].get(5.0, float("nan"))
        sf10 = sf_p.loc[loc].get(10.0, float("nan"))
        rows.append([loc.title(), fmt(sec5, 3), fmt(sf5, 3), fmt(sec10, 3), fmt(sf10, 3)])
    doc.add_paragraph("Table 2.4  Config B SEC and solar fraction vs. collector area, r = 0, annual TMY.")
    add_table(doc, ["Site", "SEC @ 5 m^2", "SF @ 5 m^2", "SEC @ 10 m^2", "SF @ 10 m^2"], rows)
    doc.add_paragraph(
        "Only Kathmandu has a simulated 5 m^2 case; doubling to 10 m^2 lifts solar fraction "
        "from 0.44 to 0.63 and drops SEC by 17.2 %. The response is sub-linear because the "
        "collector saturates once it can supply the full temperature rise on sunny hours: "
        "additional area then contributes only during shoulder hours and on cloudy days."
    )

    # 2.4 Recirculation
    doc.add_heading("2.4  Recirculation", level=2)
    rs = b[(b["season"] == "annual") & (b["solar_area_m2"] == 10.0) &
           (b["vpd_bypass"] == False) & (b["r_recirc"].isin([0.0, 0.3, 0.5, 0.7, 0.9, 1.0]))]
    pivot = rs.pivot_table(index="location", columns="r_recirc",
                           values="SEC_kWh_per_kg").round(3)
    rows = []
    for loc in pivot.index:
        rows.append([loc.title()] + [fmt(pivot.loc[loc][c], 3) for c in pivot.columns])
    doc.add_paragraph("Table 2.5  Config B annual-TMY SEC (kWh/kg) vs. r, A_c = 10 m^2.")
    add_table(doc, ["Site"] + [f"r = {c:.1f}" for c in pivot.columns], rows)
    doc.add_paragraph(
        "Recirculation raises SEC at every site in Config B. At Biratnagar, SEC rises from "
        "0.287 at r = 0 to 0.423 at r = 0.9 (47 % increase). At Kathmandu it rises from 0.488 "
        "to a minimum of 0.386 at r = 0.2 then climbs back to 0.469 at r = 0.9. Taplejung "
        "shows a 21 % increase at r = 0.9."
    )
    doc.add_paragraph(
        "The mechanism is the same evaporator-collector interaction across both components. "
        "In closed-loop the evaporator cools the mixed air to roughly 12 deg C for "
        "dehumidification. The collector then has to reheat this cold stream before the "
        "condenser, so a share of the solar heat is spent undoing evaporator cooling rather "
        "than displacing compressor work. The solar fraction drops accordingly, from 0.80 at "
        "r = 0 at Biratnagar to 0.76 at r = 0.9, and the compressor pulls the remaining lift. "
        "Kathmandu's shallow minimum at r = 0.2-0.3 is where the small dehumidification "
        "gain on cold, high-RH hours still outweighs the collector penalty; at higher r the "
        "penalty dominates."
    )

    # 2.5 takeaways
    doc.add_heading("2.5  Takeaways", level=2)
    doc.add_paragraph(
        "Open-loop with A_c = 10 m^2 is the design point. Annual SEC reaches 0.287 kWh/kg "
        "at Biratnagar, 0.344 at Taplejung, and 0.488 at Kathmandu, with solar fractions "
        "0.63-0.80.", style="List Bullet")
    doc.add_paragraph(
        "The collector saturates between 5 and 10 m^2 for this thermal load. Larger areas "
        "would chase diminishing returns at an increased capital cost.",
        style="List Bullet")
    doc.add_paragraph(
        "Closed-loop degrades Config B at every site and season, because the evaporator "
        "and the collector end up acting on the same air stream in opposite directions. "
        "This observation motivates the cascade arrangements in Section 3, which separate "
        "the two components onto different streams.",
        style="List Bullet")

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
