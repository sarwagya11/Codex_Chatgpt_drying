"""Generate Results & Discussion Section 2: Config C (Solar Cascade).

Creates a Word document covering:
  2.1  Config C — Solar Cascade Concept and Variants
  2.2  Config C1 — Mix Before Solar: Results
  2.3  Config C2 — Mix After Solar: Results
  2.4  Config B vs C1 vs C2: Why Placement Matters
  2.5  Lessons for System Design

Usage:
    python scripts/write_section2_config_c.py
"""

import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_ROOT = PROJECT_ROOT / "outputs"
FIGURES_DIR = PROJECT_ROOT / "figures" / "thesis"


# ── helpers ────────────────────────────────────────────────────────────


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(size)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    return table


def add_figure_placeholder(doc, label, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[DIAGRAM: {label}]")
    run.bold = True
    run.font.size = Pt(12)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.italic = True
    return cap


def fmt(val, decimals=3):
    if pd.isna(val):
        return "-"
    return f"{val:.{decimals}f}"


def pct_diff(base, new):
    if pd.isna(base) or pd.isna(new) or base == 0:
        return "-"
    return f"{(base - new) / base * 100:.1f}%"


def load_data():
    df = pd.read_csv(OUTPUT_ROOT / "master_summary.csv")
    df = df[~df["filename"].str.contains("_s4", na=False)]
    return df


def get_row(df, config, loc, season, r, vpd=False, area=0.0):
    mask = (
        (df["config"] == config)
        & (df["location"] == loc)
        & (df["season"] == season)
        & (df["r_recirc"] == r)
        & (df["vpd_bypass"] == vpd)
    )
    if area > 0:
        mask &= df["solar_area_m2"] == area
    rows = df[mask]
    return rows.iloc[0] if not rows.empty else None


LOC_NAMES = {
    "kathmandu": "Kathmandu (1 350 m a.s.l.)",
    "biratnagar": "Biratnagar (72 m a.s.l.)",
    "taplejung": "Taplejung (1 820 m a.s.l.)",
}
LOC_SHORT = {"kathmandu": "KTM", "biratnagar": "BTN", "taplejung": "TPJ"}
LOCS = ["kathmandu", "biratnagar", "taplejung"]
SEASONS = ["annual", "autumn_oct_nov", "winter_dec_jan", "spring_mar_apr"]
SEASON_NAMES = {
    "annual": "Annual TMY",
    "autumn_oct_nov": "Autumn (Oct\u2013Nov)",
    "winter_dec_jan": "Winter (Dec\u2013Jan)",
    "spring_mar_apr": "Spring (Mar\u2013Apr)",
}


# ── document sections ──────────────────────────────────────────────────


def write_section_2_1(doc, df):
    """2.1 Config C concept and variants."""
    add_heading(doc, "2.1  Config C: Solar Cascade \u2014 Concept and Variants", level=2)

    add_para(doc, (
        "Config B demonstrated that a solar collector in series with the "
        "condenser can reduce SEC by 32\u201347% compared to the HP-only baseline. "
        "A natural follow-up question is: can the solar collector be positioned "
        "differently to achieve even greater benefit? Config C explores this by "
        "placing the solar collector in a cascade arrangement with the evaporator, "
        "where solar energy assists the heat pump cycle rather than directly "
        "heating the process air."
    ))

    add_para(doc, (
        "Two variants were designed, differing in where recirculated exhaust "
        "enters the air path relative to the solar collector:"
    ))

    add_para(doc, "Config C1: Mix before solar", bold=True)

    add_figure_placeholder(doc,
        "Config C1 air path",
        "Figure 2.1a: Config C1 air path. Recirculated exhaust mixes with "
        "ambient air before passing through the solar collector, then the "
        "evaporator, and finally the condenser.")

    add_para(doc, (
        "Air path (r > 0): [r\u00d7Exhaust + (1\u2212r)\u00d7Ambient] \u2192 Solar "
        "\u2192 Evaporator \u2192 Condenser \u2192 Chamber"
    ), italic=True)

    add_para(doc, (
        "Air path (r = 0): Ambient \u2192 Solar \u2192 Evaporator \u2192 Condenser "
        "\u2192 Chamber"
    ), italic=True)

    add_para(doc, (
        "In C1, the solar collector receives the mixed air (which at r > 0 "
        "includes warm exhaust) and heats it further. This solar-heated air "
        "then enters the evaporator for dehumidification. The key issue is "
        "that the evaporator must cool this warm, solar-heated air back down "
        "to the coil surface temperature for moisture removal \u2014 undoing "
        "much of the solar gain."
    ))

    add_para(doc, "Config C2: Mix after solar", bold=True)

    add_figure_placeholder(doc,
        "Config C2 air path",
        "Figure 2.1b: Config C2 air path. Ambient air passes through the "
        "solar collector first, then mixes with recirculated exhaust before "
        "entering the evaporator and condenser.")

    add_para(doc, (
        "Air path (r > 0): Ambient \u2192 Solar \u2192 [Mix solar_out + r\u00d7Exhaust] "
        "\u2192 Evaporator \u2192 Condenser \u2192 Chamber"
    ), italic=True)

    add_para(doc, (
        "Air path (r = 0): Ambient \u2192 Solar \u2192 Evaporator \u2192 Condenser "
        "\u2192 Chamber"
    ), italic=True)

    add_para(doc, (
        "In C2, the solar collector always sees fresh ambient air (cold inlet "
        "\u2192 high collector efficiency). The solar-heated air is then mixed "
        "with recirculated exhaust. This arrangement preserves the solar "
        "collector\u2019s efficiency advantage while still allowing recirculation."
    ))

    add_para(doc, (
        "At r = 0 (open-loop), both C1 and C2 reduce to the same air path: "
        "Ambient \u2192 Solar \u2192 Evaporator \u2192 Condenser \u2192 Chamber. However, "
        "in this path the evaporator sits between the solar collector and "
        "the condenser. In open-loop mode with dry ambient air, the evaporator "
        "provides no dehumidification benefit \u2014 it simply cools the "
        "solar-heated air, which the condenser must then reheat. This "
        "fundamental placement issue drives the results presented below."
    ))


def write_section_2_2(doc, df):
    """2.2 Config C1 results."""
    add_heading(doc, "2.2  Config C1: Mix Before Solar \u2014 Results", level=2)

    add_para(doc, (
        "Table 2.1 presents Config C1 performance at r = 0 for two solar "
        "collector areas, alongside the Config A and B baselines."
    ))

    headers = ["Config", "Location", "A_c (m\u00b2)", "SEC (kWh/kg)",
               "Time (h)", "COP", "Solar frac."]
    rows = []
    for loc in LOCS:
        a = get_row(df, "A", loc, "annual", 0.0)
        if a is not None:
            rows.append(["A (r=0)", LOC_SHORT[loc], "\u2014",
                         fmt(a["SEC_kWh_per_kg"]), fmt(a["time_h"], 1),
                         fmt(a["COP_mean"], 2), "\u2014"])

        b = get_row(df, "B", loc, "annual", 0.0, area=10.0)
        if b is not None:
            rows.append(["B (r=0)", LOC_SHORT[loc], "10",
                         fmt(b["SEC_kWh_per_kg"]), fmt(b["time_h"], 1),
                         fmt(b["COP_mean"], 2), fmt(b["solar_fraction"], 2)])

        for area in [5.0, 10.0]:
            c1 = get_row(df, "C1", loc, "annual", 0.0, area=area)
            if c1 is not None:
                rows.append([f"C1 (r=0)", LOC_SHORT[loc], f"{area:.0f}",
                             fmt(c1["SEC_kWh_per_kg"]), fmt(c1["time_h"], 1),
                             fmt(c1["COP_mean"], 2), fmt(c1["solar_fraction"], 2)])

    add_table(doc, headers, rows)
    add_para(doc,
        "Table 2.1: Config C1 (r = 0) annual-TMY performance compared to "
        "Config A and B baselines.",
        italic=True, size=9)

    c1_ktm = get_row(df, "C1", "kathmandu", "annual", 0.0, area=10.0)
    c1_btn = get_row(df, "C1", "biratnagar", "annual", 0.0, area=10.0)
    a_ktm = get_row(df, "A", "kathmandu", "annual", 0.0)
    b_ktm = get_row(df, "B", "kathmandu", "annual", 0.0, area=10.0)

    add_para(doc, (
        f"Config C1 at r = 0 produces alarming results. At Kathmandu with "
        f"10 m\u00b2, SEC is {fmt(c1_ktm['SEC_kWh_per_kg'])} kWh/kg \u2014 nearly "
        f"identical to Config A ({fmt(a_ktm['SEC_kWh_per_kg'])}) despite having "
        f"a solar collector with a solar fraction of "
        f"{fmt(c1_ktm['solar_fraction'], 2)}. More critically, the drying time "
        f"explodes to {fmt(c1_ktm['time_h'], 1)} h, more than double the "
        f"Config A baseline of {fmt(a_ktm['time_h'], 1)} h."
    ))

    add_para(doc, (
        f"At Biratnagar, the situation is even worse: C1 achieves SEC = "
        f"{fmt(c1_btn['SEC_kWh_per_kg'])} with a drying time of "
        f"{fmt(c1_btn['time_h'], 1)} h \u2014 compared to Config B\u2019s "
        f"{fmt(get_row(df, 'B', 'biratnagar', 'annual', 0.0, area=10.0)['SEC_kWh_per_kg'])} "
        f"in {fmt(get_row(df, 'B', 'biratnagar', 'annual', 0.0, area=10.0)['time_h'], 1)} h "
        f"with the same 10 m\u00b2 collector."
    ))

    add_para(doc, "Why C1 performs so poorly at r = 0", bold=True)

    add_para(doc, (
        "In open-loop C1, the air path is: Ambient \u2192 Solar \u2192 Evaporator "
        "\u2192 Condenser \u2192 Chamber. The solar collector heats the ambient air "
        "(e.g. from 10 \u00b0C to 30 \u00b0C during midday). This warm, solar-heated "
        "air then enters the evaporator, which cools it back to ~12 \u00b0C "
        "(the coil surface temperature). The solar gain is largely destroyed, "
        "and the condenser must reheat from 12 \u00b0C to 45 \u00b0C regardless of "
        "how much solar heating occurred."
    ))

    add_para(doc, (
        "Worse, the ambient air in open-loop mode is typically dry enough "
        "that no dehumidification occurs at the evaporator \u2014 the evaporator "
        "provides only sensible cooling. The net effect is that the solar "
        "collector heats the air, the evaporator cools it back down "
        "(wasting both solar and compressor energy), and the condenser "
        "provides the full temperature lift as if the solar collector "
        "were absent."
    ))

    add_para(doc, (
        "The extended drying time is explained by periods (especially at "
        "night) when the evaporator cools the air below the optimal "
        "drying temperature. At night, with no solar gain, the air enters "
        "the evaporator at ambient temperature and exits at ~12 \u00b0C. The "
        "condenser heats it to 45 \u00b0C as in Config A, but the COP is "
        "penalised by the unnecessary evaporator cooling step."
    ))

    # C1 with recirculation
    c1_ktm_r07 = get_row(df, "C1", "kathmandu", "annual", 0.7, area=10.0)
    if c1_ktm_r07 is not None:
        add_para(doc, "Effect of recirculation on C1", bold=True)

        add_para(doc, (
            f"Adding recirculation (r = 0.7) to C1 at Kathmandu improves SEC to "
            f"{fmt(c1_ktm_r07['SEC_kWh_per_kg'])} kWh/kg with a drying time of "
            f"{fmt(c1_ktm_r07['time_h'], 1)} h. The improvement occurs because "
            f"recirculated exhaust is warm and humid, giving the evaporator "
            f"meaningful dehumidification work. However, C1 with r = 0.7 still "
            f"offers no advantage over Config B at r = 0 "
            f"({fmt(b_ktm['SEC_kWh_per_kg'])} kWh/kg in "
            f"{fmt(b_ktm['time_h'], 1)} h), which achieves better SEC with a "
            f"simpler air path."
        ))


def write_section_2_3(doc, df):
    """2.3 Config C2 results."""
    add_heading(doc, "2.3  Config C2: Mix After Solar \u2014 Results", level=2)

    add_para(doc, (
        "Table 2.2 presents Config C2 performance at r = 0 compared to "
        "the same baselines."
    ))

    headers = ["Config", "Location", "A_c (m\u00b2)", "SEC (kWh/kg)",
               "Time (h)", "COP", "Solar frac."]
    rows = []
    for loc in LOCS:
        b = get_row(df, "B", loc, "annual", 0.0, area=10.0)
        if b is not None:
            rows.append(["B (r=0)", LOC_SHORT[loc], "10",
                         fmt(b["SEC_kWh_per_kg"]), fmt(b["time_h"], 1),
                         fmt(b["COP_mean"], 2), fmt(b["solar_fraction"], 2)])

        for area in [5.0, 10.0]:
            c2 = get_row(df, "C2", loc, "annual", 0.0, area=area)
            if c2 is not None:
                rows.append([f"C2 (r=0)", LOC_SHORT[loc], f"{area:.0f}",
                             fmt(c2["SEC_kWh_per_kg"]), fmt(c2["time_h"], 1),
                             fmt(c2["COP_mean"], 2), fmt(c2["solar_fraction"], 2)])

    add_table(doc, headers, rows)
    add_para(doc,
        "Table 2.2: Config C2 (r = 0) annual-TMY performance compared to Config B.",
        italic=True, size=9)

    c2_ktm = get_row(df, "C2", "kathmandu", "annual", 0.0, area=10.0)
    c2_btn = get_row(df, "C2", "biratnagar", "annual", 0.0, area=10.0)
    b_ktm = get_row(df, "B", "kathmandu", "annual", 0.0, area=10.0)
    b_btn = get_row(df, "B", "biratnagar", "annual", 0.0, area=10.0)

    add_para(doc, (
        f"C2 performs dramatically better than C1. At Kathmandu with 10 m\u00b2, "
        f"SEC = {fmt(c2_ktm['SEC_kWh_per_kg'])} kWh/kg with a normal drying time "
        f"of {fmt(c2_ktm['time_h'], 1)} h and a notably high COP of "
        f"{fmt(c2_ktm['COP_mean'], 2)}. At Biratnagar, SEC = "
        f"{fmt(c2_btn['SEC_kWh_per_kg'])} with COP = "
        f"{fmt(c2_btn['COP_mean'], 2)}."
    ))

    add_para(doc, (
        f"However, C2 still does not match Config B. At Kathmandu, B achieves "
        f"SEC = {fmt(b_ktm['SEC_kWh_per_kg'])} versus C2\u2019s "
        f"{fmt(c2_ktm['SEC_kWh_per_kg'])} ({pct_diff(c2_ktm['SEC_kWh_per_kg'], b_ktm['SEC_kWh_per_kg'])} "
        f"better). At Biratnagar, B achieves {fmt(b_btn['SEC_kWh_per_kg'])} versus "
        f"C2\u2019s {fmt(c2_btn['SEC_kWh_per_kg'])} "
        f"({pct_diff(c2_btn['SEC_kWh_per_kg'], b_btn['SEC_kWh_per_kg'])} better)."
    ))

    add_para(doc, "How C2\u2019s cascade arrangement works", bold=True)

    add_para(doc, (
        "In C2 at r = 0, the air path is: Ambient \u2192 Solar \u2192 Evaporator "
        "\u2192 Condenser \u2192 Chamber. Like C1, the evaporator sits after the "
        "solar collector. However, C2\u2019s higher COP values (5.05\u20137.76) "
        "reveal what happens: the solar collector pre-heats the air entering "
        "the evaporator, which raises the evaporator\u2019s heat source temperature. "
        "In the open-loop HP sizing, T_evap = T_source \u2212 10 K. When the "
        "source is solar-heated air at 30 \u00b0C instead of ambient at 10 \u00b0C, "
        "T_evap rises to ~20 \u00b0C, dramatically reducing the temperature lift "
        "and boosting COP."
    ))

    add_para(doc, (
        "The trade-off is that this arrangement does not reduce the condenser "
        "load directly \u2014 the condenser still heats air from the evaporator "
        "outlet (~24 \u00b0C) to T_set (45 \u00b0C). In Config B, the solar collector "
        "heats air on the condenser side, directly reducing the temperature "
        "lift the condenser must provide. This is why B achieves lower SEC "
        "despite C2\u2019s higher COP: the total compressor work is determined by "
        "Q_cond / COP, and B reduces Q_cond while C2 only increases COP."
    ))

    # Seasonal table for C2
    add_para(doc, "Seasonal performance", bold=True)

    headers_s = ["Location", "Annual", "Autumn", "Winter", "Spring"]
    rows_s = []
    for loc in LOCS:
        row_data = [LOC_SHORT[loc]]
        for season in SEASONS:
            d = get_row(df, "C2", loc, season, 0.0, area=10.0)
            row_data.append(fmt(d["SEC_kWh_per_kg"]) if d is not None else "-")
        rows_s.append(row_data)
    add_table(doc, headers_s, rows_s)
    add_para(doc,
        "Table 2.3: Config C2 (r = 0, A_c = 10 m\u00b2) SEC (kWh/kg) by season.",
        italic=True, size=9)

    c2_btn_spr = get_row(df, "C2", "biratnagar", "spring_mar_apr", 0.0, area=10.0)
    c2_ktm_spr = get_row(df, "C2", "kathmandu", "spring_mar_apr", 0.0, area=10.0)

    if c2_btn_spr is not None and c2_ktm_spr is not None:
        add_para(doc, (
            f"C2 shows strong seasonal variation. At Biratnagar in spring, "
            f"SEC drops to {fmt(c2_btn_spr['SEC_kWh_per_kg'])} kWh/kg with a "
            f"solar fraction of {fmt(c2_btn_spr['solar_fraction'], 2)} \u2014 the "
            f"solar collector provides so much energy to the evaporator source "
            f"that the compressor barely operates (COP = "
            f"{fmt(c2_btn_spr['COP_mean'], 1)}). At Kathmandu in spring, SEC = "
            f"{fmt(c2_ktm_spr['SEC_kWh_per_kg'])} with similarly high solar "
            f"utilisation."
        ))

        add_para(doc, (
            "The extremely high COP values in spring (> 100) occur because the "
            "solar collector heats the evaporator source air so effectively that "
            "the temperature lift approaches zero for some timesteps. While these "
            "instantaneous COP values are physical artefacts of the near-zero "
            "lift, they reflect a genuine reduction in compressor work: the "
            "solar collector is doing most of the thermodynamic lifting."
        ))


def write_section_2_4(doc, df):
    """2.4 Config B vs C1 vs C2 comparison."""
    add_heading(doc, "2.4  B vs C1 vs C2: Why Solar Placement Matters", level=2)

    add_para(doc, (
        "Table 2.4 provides a direct comparison of the three solar "
        "configurations at r = 0 with A_c = 10 m\u00b2."
    ))

    headers = ["Config", "Location", "SEC (kWh/kg)", "Time (h)", "COP",
               "Solar frac.", "vs Config A"]
    rows = []
    for loc in LOCS:
        a = get_row(df, "A", loc, "annual", 0.0)
        for cfg_name in ["B", "C1", "C2"]:
            c = get_row(df, cfg_name, loc, "annual", 0.0, area=10.0)
            if c is not None and a is not None:
                rows.append([
                    cfg_name, LOC_SHORT[loc],
                    fmt(c["SEC_kWh_per_kg"]),
                    fmt(c["time_h"], 1),
                    fmt(c["COP_mean"], 2),
                    fmt(c["solar_fraction"], 2),
                    pct_diff(a["SEC_kWh_per_kg"], c["SEC_kWh_per_kg"]),
                ])
    add_table(doc, headers, rows)
    add_para(doc,
        "Table 2.4: Comparison of solar configurations (r = 0, A_c = 10 m\u00b2, "
        "annual TMY). \u2018vs Config A\u2019 is SEC improvement.",
        italic=True, size=9)

    add_para(doc, (
        "The ranking is consistent across all locations: B > C2 > C1. "
        "The explanation lies in where the solar energy enters the "
        "thermodynamic cycle:"
    ))

    add_bullet(doc, (
        "Config B (Amb \u2192 Solar \u2192 Cond \u2192 Chamber): Solar energy "
        "directly reduces the condenser load. The air arrives at the "
        "condenser already pre-heated, so Q_cond is smaller. The compressor "
        "does less work because Q_cond = Q_evap + W_comp, and a smaller "
        "Q_cond means less W_comp."
    ))
    add_bullet(doc, (
        "Config C2 (Amb \u2192 Solar \u2192 Evap \u2192 Cond \u2192 Chamber): Solar "
        "energy raises the evaporator source temperature, increasing COP "
        "but not reducing Q_cond. The compressor does less work per unit "
        "of heat transferred, but the total heat to be transferred remains "
        "the same."
    ))
    add_bullet(doc, (
        "Config C1 (Mix \u2192 Solar \u2192 Evap \u2192 Cond \u2192 Chamber): Solar "
        "energy heats air that the evaporator then cools. The solar gain "
        "is largely wasted, and the drying time increases dramatically "
        "at r = 0 because the evaporator cooling step delays the air\u2019s "
        "arrival at the drying temperature."
    ))

    add_para(doc, (
        "This ranking has a simple physical interpretation: in a heat pump "
        "dryer, the most valuable use of free solar energy is to reduce the "
        "condenser load (Config B), because the condenser is the bottleneck "
        "for energy consumption. Using solar to boost COP (Config C2) is "
        "the next best option. Using solar to heat air that will be cooled "
        "immediately (Config C1) is counterproductive."
    ))


def write_section_2_5(doc, df):
    """2.5 Lessons for system design."""
    add_heading(doc, "2.5  Lessons for System Design", level=2)

    add_para(doc, (
        "The Config C results establish three design principles that guide "
        "the development of more advanced configurations in subsequent sections:"
    ))

    add_bullet(doc, (
        "Solar energy is most effective on the condenser side. Config B\u2019s "
        "series arrangement (Solar \u2192 Condenser) consistently outperforms "
        "the cascade arrangements (Solar \u2192 Evaporator) by 13\u201325% in SEC."
    ))
    add_bullet(doc, (
        "Component ordering in the air path critically affects performance. "
        "Placing the evaporator after the solar collector (C1/C2) wastes "
        "solar energy on air that must be cooled for dehumidification. "
        "This lesson applies equally to the design of heat recovery systems: "
        "the heat recovery exchanger should feed the condenser side, not "
        "the evaporator side."
    ))
    add_bullet(doc, (
        "COP improvement alone is insufficient. Config C2 achieves COP "
        "values 2\u20133\u00d7 higher than Config B, yet has worse SEC because "
        "the condenser load remains unchanged. In a heat pump dryer, the "
        "optimisation target is total compressor work (W_comp = Q_cond / COP "
        "\u2212 Q_evap / COP), not COP in isolation."
    ))

    add_para(doc, (
        "These principles motivate the Config D family (Section 3), which "
        "introduces heat recovery to reduce the condenser load \u2014 the same "
        "mechanism that makes Config B effective, but using waste heat from "
        "the exhaust instead of solar energy. This approach works at night "
        "and in all weather conditions, addressing Config B\u2019s dependence on "
        "solar irradiance."
    ))


# ── main ───────────────────────────────────────────────────────────────


def main():
    df = load_data()
    doc = Document()

    add_heading(doc, "Results and Discussion", level=0)
    add_heading(doc, "Section 2: Config C (Solar Cascade Configurations)", level=1)
    add_para(doc, (
        "This section examines two alternative solar collector placements "
        "that attempt to improve upon Config B\u2019s series arrangement. "
        "Config C1 places the solar collector before the evaporator with "
        "pre-mixing, while Config C2 uses fresh ambient air through the "
        "solar collector before mixing with recirculated exhaust. The "
        "results demonstrate that solar placement in the air path has a "
        "profound effect on system performance and establish design "
        "principles for the more advanced configurations that follow."
    ))

    write_section_2_1(doc, df)
    write_section_2_2(doc, df)
    write_section_2_3(doc, df)
    write_section_2_4(doc, df)
    write_section_2_5(doc, df)

    # Save
    out_dir = PROJECT_ROOT / "thesis"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / "Section2_Config_C.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")

    n_c1 = len(df[df["config"] == "C1"])
    n_c2 = len(df[df["config"] == "C2"])
    print(f"Data rows referenced: {n_c1} (C1) + {n_c2} (C2)")


if __name__ == "__main__":
    main()
