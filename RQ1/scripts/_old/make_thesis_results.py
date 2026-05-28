"""Generate thesis/Results_Discussion_v3.docx - full Results chapter, §1-§8.

Storytelling structure, concise, table-heavy, objective language.
D3 is excluded per scope decision. Data pulled live from master_summary.csv.
"""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL

ROOT = Path(__file__).resolve().parents[1]
SUM = ROOT / "outputs" / "master_summary.csv"
OUT = ROOT / "thesis" / "Results_Discussion_v3.docx"


def fmt(x, d=3):
    try:
        if x != x: return "-"
        return f"{float(x):.{d}f}"
    except Exception:
        return "-"


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"; t.autofit = True
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            c = t.rows[i].cells[j]; c.text = str(v)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9.5)
    return t


def para(doc, text, style=None):
    p = doc.add_paragraph(text)
    if style: p.style = style
    return p


def section1(doc, df):
    a = df[df["config"] == "A"]
    doc.add_heading("Section 1  Config A: Heat-Pump-Only Baseline", level=1)
    para(doc, "Config A is the reference case. A single-stage vapour-compression heat pump "
         "(R134a, eta_is = 0.75) heats ambient air to T_set = 45 deg C and dries 3.0 kg of "
         "apple slices (X0 = 6.5, X_f = 0.10 db) across ten trays. Every later configuration "
         "adds free-energy components on top of this baseline.")

    doc.add_heading("1.1  Air paths", level=2)
    para(doc, "Open-loop (r = 0).", style="List Bullet")
    para(doc, "Ambient -> Condenser -> Chamber -> Exhaust. Evaporator draws heat from a "
         "separate ambient stream; T_evap ~ T_amb - 10 K.", style="List Bullet")
    para(doc, "Closed-loop (r > 0).", style="List Bullet")
    para(doc, "Fraction r of exhaust mixes with fresh ambient, is dehumidified at the "
         "evaporator (fixed T_evap = 5 deg C), then reheated at the condenser. First law "
         "enforced: Q_cond = Q_evap + W_comp.", style="List Bullet")

    doc.add_heading("1.2  Open-loop baseline (r = 0)", level=2)
    ann0 = a[(a["season"] == "annual") & (a["r_recirc"] == 0.0) & (a["vpd_bypass"] == False)]
    rows = [[r["location"].title(), fmt(r["time_h"], 2), fmt(r["SEC_kWh_per_kg"], 3),
             fmt(r["SMER_kg_per_kWh"], 3), fmt(r["W_comp_kWh"], 2), fmt(r["COP_mean"], 2)]
            for _, r in ann0.iterrows()]
    para(doc, "Table 1.1  Config A annual-TMY performance, r = 0.")
    add_table(doc, ["Site", "Time (h)", "SEC (kWh/kg)", "SMER (kg/kWh)",
                    "W_comp (kWh)", "COP_mean"], rows)
    para(doc, "Annual SEC spans 0.543 kWh/kg at Biratnagar to 0.717 at Kathmandu, a 32 % "
         "range. Biratnagar's warmer ambient (T_amb ~ 25 deg C) gives a smaller condenser-"
         "to-evaporator lift and lifts COP from 3.63 to 4.43. The lower altitude (72 m vs. "
         "1 350 m) also raises air density and reduces time-integrated compressor load. "
         "Drying time sits in a narrow 13.9-14.5 h band because T_to_chamber is fixed at "
         "45 deg C, so the kinetic driving force is identical across sites.")

    seas = a[(a["r_recirc"] == 0.0) & (a["vpd_bypass"] == False) &
             (a["season"].isin(["autumn_oct_nov", "winter_dec_jan", "spring_mar_apr"]))]
    pivot = seas.pivot_table(index="location", columns="season",
                             values="SEC_kWh_per_kg").round(3)
    pivot = pivot[["winter_dec_jan", "spring_mar_apr", "autumn_oct_nov"]]
    rows = [[loc.title(), fmt(r["winter_dec_jan"], 3), fmt(r["spring_mar_apr"], 3),
             fmt(r["autumn_oct_nov"], 3)] for loc, r in pivot.iterrows()]
    para(doc, "Table 1.2  Config A seasonal SEC (kWh/kg), r = 0.")
    add_table(doc, ["Site", "Winter", "Spring", "Autumn"], rows)
    para(doc, "Seasonal SEC is 15-30 % below the annual-TMY value at every site. The "
         "annual TMY averages in monsoon months (high RH, alpha_RH suppression) which "
         "are absent from the three drying seasons (autumn, winter, spring) used here. "
         "These seasons match the post-harvest window for apple in Nepal.")

    doc.add_heading("1.3  Recirculation sweep", level=2)
    rs = a[(a["season"] == "annual") & (a["vpd_bypass"] == False) &
           (a["r_recirc"].isin([0.0, 0.3, 0.5, 0.7, 0.9, 1.0]))]
    pivot = rs.pivot_table(index="location", columns="r_recirc",
                           values="SEC_kWh_per_kg").round(3)
    rows = [[loc.title()] + [fmt(pivot.loc[loc][c], 3) for c in pivot.columns]
            for loc in pivot.index]
    para(doc, "Table 1.3  Config A annual-TMY SEC (kWh/kg) vs. recirculation fraction r.")
    add_table(doc, ["Site"] + [f"r = {c:.1f}" for c in pivot.columns], rows)
    para(doc, "At Kathmandu SEC falls from 0.717 (r = 0) to 0.669 (r = 0.9), a 6.7 % "
         "reduction. At Biratnagar SEC rises from 0.543 to 0.753 (39 % increase). "
         "Taplejung shows a 14 % increase at r = 0.9.")
    para(doc, "The crossover is set by the evaporator operating point. Open-loop T_evap "
         "tracks ambient (15 deg C at Biratnagar, 5 deg C at Kathmandu). Closed-loop "
         "T_evap is fixed at 5 deg C. At warm sites the fixed setpoint forces a lower "
         "saturation temperature than ambient requires, raising the pressure ratio and "
         "compressor work. At cool sites the fixed setpoint is close to the open-loop "
         "value, and the dehumidification benefit of mixing cooler air into the "
         "condenser inlet produces a net gain. Between r = 0.9 and r = 1.0 the change is "
         "under 1 %: full closure has no fresh-air dilution, so steady-state SEC is set "
         "by the evaporator's dehumidification rate.")

    doc.add_heading("1.4  Takeaways", level=2)
    para(doc, "Drying time is fixed by T_to_chamber and ambient humidity, not by "
         "configuration. Every configuration reaching 45 deg C chamber inlet dries in "
         "14-15 h under annual TMY. Differences between configurations appear in SEC "
         "and SMER.", style="List Bullet")
    para(doc, "SEC is site-driven. Biratnagar 0.543, Taplejung 0.566, Kathmandu 0.717. "
         "Any configuration's SEC must be read against its site.", style="List Bullet")
    para(doc, "Open-loop is the strongest baseline at warm sites. Closed-loop "
         "recirculation only reduces SEC when ambient is cold enough for the fixed "
         "T_evap = 5 deg C to sit at or below the open-loop evaporator setpoint.",
         style="List Bullet")


def section2(doc, df):
    a = df[df["config"] == "A"]
    b = df[df["config"] == "B"]
    doc.add_heading("Section 2  Config B: Solar + Heat Pump in Series", level=1)
    para(doc, "Config B places a flat-plate solar air collector in series upstream of the "
         "condenser. Open-loop path: Ambient -> Collector -> Condenser -> Chamber. The "
         "heat pump provides only the residual temperature lift that the collector "
         "cannot. Collector model: Hottel-Whillier-Bliss, F_R(tau-alpha) = 0.68, "
         "F_R U_L = 4.2 W/m^2/K.")

    doc.add_heading("2.1  Open-loop baseline, r = 0, A_c = 10 m^2", level=2)
    a0 = a[(a["season"] == "annual") & (a["r_recirc"] == 0.0) &
           (a["vpd_bypass"] == False)].set_index("location")
    b0 = b[(b["season"] == "annual") & (b["r_recirc"] == 0.0) &
           (b["solar_area_m2"] == 10.0) & (b["vpd_bypass"] == False)].set_index("location")
    rows = []
    for loc in ["kathmandu", "biratnagar", "taplejung"]:
        sec_a = a0.loc[loc, "SEC_kWh_per_kg"]
        sec_b = b0.loc[loc, "SEC_kWh_per_kg"]
        red = (sec_a - sec_b) / sec_a * 100
        rows.append([loc.title(), fmt(sec_a, 3), fmt(sec_b, 3), f"-{red:.1f} %",
                     fmt(b0.loc[loc, "solar_fraction"], 3),
                     fmt(b0.loc[loc, "Q_solar_kWh"], 2),
                     fmt(b0.loc[loc, "W_comp_kWh"], 2)])
    para(doc, "Table 2.1  Config B vs. Config A, annual TMY.")
    add_table(doc, ["Site", "SEC_A", "SEC_B", "Reduction", "Solar frac",
                    "Q_solar (kWh)", "W_comp (kWh)"], rows)
    para(doc, "The collector offsets 63-80 % of the condenser duty. SEC reduction is "
         "largest at Biratnagar (47.2 %) where the solar resource and warm ambient "
         "compound. Kathmandu 32.0 %, Taplejung 39.2 %.")

    doc.add_heading("2.2  Seasonal performance", level=2)
    bs = b[(b["r_recirc"] == 0.0) & (b["solar_area_m2"] == 10.0) & (b["vpd_bypass"] == False)]
    pivot = bs.pivot_table(index="location", columns="season", values="SEC_kWh_per_kg").round(3)
    pivot = pivot[["winter_dec_jan", "spring_mar_apr", "autumn_oct_nov"]]
    sfp = bs.pivot_table(index="location", columns="season", values="solar_fraction").round(3)
    sfp = sfp[["winter_dec_jan", "spring_mar_apr", "autumn_oct_nov"]]
    rows = [[loc.title(),
             f'{fmt(pivot.loc[loc]["winter_dec_jan"],3)} ({fmt(sfp.loc[loc]["winter_dec_jan"],2)})',
             f'{fmt(pivot.loc[loc]["spring_mar_apr"],3)} ({fmt(sfp.loc[loc]["spring_mar_apr"],2)})',
             f'{fmt(pivot.loc[loc]["autumn_oct_nov"],3)} ({fmt(sfp.loc[loc]["autumn_oct_nov"],2)})']
            for loc in pivot.index]
    para(doc, "Table 2.2  Config B seasonal SEC kWh/kg (solar fraction), r = 0, A_c = 10 m^2.")
    add_table(doc, ["Site", "Winter", "Spring", "Autumn"], rows)
    para(doc, "Spring gives the lowest SEC at every site (0.128 Biratnagar, 0.182 "
         "Kathmandu, 0.212 Taplejung) with solar fractions 0.88-0.93.")

    doc.add_heading("2.3  Collector area (Kathmandu, annual)", level=2)
    ka = b[(b["season"] == "annual") & (b["location"] == "kathmandu") &
           (b["r_recirc"] == 0.0) & (b["vpd_bypass"] == False) &
           (b["solar_area_m2"].isin([5.0, 10.0]))]
    rows = [[fmt(r["solar_area_m2"], 1), fmt(r["SEC_kWh_per_kg"], 3),
             fmt(r["solar_fraction"], 3)] for _, r in ka.iterrows()]
    para(doc, "Table 2.3  Config B Kathmandu SEC vs. collector area.")
    add_table(doc, ["A_c (m^2)", "SEC (kWh/kg)", "Solar frac"], rows)
    para(doc, "Doubling from 5 to 10 m^2 lifts solar fraction from 0.44 to 0.63 and "
         "drops SEC by 17.2 %. The response is sub-linear because the collector "
         "saturates once it can supply the full temperature rise on sunny hours.")

    doc.add_heading("2.4  Recirculation", level=2)
    rs = b[(b["season"] == "annual") & (b["solar_area_m2"] == 10.0) &
           (b["vpd_bypass"] == False) & (b["r_recirc"].isin([0.0, 0.3, 0.5, 0.7, 0.9]))]
    pivot = rs.pivot_table(index="location", columns="r_recirc", values="SEC_kWh_per_kg").round(3)
    rows = [[loc.title()] + [fmt(pivot.loc[loc][c], 3) for c in pivot.columns]
            for loc in pivot.index]
    para(doc, "Table 2.4  Config B annual-TMY SEC (kWh/kg) vs. r, A_c = 10 m^2.")
    add_table(doc, ["Site"] + [f"r = {c:.1f}" for c in pivot.columns], rows)
    para(doc, "Recirculation raises SEC at Biratnagar (+47 % at r = 0.9) and Taplejung "
         "(+21 %). Kathmandu shows a shallow minimum near r = 0.2-0.3 before climbing "
         "back up. The mechanism is that the evaporator cools the mixed air to ~12 "
         "deg C for dehumidification, and the collector then spends heat undoing that "
         "cooling rather than displacing compressor work. Solar fraction drops from "
         "0.80 to 0.76 at Biratnagar as r rises.")

    doc.add_heading("2.5  Takeaways", level=2)
    para(doc, "Open-loop with A_c = 10 m^2 is the design point: 0.287 Biratnagar, "
         "0.344 Taplejung, 0.488 Kathmandu.", style="List Bullet")
    para(doc, "The collector saturates between 5 and 10 m^2 for this thermal load.",
         style="List Bullet")
    para(doc, "Closed-loop degrades Config B because the evaporator and collector act "
         "on the same stream in opposite directions. This motivates the cascade "
         "arrangements in Section 3.", style="List Bullet")


def section3(doc, df):
    doc.add_heading("Section 3  Config C: Solar Cascade (C1 and C2)", level=1)
    para(doc, "Configs C1 and C2 separate the solar and heat-pump streams. In both, "
         "the solar collector heats air directly toward T_set; the heat pump only "
         "finishes the residual lift. The two variants differ in where ambient mixes "
         "with the recirculated exhaust.")
    para(doc, "C1: Mix -> Solar -> Evap -> Cond -> Chamber. Mixing happens before the "
         "collector, so the collector always operates on the recirculated (warmer) "
         "stream.", style="List Bullet")
    para(doc, "C2: Solar -> Mix -> Evap -> Cond -> Chamber. The collector sees ambient "
         "air only; mixing with exhaust happens after the collector, before the "
         "evaporator.", style="List Bullet")
    para(doc, "At r = 0 the two configurations are equivalent in topology (no "
         "recirculation to mix) and reduce to a solar-priority cascade with the "
         "evaporator drawing from the solar-heated stream.")

    doc.add_heading("3.1  Open-loop baseline (r = 0, A_c = 10 m^2)", level=2)
    c1 = df[(df["config"] == "C1") & (df["season"] == "annual") & (df["r_recirc"] == 0.0) &
            (df["solar_area_m2"] == 10.0) & (df["vpd_bypass"] == False)]
    c2 = df[(df["config"] == "C2") & (df["season"] == "annual") & (df["r_recirc"] == 0.0) &
            (df["solar_area_m2"] == 10.0) & (df["vpd_bypass"] == False)]
    rows = []
    for loc in ["kathmandu", "biratnagar", "taplejung"]:
        r1 = c1[c1["location"] == loc].iloc[0]
        r2 = c2[c2["location"] == loc].iloc[0]
        rows.append([loc.title(),
                     fmt(r1["SEC_kWh_per_kg"], 3), fmt(r1["time_h"], 1), fmt(r1["solar_fraction"], 3),
                     fmt(r2["SEC_kWh_per_kg"], 3), fmt(r2["time_h"], 1), fmt(r2["solar_fraction"], 3)])
    para(doc, "Table 3.1  C1 vs. C2 at r = 0, A_c = 10 m^2, annual TMY.")
    add_table(doc, ["Site", "C1 SEC", "C1 time (h)", "C1 SF",
                    "C2 SEC", "C2 time (h)", "C2 SF"], rows)
    para(doc, "C1 and C2 diverge at r = 0 because of how the evaporator is fed. In C1 "
         "the solar-heated air passes through the evaporator before the condenser. When "
         "solar irradiance is high the evaporator sees an already-hot stream it cannot "
         "effectively absorb heat from (the evap saturation temperature tracks its own "
         "supply), which stalls the cycle and extends drying time to 28-32 h. In C2 "
         "the solar stream goes straight to the chamber through the condenser while "
         "the evaporator draws from ambient; the heat pump runs on its normal envelope.")
    para(doc, "C2 delivers 0.381 kWh/kg at Biratnagar, 0.424 at Taplejung, 0.561 at "
         "Kathmandu, beating Config B at Kathmandu but not at Biratnagar or Taplejung. "
         "Solar fraction is slightly lower than Config B (0.59-0.76) because the "
         "evaporator rejects its heat to ambient instead of displacing condenser duty.")

    doc.add_heading("3.2  Takeaways", level=2)
    para(doc, "C1 is not operationally viable in this thermal envelope. Cascade "
         "placement that forces the evaporator to draw from a solar-preheated stream "
         "stalls the heat-pump cycle.", style="List Bullet")
    para(doc, "C2 preserves an independent evaporator envelope and performs similarly "
         "to Config B on an annual basis. It trades lower solar fraction for decoupled "
         "component operation.", style="List Bullet")
    para(doc, "Neither C-variant beats Config B at Biratnagar, which remains the "
         "benchmark at warm sites.", style="List Bullet")


def section4(doc, df):
    doc.add_heading("Section 4  Config D: Heat-Recovery Exchanger", level=1)
    para(doc, "Config D adds a counter-flow air-to-air plate heat exchanger (HRX, "
         "eps = 0.70, sensible-only) between exhaust and supply. No solar component. "
         "Two variants differ in how the evaporator is fed.")
    para(doc, "D1: Ambient -> HRX -> Condenser -> Chamber; Exhaust -> HRX -> expelled. "
         "Evaporator draws from ambient.", style="List Bullet")
    para(doc, "D2: same supply side as D1; evaporator draws from a dynamic mix of "
         "exhaust (post-HRX) and ambient, which raises its saturation temperature "
         "toward the exhaust-enthalpy side.", style="List Bullet")

    doc.add_heading("4.1  Open-loop baseline, r = 0", level=2)
    for cfg in ["D1", "D2"]:
        s = df[(df["config"] == cfg) & (df["season"] == "annual") & (df["vpd_bypass"] == False)]
        rows = [[r["location"].title(), fmt(r["time_h"], 2), fmt(r["SEC_kWh_per_kg"], 3),
                 fmt(r["W_comp_kWh"], 2), fmt(r["Q_HRX_kWh"], 2), fmt(r["COP_mean"], 2)]
                for _, r in s.iterrows()]
        para(doc, f"Table 4.{1 if cfg=='D1' else 2}  Config {cfg} annual-TMY, r = 0.")
        add_table(doc, ["Site", "Time (h)", "SEC", "W_comp (kWh)",
                        "Q_HRX (kWh)", "COP_mean"], rows)

    para(doc, "D1 achieves 0.293 at Biratnagar, 0.310 at Taplejung, 0.365 at Kathmandu. "
         "The HRX recovers 20-25 kWh of exhaust enthalpy per batch, offsetting a "
         "comparable share of the condenser duty. SEC reduction vs. Config A is "
         "46-49 % without any solar input.")
    para(doc, "D2 improves on D1 by 3-4 % across sites by feeding warmer air to the "
         "evaporator, raising T_evap_sat and reducing the pressure ratio. SEC at "
         "Biratnagar is 0.282; Kathmandu 0.354; Taplejung 0.299.")

    doc.add_heading("4.2  Seasonal (D2)", level=2)
    d2 = df[(df["config"] == "D2") & (df["vpd_bypass"] == False)]
    pivot = d2.pivot_table(index="location", columns="season", values="SEC_kWh_per_kg").round(3)
    cols = [c for c in ["winter_dec_jan", "spring_mar_apr", "autumn_oct_nov"] if c in pivot.columns]
    pivot = pivot[cols]
    rows = [[loc.title()] + [fmt(pivot.loc[loc][c], 3) for c in cols] for loc in pivot.index]
    para(doc, "Table 4.3  Config D2 seasonal SEC (kWh/kg), r = 0.")
    add_table(doc, ["Site", "Winter", "Spring", "Autumn"], rows)

    doc.add_heading("4.3  Takeaways", level=2)
    para(doc, "HRX alone (no solar) drops SEC by 46-49 % vs. Config A. This is the "
         "largest single-component gain in the study.", style="List Bullet")
    para(doc, "D2 improves on D1 by warming the evaporator supply; the gain is small "
         "(3-4 %) but consistent.", style="List Bullet")
    para(doc, "D1 and D2 preserve an open-loop path and do not trigger the "
         "evaporator-collector conflict that limits closed-loop Configs A and B.",
         style="List Bullet")


def section5(doc, df):
    doc.add_heading("Section 5  VPD Bypass Control", level=1)
    para(doc, "The vapour-pressure-deficit (VPD) bypass diverts a share of the process "
         "air around the drying chamber when the chamber humidity has fallen enough "
         "that continued drying becomes kinetically slow (VPD exceeds 0.05 kPa). The "
         "bypassed stream is expelled (in D/E, exhaust-side bypass) or returned to the "
         "condenser inlet (in A/B, condenser-direct bypass), cutting fan and compressor "
         "work during the tail of the batch. Activation oscillates with 3x hysteresis "
         "and 600 s dwell to avoid chatter.")

    doc.add_heading("5.1  Annual-TMY SEC reduction with VPD bypass", level=2)
    configs = [("A", 0.0, 0.9), ("B", 10.0, 0.9), ("D1", 0.0, 0.0),
               ("D2", 0.0, 0.0), ("E1", 10.0, 0.0), ("E2", 10.0, 0.0)]
    rows = []
    for cfg, area, r in configs:
        sub = df[(df["config"] == cfg) & (df["season"] == "annual") &
                 (df["r_recirc"] == r) & (df["solar_area_m2"] == area)]
        for loc in ["kathmandu", "biratnagar", "taplejung"]:
            off = sub[(sub["location"] == loc) & (sub["vpd_bypass"] == False)]
            on = sub[(sub["location"] == loc) & (sub["vpd_bypass"] == True)]
            if len(off) == 0 or len(on) == 0: continue
            o = off.iloc[0]["SEC_kWh_per_kg"]
            n = on.iloc[0]["SEC_kWh_per_kg"]
            red = (o - n) / o * 100
            label = cfg if r == 0.0 else f"{cfg} r={r}"
            rows.append([label, loc.title(), fmt(o, 3), fmt(n, 3), f"-{red:.1f} %"])
    para(doc, "Table 5.1  Annual-TMY SEC (kWh/kg) with and without VPD bypass.")
    add_table(doc, ["Config", "Site", "SEC off", "SEC on", "Reduction"], rows)

    para(doc, "VPD bypass reduces SEC by 15-32 % depending on configuration and site. "
         "The gain is largest where the tail-end compressor duty is a large share of "
         "total W_comp (HP-only and HP-dominant configurations). In the E family the "
         "absolute SEC reduction is smaller because compressor work is already low, "
         "but the relative gain remains 27-30 %.")
    para(doc, "Drying time extends by 5-10 % with bypass active (e.g., Config A r = 0.9 "
         "Kathmandu: 14.9 h -> 15.6 h). The bypass trades wall-clock time for energy; "
         "it is worth activating whenever the energy reduction outweighs the time "
         "cost for the downstream operation (packaging, cooling).")

    doc.add_heading("5.2  Takeaways", level=2)
    para(doc, "VPD bypass is a control-only modification with no hardware cost beyond "
         "a damper and sensor. It stacks multiplicatively with every other "
         "configuration improvement.", style="List Bullet")
    para(doc, "The largest absolute SEC wins are in HP-only (Config A) and HP+HRX "
         "(D1, D2) where compressor work dominates.", style="List Bullet")
    para(doc, "The largest relative wins are in the E family (see Section 6), where "
         "the bypass removes the last share of compressor work during solar-rich hours.",
         style="List Bullet")


def section6(doc, df):
    doc.add_heading("Section 6  Config E: HRX + Solar + Heat Pump", level=1)
    para(doc, "Config E combines the HRX recovery of Section 4 with the solar collector "
         "of Section 2. Three variants differ in the air-path topology:")
    para(doc, "E1: Ambient -> HRX -> Solar -> Cond -> Chamber. Evaporator on ambient. "
         "Exhaust -> HRX -> expelled.", style="List Bullet")
    para(doc, "E2: same supply path as E1; evaporator on iterated exhaust-ambient mix "
         "(like D2 vs. D1).", style="List Bullet")
    para(doc, "E3: Ambient -> HRX -> Cond -> Solar -> Chamber (solar after condenser). "
         "Solar-priority control: heat pump switches off when solar alone from HRX "
         "output meets T_set; otherwise HP provides partial lift with a variable T_cond.",
         style="List Bullet")

    doc.add_heading("6.1  Open-loop baseline (r = 0, A_c = 10 m^2)", level=2)
    rows = []
    for cfg in ["E1", "E2", "E3"]:
        s = df[(df["config"] == cfg) & (df["season"] == "annual") &
               (df["solar_area_m2"] == 10.0) & (df["vpd_bypass"] == False)]
        for loc in ["kathmandu", "biratnagar", "taplejung"]:
            row = s[s["location"] == loc]
            if len(row) == 0: continue
            r = row.iloc[0]
            rows.append([cfg, loc.title(), fmt(r["SEC_kWh_per_kg"], 3),
                         fmt(r["solar_fraction"], 3), fmt(r["W_comp_kWh"], 2),
                         fmt(r["time_h"], 2)])
    para(doc, "Table 6.1  Config E annual-TMY performance, r = 0, A_c = 10 m^2.")
    add_table(doc, ["Config", "Site", "SEC", "Solar frac", "W_comp (kWh)", "Time (h)"], rows)
    para(doc, "E2 is the strongest variant at every site: 0.129 Biratnagar, 0.163 "
         "Taplejung, 0.197 Kathmandu. E1 trails E2 by 6-10 % and E3 matches E1. Solar "
         "fractions reach 0.73-0.88. Compared to Config A these are 74-76 % SEC "
         "reductions on annual TMY.")
    para(doc, "E2's advantage over E1 comes from the iterative evaporator mix: warming "
         "the evap supply raises T_evap_sat and lowers pressure ratio, which shows up "
         "as lower W_comp at the same thermal load. E3's solar-after-condenser "
         "placement gives a comparable SEC to E1 but with a lower solar fraction, "
         "because in hours when solar alone meets T_set the HP idles and the HRX "
         "benefit is the dominant contributor.")

    doc.add_heading("6.2  E2 collector-area sweep", level=2)
    for loc in ["kathmandu", "biratnagar"]:
        s = df[(df["config"] == "E2") & (df["season"] == "annual") &
               (df["location"] == loc) & (df["vpd_bypass"] == False)]
        s = s.sort_values("solar_area_m2")
        rows = [[fmt(r["solar_area_m2"], 1), fmt(r["SEC_kWh_per_kg"], 3),
                 fmt(r["solar_fraction"], 3)] for _, r in s.iterrows()]
        para(doc, f"Table 6.{2 if loc=='kathmandu' else 3}  E2 SEC vs. A_c, {loc.title()} annual.")
        add_table(doc, ["A_c (m^2)", "SEC (kWh/kg)", "Solar frac"], rows)
    para(doc, "The E2 area response is steeper than Config B because the HRX has "
         "already reduced the thermal load that the collector must meet. At Biratnagar "
         "a 5 m^2 collector already gives 0.155 kWh/kg (solar fraction 0.77); doubling "
         "to 10 m^2 takes this to 0.129 and 0.88. Beyond 10 m^2 returns diminish: 15 "
         "m^2 drops SEC to 0.120 and 20 m^2 to 0.116 at Biratnagar. The knee of the "
         "curve sits at 8-10 m^2 for both sites tested.")

    doc.add_heading("6.3  E2 seasonal with VPD bypass (A_c = 10 m^2)", level=2)
    s = df[(df["config"] == "E2") & (df["solar_area_m2"] == 10.0) & (df["vpd_bypass"] == True)]
    pivot = s.pivot_table(index="location", columns="season", values="SEC_kWh_per_kg").round(3)
    cols = [c for c in ["winter_dec_jan", "spring_mar_apr", "autumn_oct_nov"] if c in pivot.columns]
    pivot = pivot[cols] if cols else pivot
    rows = [[loc.title()] + [fmt(pivot.loc[loc][c], 3) for c in pivot.columns]
            for loc in pivot.index]
    para(doc, "Table 6.4  E2 + VPD bypass seasonal SEC (kWh/kg).")
    add_table(doc, ["Site"] + [c.replace("_", " ").title() for c in pivot.columns], rows)
    para(doc, "With VPD bypass active, E2 at Biratnagar in spring reaches 0.072 kWh/kg "
         "with solar fraction 0.96, the lowest value recorded across the entire study. "
         "Annual-TMY values drop to 0.097 Biratnagar, 0.129 Taplejung, 0.144 Kathmandu.")

    doc.add_heading("6.4  Takeaways", level=2)
    para(doc, "E2 is the best-performing configuration in this study. The combination "
         "of HRX recovery, solar pre-heat, and iterative evaporator mixing cuts SEC to "
         "19-26 % of the Config A baseline.", style="List Bullet")
    para(doc, "VPD bypass on E2 gives a further 27-30 % relative reduction, reaching "
         "0.097 kWh/kg annual at Biratnagar.", style="List Bullet")
    para(doc, "Collector area beyond 10 m^2 shows diminishing returns; 8-10 m^2 is a "
         "reasonable design point for the 3 kg batch size used here.",
         style="List Bullet")


def section7(doc, df):
    doc.add_heading("Section 7  Grand Synthesis", level=1)
    para(doc, "This section compiles the best annual-TMY result from each configuration "
         "family and traces the Pareto frontier across SEC, capital complexity, and "
         "solar fraction.")

    picks = [
        ("A (r=0)", "A", {"season": "annual", "r_recirc": 0.0,
                           "solar_area_m2": 0.0, "vpd_bypass": False}),
        ("A (r=0.9, VPD)", "A", {"season": "annual", "r_recirc": 0.9,
                                  "solar_area_m2": 0.0, "vpd_bypass": True}),
        ("B (r=0, 10m^2)", "B", {"season": "annual", "r_recirc": 0.0,
                                  "solar_area_m2": 10.0, "vpd_bypass": False}),
        ("C2 (r=0, 10m^2)", "C2", {"season": "annual", "r_recirc": 0.0,
                                    "solar_area_m2": 10.0, "vpd_bypass": False}),
        ("D2 (r=0)", "D2", {"season": "annual", "r_recirc": 0.0,
                             "solar_area_m2": 0.0, "vpd_bypass": False}),
        ("D2 + VPD", "D2", {"season": "annual", "r_recirc": 0.0,
                             "solar_area_m2": 0.0, "vpd_bypass": True}),
        ("E2 (r=0, 10m^2)", "E2", {"season": "annual", "r_recirc": 0.0,
                                    "solar_area_m2": 10.0, "vpd_bypass": False}),
        ("E2 + VPD (10m^2)", "E2", {"season": "annual", "r_recirc": 0.0,
                                     "solar_area_m2": 10.0, "vpd_bypass": True}),
    ]
    rows = []
    for label, cfg, flt in picks:
        s = df[df["config"] == cfg].copy()
        for k, v in flt.items():
            s = s[s[k] == v]
        sk = s[s["location"] == "kathmandu"]
        sb = s[s["location"] == "biratnagar"]
        st = s[s["location"] == "taplejung"]
        rows.append([label,
                     fmt(sk.iloc[0]["SEC_kWh_per_kg"], 3) if len(sk) else "-",
                     fmt(sb.iloc[0]["SEC_kWh_per_kg"], 3) if len(sb) else "-",
                     fmt(st.iloc[0]["SEC_kWh_per_kg"], 3) if len(st) else "-"])
    para(doc, "Table 7.1  Annual-TMY SEC (kWh/kg) by configuration and site.")
    add_table(doc, ["Configuration", "Kathmandu", "Biratnagar", "Taplejung"], rows)

    doc.add_heading("7.1  Ranking", level=2)
    para(doc, "At every site the ranking is the same: E2+VPD < E2 < D2+VPD < D2 < "
         "B < C2 < A+VPD < A. The addition of HRX (D family) delivers a 46-49 % SEC "
         "reduction; adding solar on top of HRX (E family) delivers a further 45-55 % "
         "reduction; adding VPD bypass on top of E2 delivers another 27-30 %. Each "
         "lever is independent and stacks.")

    doc.add_heading("7.2  Component contribution summary", level=2)
    contrib = [
        ("HRX (sensible-only, eps = 0.70)", "46-49 %", "D1/D2 vs A", "Recovers exhaust enthalpy"),
        ("Solar collector (10 m^2)", "32-47 %", "B vs A", "Pre-heats supply air"),
        ("HRX + solar in series", "71-76 %", "E2 vs A", "Levers stack"),
        ("VPD bypass (control only)", "15-32 %", "any cfg on/off", "Cuts tail-end duty"),
        ("Closed-loop recirculation", "-39 % to +7 %", "r=0.9 vs r=0", "Site-dependent"),
    ]
    para(doc, "Table 7.2  Annual-TMY SEC contribution of each lever.")
    add_table(doc, ["Component / lever", "SEC reduction", "Comparison", "Mechanism"], contrib)

    doc.add_heading("7.3  Best-of-site", level=2)
    para(doc, "The lowest annual-TMY SEC recorded in this study is 0.097 kWh/kg at "
         "Biratnagar (E2 + VPD, A_c = 10 m^2). Taplejung's best is 0.129, Kathmandu's "
         "is 0.144. The 33 % spread between best and worst site in E2 mirrors the "
         "32 % spread seen in Config A, confirming that site-driven climate effects "
         "persist through every layer of the configuration stack.")


def section8(doc, df):
    doc.add_heading("Section 8  Limitations and Scope", level=1)
    para(doc, "The results above rest on a set of modelling choices whose validity "
         "bounds the applicability of the conclusions.")

    doc.add_heading("8.1  Kinetic model", level=2)
    para(doc, "Drying kinetics use a parametric Midilli form (K_ref = 1.63e-4 /s, "
         "Ea/R = 2711 K, alpha_RH = 1.75) fit to apple-slice data at R^2 = 0.90. The "
         "humidity-suppression exponent alpha_RH is the single largest sensitivity in "
         "the model. Under temperature-limited kinetics (alpha_RH < 1) the evaporator "
         "strategy and VPD bypass conclusions in Sections 1 and 5 would not hold "
         "unchanged.")

    doc.add_heading("8.2  Heat-pump envelope", level=2)
    para(doc, "The R134a cycle assumes eta_is = 0.75, a fixed T_cond = 55 deg C in the "
         "non-variable-T_cond configurations, and no frost on the evaporator. Operation "
         "was restricted to T_evap >= -5 deg C and pressure ratio <= 10. No cases in "
         "the reported results violated these bounds, but operation outside this "
         "envelope (colder ambient, very low solar resource) would require reshaping "
         "the cycle.")

    doc.add_heading("8.3  HRX assumptions", level=2)
    para(doc, "The HRX is modelled as sensible-only with eps = 0.70 (counter-flow "
         "plate). No condensation is allowed on the exhaust side. In practice a "
         "humidity-heavy exhaust would condense on the HRX plates at Kathmandu and "
         "Taplejung winter conditions. This would raise the effective UA and slightly "
         "improve D/E performance beyond the reported values but also introduce water "
         "management complexity not modelled here.")

    doc.add_heading("8.4  Solar resource", level=2)
    para(doc, "Irradiance comes from TMY data at the three sites (NASA POWER, 2024 "
         "retrieval). Day-to-day variability is captured within each TMY, but "
         "multi-year variability and cloud-enhancement events are not. Collector "
         "thermal model is steady-state Hottel-Whillier-Bliss without explicit heat "
         "capacity, so transient performance on partly-cloudy hours is smoothed.")

    doc.add_heading("8.5  Experimental validation", level=2)
    para(doc, "No experimental validation has been performed to date. The sizing "
         "document (SAHPD_Sizing.docx) specifies a 1.5-ton reversible inverter split "
         "AC and a 400 m^3/h counter-flow HRX as the reference build. Experimental "
         "comparison will be restricted to Config A (HP-only) and Config D2 (HRX) in "
         "the first phase; Config E requires the solar collector to be commissioned.")

    doc.add_heading("8.6  Scope notes", level=2)
    para(doc, "This chapter does not report Config D3 (HRX with swapped hot-cold "
         "streams). D3 carries a humidity-transfer risk not present in D1/D2 and was "
         "excluded from the comparative analysis. A separate internal note retains the "
         "D3 results for reference.", style="List Bullet")
    para(doc, "Thermal storage, an obvious extension to the E family, is not "
         "modelled. Storage would shift the seasonal SEC floor and is flagged for "
         "a future iteration of the study.", style="List Bullet")
    para(doc, "The study uses a 3.0 kg batch on ten trays. Scale-up to industrial "
         "throughput would require reassessment of chamber geometry, airflow "
         "distribution, and HRX face area.", style="List Bullet")


def main():
    df = pd.read_csv(SUM)
    df = df[~df["filename"].str.contains("_s4", na=False)]

    doc = Document()
    for s in doc.sections:
        s.left_margin = Cm(2); s.right_margin = Cm(2)

    doc.add_heading("Results and Discussion", level=0)
    para(doc, "This chapter reports the simulation results for ten SAHPD configurations "
         "across three Nepalese sites (Kathmandu, Biratnagar, Taplejung) under annual-"
         "TMY and three drying-season windows. Each section builds on the previous: "
         "Config A establishes the heat-pump-only baseline; Config B adds a solar "
         "collector in series; Config C explores two solar cascade topologies; "
         "Config D adds a heat-recovery exchanger; Section 5 introduces the VPD-based "
         "bypass control used across configurations; Config E combines HRX and solar; "
         "Section 7 synthesises across configurations; Section 8 states limitations. "
         "All tables draw live from outputs/master_summary.csv. Operating parameters: "
         "T_set = 45 deg C, m_p_dry = 3.0 kg, X0 = 6.5, X_f = 0.10 db, v = 1.1 m/s, "
         "m_da = 0.098 kg/s.")

    section1(doc, df)
    section2(doc, df)
    section3(doc, df)
    section4(doc, df)
    section5(doc, df)
    section6(doc, df)
    section7(doc, df)
    section8(doc, df)

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
