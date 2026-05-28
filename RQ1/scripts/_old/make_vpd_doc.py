"""Generate VPD strategy explanation Word document."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading("VPD-Based Condenser-Direct Bypass Strategy", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("Technical Document - SAHPD RQ1 Project")
doc.add_paragraph("Author: Wasti Sarwar")
doc.add_paragraph("Date: March 2026")
doc.add_paragraph("")

# ==== Section 1 ====
doc.add_heading("1. Problem Statement", level=1)
doc.add_paragraph(
    "In a closed-loop heat pump dryer (HPCD) with recirculation, exhaust air from the drying chamber "
    "is mixed with fresh ambient air and passed through an evaporator coil for dehumidification before "
    "being reheated by the condenser to T_set before entering the chamber. This evaporator-based "
    "dehumidification is essential when the recirculated air carries significant moisture from the product."
)
doc.add_paragraph(
    "However, as drying progresses, the product moisture content decreases and the drying rate falls "
    "dramatically. In the late stages of drying (the falling-rate period), the exhaust air is already "
    "quite dry - its humidity ratio approaches that of ambient air. At this point, running the air through "
    "the evaporator provides minimal dehumidification benefit, yet the compressor still consumes "
    "significant energy to cool and reheat the air. Analysis of simulation data shows that the final "
    "4-6 hours of drying (representing only ~4% of total water removal) can consume 30-44% of total "
    "compressor energy."
)

doc.add_heading("1.1 The Condenser Penalty Concept", level=2)
doc.add_paragraph(
    "To quantify when the evaporator becomes energetically wasteful, we define the Condenser Penalty "
    "Fraction (CPF):"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CPF = (VPD_post_evap - VPD_exhaust) / VPD_post_evap")
run.italic = True

doc.add_paragraph(
    "where VPD (Vapour Pressure Deficit) is the difference between the saturation vapour pressure "
    "at the chamber temperature T_set and the actual vapour pressure of the air stream, "
    "evaluated at the drying temperature. VPD directly measures the air's drying potential - "
    "its ability to absorb moisture from the product surface."
)
doc.add_paragraph(
    "VPD_post_evap is the drying potential of air after evaporator dehumidification (the normal path), "
    "while VPD_exhaust is the drying potential of the exhaust air if sent directly to the condenser "
    "(bypassing the evaporator). When CPF is small (< 5%), the evaporator barely improves the air's "
    "drying potential, and the compressor energy used for evaporator cooling is largely wasted."
)

# ==== Section 2 ====
doc.add_heading("2. The VPD-Based Condenser-Direct Bypass Strategy", level=1)
doc.add_paragraph(
    "The strategy introduces an alternative air path that bypasses the evaporator when its benefit "
    "is minimal:"
)

doc.add_heading("2.1 Normal Mode (Evaporator Active)", level=2)
doc.add_paragraph(
    "Exhaust air is mixed with ambient air at recirculation ratio r, passed through the evaporator "
    "coil (which cools and dehumidifies it), then reheated by the condenser to T_set before entering "
    "the chamber. The evaporator operates as a dehumidifier, and the heat pump cycle runs with "
    "T_evap determined by the dewpoint of the mixed air."
)

doc.add_heading("2.2 Condenser-Direct Mode", level=2)
doc.add_paragraph(
    "When CPF falls below a threshold (default: 0.05 or 5%), the system switches to condenser-direct "
    "mode. In this mode:"
)
for item in [
    "Exhaust air (warm, ~43 C, relatively dry) is routed directly to the HP condenser, bypassing the evaporator.",
    "The condenser only needs to reheat the air by ~2 C (from ~43 C to 45 C), requiring very small Q_cond (~0.2 kW vs ~2.5 kW in normal mode).",
    "The evaporator draws fresh ambient air as a sensible heat source (not for dehumidification).",
    "The compressor work drops dramatically because Q_cond is tiny.",
    "Humidity in the closed loop builds naturally as the product continues releasing moisture.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("2.3 Switching Criteria", level=2)
doc.add_paragraph("The system oscillates between modes based on two thresholds:")
for item in [
    "Evap to Cond-Direct: When CPF < 0.05 (evaporator benefit < 5%)",
    "Cond-Direct to Evap: When CPF > 0.15 (= 3 x 0.05; humidity has accumulated sufficiently that the evaporator would provide meaningful dehumidification)",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "The 3x hysteresis ratio prevents chattering at the switching boundary. The asymmetric band "
    "(5% to 15%) ensures that once the system enters condenser-direct mode, it stays there long "
    "enough for humidity to accumulate meaningfully before switching back."
)

# ==== Section 3 ====
doc.add_heading("3. Physics-Based Minimum Dwell Time", level=1)
doc.add_paragraph(
    "Even with hysteresis, rapid drying rates (especially at lower recirculation ratios) can cause "
    "frequent mode transitions. Real heat pump compressors have minimum on/off times of 3-5 minutes "
    "to prevent mechanical damage from short-cycling. To enforce physically realistic switching, "
    "we compute a humidity accumulation dwell time tau_humidity."
)

doc.add_heading("3.1 Derivation", level=2)
doc.add_paragraph(
    "The key insight is that in condenser-direct mode, the product continues releasing water vapour "
    "into the closed air loop. The rate of humidity increase in the air stream is:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("d(omega)/dt = m_dot_water / m_dot_da").italic = True

doc.add_paragraph(
    "where m_dot_water = dm_w_total / dt is the instantaneous drying rate (kg water per second) "
    "measured from the previous simulation timestep, and m_dot_da is the dry air mass flow rate."
)
doc.add_paragraph(
    "The condenser penalty fraction CPF depends on the exhaust humidity omega_exhaust. "
    "The sensitivity of CPF to changes in omega_exhaust is:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("d(CPF)/d(omega) = [CPF(omega+delta) - CPF(omega-delta)] / (2*delta)").italic = True

doc.add_paragraph(
    "evaluated by central finite difference with delta = 0.01*omega. "
    "Combining these gives the rate of change of CPF over time:"
)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("d(CPF)/dt = |d(CPF)/d(omega)| x d(omega)/dt").italic = True

doc.add_paragraph("The minimum dwell time to cross the switching threshold is:")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("tau_humidity = |CPF_target - CPF_current| / |d(CPF)/dt|").italic = True

doc.add_paragraph(
    "This tau is clamped to [tau_min, tau_max] = [300 s, 7200 s] (5 min to 2 hours). "
    "The 300-second floor is the compressor short-cycle protection limit. "
    "The 7200-second ceiling prevents indefinite lock-in when the drying rate approaches zero."
)

doc.add_heading("3.2 Physical Interpretation", level=2)
doc.add_paragraph(
    "This dwell time is self-adapting to the drying physics:"
)
for item in [
    "Early in drying (high drying rate): humidity builds/depletes quickly, so tau is short (5-15 min). "
    "The system can switch modes frequently because the air state genuinely changes fast.",
    "Late in drying (low drying rate): humidity changes very slowly, so tau is long (1-2 hours). "
    "The system stays in condenser-direct mode for extended periods because there is barely any water "
    "to accumulate. This is physically correct - there is no need to dehumidify when the product "
    "is nearly dry.",
    "The dwell time tracks the actual measured drying rate from the simulation, not an assumed model. "
    "This makes it robust to different products, temperatures, and operating conditions.",
]:
    doc.add_paragraph(item, style="List Bullet")

# ==== Section 4 ====
doc.add_heading("4. Air Flow Model in the Drying Chamber", level=1)
doc.add_paragraph(
    "The drying chamber uses a cross-flow, trays-in-series configuration. "
    "Hot dry air enters the chamber and flows horizontally across each tray in sequence:"
)
for item in [
    "Air enters at T_set (45 C) with humidity omega_to_chamber.",
    "It flows horizontally across Tray 0, exchanging heat and moisture with the product.",
    "The same air stream (now cooler and more humid) enters Tray 1, then Tray 2, etc.",
    "After passing through all 10 trays in series, the air exits as exhaust.",
    "Each tray is subdivided into n_sections = 4 along the airflow direction.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "With n_sections = 4, each tray is divided into 4 strips along the airflow direction. "
    "Section 0 (air inlet side) sees the hottest, driest air and dries fastest. "
    "Section 3 (air outlet side) sees cooler, more humid air and dries slowest. "
    "This creates a realistic intra-tray moisture gradient that better represents cross-flow behaviour. "
    "The exhaust humidity (which drives VPD switching) is computed from the air exiting the last "
    "section of the last tray, providing an accurate representation of the loop humidity state."
)

# ==== Section 5 ====
doc.add_heading("5. Results Summary", level=1)

doc.add_heading("5.1 Config A - HP Only (n_sections = 4)", level=2)

table = doc.add_table(rows=11, cols=6)
table.style = "Light Shading Accent 1"
headers = ["Location", "r", "Baseline SEC", "VPD SEC", "Savings %", "Drying Time (h)"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data_a = [
    ["Kathmandu", "0.3", "0.785", "0.547", "30%", "15.4"],
    ["Kathmandu", "0.5", "0.724", "0.505", "30%", "16.2"],
    ["Kathmandu", "0.7", "0.643", "0.448", "30%", "17.7"],
    ["Kathmandu", "0.9", "0.515", "0.401", "22%", "21.6"],
    ["Kathmandu", "0 (ref)", "0.784", "-", "-", "14.2"],
    ["Biratnagar", "0.3", "0.843", "0.576", "32%", "15.6"],
    ["Biratnagar", "0.5", "0.804", "0.548", "32%", "16.2"],
    ["Biratnagar", "0.7", "0.747", "0.530", "29%", "17.2"],
    ["Biratnagar", "0.9", "0.659", "0.582", "12%", "18.1"],
    ["Biratnagar", "0 (ref)", "0.597", "-", "-", "14.8"],
]
for row_idx, row_data in enumerate(data_a):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = val

doc.add_paragraph("")
doc.add_heading("5.2 Config B - Solar + HP (n_sections = 4)", level=2)

table2 = doc.add_table(rows=11, cols=6)
table2.style = "Light Shading Accent 1"
for i, h in enumerate(headers):
    table2.rows[0].cells[i].text = h

data_b = [
    ["Kathmandu", "0.3", "0.597", "0.333", "44%", "15.2"],
    ["Kathmandu", "0.5", "0.575", "0.310", "46%", "15.9"],
    ["Kathmandu", "0.7", "0.543", "0.285", "48%", "16.9"],
    ["Kathmandu", "0.9", "0.487", "0.330", "32%", "19.2"],
    ["Kathmandu", "0 (ref)", "0.541", "-", "-", "14.2"],
    ["Biratnagar", "0.3", "0.479", "0.245", "49%", "15.6"],
    ["Biratnagar", "0.5", "0.466", "0.226", "51%", "16.1"],
    ["Biratnagar", "0.7", "0.454", "0.228", "50%", "17.1"],
    ["Biratnagar", "0.9", "0.438", "0.327", "25%", "18.2"],
    ["Biratnagar", "0 (ref)", "0.323", "-", "-", "14.8"],
]
for row_idx, row_data in enumerate(data_b):
    for col_idx, val in enumerate(row_data):
        table2.rows[row_idx + 1].cells[col_idx].text = val

# ==== Section 6 ====
doc.add_heading("6. Key Findings", level=1)
for item in [
    "The VPD-based condenser-direct strategy achieves 12-51% SEC reduction compared to baseline recirculation across all tested configurations.",
    "Config B (Solar + HP) benefits more from the strategy (25-51% savings) than Config A (HP-only, 12-32% savings), because solar heating provides free energy during condenser-direct periods.",
    "The optimal recirculation ratio with VPD strategy is r = 0.7 for Kathmandu (cold, dry climate) and r = 0.5 for Biratnagar (warm, humid climate).",
    "The humidity accumulation dwell time naturally adapts: 5-15 min in the constant-rate period (high drying rate) growing to 1-2 hours in the falling-rate period (low drying rate).",
    "The 3x hysteresis band (5% CPF on, 15% CPF off) combined with the physics-based dwell produces clean, non-chattering transitions that are realistic for residential AC compressors.",
    "Using n_sections = 4 per tray (instead of n = 1) captures intra-tray moisture gradients, providing more accurate exhaust humidity for VPD switching decisions.",
]:
    doc.add_paragraph(item, style="List Bullet")

# ==== Section 7 ====
doc.add_heading("7. Implementation Notes", level=1)
doc.add_paragraph(
    "The strategy is controlled by a single parameter: cond_penalty_thresh (default: 0.05). "
    "Setting it to 0.0 disables the strategy entirely. The humidity dwell time is computed "
    "automatically from the simulation state - no tuning parameters are needed beyond the threshold."
)
doc.add_paragraph(
    "The implementation uses a 3-way valve concept: one position routes air through the evaporator "
    "(normal mode), the other routes exhaust directly to the condenser (condenser-direct mode). "
    "The evaporator always operates - in condenser-direct mode it draws ambient air as a sensible "
    "heat source. This maintains continuous compressor operation and avoids start/stop cycling."
)

# Save
output_path = r"D:\Masters\SAHPD papers\VPD_Condenser_Direct_Bypass_Strategy.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
