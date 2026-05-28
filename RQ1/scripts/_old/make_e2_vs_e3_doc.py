"""Generate a thesis-ready comparison document for Config E2 vs E3.

Explains why E2 (Solar -> Condenser) outperforms E3 (Condenser -> Solar)
despite E3 exhibiting higher peak COP, by linking heat-pump and solar-
collector physics to the topology choice.

Output: Compare_E2_vs_E3.docx in the RQ1 root.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = PROJECT_ROOT / "figures" / "compare_E2_E3"
PLOT_DIR = PROJECT_ROOT / "plots" / "compare_E_autumn_ktm"


# --------------------------------------------------------------------------
# Styling helpers
# --------------------------------------------------------------------------
def set_default_font(doc, name="Times New Roman", size_pt=11):
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), name)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, bold=False, italic=False, justify=True, size=11):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r.font.size = Pt(11)
    r.italic = True
    return p


def add_figure(doc, path, caption, width_cm=15.5):
    if not path.exists():
        add_para(doc, f"[Missing figure: {path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.font.name = "Times New Roman"
    cr.font.size = Pt(10)
    cr.italic = True


def add_table(doc, header, rows, col_widths_cm=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header
    for j, text in enumerate(header):
        c = tbl.rows[0].cells[j]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Body
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = tbl.rows[i + 1].cells[j]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)
            if j == 0:
                r.bold = True

    if col_widths_cm:
        for j, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[j].width = Cm(w)


# --------------------------------------------------------------------------
# Document body
# --------------------------------------------------------------------------
def build():
    doc = Document()
    set_default_font(doc)

    for sec in doc.sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # -------- Title --------
    title = doc.add_heading(
        "Comparison of Configurations E2 and E3: "
        "Solar–Heat Pump Coupling Topology",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = subtitle.add_run(
        "Case study: Kathmandu, Autumn (Oct–Nov), A_c = 10 m², ε_HRX = 0.70"
    )
    s.font.name = "Times New Roman"
    s.font.size = Pt(11)
    s.italic = True

    # -------- 1. Introduction --------
    add_heading(doc, "1. Objective", level=1)
    add_para(
        doc,
        "Configurations E2 and E3 both combine a sensible-counterflow heat-recovery "
        "exchanger (HRX), a flat-plate solar air collector, and a vapour-compression "
        "heat pump (HP) to raise fresh ambient air to the drying setpoint "
        "T_set = 45 °C. They differ only in the relative ordering of the solar "
        "collector and the condenser on the air path that feeds the drying chamber. "
        "This document quantifies the resulting performance gap under identical "
        "boundary conditions and explains the difference from first principles so "
        "that the topology choice can be defended in the thesis discussion.",
    )

    # -------- 2. Topology --------
    add_heading(doc, "2. Topology and Control", level=1)

    add_heading(doc, "2.1 Air paths", level=2)
    add_para(
        doc,
        "In Config E2 the solar collector is placed upstream of the condenser, "
        "so the air entering the collector is cold HRX-preheated ambient air. "
        "In Config E3 the order is reversed: the condenser heats the air first, "
        "and the collector performs a final top-up before the chamber inlet. "
        "The evaporator side and the HRX are identical in both configurations.",
    )

    add_table(
        doc,
        header=["Config", "Condenser stream", "Evaporator stream", "Control strategy"],
        rows=[
            [
                "E2",
                "Amb → HRX → Solar → Cond → Chamber",
                "Exhaust (post-HRX) + iterative ambient make-up",
                "Fixed T_cond = 55 °C; HP idles when T_after_solar ≥ T_set",
            ],
            [
                "E3",
                "Amb → HRX → Cond → Solar → Chamber",
                "Exhaust (post-HRX) + iterative ambient make-up",
                "Variable T_cond ≤ 55 °C, solar-priority: HP OFF when solar alone meets T_set, partial lift otherwise",
            ],
        ],
        col_widths_cm=[1.7, 5.8, 4.6, 4.4],
    )

    add_heading(doc, "2.2 Solar-priority control in E3", level=2)
    add_para(
        doc,
        "E3 exploits the freedom to let the condenser operate below its nominal "
        "setpoint. At each time step the simulator back-calculates the highest "
        "collector inlet temperature T_cond,out that still allows the downstream "
        "collector to reach T_set, using the Hottel–Whillier–Bliss relation",
    )
    add_equation(
        doc,
        "T_out  =  T_in  +  (F_R · (τα) · G · A_c  −  F_R · U_L · A_c · (T_in − T_amb)) / (ṁ · c_p)",
    )
    add_para(
        doc,
        "If the required T_cond,out is lower than T_set, the HP is asked to "
        "deliver a smaller lift; if the HRX output alone satisfies the equation "
        "above, the compressor is switched off entirely and the collector "
        "supplies the whole heating duty. E2 has no such degree of freedom — its "
        "condenser always targets 55 °C and the excess solar heat, when "
        "T_after_solar exceeds 45 °C, is simply clipped at the chamber inlet.",
    )

    # -------- 3. Quantitative comparison --------
    add_heading(doc, "3. Numerical Results", level=1)
    add_para(
        doc,
        "Both configurations were simulated on the Kathmandu PVGIS-TMY autumn "
        "subset (October + November, 1464 hourly records). The drying batch is "
        "3.0 kg of apple on 10 trays, initial moisture 5.20 kg/kg d.b., target "
        "0.10 kg/kg d.b. The solar collector area is 10 m² and the HRX "
        "effectiveness is 0.70. Fan power and drying kinetics are identical "
        "across both runs.",
    )

    add_table(
        doc,
        header=["Quantity", "E2 (Solar→Cond)", "E3 (Cond→Solar)", "E3 − E2"],
        rows=[
            ["Drying time [h]", "16.93", "16.93", "0.00"],
            ["Water removed [kg]", "19.31", "19.31", "0.00"],
            ["W_comp total [kWh]", "1.956", "2.157", "+0.201 (+10.3 %)"],
            ["W_fan total [kWh]", "0.415", "0.415", "0.000"],
            ["SEC [kWh/kg_water]", "0.1228", "0.1333", "+0.0105 (+8.6 %)"],
            ["Q_cond total [kWh]", "9.95", "11.16", "+1.21"],
            ["Q_solar useful [kWh]", "20.95", "19.78", "−1.18 (−5.6 %)"],
            ["HP on-fraction", "0.618", "0.618", "0.000"],
            ["T_cond mean (HP-on) [°C]", "55.0", "53.5", "−1.5"],
            ["T_cond minimum [°C]", "55.0", "36.2", "−18.8"],
            ["Lift mean (HP-on) [K]", "39.74", "38.28", "−1.46"],
            ["Lift minimum [K]", "36.38", "23.49", "−12.90"],
            ["COP mean (HP-on)", "5.22", "5.47", "+0.25 (+4.8 %)"],
            ["COP peak", "5.63", "8.56", "+2.93 (+52 %)"],
            ["Collector η energy-weighted", "0.467", "0.442", "−0.025"],
            ["Collector T_in mean (day) [°C]", "30.2", "34.9", "+4.7"],
            ["Collector T_out mean (day) [°C]", "50.5", "54.1", "+3.6"],
        ],
        col_widths_cm=[6.0, 3.4, 3.4, 3.7],
    )

    add_para(
        doc,
        "The headline result is that E2 dries the same batch with 8.6 % less "
        "specific energy consumption, even though E3 attains a 52 % higher "
        "peak COP and a lower mean thermodynamic lift. The remainder of this "
        "document resolves that apparent paradox.",
        italic=False,
    )

    add_figure(
        doc,
        FIG_DIR / "hp_comparison.png",
        "Figure 1. Heat-pump performance of Config E2 (blue) and Config E3 (red) "
        "over a representative autumn drying batch in Kathmandu. "
        "Top-left: saturation temperatures. Top-right: thermodynamic lift. "
        "Bottom-left: instantaneous COP when the compressor is running. "
        "Bottom-right: compressor work and useful solar heat.",
    )

    # -------- 4. Physics --------
    add_heading(doc, "4. Physical Explanation", level=1)

    add_heading(doc, "4.1 Carnot framing of the heat pump", level=2)
    add_para(
        doc,
        "For a vapour-compression cycle with fixed isentropic efficiency, "
        "the actual COP tracks the Carnot bound",
    )
    add_equation(
        doc,
        "COP_Carnot  =  T_cond  /  (T_cond − T_evap)",
    )
    add_para(
        doc,
        "with an approximately constant second-law efficiency "
        "(0.61–0.62 in our validation set for η_is = 0.75). E3’s solar-priority "
        "control reduces T_cond during peak insolation hours, shrinking the "
        "denominator and producing the dramatic COP spike visible in Figure 1 "
        "(peaks at 8.56 when T_cond dips to 36 °C). E2’s fixed condenser "
        "target keeps the lift nearly flat at ~40 K, so its COP only varies "
        "through the small swing in T_evap.",
    )

    add_heading(doc, "4.2 Hottel–Whillier–Bliss collector efficiency", level=2)
    add_para(
        doc,
        "The flat-plate collector efficiency is the ratio of useful air-side "
        "enthalpy gain to incident radiation:",
    )
    add_equation(
        doc,
        "η_c  =  F_R · (τα)  −  F_R · U_L · (T_in − T_amb) / G",
    )
    add_para(
        doc,
        "The second term — the loss term — is linear in the inlet-to-ambient "
        "temperature difference (T_in − T_amb). Placing the collector upstream "
        "of the condenser (E2) gives it cold HRX-warmed ambient air, typically "
        "22–35 °C, so (T_in − T_amb) rarely exceeds 10 K and collector losses "
        "are minimised. Placing the collector downstream of the condenser (E3) "
        "forces it to operate with air that has already been lifted part of "
        "the way to T_set, so (T_in − T_amb) is 5–15 K larger at every moment. "
        "The measured energy-weighted efficiency drops from 0.467 (E2) to "
        "0.442 (E3), a 2.5 percentage-point penalty.",
    )

    add_figure(
        doc,
        FIG_DIR / "solar_comparison.png",
        "Figure 2. Solar collector operating point. "
        "Left: air temperatures at the collector inlet (solid) and outlet "
        "(dashed). E3’s inlet is consistently warmer because the condenser "
        "precedes the collector. Right: instantaneous collector efficiency η_c. "
        "Both configurations approach the same peak value at midday, but E3 "
        "starts producing useful heat later and E2 holds its efficiency over a "
        "longer window.",
    )

    add_heading(doc, "4.3 Why higher peak COP does not translate to lower SEC", level=2)
    add_para(
        doc,
        "The apparent paradox is resolved by noting that the COP gain and the "
        "collector-efficiency loss act on different quantities. E3 improves the "
        "quality of a smaller quantity of heat-pump output: during peak sun the "
        "HP provides only a partial lift, so although each kWh of compressor "
        "work buys more condenser heat, the compressor is supplying less of the "
        "total thermal demand. E2 instead improves the quantity of useful solar "
        "heat delivered to the air stream by keeping the collector on its "
        "efficient operating branch. Over the whole batch the balance favours "
        "E2:",
    )
    add_equation(
        doc,
        "Q_solar,E2 − Q_solar,E3  ≈  +1.18 kWh   >   "
        "(W_comp,E3 − W_comp,E2) · COP_E2  ≈  +1.05 kWh",
    )
    add_para(
        doc,
        "In other words, the extra solar heat that E2 extracts is larger than "
        "the thermal equivalent of the extra compressor work E3 needs, so the "
        "net electrical energy consumption is lower for E2 despite its lower "
        "peak COP.",
    )

    add_heading(doc, "4.4 Solar clipping in E2 is a small effect", level=2)
    add_para(
        doc,
        "An obvious concern with E2 is that on high-insolation hours the "
        "collector outlet can exceed T_set. The simulator enforces "
        "T_air_in,cond = min(T_after_solar, T_set), so any surplus is "
        "effectively wasted. In the autumn Kathmandu run this clipping is "
        "active for ~6 h per batch, but the collector efficiency during those "
        "hours is already near its Hottel–Whillier ceiling (η ≈ 0.49). Moving "
        "the same collector behind the condenser (E3 topology) to avoid "
        "clipping does not reclaim that energy — instead it degrades the "
        "collector’s loss term and the HP has to make up the shortfall.",
    )

    # -------- 5. Discussion --------
    add_heading(doc, "5. Discussion and Design Implication", level=1)
    add_para(
        doc,
        "The E2-versus-E3 comparison isolates a design principle that is "
        "easily obscured by heat-pump-centric reasoning: in a hybrid "
        "solar-HP dryer the dominant optimisation target is the collector "
        "loss term, not the HP lift. Within the range of conditions studied "
        "here (dry-bulb 15–28 °C, GHI up to 900 W m⁻²), the flat-plate "
        "collector’s heat-loss coefficient U_L is large enough that a 5 °C "
        "increase in inlet temperature costs about the same amount of "
        "delivered heat as a 1 K reduction in HP lift buys. The correct "
        "design move is therefore to place the collector where the air is "
        "coldest, which in this hybrid family means upstream of the "
        "condenser (Config E2).",
    )
    add_para(
        doc,
        "Three caveats are worth stating explicitly for the thesis:",
    )
    add_para(
        doc,
        "(i) The result is specific to air-based collectors with relatively "
        "high U_L (~6 W m⁻² K⁻¹ in our model). Evacuated-tube or selective-"
        "surface collectors with lower U_L would narrow the gap because the "
        "loss penalty of the E3 topology scales directly with U_L.",
    )
    add_para(
        doc,
        "(ii) E3’s variable-T_cond strategy is still valuable as a control "
        "concept: it is the mechanism that lets the HP decouple from the "
        "setpoint when renewable heat is abundant. The present result only "
        "says that placing the collector behind the condenser is the wrong "
        "physical location to realise that benefit in an air system.",
    )
    add_para(
        doc,
        "(iii) At very low insolation, both configurations degenerate to a "
        "Configuration D (HRX-only) solution, so the E2/E3 gap closes at "
        "night. The quoted 8.6 % SEC advantage is therefore season- and "
        "site-specific; the physical direction of the inequality, however, "
        "is not — it follows directly from the sign of ∂η_c / ∂T_in.",
    )

    # -------- 6. Conclusion --------
    add_heading(doc, "6. Conclusion", level=1)
    add_para(
        doc,
        "Under identical weather, kinetics, HRX, and heat-pump parameters, "
        "Config E2 (Solar → Condenser) reaches the same dryness with 8.6 % "
        "less specific energy consumption than Config E3 (Condenser → Solar) "
        "on a representative Kathmandu autumn batch. The advantage arises "
        "because the solar collector in E2 operates with a cooler inlet "
        "stream (~30 °C) than in E3 (~35 °C), reducing its heat-loss term "
        "and increasing the energy-weighted collection efficiency from "
        "0.442 to 0.467. E3’s higher peak COP (up to 8.56 versus 5.63 in E2) "
        "is a real but smaller benefit that does not compensate for the lost "
        "solar yield. The design recommendation for air-based solar-assisted "
        "heat-pump dryers is therefore to place the collector upstream of "
        "the condenser whenever the collector is the dominant source of "
        "thermal losses.",
    )

    out_path = PROJECT_ROOT / "Compare_E2_vs_E3.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")
    print(f"Size: {out_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build()
