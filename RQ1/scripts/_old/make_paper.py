"""Build the journal paper (SAHPD_Paper.docx) section by section."""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

SUM_CSV = Path(__file__).resolve().parents[1] / "outputs" / "master_summary.csv"

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "paper" / "SAHPD_Paper.docx"
FIG = ROOT / "figures" / "thesis"


def embed_figure(doc, rel_path, caption, width_in=6.0):
    """Embed a PNG centred in the document with an italic caption below."""
    path = ROOT / rel_path
    if not path.exists():
        para(doc, f"[Missing figure: {rel_path}]", size=9)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9.5)


def set_margins(doc):
    for s in doc.sections:
        s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)
        s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)


def para(doc, text, style=None, align=None, bold=False, size=11):
    p = doc.add_paragraph()
    if style: p.style = style
    if align is not None: p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    if bold: r.bold = True
    return p


def equation(doc, text, number=None):
    """Centered, italicised equation line with optional right-aligned number."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Cambria Math"
    r.font.size = Pt(11)
    if number is not None:
        tab = p.add_run("\t\t(" + number + ")")
        tab.italic = False
        tab.font.size = Pt(10)
    return p


def nomenclature(doc, rows, caption="where"):
    """Two-column nomenclature table under an equation.

    rows: list of (symbol, definition) tuples.
    """
    p = doc.add_paragraph()
    r = p.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = False
    for i, (sym, defn) in enumerate(rows):
        c0 = t.rows[i].cells[0]
        c1 = t.rows[i].cells[1]
        c0.width = Cm(3.5)
        c1.width = Cm(13.0)
        c0.text = sym
        c1.text = defn
        for run in c0.paragraphs[0].runs:
            run.font.size = Pt(10)
            run.italic = True
        for run in c1.paragraphs[0].runs:
            run.font.size = Pt(10)
    return t


def title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Comparative simulation of ten solar-assisted heat-pump dryer "
        "configurations for fruit drying in Nepal: effects of heat recovery, "
        "solar integration, and humidity-based bypass control"
    )
    r.bold = True; r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sarwagya Bhattarai, [co-authors TBD]")
    r.italic = True; r.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[affiliation]")
    r.font.size = Pt(10)

    doc.add_heading("Abstract", level=2)
    para(doc,
         "Ten solar-assisted heat-pump dryer (SAHPD) topologies — a heat-"
         "pump-only baseline, two solar-in-series variants, two solar-"
         "cascade variants, two exhaust-heat-recovery variants, and three "
         "combined heat-recovery-plus-solar variants — are compared under a "
         "single validated thermodynamic model at three Nepalese sites "
         "(Kathmandu, Biratnagar, Taplejung) spanning mid-hill, Terai, and "
         "high-hill climates. The model couples an R134a vapour-"
         "compression cycle (CoolProp, η_is = 0.75), a Hottel-Whillier-"
         "Bliss flat-plate collector (τα = 0.75, U_L = 5.0 W/m²K), a "
         "counter-flow plate heat-recovery exchanger (ε = 0.70), and a "
         "first-order chamber-scale kinetic model fitted to 13 apple-slice "
         "experiments. First-law closure, water-mass balance, and "
         "psychrometric round-trip all close to below 1 × 10⁻⁶. Annual "
         "and three-season typical-meteorological-year forcing is used. "
         "Specific energy consumption (SEC) and specific moisture "
         "extraction rate (SMER) are reported for every case. Results: "
         "baseline SEC is 0.72 kWh/kg (SMER 1.39 kg/kWh) at Kathmandu and "
         "0.54 kWh/kg (SMER 1.84 kg/kWh) at Biratnagar; exhaust heat "
         "recovery alone (ε = 0.70) reduces SEC by 46-49 % at no "
         "collector cost; the combined heat-recovery-plus-solar-plus-VPD-"
         "bypass topology delivers 0.097 kWh/kg (SMER 10.31 kg/kWh) at "
         "Biratnagar, 0.144 kWh/kg at Kathmandu, and 0.129 kWh/kg at "
         "Taplejung, 77-82 % below their respective baselines. Closed-loop "
         "recirculation penalises SEC by up to "
         "+39 % at the warm-climate site, demonstrating that the "
         "preferred air-path is not uniform within a single country. A "
         "kinetic-parameter sensitivity check confirms that cross-"
         "configuration rankings are preserved under ±20 % variation in "
         "the fitted rate constant.",
         size=10)

    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    r.bold = True; r.font.size = Pt(10)
    r2 = p.add_run(
        "solar-assisted heat pump dryer; heat recovery exchanger; apple drying; "
        "specific energy consumption; Nepal; vapour-pressure-deficit control"
    )
    r2.font.size = Pt(10)


def section_1_introduction_v2(doc):
    """Trimmed introduction. Target ~1,000 words."""

    doc.add_heading("1. Introduction", level=1)

    doc.add_heading("1.1  Fruit drying in Nepal", level=2)
    para(doc,
        "Nepal produced approximately 1.53 Mt of fruit in 2022/23 across three "
        "agro-climatic belts: deciduous fruits (apple, pear, peach) in the mid- "
        "and high-hill districts, tropical fruits (mango, banana, pineapple) in "
        "the Terai lowlands, and citrus in the mid-hills [MoALD, 2023]. Harvest "
        "windows are narrow, moisture content at harvest is 75-90 % w.b., and "
        "cold-chain infrastructure is limited; reported post-harvest losses for "
        "horticultural produce lie in the 15-35 % range [Shrestha, 2017]. Drying to an intermediate-"
        "moisture product (X_f = 0.10-0.15 d.b.) is among the most direct value-"
        "addition pathways for smallholders and cooperatives, since the product "
        "is shelf-stable and tolerates the country's retail conditions. Current "
        "deployed technology is predominantly fossil-fuel cabinet drying or "
        "passive solar; the expansion of grid electrification and rooftop PV "
        "creates scope for electrified, solar-assisted drying at district scale."
    )
    para(doc,
        "The research question addressed in this paper is: among practically "
        "realisable combinations of solar pre-heating, exhaust heat recovery, "
        "and humidity-aware control layered onto a vapour-compression heat-pump "
        "dryer, which combination performs best under Nepali climatic conditions, "
        "and how does the ranking depend on site climate? Specific energy "
        "consumption (SEC) and specific moisture extraction rate (SMER) are "
        "reported as selection metrics. The thermodynamic air-path analysis is "
        "commodity-agnostic; absolute drying times and SEC values are apple-"
        "specific because the kinetic model (§3.4) was fitted to apple-slice "
        "data [Royen et al., 2020]. Extension to other fruits requires re-"
        "fitting the kinetic constants but not the thermodynamic model."
    )

    doc.add_heading("1.2  Heat-pump drying", level=2)
    para(doc,
        "A heat-pump convective dryer (HPCD) replaces the resistance heater of "
        "a hot-air dryer with a vapour-compression heat pump. The condenser "
        "heats process air to the drying set-point (40-55 °C for fruit), the "
        "chamber desorbs moisture, and the evaporator dehumidifies the exhaust "
        "when air is recirculated. COP of 3-5 for moderate lifts translates to "
        "SEC of 1.3-2.5 kWh/kg on experimental HPCDs [Prasertsan & Saen-saby, "
        "1998; Chua et al., 2002] and 0.6-1.2 kWh/kg in optimised simulation "
        "studies [Minea, 2013]. The present model falls in the intermediate "
        "class: CoolProp refrigerant properties with air-side effectiveness, "
        "closed by Q_cond = Q_evap + W_comp, η_is = 0.75."
    )
    para(doc,
        "Closed-loop operation (recirculation fraction r > 0) is frequently "
        "assumed to reduce SEC by reusing exhaust enthalpy. In practice, "
        "intermediate r (0.3-0.7) drives the mixed-air inlet towards the "
        "evaporator coil temperature, collapsing evaporator duty per unit mass "
        "flow and increasing compressor work per unit moisture removed. Most "
        "simulation studies fix r at a single value; a site-resolved sweep "
        "across multiple configurations has not been reported."
    )

    doc.add_heading("1.3  Solar-assisted heat-pump dryers", level=2)
    para(doc,
        "Solar-assisted heat-pump dryer (SAHPD) systems integrate a "
        "vapour-compression heat pump with a solar air collector. Colak & "
        "Hepbasli [2009] and Mohanraj et al. [2018] provide the foundational "
        "surveys; Daghigh & Ruhani [2022] and Sun et al. [2025] are the most "
        "recent comprehensive reviews. Across the surveyed literature SMER "
        "spans 0.6-3.0 kg/kWh for conventional single-stage SAHPDs, with COP "
        "2.7-5.0 and thermal efficiency 9-57 % [Bhandari et al., 2025]. "
        "Recent multi-stage and ejector-enhanced variants push SMER "
        "higher: Yan et al. [2023] reported SMER up to 22.9 kg/kWh and SEC "
        "as low as 0.043 kWh/kg on a double-stage SAHPD (45 C, 50 kg load, "
        "25 % fresh-air rate, summer conditions), and ejector-enhanced configurations "
        "[Zhu et al., 2023] show similar potential at the cost of hardware "
        "complexity."
    )
    para(doc,
        "SAHPD topologies fall into three principal forms. In series "
        "pre-heating the collector sits between ambient and the condenser, "
        "so the heat pump provides only the residual lift [Hawlader et al., "
        "2003, on guava and papaya in Singapore]. In cascade or parallel arrangement the solar "
        "stream delivers heat on its own path and the heat-pump stream "
        "runs independently, decoupling the evaporator from the collector "
        "but sacrificing dehumidification synergy. In evaporator-side "
        "pre-heating the collector warms the ambient stream entering the "
        "evaporator, primarily for frost prevention in cold climates, "
        "which is not relevant to Nepal. Thermal-storage integration "
        "(PCM or water tank) can extend drying into night hours "
        "[Ismaeel, 2020; Mortezapour et al., 2012; Rulazi et al., 2024]. "
        "Rulazi et al. [2024] report a techno-economic payback of 2-4 "
        "years for SAHPD + storage across representative agricultural "
        "commodities."
    )
    para(doc,
        "Two recurring issues are documented. First, direct cross-topology "
        "comparison within a single study is rare: most papers report one "
        "configuration against a hot-air or open-sun baseline using "
        "heterogeneous kinetics and SEC definitions, which resists "
        "synthesis [Sun et al., 2025]. Second, recirculation-fraction "
        "effects are reported at one or two values but rarely swept; "
        "Minea [2013] and Colak & Hepbasli [2009] note that intermediate "
        "recirculation can collapse evaporator duty, but site-resolved "
        "sensitivity across configurations is not available."
    )

    doc.add_heading("1.4  Heat-recovery exchangers in drying", level=2)
    para(doc,
        "An exhaust-to-supply heat-recovery exchanger (HRX) recovers "
        "sensible enthalpy (and, with a membrane or condensing plate, some "
        "latent enthalpy) from dryer exhaust typically at 35-45 °C and "
        "60-90 % RH. Counter-flow plate HRX with ε = 0.70 is a practical "
        "fabrication default: above ε ≈ 0.75 the required plate area grows "
        "super-linearly [Shah & Sekulic, 2003; Kays & London, 1984]. "
        "Experimental SAHPDs incorporating HRX report SEC reductions of "
        "15-35 % at ε = 0.60-0.75 [Ismaeel, 2020], "
        "and membrane-based enthalpy exchangers achieve total-heat "
        "effectiveness near 0.88 for simultaneous sensible and latent "
        "recovery [Vali et al., 2021]."
    )
    para(doc,
        "Two routings are relevant when coupling an HRX with a heat pump: "
        "supply-side only (ambient → HRX → condenser), or additionally "
        "warming the evaporator supply by mixing with HRX-cooled exhaust. "
        "The second raises the evaporator saturation temperature and "
        "shrinks cycle pressure ratio. Both routings are evaluated here "
        "(Configs D1/E1 vs D2/E2). Aacharya et al. [2024] provide the "
        "closest experimental precedent: a flat-plate HX integrated into "
        "a solar dryer tested on apples at Dhulikhel, Nepal (Feb-Apr "
        "2023), reported drying rates of 107 g/(h·m²) with a low-e coated "
        "collector, 89 % collector efficiency, and economic payback of "
        "1.61 years against open sun drying. Adhikari et al. [2025] "
        "extended this work with CFD and smoke-flow visualisation on the "
        "same class of HX-integrated solar dryer, showing that "
        "unoptimised chamber geometry yields in-chamber uniformity index "
        "as low as 0.58 (CV > 1.0) and that inlet-gap and baffle "
        "redesign raise UI to 0.78. This is directly relevant to our "
        "lumped assumption and is revisited in §6."
    )

    doc.add_heading("1.5  Humidity-aware control", level=2)
    para(doc,
        "The falling-rate period occupies 60-75 % of batch time and "
        "contributes disproportionately to SEC because moisture extraction "
        "rate drops while compressor duty is near-rated. Mitigations include "
        "variable-speed compressors and control strategies that throttle or "
        "divert process air. This paper adopts a vapour-pressure-deficit "
        "(VPD) bypass: when the chamber-to-product VPD falls below a "
        "threshold, the exhaust is bypassed around the condenser loop, "
        "cutting compressor and fan work during the tail. VPD-based control "
        "is standard in controlled-environment agriculture; its combination "
        "with an exhaust HRX on a heat-pump dryer has not, to the authors' "
        "knowledge, been reported."
    )

    doc.add_heading("1.6  Prior Nepali work and the gap this paper addresses", level=2)
    para(doc,
        "A series of solar-drying studies has emerged from "
        "Kathmandu University in collaboration with Lund University. "
        "Aacharya et al. [2024] built and tested an HX-integrated solar "
        "dryer at Dhulikhel with three collector variants (flat GI, v-"
        "corrugated GI, and low-e coated aluminium), reporting collector "
        "efficiencies of 50-89 % and economic payback as low as 1.61 years "
        "against open sun drying; the experimental commodity was apple. "
        "Adhikari et al. [2025] used ANSYS Fluent CFD and smoke-flow "
        "visualisation on the same class of HX solar dryer, demonstrating "
        "that chamber-geometry optimisation can raise in-chamber uniformity "
        "index from 0.58 to 0.78. Earlier Nepali work reports passive and "
        "active solar cabinet dryers from RECAST at Tribhuvan University "
        "and rack-type dryers at high-altitude sites. Critically, however, "
        "no peer-reviewed Nepali study reports a heat-pump dryer, solar-"
        "assisted or otherwise. The present paper therefore extends the "
        "Kathmandu University + Lund line of work by adding vapour-"
        "compression cycle integration, humidity-aware control, and "
        "seasonal TMY analysis across three climate regimes."
    )
    para(doc,
        "Internationally, comparative SAHPD modelling studies are typically "
        "limited to two or three topologies. The combination of HRX + "
        "solar + VPD-bypass on a single air path, with a systematic "
        "recirculation sweep and cross-site ranking, has not, to the "
        "authors' knowledge, been previously reported. This paper makes "
        "three contributions:"
    )
    para(doc,
        "(i) Ten configurations (HP-only, two solar-in-series, two solar-"
        "cascade, two HRX-only, three combined HRX + solar) are compared "
        "under one validated thermodynamic model, one kinetic model, one TMY "
        "dataset, and one set of operating points. Fair cross-topology "
        "ranking is the primary deliverable.",
        style="List Number")
    para(doc,
        "(ii) The HRX + solar + VPD-bypass combination (Config E2 with "
        "bypass, 10 m² collector) delivers annual-TMY SEC of 0.097 kWh/kg at "
        "the Terai site. Individual component contributions are decomposed.",
        style="List Number")
    para(doc,
        "(iii) The closed-loop recirculation penalty is mapped across three "
        "Nepalese sites spanning Terai, mid-hill, and high-hill climates, "
        "showing that the preferred configuration is not uniform within the "
        "country.",
        style="List Number")
    para(doc,
        "Section 2 describes the ten configurations. Section 3 presents the "
        "mathematical model. Section 4 describes the simulation setup and "
        "model validation. Section 5 reports results; Section 6 concludes; "
        "Section 7 lists references."
    )


def section_2_system_description(doc):
    """Section 2: System description. Describes the ten air-path configurations
    (A, B, C1, C2, D1, D2, E1, E2, E3; D3 excluded) and the VPD bypass control.
    ~1,200 words. Figure placeholders inserted at the points where the final
    paper will carry schematic diagrams of each air path."""

    doc.add_heading("2. System description", level=1)

    para(doc,
        "All ten configurations share a common dryer chamber, common set-point "
        "(T_set = 45 deg C, m_p_dry = 3.0 kg apple slices on ten trays, v_air = "
        "1.1 m/s superficial), and a common R134a vapour-compression heat pump "
        "with isentropic efficiency eta_is = 0.75. They differ only in the "
        "routing of the process air and in the placement of auxiliary "
        "components (flat-plate solar air collector, counter-flow plate HRX, "
        "VPD-triggered exhaust bypass). This section describes the air paths "
        "at a functional level; the component equations are deferred to "
        "Section 3."
    )

    # 2.1 Config A
    doc.add_heading("2.1  Config A: heat-pump-only reference", level=2)
    para(doc,
        "Config A is the reference topology. In open-loop operation (recirculation "
        "fraction r = 0) ambient air passes through the condenser, is heated to "
        "T_set, traverses the chamber, and is exhausted. The evaporator draws a "
        "separate ambient stream, sized to absorb Q_evap = Q_cond minus W_comp. "
        "In closed-loop operation (r > 0) a fraction r of the chamber exhaust is "
        "recirculated to the evaporator inlet, where it is cooled below its dew "
        "point, dehumidified, mixed with the make-up fraction (1 - r) of fresh "
        "ambient, and reheated at the condenser. This places the evaporator and "
        "the condenser in series on the same stream, which is the conventional "
        "closed-loop HPCD arrangement. No solar or heat-recovery component is "
        "present."
    )
    para(doc, "[Figure 1: Config A schematic, open- and closed-loop.]", size=9)

    # 2.2 Solar-in-series and cascade
    doc.add_heading("2.2  Configs B, C1, C2: solar-augmented variants", level=2)
    para(doc,
        "Config B places a flat-plate solar air collector in series between the "
        "evaporator-mix point and the condenser. The collector delivers partial "
        "temperature lift; the condenser trims the remaining lift to T_set. This "
        "arrangement gives the highest solar fraction when ambient is cold and "
        "insolation is high, because the collector absorbs the bulk of the heat "
        "demand; it gives a lower solar fraction in warm, cloudy conditions, "
        "because the collector outlet already exceeds T_set and must be "
        "throttled or bypassed."
    )
    para(doc,
        "Configs C1 and C2 are cascade arrangements in which the solar stream is "
        "thermally decoupled from the evaporator. In C1 (mix-before-solar) the "
        "recirculated exhaust and fresh ambient are mixed upstream of the "
        "collector and the mixed stream is then split between the solar path "
        "and the heat-pump path. In C2 (mix-after-solar) fresh ambient is "
        "pre-heated by the collector first, then mixed with the recirculated "
        "exhaust, and the mixed stream drives the heat pump. Both C1 and C2 are "
        "open-loop only in the results of this paper (r = 0), because the "
        "cascade benefit disappears when r > 0 and the evaporator must still "
        "dehumidify the combined stream."
    )
    para(doc, "[Figure 2: Configs B, C1, C2 schematics.]", size=9)

    # 2.3 HRX-only
    doc.add_heading("2.3  Configs D1, D2: heat-recovery variants", level=2)
    para(doc,
        "The D-family adds a counter-flow plate HRX (eps = 0.70) that transfers "
        "sensible heat from the chamber exhaust to an ambient supply stream, "
        "without refrigerant-side modification. Config D1 routes only the "
        "condenser-supply stream through the HRX: ambient enters the cold side, "
        "is pre-heated, and is delivered to the condenser; exhaust enters the "
        "hot side, is cooled, and is expelled. The evaporator takes a separate "
        "ambient stream unaffected by the HRX. Config D2 additionally warms the "
        "evaporator inlet by mixing a controlled fraction of the HRX-cooled "
        "exhaust with fresh ambient. This raises the evaporator saturation "
        "temperature and reduces the cycle pressure ratio at the cost of a "
        "small loss of sensible recovery on the supply side."
    )
    para(doc,
        "A third variant (D3) with the HRX streams swapped (exhaust to condenser, "
        "ambient to evaporator) was simulated during the exploratory phase but "
        "is excluded from the present paper. At Taplejung, where the cold "
        "ambient drives the supply-side dew point close to the HRX plate "
        "temperature, D3 showed a water-mass-balance residual of 2.3 % over "
        "the batch (vs < 10⁻⁶ for all other configs). The physical cause is "
        "condensation on the supply-side plate that is not captured by the "
        "sensible-only HRX model (Section 3.3); rather than add a latent-"
        "transfer sub-model for a single variant, D3 is excluded."
    )
    para(doc, "[Figure 3: Configs D1, D2 schematics.]", size=9)

    # 2.4 Combined HRX + solar
    doc.add_heading("2.4  Configs E1, E2, E3: combined HRX and solar", level=2)
    para(doc,
        "The E-family carries both an HRX and a solar collector on the supply "
        "stream. In Config E1, ambient passes through the HRX cold side and is "
        "then pre-heated further by the collector before reaching the condenser; "
        "the evaporator draws a separate ambient stream as in D1. Config E2 "
        "retains the same supply-side sequence (ambient -> HRX -> collector -> "
        "condenser) but the evaporator supply is a mixed stream of HRX-cooled "
        "exhaust and fresh ambient, sized iteratively such that the evaporator "
        "heat duty closes the first law for the specified T_cond."
    )
    para(doc,
        "Config E3 differs in the placement of the collector on the supply side: "
        "the HRX pre-heats ambient, the condenser brings the stream up, and the "
        "collector is then used as a post-heater downstream of the condenser. "
        "When the combined HRX + collector delivery already exceeds T_set, the "
        "heat pump is switched off and the chamber is supplied by solar alone; "
        "when it does not, the heat pump provides a reduced condenser lift and "
        "the collector finishes the delivery. This solar-priority control "
        "reduces compressor run time on high-insolation hours. Evaporator "
        "supply is mixed as in E2."
    )
    para(doc, "[Figure 4: Configs E1, E2, E3 schematics.]", size=9)

    # 2.5 VPD bypass
    doc.add_heading("2.5  VPD-triggered exhaust bypass", level=2)
    para(doc,
        "The VPD bypass is an optional control overlay that can be applied to "
        "Configs A, B, D1, D2, E1, E2, and E3. At every simulation time-step the "
        "vapour-pressure deficit between chamber air and product surface is "
        "evaluated using the GAB sorption isotherm for apple at the local "
        "product moisture content. When VPD falls below a threshold of 0.05 kPa "
        "the bypass is engaged: a portion of the condenser outlet is routed "
        "around the chamber and expelled, and the compressor is unloaded in "
        "proportion. A three-fold hysteresis (engagement at VPD < 0.05 kPa, "
        "release at VPD > 0.15 kPa) and a 600-second minimum dwell time prevent "
        "oscillation on the shoulder of the transition. The bypass thus acts as "
        "a rate-aware duty cut that is operationally justified: compressor "
        "energy spent during the low-VPD tail removes very little additional "
        "moisture per kWh (Section 1.5)."
    )
    para(doc,
        "The bypass is treated as a separate design variable in Section 5, "
        "reported in two columns per configuration (bypass off, bypass on), so "
        "that its marginal contribution to SEC can be read directly from the "
        "results tables."
    )

    # 2.6 Summary of configuration space
    doc.add_heading("2.6  Summary of the configuration space", level=2)
    para(doc,
        "The ten configurations span three axes of design choice: solar "
        "integration (absent in A, D1, D2; present in B, C1, C2, E1, E2, E3), "
        "heat recovery (absent in A, B, C1, C2; present in D1, D2, E1, E2, E3), "
        "and recirculation fraction r (scanned from 0 to 1 on Configs A and B). "
        "The VPD bypass is orthogonal and toggled on or off for each "
        "configuration that admits it. This three-axis scope is the basis for "
        "the comparative ranking in Section 5."
    )


def section_3_mathematical_model(doc):
    """Section 3: Mathematical model. Heavy equations. ~1,800 words."""

    doc.add_heading("3. Mathematical model", level=1)
    para(doc,
        "This section presents the component models, the thermodynamic closure "
        "equations that link them, and the drying and sorption submodels. All "
        "quantities are in SI units. Vector state variables are written without "
        "ornamentation; specific quantities carry a lower-case symbol (h for "
        "specific enthalpy, cp for specific heat). Air-side properties are "
        "moist-air properties on the dry-air basis."
    )

    # 3.1 Vapour-compression cycle
    doc.add_heading("3.1  Vapour-compression cycle", level=2)
    para(doc,
        "The heat pump uses R134a and is implemented as a four-process cycle on "
        "refrigerant-property tables (CoolProp PropsSI [Bell et al., 2014]). Compression 1 -> 2 is "
        "non-isentropic with fixed isentropic efficiency eta_is and an "
        "additional mechanical-efficiency factor eta_mech on shaft power; "
        "condensation 2 -> 3 is at fixed high-side pressure with fixed "
        "sub-cooling dT_sc; throttling 3 -> 4 is isenthalpic; evaporation "
        "4 -> 1 is at fixed low-side pressure with fixed superheat dT_sh. The "
        "cycle duties are:"
    )
    equation(doc, "h₂ = h₁ + (h₂ₛ − h₁) / η_is", "1")
    equation(doc, "Q̇_cond = ṁ_ref · (h₂ − h₃)", "2")
    equation(doc, "Q̇_evap = ṁ_ref · (h₁ − h₄)", "3")
    equation(doc, "Ẇ_comp = ṁ_ref · (h₂ − h₁) / η_mech", "4")
    equation(doc, "COP = Q̇_cond / Ẇ_comp", "5")
    nomenclature(doc, [
        ("h₁",     "refrigerant specific enthalpy at evaporator outlet (superheated vapour at T_evap + dT_sh, p_evap) [kJ/kg]"),
        ("h₂",     "refrigerant specific enthalpy at compressor outlet (actual) [kJ/kg]"),
        ("h₂ₛ",    "refrigerant specific enthalpy at end of an isentropic compression to p_cond from state 1 [kJ/kg]"),
        ("h₃",     "refrigerant specific enthalpy at condenser outlet (subcooled liquid at T_cond − dT_sc) [kJ/kg]"),
        ("h₄",     "refrigerant specific enthalpy after the throttle (h₄ = h₃) [kJ/kg]"),
        ("η_is",   "compressor isentropic efficiency; fixed at 0.75 in this work"),
        ("η_mech", "compressor mechanical efficiency on shaft power; fixed at 0.95"),
        ("ṁ_ref",  "refrigerant mass flow rate, solved from the condenser-duty target [kg/s]"),
        ("Q̇_cond", "condenser thermal duty [kW]"),
        ("Q̇_evap", "evaporator thermal duty [kW]"),
        ("Ẇ_comp", "compressor shaft power [kW]; motor losses are not modelled separately"),
        ("COP",    "coefficient of performance, Q̇_cond / Ẇ_comp"),
        ("dT_sh",  "evaporator superheat; 5 K"),
        ("dT_sc",  "condenser subcooling; 5 K"),
    ])
    para(doc,
        "The condensing and evaporating temperatures are set by a pinch model "
        "on the refrigerant side: the condensing saturation temperature is "
        "T_cond = T_set + dT_pinch (dT_pinch = 10 K), and the evaporating "
        "saturation temperature is placed 10 K below the evaporator-source "
        "air temperature. The air-side coil effectiveness (eq. 6) then "
        "determines the actual air outlet temperature, which can exceed "
        "T_set when the inlet is cold. In the code the condenser air outlet "
        "is capped at T_set; any residual capacity above T_set is unused. "
        "This means the heat pump is slightly oversized relative to the "
        "minimum required lift, which is conservative for SEC. The air-side "
        "equations are:"
    )
    equation(doc, "T_air,cond,out = T_air,cond,in + ε_cond · (T_cond − T_air,cond,in)", "6")
    equation(doc, "T_air,evap,out = T_air,evap,in − ε_evap · (T_air,evap,in − T_evap,coil)", "7")
    nomenclature(doc, [
        ("T_cond",          "refrigerant saturation temperature on the condenser side [°C]"),
        ("T_evap",          "refrigerant saturation temperature on the evaporator side [°C]"),
        ("T_evap,coil",     "evaporator coil-surface temperature, T_evap + dT_approach with dT_approach = 3 K"),
        ("ε_cond, ε_evap",  "air-side sensible effectiveness of the condenser / evaporator coil; fixed at 0.85"),
    ])

    # 3.2 Collector
    doc.add_heading("3.2  Solar air collector (Hottel-Whillier-Bliss)", level=2)
    para(doc,
        "The flat-plate solar air collector is modelled with the steady-state "
        "Hottel-Whillier-Bliss formulation [Hottel and Whillier, 1955; Bliss, 1959; Duffie and Beckman, 2013]. The implementation splits the "
        "useful gain into absorbed-heat and loss terms, then pre-multiplies by "
        "the heat-removal factor F_R:"
    )
    equation(doc, "Q̇_abs = A_c · (τα) · K_θ · G_T", "8")
    equation(doc, "Q̇_loss = A_c · U_L · (T_in − T_amb)", "9")
    equation(doc, "Q̇_sol = F_R · (Q̇_abs − Q̇_loss)   ,   Q̇_sol ≥ 0", "10")
    equation(doc, "F_R = (C_min / UA) · [1 − exp(−UA · F′ / C_min)]", "11")
    equation(doc, "η_c = Q̇_sol / (A_c · G_T)", "12")
    equation(doc, "T_out = T_in + Q̇_sol / (ṁ_a · cp_a)", "13")
    nomenclature(doc, [
        ("A_c",   "gross collector area [m²]; swept from 2 to 10 m² in the area sweep (Section 5)"),
        ("τα",    "transmittance-absorptance product (optical efficiency at normal incidence); 0.75"),
        ("K_θ",   "incidence-angle modifier; 1.0 in this work (annual-average normal-equivalent)"),
        ("U_L",   "overall thermal-loss coefficient [W m⁻² K⁻¹]; 5.0 (single-glazed flat plate)"),
        ("F′",    "collector efficiency factor; 0.90"),
        ("F_R",   "heat-removal factor (derived from F′, UA, and C_min via eq. 11)"),
        ("G_T",   "in-plane irradiance from the TMY file [W/m²]"),
        ("T_in",  "collector air inlet temperature [°C]"),
        ("T_amb", "ambient dry-bulb temperature [°C]"),
        ("ṁ_a",   "air mass flow rate through the collector [kg/s]"),
        ("cp_a",  "specific heat of dry air at constant pressure; 1.006 kJ/(kg·K)"),
        ("C_min", "minimum heat-capacity rate, C_min = ṁ_a · cp_a [kW/K]"),
        ("UA",    "A_c · U_L [kW/K]"),
    ])
    para(doc,
        "When Q̇_sol > 0 but the resulting T_out exceeds T_set, the collector is "
        "bypassed rather than throttled: the air is routed around the plate and "
        "no useful gain is credited, because any clipped gain would otherwise "
        "have to be dumped by the heat pump running in reverse, which is not "
        "modelled. A stagnation limit T_stag,max = 150 °C is imposed for "
        "diagnostic purposes and is never approached in the TMY data used."
    )

    # 3.3 HRX
    doc.add_heading("3.3  Counter-flow plate HRX", level=2)
    para(doc,
        "The exhaust-to-supply heat-recovery exchanger is modelled with a fixed "
        "sensible-heat effectiveness ε_HRX under balanced flow (ṁ_s = ṁ_e, "
        "C_r = 1). The supply and exhaust outlet temperatures are:"
    )
    equation(doc, "T_s,out = T_s,in + ε_HRX · (T_e,in − T_s,in)", "14")
    equation(doc, "T_e,out = T_e,in − ε_HRX · (T_e,in − T_s,in)", "15")
    equation(doc, "Q̇_HRX = ṁ · cp_a · (T_s,out − T_s,in)", "16")
    para(doc,
        "Equations (14)-(15) are equivalent to the general effectiveness "
        "relation Q̇ = ε_HRX · C_min · (T_e,in − T_s,in) under balanced flow. "
        "For a counter-flow plate exchanger operating at C_r = 1, the closed-"
        "form ε-NTU result is ε = NTU / (1 + NTU); the present study fixes "
        "ε_HRX = 0.70, which corresponds to NTU ≈ 2.3 and is a representative "
        "value for brazed-aluminium or corrugated-polypropylene cores of "
        "practical size."
    )
    nomenclature(doc, [
        ("ε_HRX",           "sensible-heat effectiveness of the HRX; fixed at 0.70"),
        ("T_s,in, T_s,out", "supply-side inlet and outlet air temperatures [°C]"),
        ("T_e,in, T_e,out", "exhaust-side inlet and outlet air temperatures [°C]"),
        ("ṁ",               "common air mass flow rate on both sides (balanced) [kg/s]"),
    ])
    para(doc,
        "Condensation on the exhaust-side plate is permitted when T_e,out falls "
        "below the exhaust dew-point; in that case the supply-side moisture "
        "ratio is unchanged and the exhaust-side moisture ratio is set to "
        "saturation at T_e,out. Any resulting condensate mass enters the "
        "water-mass balance of Section 3.7 as a negative source term."
    )

    # 3.4 Drying kinetics
    doc.add_heading("3.4  Drying kinetics (three-stage fit on apple slice data)", level=2)
    para(doc,
        "The drying-kinetics sub-model used at chamber scale was obtained from "
        "a three-stage fit on thirteen constant-condition apple-slice drying "
        "experiments spanning temperature, relative humidity, air velocity, "
        "and slice-thickness sweeps. The three stages are:"
    )
    para(doc,
        "Phase 1 — per-curve recursive piecewise Page/Midilli fit. Each "
        "experimental MR(t) curve is first smoothed with a LOWESS filter, "
        "then split at an optimal t_split into a left (falling-rate-dominated) "
        "segment and a right (tail) segment. On each segment one of four "
        "candidate forms is fitted with scipy.optimize.curve_fit and selected "
        "by corrected Akaike information criterion (AICc):",
        style="List Number")
    equation(doc, "MR_Page(t) = exp(−k · tⁿ)", "17a")
    equation(doc, "MR_Midilli(t) = exp(−k · tⁿ) + b · t    [Midilli et al., 2002]", "17b")
    para(doc,
        "with optional left-shift τ (Page_shift, Midilli_shift). Monotonicity, "
        "join-continuity, and tail-slope continuity are enforced by soft "
        "penalties; the output of Phase 1 for each dataset is the tuple "
        "(family_L, k_L, n_L, b_L, t_split, family_R, k_R, n_R, b_R, "
        "join_offset).",
        style="List Number")
    para(doc,
        "Phase 2 — regression of the piecewise parameters on operating "
        "conditions. The Phase-1 parameter tuples are regressed on "
        "(T, RH, v, d) using scikit-learn pipelines (SimpleImputer, "
        "StandardScaler, then Ridge / Lasso / LinearRegression / "
        "RandomForest), with log or signed-log target transforms on the "
        "positive and signed coefficients respectively. Model selection is "
        "by five-fold cross-validated RMSE. The Phase-2 output is a set of "
        "regressors g(T, RH, v, d) → {k_L, n_L, b_L, t_split, k_R, n_R, b_R, "
        "offset} that predict the piecewise Midilli/Page parameters at any "
        "operating point in the validity box.",
        style="List Number")
    para(doc,
        "Phase 3 — first-order chamber update for TMY-scale simulation. The "
        "piecewise-Midilli reconstruction is accurate on any single batch but "
        "is expensive to integrate under transient boundary conditions "
        "(TMY-driven T and RH). For the ten-configuration sweep of this "
        "paper, a four-factor log-linear effective coefficient is fitted by "
        "ordinary least squares (in log space) across all thirteen "
        "experiments:",
        style="List Number")
    equation(doc,
        "ln k_eff = ln K_ref + (E_a/R)·(1/T_ref − 1/T) − α_RH · RH + γ_v · ln(v/v_ref) + δ_d · ln(d_ref/d)",
        "17")
    para(doc,
        "and used in a first-order relaxation of X toward the GAB equilibrium "
        "moisture content X_e:"
    )
    equation(doc, "dX / dt = −k_eff(T, RH, v, d) · (X − X_e)", "18")
    equation(doc, "ṁ_w = m_p,dry · k_eff · (X − X_e)", "19")
    nomenclature(doc, [
        ("k_eff",   "effective first-order drying coefficient [1/s]"),
        ("X",       "instantaneous dry-basis moisture content [kg water / kg dry solid]"),
        ("X_e",     "equilibrium dry-basis moisture content from the GAB isotherm (eq. 20)"),
        ("m_p,dry", "dry mass of product in the chamber; 3.0 kg"),
        ("ṁ_w",     "instantaneous moisture-removal mass rate [kg/s]"),
        ("K_ref",   "kinetic coefficient at (T_ref, RH = 0, v_ref, d_ref); 1.63 × 10⁻⁴ s⁻¹"),
        ("E_a/R",   "effective activation energy divided by R; 2711 K"),
        ("α_RH",    "humidity-suppression coefficient; 1.75 (dimensionless; RH is expressed as a fraction 0-1 in eq. 17)"),
        ("γ_v",     "air-velocity exponent, fitted from the v ∈ {0.60, 0.85, 1.10} m/s sweep"),
        ("δ_d",     "slice-thickness exponent, fitted from the d ∈ {4, 6, 8, 10} mm sweep"),
        ("T_ref",   "reference temperature for the fit; 318.15 K (45 °C)"),
        ("v_ref",   "reference velocity; 1.1 m/s"),
        ("d_ref",   "reference slice thickness; 6 mm"),
    ])
    para(doc,
        "The Phase-3 log-linear fit on all thirteen data points delivers "
        "R² ≈ 0.90 on ln k_eff. Equations (17)-(19) form the kinetic sub-model "
        "actually coupled to the thermodynamic model in the TMY runs of "
        "Section 5. The Phase-2 piecewise-parameter regressors are retained "
        "as a higher-fidelity alternative for single-batch reconstructions "
        "(Section 4.3 uses the Phase-3 form; the Phase-1/2 pipeline "
        "validation is documented in the authors' companion work on "
        "kinetic identification)."
    )

    # 3.5 GAB
    doc.add_heading("3.5  Equilibrium moisture content (GAB isotherm)", level=2)
    para(doc,
        "The equilibrium moisture content for apple is computed from the "
        "three-parameter Guggenheim-Anderson-de Boer isotherm with temperature-"
        "dependent parameters. The GAB constants used here are fitted to "
        "apple sorption data from Vega-Gálvez et al. [2012] at 40-60 °C:"
    )
    equation(doc, "X_e = (X_m · C · K · a_w) / [(1 − K·a_w) · (1 − K·a_w + C·K·a_w)]", "20")
    equation(doc, "X_m(T) = X_m,0 · exp(ΔH_xm / (R·T))", "21")
    equation(doc, "C(T)   = C_0   · exp(ΔH_C  / (R·T))", "22")
    equation(doc, "K(T)   = K_0   · exp(ΔH_K  / (R·T))", "23")
    nomenclature(doc, [
        ("X_e",   "equilibrium moisture content on dry basis [kg/kg db]"),
        ("a_w",   "water activity, taken equal to chamber RH (fraction), clamped at 0.95"),
        ("X_m",   "monolayer capacity [kg/kg db]"),
        ("C",     "Guggenheim energy constant"),
        ("K",     "multilayer factor"),
        ("X_m,0", "monolayer pre-exponential; 3.141 × 10⁻³ kg/kg db"),
        ("ΔH_xm", "monolayer sorption enthalpy; +8057 J/mol (positive, so X_m decreases with increasing T as the exponent +ΔH/(RT) shrinks)"),
        ("C_0",   "Guggenheim pre-exponential; 4.923 × 10⁻³"),
        ("ΔH_C",  "Guggenheim enthalpy; 17241 J/mol"),
        ("K_0",   "multilayer pre-exponential; 0.9904"),
        ("ΔH_K",  "multilayer enthalpy; ≈ 0 J/mol (K effectively temperature-independent)"),
        ("R",     "universal gas constant; 8.314 J/(mol·K)"),
        ("T",     "temperature [K]"),
    ])

    # 3.6 Psychrometrics and control
    doc.add_heading("3.6  Psychrometrics, VPD, and bypass control", level=2)
    para(doc,
        "The moist-air state is described on the dry-air basis. The humidity "
        "ratio, specific enthalpy, and saturation vapour pressure follow the "
        "standard HVAC relations, with saturation pressure from the Magnus "
        "form as improved by Alduchov and Eskridge [1996]:"
    )
    equation(doc, "ω = 0.62198 · p_w / (p − p_w)", "24")
    equation(doc, "h = cp_a · T + ω · (h_fg + cp_v · T)", "25")
    equation(doc, "p_w,sat(T) = 610.94 · exp[17.625 · T / (T + 243.04)]", "26")
    nomenclature(doc, [
        ("ω",         "humidity ratio [kg water / kg dry air]"),
        ("p",         "total atmospheric pressure [Pa]; 101,325 (sea level) or site-corrected"),
        ("p_w",       "partial pressure of water vapour [Pa]"),
        ("p_w,sat",   "saturation vapour pressure of water [Pa]"),
        ("h",         "moist-air specific enthalpy on dry-air basis [kJ/kg]"),
        ("cp_a",      "specific heat of dry air; 1.006 kJ/(kg·K)"),
        ("cp_v",      "specific heat of water vapour; 1.86 kJ/(kg·K)"),
        ("h_fg",      "reference latent heat of water; 2501 kJ/kg"),
        ("T",         "air temperature [°C] in eqs. (25)-(26)"),
    ])
    para(doc,
        "The vapour-pressure deficit driving drying is the gap between the "
        "saturation vapour pressure at the product surface temperature and the "
        "actual partial pressure of water vapour in the chamber air:"
    )
    equation(doc, "VPD = p_w,sat(T_prod) − p_w(chamber)", "27")
    para(doc,
        "The bypass control engages and releases on VPD, with a three-fold "
        "hysteresis and a minimum dwell time to suppress chattering:"
    )
    equation(doc, "bypass ON   if  VPD < VPD_on   and   Δt_since_switch > τ_dwell", "28")
    equation(doc, "bypass OFF  if  VPD > VPD_off  and   Δt_since_switch > τ_dwell", "29")
    nomenclature(doc, [
        ("T_prod",          "product-surface temperature (taken equal to chamber air temperature at quasi-steady conditions) [°C]"),
        ("VPD_on",          "engagement threshold; 0.05 kPa"),
        ("VPD_off",         "release threshold; 0.15 kPa (three-fold margin over VPD_on)"),
        ("τ_dwell",         "minimum dwell time between switch events; 600 s"),
        ("Δt_since_switch", "elapsed time since the last state change [s]"),
    ])

    # 3.7 First-law closure
    doc.add_heading("3.7  First-law and water-mass closure", level=2)
    para(doc,
        "At every integration step the refrigerant mass flow rate and the "
        "saturation temperatures are solved jointly with the air-side energy "
        "balance so that the cycle first law is satisfied to machine "
        "precision:"
    )
    equation(doc, "Q̇_cond − Q̇_evap − Ẇ_comp = 0", "30")
    para(doc,
        "The chamber water-mass balance over one time-step Δt is:"
    )
    equation(doc, "Δm_w,chamber = (ṁ_s · ω_in − ṁ_s · ω_out) · Δt + ṁ_w · Δt", "31")
    nomenclature(doc, [
        ("ω_in, ω_out", "humidity ratio at chamber inlet and outlet [kg/kg da]"),
        ("ṁ_s",         "supply-air mass flow rate on dry-air basis [kg/s]"),
        ("ṁ_w",         "moisture-removal rate from the product, eq. (18)"),
        ("Δm_w,chamber","net water mass accumulated in the chamber over Δt [kg] (~ 0 at quasi-steady)"),
    ])
    para(doc,
        "Equations (30) and (31) are enforced at each time-step by a fixed-"
        "point iteration (for E2 and E3, an outer iteration on the evaporator "
        "mix ratio is additionally required). Section 4.3 demonstrates that "
        "the residual on both (30) and (31) is below 1 × 10⁻⁶ on all six "
        "validation configurations."
    )

    # 3.8 Integrated SEC
    doc.add_heading("3.8  Performance metric (SEC)", level=2)
    para(doc,
        "The specific energy consumption over a batch is the ratio of "
        "electrical input to water removed:"
    )
    equation(doc, "SEC = ∫₀ᵗᶠ (Ẇ_comp + Ẇ_fan) dt  /  ∫₀ᵗᶠ ṁ_w dt", "32")
    nomenclature(doc, [
        ("Ẇ_fan",     "circulating-fan electrical input [kW]; computed from dP × V̇ / η_fan (η_fan = 0.60), typically 25 W (2-5 % of Ẇ_comp depending on configuration)"),
        ("t_f",       "batch end time, defined by X(t_f) = X_target with X_target = 0.10 kg/kg db"),
        ("SEC",       "specific energy consumption [kWh per kg water removed]"),
    ])
    para(doc,
        "Solar gain Q̇_sol is not credited against the numerator of equation "
        "(32); it enters through the reduced condenser duty that the heat "
        "pump has to supply, and therefore through the reduced Ẇ_comp."
    )


def section_4_simulation_setup(doc):
    """Section 4: Simulation setup and model validation. ~700 words."""

    doc.add_heading("4. Simulation setup and model validation", level=1)

    # 4.1 Sites
    doc.add_heading("4.1  Sites and typical-meteorological-year data", level=2)
    para(doc,
        "Three Nepalese sites are simulated. They span the principal climatic "
        "regimes of the country and cover the main horticultural production "
        "zones. For each site the typical-meteorological-year (TMY) hourly "
        "series of ambient dry-bulb temperature, relative humidity, global "
        "horizontal irradiance, and in-plane irradiance (south-facing, tilt "
        "equal to site latitude) is drawn from the PVGIS SARAH-2 dataset [PVGIS JRC, 2024]. "
        "Atmospheric pressure at each site is computed from elevation using "
        "the standard atmosphere (p_atm = 101,325 · (1 − 0.0065 · z / 288.15)"
        "^5.2561), then held fixed in the psychrometric calculations of "
        "Section 3.6. The PVGIS SARAH-2 dataset has a stated mean bias of "
        "±2 % on GHI for the South Asian region; propagated through the "
        "collector model this corresponds to ~±1-2 % uncertainty on SEC "
        "for solar-integrated configurations. The TMY construction averages "
        "inter-annual variability, so the reported SEC values represent a "
        "climatological mean rather than any specific year."
    )

    # Table 4.1 — site parameters
    t = doc.add_table(rows=4, cols=5)
    t.style = "Light Grid Accent 1"
    hdr = ["Site", "Elevation [m]", "p_atm [Pa]", "Lat. [°N]", "Climate regime"]
    rows = [
        ["Kathmandu",  "1,350", "86,120",  "27.7", "Mid-hill, monsoonal, mild winter"],
        ["Biratnagar", "72",    "100,460", "26.5", "Terai lowland, hot-humid tropical"],
        ["Taplejung",  "1,820", "81,400",  "27.4", "High-hill, cool, high irradiance"],
    ]
    for j, h in enumerate(hdr):
        cell = t.rows[0].cells[j]; cell.text = h
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            cell = t.rows[i].cells[j]; cell.text = v
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(10)
    para(doc, "Table 4.1  Site parameters used in the TMY simulations.", size=9)

    para(doc,
        "Kathmandu represents the mid-hill urban cluster where a large fraction "
        "of retail processing capacity sits. Biratnagar is the Terai benchmark "
        "for tropical-fruit drying (mango, pineapple, banana). Taplejung is the "
        "high-hill benchmark with the highest annual irradiance of the three "
        "and the coldest winter ambient. Seasonal sub-series (winter Dec-Jan, "
        "spring Mar-Apr, autumn Oct-Nov) are extracted from the same TMY "
        "files for the seasonal analyses of Section 5; the monsoon months "
        "(Jun-Sep) are excluded from the seasonal tables because the dryer is "
        "not typically operated through the monsoon and the TMY-derived annual "
        "results already carry the monsoon penalty in their full-year averages."
    )

    para(doc, "[Figure 5: annual-TMY profiles of T_amb, RH_amb, and G_T "
        "for the three sites.]", size=9)

    # 4.2 Operating inputs
    doc.add_heading("4.2  Operating inputs and design choices", level=2)
    para(doc,
        "All simulations use a single set of operating inputs, fixed across "
        "configurations and sites so that inter-configuration differences are "
        "attributable to the air-path routing alone:"
    )
    nomenclature(doc, [
        ("T_set",       "chamber air set-point temperature; 45 °C"),
        ("m_p,dry",     "dry mass of apple product per batch; 3.0 kg (≈ 22.5 kg wet at X₀ = 6.5 kg/kg db)"),
        ("N_trays",     "number of trays in the chamber; 10"),
        ("X₀",          "initial dry-basis moisture content; 6.5 kg/kg db"),
        ("X_target",    "batch-end dry-basis moisture content; 0.10 kg/kg db"),
        ("v_air",       "superficial air velocity through the tray stack; 1.1 m/s"),
        ("d_slice",     "apple slice thickness; 5 mm"),
        ("Refrigerant", "R134a (GWP = 1430; selected because T_crit = 101 °C permits T_cond up to 70 °C with adequate sub-cooling margin)"),
        ("η_is",        "compressor isentropic efficiency; 0.75"),
        ("η_mech",      "compressor mechanical efficiency; 0.95"),
        ("ε_cond, ε_evap", "air-side coil effectiveness; 0.85"),
        ("ε_HRX",       "HRX sensible effectiveness; 0.70"),
        ("A_c",         "collector area; 10 m² (baseline), swept 2-10 m² in the area-sweep runs"),
        ("F_R·(τα), U_L", "collector parameters as Section 3.2 (τα = 0.75, U_L = 5.0 W/m²K)"),
        ("r",           "recirculation fraction; 0 (open-loop) baseline, swept 0.0-1.0 on A and B"),
        ("VPD_on, VPD_off", "bypass thresholds; 0.05 and 0.15 kPa"),
        ("Δt",          "simulation time-step; 60 s"),
    ])
    para(doc,
        "The condenser and evaporator are sized at each time-step to meet the "
        "instantaneous thermal demand at T_set, rather than fixed at a "
        "nameplate rating. The reported SEC therefore reflects an "
        "idealised inverter-driven compressor behaviour; fixed-speed "
        "compressor effects (cycling losses, part-load degradation) are not "
        "included and are flagged in Section 6 as the most important "
        "experimental extension."
    )

    # 4.3 Validation
    doc.add_heading("4.3  Model validation and energy-balance check", level=2)
    para(doc,
        "Three internal consistency checks are imposed on every simulation and "
        "were completed on 2026-04-09 for a representative batch of six "
        "configurations (A r = 0, A r = 0.9, D1, D2, E1, E2) at Kathmandu and "
        "Biratnagar. The three checks are:"
    )
    para(doc,
        "First-law residual on the vapour-compression cycle, |Q̇_cond − Q̇_evap "
        "− Ẇ_comp| / Q̇_cond, maximum across all time-steps of each batch. The "
        "measured residual was below 1 × 10⁻⁶ on all six configurations; a "
        "numerical target of 1 × 10⁻⁴ was set a priori.",
        style="List Number")
    para(doc,
        "Water-mass closure on the chamber, |Σ Δm_w,chamber − m_w,removed,"
        "product| / m_w,removed, summed over the batch. The measured residual "
        "was below 1 × 10⁻⁶ on all six configurations.",
        style="List Number")
    para(doc,
        "Psychrometric round-trip consistency, (ω → T, RH) → (ω′) with |ω − ω′| "
        "as the residual. The measured residual was below 4 × 10⁻⁶ across all "
        "time-steps, driven by the bisection tolerance on the dew-point solver.",
        style="List Number")
    para(doc,
        "In addition to the numerical closures, the following physical-"
        "plausibility checks were verified on every run and failed none:"
    )
    para(doc,
        "COP in the range 3.5-4.8 across all sites and seasons. The Carnot "
        "COP for the typical lift (T_evap ≈ 10 °C, T_cond ≈ 55 °C) is "
        "T_cond/(T_cond - T_evap) ≈ 7.3; the simulated COP/COP_Carnot ratio "
        "of 0.61-0.62 is consistent with η_is = 0.75 and η_mech = 0.95 "
        "(the product η_is × η_mech = 0.71 maps to ~0.61 Carnot ratio after "
        "accounting for superheat and subcooling losses in the real-gas cycle).",
        style="List Bullet")
    para(doc,
        "No frost flag on the evaporator in any run (T_evap,coil ≥ 0 °C in all "
        "configurations under the Nepalese TMY envelope).",
        style="List Bullet")
    para(doc,
        "All configurations deliver T_to_chamber = T_set at every time-step, "
        "so the drying kinetics (and the drying time) are identical across "
        "configurations; the differences are entirely in energy input.",
        style="List Bullet")
    para(doc,
        "Against published experimental work, the open-loop Config A annual "
        "SEC of 0.72 kWh/kg at Kathmandu and 0.54 kWh/kg at Biratnagar is "
        "within the 0.6-1.2 kWh/kg range reported for optimised HPCD "
        "simulations (Minea 2013) and is broadly consistent with the "
        "experimental 1.3-2.5 kWh/kg envelope reported for smaller lab-scale "
        "systems by Prasertsan and Saen-saby (1998) and Chua et al. (2002). "
        "The simulated values sit on the lower end of the published range, "
        "which is expected for an idealised-compressor model operating at a "
        "relatively low set-point (45 °C) and against warm ambient for a "
        "substantial fraction of the year."
    )
    para(doc,
        "Kinetic-sub-model sensitivity. The thin-layer kinetic parameters "
        "were calibrated on a single apple-slice dataset at 82 kPa ambient "
        "[Royen et al., 2020]. The Phase-3 log-linear fit on these data (Section 3.4) gives R² = 0.90 on ln k_eff. Two "
        "sensitivity checks were run. First, K_ref was perturbed by ±20 %; "
        "the resulting change in batch time is ±12-15 % and SEC changes "
        "in proportion (−11 % to +14 %). Crucially, the relative ranking "
        "of the ten configurations is preserved under either perturbation: "
        "E2 + bypass remains the lowest-SEC topology at both sites, and "
        "the relative gaps between configurations change by less than "
        "3 percentage points. Second, the humidity-suppression exponent "
        "α_RH was perturbed between 1.5 and 2.0 (the 95 % CI of the fit); "
        "the absolute SEC shifts by ±6 % but cross-configuration rankings "
        "again hold. This supports interpreting the absolute SEC values as "
        "a best-case bound and the comparative rankings as robust. "
        "Cross-dataset validation was not performed, which is stated as a "
        "limitation in §6; the reported ranges for apple drying by "
        "Sacilik & Elicin [2006], Vega-Gálvez et al. [2012] and Velic et "
        "al. [2004] give K_ref within ±20 % of the value used here, "
        "consistent with the perturbation envelope tested."
    )
    embed_figure(doc, "outputs/config_A/biratnagar/plots_r0.0/baseline_r0.0_overview.png",
        "Figure 5b. Example simulation output: Config A (baseline, r = 0) at Biratnagar. Six-panel overview of a single 19.2-kg-water batch showing (top left) tray-level moisture content evolution, (top right) air-side temperature profile, (middle left) instantaneous COP, (middle right) compressor/fan power and cumulative electrical energy, (bottom left) relative humidity at ambient, chamber inlet, and exhaust, and (bottom right) cumulative water removed against target. Drying time = 14.8 h, mean COP = 4.07.",
        width_in=6.0)

    # 4.4 Effectiveness sensitivity and time-step convergence
    doc.add_heading("4.4  Effectiveness sensitivity and time-step convergence", level=2)
    para(doc,
        "Two reviewer-requested robustness checks were performed on the "
        "heat-exchanger effectiveness parameters and the integration "
        "time-step."
    )
    para(doc,
        "Effectiveness sensitivity. Three parameters were swept independently "
        "at Kathmandu: condenser effectiveness ε_cond (0.75, 0.85, 0.95), "
        "evaporator effectiveness ε_evap (0.75, 0.85, 0.95), and HRX "
        "effectiveness ε_HRX (0.60, 0.70, 0.80). The results are summarised "
        "below."
    )
    # F1 sensitivity table
    _add_table(doc,
        ["Parameter", "Value", "Config", "SEC (kWh/kg)"],
        [
            ["ε_cond", "0.75", "A", "0.756"],
            ["ε_cond", "0.85", "A", "0.717"],
            ["ε_cond", "0.95", "A", "0.717"],
            ["ε_evap", "0.75", "E2", "0.199"],
            ["ε_evap", "0.85", "E2", "0.197"],
            ["ε_evap", "0.95", "E2", "0.196"],
            ["ε_HRX",  "0.60", "E2", "0.225"],
            ["ε_HRX",  "0.70", "E2", "0.197"],
            ["ε_HRX",  "0.80", "E2", "0.172"],
            ["ε_HRX",  "0.60", "D1", "0.416"],
            ["ε_HRX",  "0.70", "D1", "0.365"],
            ["ε_HRX",  "0.80", "D1", "0.314"],
        ],
        "Table 4.2  Effectiveness sensitivity at Kathmandu (annual TMY, r = 0). "
        "ε_HRX is the dominant lever: a ±0.10 change in ε_HRX shifts SEC by "
        "~24 % on D1 and ~13 % on E2. Condenser and evaporator effectiveness "
        "have minor effects (< 5 %) because the pinch model already sets the "
        "saturation temperatures.",
    )
    para(doc,
        "ε_HRX is by far the most influential parameter: raising it from 0.60 "
        "to 0.80 reduces D1 SEC by 24 % and E2 SEC by 24 %. Condenser and "
        "evaporator effectiveness have minor effects (< 5 %) because the "
        "fixed-pinch model determines the saturation temperatures independently "
        "of air-side effectiveness. This means accurate characterisation of the "
        "HRX is the most important experimental input for D- and E-family "
        "predictions."
    )
    # F2 convergence table
    para(doc,
        "Time-step convergence. Configs A and E2 were re-run at Kathmandu "
        "with Δt = 30 s, 60 s, and 120 s:"
    )
    _add_table(doc,
        ["Config", "Δt = 30 s", "Δt = 60 s", "Δt = 120 s"],
        [
            ["A",  "0.717", "0.717", "0.715"],
            ["E2", "0.197", "0.197", "0.197"],
        ],
        "Table 4.3  Time-step convergence: SEC [kWh/kg] at Kathmandu, annual TMY. "
        "The maximum variation across a 4× range in Δt is 0.3 % (Config A), "
        "confirming that Δt = 60 s is adequate.",
    )
    para(doc,
        "The maximum SEC variation across the 4× time-step range is 0.3 % "
        "(Config A) and < 0.1 % (E2), confirming that Δt = 60 s provides "
        "sufficient temporal resolution for the quasi-steady boundary "
        "conditions of this study."
    )


def _fmt(x, d=3):
    try:
        if pd.isna(x):
            return "-"
        return f"{float(x):.{d}f}"
    except Exception:
        return "-"


def _add_table(doc, headers, rows, caption=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = str(h)
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            c = t.rows[i].cells[j]; c.text = str(v)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9.5)
    if caption:
        para(doc, caption, size=9)
    return t


def section_5_results(doc):
    """Section 5: Results and discussion. Pulled from master_summary.csv."""
    df = pd.read_csv(SUM_CSV)
    df = df[~df["filename"].astype(str).str.contains("_s4", na=False)]

    doc.add_heading("5. Results and discussion", level=1)
    para(doc,
        "Two metrics are reported throughout: specific energy consumption "
        "SEC = (W_comp + W_fan) / m_water [kWh/kg] (lower is better), and "
        "specific moisture extraction rate SMER = m_water / (W_comp + W_fan) "
        "[kg/kWh] (higher is better); SMER = 1/SEC and is reported "
        "alongside SEC to aid comparison with prior SAHPD literature, where "
        "SMER is the more common headline figure. Solar fraction SF = "
        "Q_solar_usable / Q_cond. Results follow: baseline and site-dependence "
        "(§5.1), solar integration (§5.2), heat recovery (§5.3), VPD bypass "
        "(§5.4), combined topologies (§5.5), seasonal variation (§5.6), and "
        "synthesis with a configuration-ranking table (§5.7)."
    )

    # 5.1 Baseline
    doc.add_heading("5.1  Baseline (Config A) and site-dependence of the closed-loop penalty", level=2)
    a = df[(df["config"] == "A") & (df["season"] == "annual") & (df["vpd_bypass"] == False)]
    rows = []
    for loc in ["kathmandu", "biratnagar", "taplejung"]:
        row = a[a["location"] == loc].set_index("r_recirc")
        vals = [loc.title()]
        for r in [0.0, 0.3, 0.5, 0.7, 0.9, 1.0]:
            if r in row.index:
                vals.append(_fmt(row.loc[r, "SEC_kWh_per_kg"], 3))
            else:
                vals.append("-")
        rows.append(vals)
    _add_table(doc,
        ["Site", "r=0", "r=0.3", "r=0.5", "r=0.7", "r=0.9", "r=1.0"],
        rows,
        "Table 5.1  Config A annual-TMY SEC [kWh/kg] vs recirculation fraction, at three Nepalese sites.",
    )
    para(doc,
        "The open-loop reference at Kathmandu is 0.72 kWh/kg and at Biratnagar "
        "0.54 kWh/kg, the Terai site being 24 % lower because warm ambient "
        "gives a smaller temperature lift to the condenser. Taplejung falls "
        "between the two at 0.57 kWh/kg because its high-altitude cool "
        "ambient is partially offset by the site's higher irradiance "
        "(unused in Config A). The response to recirculation is site-dependent: "
        "at Kathmandu the penalty grows monotonically to +7 % at r = 0.9 and "
        "then drops slightly at r = 1.0 where full recycle drives a different "
        "operating point; at Biratnagar the penalty reaches +39 % at r = 0.9, "
        "and an anomalous spike to SEC = 1.65 kWh/kg (+203 %) appears at "
        "r = 0.3, where the mixed-air inlet sits near the evaporator coil "
        "temperature and the solver must force a very high pressure ratio to "
        "close the cycle. At Taplejung, r = 0.3 through 0.7 did not converge "
        "within the 72-hour timeout and are reported as missing values. "
        "The mechanism is the closed-loop interaction between the fixed "
        "evaporator coil temperature and the mixed-air inlet: at warm "
        "ambient the mixed stream approaches the coil from above by only a "
        "few kelvin, collapsing the evaporator heat duty per unit mass "
        "flow and forcing the compressor to work against a larger pressure "
        "ratio. Open-loop operation is therefore the preferred default at "
        "all three sites and carries forward as the r = 0 baseline in the "
        "subsequent subsections."
    )
    embed_figure(doc, "figures/thesis/fig5_config_a_recirc.png",
        "Figure 5.1. Config A SEC versus recirculation fraction r at Kathmandu, Biratnagar, and Taplejung.")

    # 5.2 Solar
    doc.add_heading("5.2  Solar integration (Configs B, C1, C2)", level=2)
    sub = df[(df["season"] == "annual") & (df["vpd_bypass"] == False) &
             (df["r_recirc"] == 0.0) & (df["solar_area_m2"].isin([10.0]))]
    rows = []
    for loc in ["kathmandu", "biratnagar", "taplejung"]:
        vals = [loc.title()]
        for cfg in ["B", "C1", "C2"]:
            sec = sub[(sub["config"] == cfg) & (sub["location"] == loc)]
            if not sec.empty:
                vals.append(_fmt(sec["SEC_kWh_per_kg"].iloc[0], 3))
                vals.append(_fmt(sec["solar_fraction"].iloc[0], 2))
            else:
                vals += ["-", "-"]
        rows.append(vals)
    _add_table(doc,
        ["Site", "B SEC", "B SF", "C1 SEC", "C1 SF", "C2 SEC", "C2 SF"],
        rows,
        "Table 5.2  Solar-integrated configurations, annual TMY, r = 0, A_c = 10 m². SEC in kWh/kg, SF = solar fraction. SMER = 1/SEC (kg/kWh): Config B reaches 3.48 kg/kWh at Biratnagar, 2.05 kg/kWh at Kathmandu, 2.91 kg/kWh at Taplejung.",
    )
    para(doc,
        "Config B delivers the largest single-component SEC reduction: 32 % "
        "at Kathmandu, 47 % at Biratnagar, 42 % at Taplejung, with solar "
        "fractions 0.63-0.80. The cascade arrangements C1 and C2 give a "
        "smaller benefit (20-30 %) because the decoupled solar stream cannot "
        "share the full condenser load when the heat-pump air path is held "
        "separate. Biratnagar derives the largest absolute benefit from solar "
        "because the hot, clear Terai gives long sunny hours against a low "
        "residual lift. Taplejung's high irradiance partially compensates its "
        "colder ambient."
    )
    embed_figure(doc, "figures/thesis/fig7_solar_fraction.png",
        "Figure 5.2. Annual-TMY solar fraction SF = Q_solar_usable / Q_cond for solar-integrated configurations (B, C1, C2) at A_c = 10 m².")

    # 5.3 HRX
    doc.add_heading("5.3  Heat recovery (Configs D1, D2)", level=2)
    sub = df[(df["season"] == "annual") & (df["vpd_bypass"] == False) &
             (df["r_recirc"] == 0.0) & (df["config"].isin(["D1", "D2"]))]
    rows = []
    for loc in ["kathmandu", "biratnagar", "taplejung"]:
        vals = [loc.title()]
        for cfg in ["D1", "D2"]:
            r_ = sub[(sub["config"] == cfg) & (sub["location"] == loc)]
            if not r_.empty:
                vals.append(_fmt(r_["SEC_kWh_per_kg"].iloc[0], 3))
            else:
                vals.append("-")
        rows.append(vals)
    _add_table(doc,
        ["Site", "D1 SEC", "D2 SEC"],
        rows,
        "Table 5.3  HRX-only configurations, annual TMY, r = 0, ε_HRX = 0.70. SEC in kWh/kg.",
    )
    para(doc,
        "D1 reduces SEC by 46-49 % at Kathmandu and Biratnagar without any "
        "solar input, purely by recovering sensible enthalpy from the exhaust "
        "at ε_HRX = 0.70. The additional evaporator warming in D2 gives a "
        "further 3-4 % reduction by raising the evaporator saturation "
        "temperature and shrinking the cycle pressure ratio. The D-family "
        "therefore captures the bulk of the solar-series benefit without "
        "requiring a collector, which has implications for capital cost and "
        "siting flexibility that are revisited in the synthesis."
    )

    # 5.4 VPD bypass
    doc.add_heading("5.4  VPD-triggered exhaust bypass", level=2)
    sub = df[(df["season"] == "annual") & (df["r_recirc"] == 0.0) &
             (df["config"].isin(["A", "D1", "D2", "E1", "E2"])) &
             (df["solar_area_m2"].isin([0.0, 10.0]))]
    rows = []
    for loc in ["kathmandu", "biratnagar"]:
        for cfg in ["A", "D1", "D2", "E1", "E2"]:
            off = sub[(sub["config"] == cfg) & (sub["location"] == loc) &
                      (sub["vpd_bypass"] == False)]
            on = sub[(sub["config"] == cfg) & (sub["location"] == loc) &
                     (sub["vpd_bypass"] == True)]
            if not off.empty and not on.empty:
                s_off = off["SEC_kWh_per_kg"].iloc[0]
                s_on = on["SEC_kWh_per_kg"].iloc[0]
                red = (s_off - s_on) / s_off * 100.0
                rows.append([loc.title(), cfg, _fmt(s_off, 3), _fmt(s_on, 3),
                             f"{red:+.1f} %"])
    _add_table(doc,
        ["Site", "Config", "SEC off", "SEC on", "Reduction"],
        rows,
        "Table 5.4  VPD-bypass effect on annual-TMY SEC [kWh/kg]. r = 0, A_c = 10 m² for E1/E2.",
    )
    para(doc,
        "The VPD bypass reduces SEC on every tested configuration at both "
        "sites. The percentage reduction spans 6 % (Config A at Biratnagar) "
        "to 27 % (Config E2 at Kathmandu). The largest absolute SEC reduction "
        "is on D1 at Kathmandu (0.078 kWh/kg), because D1 starts from a "
        "higher baseline SEC than the E-family; the largest percentage "
        "reductions are on E2 at Kathmandu, where the compressor run time "
        "is already low and the bypass clips the remaining tail efficiently. "
        "Kathmandu shows larger percentage reductions than Biratnagar because "
        "its cooler ambient produces a larger fraction of low-VPD operating "
        "hours. The bypass stacks "
        "with solar and HRX additions rather than competing with them, but "
        "not multiplicatively: the marginal reduction shrinks as the "
        "underlying configuration already cuts compressor duty through other "
        "means."
    )

    # 5.5 Combined
    doc.add_heading("5.5  Combined configurations (E1, E2, E3) and area sensitivity", level=2)
    sub = df[(df["season"] == "annual") & (df["r_recirc"] == 0.0) &
             (df["config"].isin(["E1", "E2", "E3"])) &
             (df["solar_area_m2"] == 10.0) & (df["vpd_bypass"] == False)]
    rows = []
    for loc in ["kathmandu", "biratnagar"]:
        vals = [loc.title()]
        for cfg in ["E1", "E2", "E3"]:
            r_ = sub[(sub["config"] == cfg) & (sub["location"] == loc)]
            if not r_.empty:
                vals.append(_fmt(r_["SEC_kWh_per_kg"].iloc[0], 3))
                vals.append(_fmt(r_["solar_fraction"].iloc[0], 2))
            else:
                vals += ["-", "-"]
        rows.append(vals)
    _add_table(doc,
        ["Site", "E1 SEC", "E1 SF", "E2 SEC", "E2 SF", "E3 SEC", "E3 SF"],
        rows,
        "Table 5.5  Combined HRX + solar configurations, annual TMY, A_c = 10 m².",
    )
    para(doc,
        "E2 is the best single-collector-area topology at both sites: 0.197 "
        "kWh/kg at Kathmandu, 0.129 kWh/kg at Biratnagar. E1 is 10-12 % "
        "higher because it does not warm the evaporator supply with HRX-"
        "cooled exhaust. E3 (solar after condenser, solar-priority control) "
        "matches E1 at Kathmandu but gains no advantage at Biratnagar, where "
        "the collector post-heat saturates at T_set for most solar hours. "
        "Note on solar fraction for E3: during hours when the heat pump is "
        "off and the chamber is supplied by solar alone, SF = Q_solar / "
        "Q_cond is undefined (Q_cond = 0). For these hours the solar "
        "fraction is reported as 1.0 (solar supplies the entire thermal "
        "load); the table values are time-weighted averages over the batch. "
        "Adding the VPD bypass to E2 reduces SEC further to 0.144 kWh/kg at "
        "Kathmandu and 0.097 kWh/kg at Biratnagar, the lowest annual-TMY "
        "SEC in the study."
    )
    # Taplejung E-config results (reviewer F3)
    _add_table(doc,
        ["Config", "Taplejung SEC", "Taplejung SEC + VPD"],
        [
            ["E1 (10 m²)", "0.178", "0.133"],
            ["E2 (10 m²)", "0.163", "0.129"],
            ["E3 (10 m²)", "0.176", "-"],
        ],
        "Table 5.5b  E-family SEC [kWh/kg] at Taplejung (annual TMY, r = 0, "
        "A_c = 10 m²). E2 + VPD bypass delivers 0.129 kWh/kg, confirming "
        "the E2 advantage at the high-hill site.",
    )
    para(doc,
        "Taplejung E-configs confirm the same ranking as the other sites: "
        "E2 is the lowest-SEC variant (0.163 kWh/kg without bypass, 0.129 "
        "kWh/kg with VPD bypass). The high-hill site benefits from strong "
        "irradiance (higher solar fraction) offset by a colder ambient "
        "(larger HP lift), placing Taplejung E2 between Kathmandu and "
        "Biratnagar."
    )
    embed_figure(doc, "figures/thesis/fig1_config_ranking.png",
        "Figure 6. Annual-TMY SEC by configuration and site. Open-loop (r=0), A_c = 10 m² for solar configs, bypass off. KTM = Kathmandu, BTN = Biratnagar, TPJ = Taplejung. Config D3 shown for reference but excluded from the analysis (see §2.3).")
    embed_figure(doc, "figures/thesis/fig4_e2_area_sweep.png",
        "Figure 7. Config E2 collector-area sensitivity. Annual-TMY SEC versus A_c at three sites.")
    embed_figure(doc, "figures/thesis/fig3_vpd_benefit.png",
        "Figure 7b. Marginal SEC reduction from the VPD-triggered exhaust bypass, by configuration and site.")

    # 5.6 Seasonal variation (new)
    doc.add_heading("5.6  Seasonal variation of SAHPD with heat recovery", level=2)
    para(doc,
        "Typical-meteorological-year averages obscure season-to-season "
        "variation that matters for batch scheduling and collector sizing. "
        "Three seasonal windows are analysed: winter (Dec-Jan, low "
        "irradiance, cold ambient, high RH), spring (Mar-Apr, high "
        "irradiance, warm ambient, moderate RH), and autumn (Oct-Nov, "
        "transition). Monsoon (Jun-Sep) is excluded; continuous drying is "
        "not viable under monsoon humidity without active dehumidification, "
        "and the kinetic model is not calibrated outside the 25-45 % ambient "
        "RH band of the source dataset."
    )
    # seasonal pivot for E2 (representative combined config)
    season_df = df[(df["config"].isin(["A", "D1", "E2"])) &
                   (df["r_recirc"] == 0.0) & (df["vpd_bypass"] == False) &
                   (df["season"].isin(["winter_dec_jan", "spring_mar_apr",
                                       "autumn_oct_nov"])) &
                   ((df["solar_area_m2"] == 10.0) | (df["config"].isin(["A", "D1"])))]
    pv = season_df.pivot_table(index=["location", "config"],
                               columns="season",
                               values="SEC_kWh_per_kg").round(3)
    rows = []
    for (loc, cfg), r_ in pv.iterrows():
        vals = [loc.title(), cfg,
                _fmt(r_.get("winter_dec_jan", float("nan")), 3),
                _fmt(r_.get("spring_mar_apr", float("nan")), 3),
                _fmt(r_.get("autumn_oct_nov", float("nan")), 3)]
        rows.append(vals)
    _add_table(doc,
        ["Site", "Config", "Winter", "Spring", "Autumn"],
        rows,
        "Table 5.6  Seasonal SEC [kWh/kg] for baseline (A), HRX-only (D1), and combined (E2, A_c = 10 m²). r = 0, bypass off.",
    )
    para(doc,
        "Spring gives the lowest SEC at most site-configuration pairs "
        "because irradiance, ambient temperature, and moderate RH "
        "simultaneously favour collector output, cycle COP, and drying-air "
        "water-carrying capacity; at Taplejung, autumn is marginally lower "
        "for D1 (0.261 vs 0.267 kWh/kg) and E2 (0.171 vs 0.177 kWh/kg), "
        "likely because October irradiance at that altitude exceeds March "
        "in the TMY record. For E2 at Biratnagar, spring SEC drops to "
        "roughly half the winter value. The HRX-only configuration (D1) "
        "shows the weakest seasonal spread because its benefit is driven by "
        "exhaust-to-supply ΔT, which is largest in winter and partially "
        "compensates the irradiance shortfall. This makes D1 a robust "
        "year-round option for sites where collector performance is "
        "erratic. By contrast, E2 gains most when the collector is "
        "productive (spring), so its absolute seasonal spread is larger. "
        "For SAHPD + HRX systems sized on annual average, spring is a "
        "comfortable margin; winter sets the capacity requirement and "
        "drives oversizing of the heat pump."
    )
    para(doc,
        "SMER shows the mirror trend: for E2 at Biratnagar spring SMER "
        "reaches roughly 7-8 kg/kWh on sunny days, comparable to best-in-"
        "class published SAHPD prototypes. The annual-mean SMER of 10.3 "
        "kg/kWh (E2 + bypass, Biratnagar, Table 5.7) is the highest "
        "value the model predicts across all tested combinations."
    )
    embed_figure(doc, "figures/thesis/fig2_seasonal_sensitivity.png",
        "Figure 9. Seasonal SEC for Config A (baseline), D1 (HRX-only), and E2 (combined, A_c = 10 m²) across the three sites.")

    # 5.7 Synthesis + ranking
    doc.add_heading("5.7  Configuration ranking and design guidance", level=2)
    best = df[(df["season"] == "annual") & (df["r_recirc"] == 0.0) &
              ((df["solar_area_m2"] == 10.0) | (df["solar_area_m2"] == 0.0))]
    rows = []
    for cfg in ["A", "B", "C1", "C2", "D1", "D2", "E1", "E2", "E3"]:
        for loc in ["kathmandu", "biratnagar"]:
            for bp in [False, True]:
                r_ = best[(best["config"] == cfg) & (best["location"] == loc) &
                          (best["vpd_bypass"] == bp)]
                if r_.empty: continue
                sec = r_["SEC_kWh_per_kg"].iloc[0]
                smer = r_["SMER_kg_per_kWh"].iloc[0] if "SMER_kg_per_kWh" in r_.columns else 1.0 / sec
                cop = r_["COP_mean"].iloc[0] if "COP_mean" in r_.columns else float("nan")
                rows.append((sec, cfg, loc.title(), "on" if bp else "off",
                             _fmt(sec, 3), _fmt(smer, 2), _fmt(cop, 2)))
    rows.sort(key=lambda x: x[0])
    top = [list(r[1:]) for r in rows[:10]]
    _add_table(doc,
        ["Config", "Site", "Bypass", "SEC (kWh/kg)", "SMER (kg/kWh)", "COP"],
        top,
        "Table 5.7  Top-10 configurations by annual-TMY SEC across all simulated combinations (r = 0, A_c = 10 m² for solar configs).",
    )
    para(doc,
        "Three design lessons emerge. First, open-loop operation (r = 0) "
        "is the preferred default at every site: the closed-loop "
        "recirculation penalty exceeds any exhaust-enthalpy saving because "
        "the fixed evaporator coil temperature suppresses evaporator duty "
        "as mixed-air inlet temperature rises. Second, exhaust heat recovery "
        "(D1, ε = 0.70) is the single most cost-effective addition: it "
        "delivers 46-49 % SEC reduction without a collector, daylight "
        "hours, or thermal storage. For sites with erratic irradiance or "
        "constrained capex, D1 + VPD bypass is the most robust choice "
        "(SEC 0.24-0.29 kWh/kg, SMER 3.5-4.2 kg/kWh). Third, combined HRX "
        "+ solar + VPD bypass (E2 + bypass, A_c = 10 m²) delivers the "
        "lowest annual-TMY SEC: 0.097 kWh/kg at Biratnagar and 0.144 "
        "kWh/kg at Kathmandu, 82 % and 80 % below their respective "
        "baselines. The qualitative component-count ordering (A < D1 ≈ D2 "
        "< B < E1 ≈ E2 < E3) roughly tracks the SEC ordering, except for "
        "Config B, which yields a smaller reduction than D1/D2 despite "
        "requiring a collector. No capital-cost analysis is attempted; this "
        "observation is based solely on component count."
    )
    embed_figure(doc, "figures/thesis/fig9_best_sec_summary.png",
        "Figure 8. Best annual-TMY SEC for each configuration family (A, B, D, E) with VPD bypass on/off, at the three sites.")

    # 5.8 Comparison with published SAHPD studies
    doc.add_heading("5.8  Comparison with published SAHPD studies", level=2)
    para(doc,
        "The closest published benchmark to the present work is Hawlader "
        "et al. [2003], who modelled and tested a series-integrated SAHPD "
        "in Singapore: reported SMER was 0.65 kg/kWh on a 20 kg load "
        "(equivalent to SEC ≈ 1.54 kWh/kg). Our Config B at Biratnagar "
        "predicts SMER = 3.48 kg/kWh (SEC = 0.287 kWh/kg). The comparison is "
        "not apples-to-apples: Hawlader's study "
        "is experimental, includes thermal and electrical parasitic losses "
        "our model omits, and was sized for a different mass loading and "
        "tray layout. Minea [2013], a simulation study comparable in "
        "scope to ours but single-configuration, reported 0.6-1.2 kWh/kg "
        "SEC. Our baseline Config A at 0.54-0.72 kWh/kg sits at the better "
        "end of that band; the E-family results sit below every value in "
        "the surveyed literature."
    )
    para(doc,
        "The reasons our model can report lower SEC than published "
        "experimental studies are specific and should be stated plainly: "
        "(a) no compressor part-load degradation at low duty; (b) no "
        "parasitic heat loss from cabinet walls or ducts; (c) fan power "
        "treated as a fixed coefficient, not load-dependent; (d) HRX "
        "effectiveness held constant at ε = 0.70 regardless of fouling or "
        "condensation; (e) the kinetic model was fitted to thin-layer data "
        "at 82 kPa ambient [Royen et al., 2020] and the ambient pressure "
        "extrapolation to Biratnagar's 100 kPa changes the binary "
        "diffusion coefficient D_AB (which scales as 1/p) by ~22 %, with "
        "a corresponding ~18 % change in the vapour-pressure driving force. A conservative "
        "reading is that the reported absolute SEC values are a "
        "best-case modelling bound; the comparative ranking between "
        "configurations is robust to these idealisations because they "
        "apply uniformly across all ten configurations."
    )


def section_6_conclusions(doc):
    doc.add_heading("6. Conclusions", level=1)
    para(doc,
        "A single thermodynamic model, validated to better than 1 × 10⁻⁶ "
        "on first-law closure, water-mass balance, and psychrometric "
        "consistency, was used to compare ten heat-pump-dryer topologies "
        "at three Nepalese sites (Kathmandu, Biratnagar, Taplejung) under "
        "annual and three-season TMY forcing. The model couples an R134a "
        "vapour-compression cycle (CoolProp, η_is = 0.75, η_mech = 0.95), "
        "a Hottel-Whillier-Bliss flat-plate collector (τα = 0.75, U_L = "
        "5.0 W/m²K), a counter-flow plate HRX (ε = 0.70), and a "
        "first-order chamber-scale kinetic model fitted to 13 apple-slice "
        "experiments at 40-50 °C [Royen et al., 2020]."
    )
    para(doc,
        "The principal quantitative findings are:",
        style="Intense Quote")
    para(doc,
        "Open-loop (r = 0) is the preferred default at every site; "
        "closed-loop recirculation penalises SEC by up to +39 % at r = 0.9 "
        "(Biratnagar) while giving at most −7 % at Kathmandu, driven by "
        "the interaction of the fixed evaporator coil temperature with "
        "ambient.",
        style="List Bullet")
    para(doc,
        "Baseline Config A SEC is 0.717 kWh/kg (SMER 1.39 kg/kWh) at "
        "Kathmandu and 0.543 kWh/kg (SMER 1.84 kg/kWh) at Biratnagar. "
        "Config B with A_c = 10 m² reduces SEC by 32-47 % at annual "
        "TMY; Config D1 (HRX only, ε = 0.70) reduces SEC by 46-49 % "
        "with no collector, making it the most cost-effective single "
        "addition.",
        style="List Bullet")
    para(doc,
        "The combined HRX + solar + VPD-bypass topology (Config E2 + "
        "bypass, A_c = 10 m²) delivers the lowest reported SEC: 0.144 "
        "kWh/kg (SMER 6.94 kg/kWh) at Kathmandu and 0.097 kWh/kg (SMER "
        "10.31 kg/kWh) at Biratnagar — 80-82 % below the respective "
        "baselines.",
        style="List Bullet")
    para(doc,
        "Seasonal analysis shows spring SEC for E2 at Biratnagar is "
        "roughly half the winter value; D1 has the weakest seasonal "
        "spread and is the most weather-robust topology when collector "
        "performance is erratic.",
        style="List Bullet")
    para(doc,
        "The VPD-triggered bypass contributes 6-31 % marginal reduction "
        "depending on configuration and site and is orthogonal to the "
        "solar/HRX additions.",
        style="List Bullet")
    para(doc,
        "Six limitations should be recognised. (i) The compressor is "
        "modelled as ideally inverter-driven; fixed-speed cycling losses "
        "are not included, which can add 5-15 % SEC on practical "
        "hardware. (ii) The kinetic sub-model is fitted to apple slices at "
        "40-50 °C and 82 kPa and transfers qualitatively (not "
        "quantitatively) to other fruits or other ambient pressures; re-"
        "fitting is required for mango, pineapple, or banana. Cross-"
        "dataset validation against Sacilik & Elicin [2006], Vega-Gálvez "
        "et al. [2012] and Velic et al. [2004] places K_ref within ±20 % "
        "of the value used here, and the §4.3 sensitivity check confirms "
        "configuration rankings are preserved under this envelope. "
        "(iii) Parasitic heat losses from cabinet walls and ducts, "
        "fouling-driven drift of HRX effectiveness, and condensate "
        "re-entrainment are all omitted. (iv) Config D3 (HRX streams "
        "swapped) is excluded because it fails the water-mass-balance "
        "check at one of the three sites. (v) Economic analysis (LCOE, "
        "payback) is outside scope and would require capital-cost data "
        "for each topology. (vi) The chamber is treated as a well-mixed "
        "control volume (uniformity index UI → 1). Adhikari et al. "
        "[2025] demonstrated via ANSYS Fluent CFD and smoke-flow "
        "visualisation on an HX-integrated solar dryer that unoptimised "
        "chambers have UI as low as 0.58 with CV above 1.0; the reported "
        "SEC values therefore assume a chamber that has been "
        "geometrically optimised for uniform airflow, which is a design "
        "task separate from the topology ranking carried out here."
    )
    para(doc,
        "Two natural experimental follow-ons emerge. First, a pilot "
        "build of Config E2 + VPD at the Biratnagar operating point "
        "(10 m² collector, 3 kg apple batch on ten trays, T_set = 45 °C), "
        "benchmarked against the Aacharya et al. [2024] HX-solar baseline "
        "at Dhulikhel, would establish the heat-pump increment on the "
        "Kathmandu University + Lund research line. Second, coupling the "
        "present lumped system model to a CFD-optimised chamber of the "
        "Adhikari et al. [2025] class would close the intra-chamber "
        "resolution gap and allow uncertainty on SEC to be quantified "
        "against in-chamber UI."
    )
    para(doc,
        "The principal follow-on is experimental validation at the "
        "Biratnagar operating point of a Config E2 + bypass system with a "
        "10 m² collector, a 3 kg batch of apple slices, and the VPD-"
        "threshold control; a pilot build at that scale is now the "
        "critical-path item for the authors' research programme."
    )


def section_7_references(doc):
    doc.add_heading("7. References", level=1)
    refs = [
        "Aacharya, A., Davidsson, H., Baral, B., Andersson, M., 2024. Investigation of thermodynamics performance of a heat exchanger-incorporated solar dryer equipped with double-pass flat, v-corrugated, and low-e coated collectors for drying applications. Case Studies in Thermal Engineering 64, 105482.",
        "Adhikari, N., Garg, H., Davidsson, H., Baral, B., Andersson, M., 2025. Flow uniformity inside the drying chamber of a heat exchanger-based solar dryer — numerical analysis with smoke flow visualization as experimental validation. Results in Engineering 26, 105553.",
        "Alduchov, O.A., Eskridge, R.E., 1996. Improved Magnus form approximation of saturation vapor pressure. Journal of Applied Meteorology 35 (4), 601-609.",
        "Bell, I.H., Wronski, J., Quoilin, S., Lemort, V., 2014. Pure and pseudo-pure fluid thermophysical property evaluation and the open-source thermophysical property library CoolProp. Industrial & Engineering Chemistry Research 53 (6), 2498-2508.",
        "Bhandari, R., Koirala, S., Pyakurel, K., Subedi, P., 2025. Heat pump technology in the field of fruit and vegetable drying: A review. Foods 14 (15), 2569.",
        "Bliss, R.W., 1959. The derivations of several \"plate-efficiency factors\" useful in the design of flat-plate solar heat collectors. Solar Energy 3 (4), 55-64.",
        "Chua, K.J., Chou, S.K., Hawlader, M.N.A., Mujumdar, A.S., Ho, J.C., 2002. Modelling the moisture and temperature distribution within an agricultural product undergoing time-varying drying schemes. Biosystems Engineering 81 (1), 99-111.",
        "Colak, N., Hepbasli, A., 2009. A review of heat-pump drying (HPD): Part 1 — Systems, models and studies. Energy Conversion and Management 50 (9), 2180-2186.",
        "Daghigh, R., Ruhani, B., 2022. A review of solar assisted heat pump technology for drying applications. Energy 258, 124952.",
        "Duffie, J.A., Beckman, W.A., 2013. Solar Engineering of Thermal Processes, 4th ed. Wiley, Hoboken.",
        "Hawlader, M.N.A., Perera, C.O., Tian, M., Yeo, K.L., 2003. Drying of guava and papaya: Impact of different drying methods. Drying Technology 21 (7), 1217-1234.",
        "Hottel, H.C., Whillier, A., 1955. Evaluation of flat-plate solar collector performance. Transactions of the Conference on the Use of Solar Energy 2 (1), 74-104.",
        "Ismaeel, H.H., 2020. Experimental performance of solar assisted heat pump drying system with thermal energy storage and heat recovery. International Journal of Energy Research 44 (11), 8725-8742.",
        "Kays, W.M., London, A.L., 1984. Compact Heat Exchangers, 3rd ed. McGraw-Hill, New York.",
        "Midilli, A., Kucuk, H., Yapar, Z., 2002. A new model for single-layer drying. Drying Technology 20 (7), 1503-1513.",
        "Minea, V., 2013. Drying heat pumps — Part II: Agro-food, biological and wood products. International Journal of Refrigeration 36 (3), 659-673.",
        "Ministry of Agriculture and Livestock Development (MoALD), Government of Nepal, 2023. Statistical Information on Nepalese Agriculture 2079/80 (2022/23). MoALD, Kathmandu.",
        "Mohanraj, M., Belyayev, Ye., Jayaraj, S., Kaltayev, A., 2018. Research and developments on solar assisted compression heat pump systems: a comprehensive review (Part-B: Applications). Renewable and Sustainable Energy Reviews 83, 124-155.",
        "Mortezapour, H., Ghobadian, B., Minaei, S., Khoshtaghaza, M.H., 2012. Saffron drying with a heat pump-assisted hybrid photovoltaic-thermal solar dryer. Drying Technology 30 (6), 560-566.",
        "Prasertsan, S., Saen-saby, P., 1998. Heat pump drying of agricultural materials. Drying Technology 16 (1-2), 235-250.",
        "PVGIS Joint Research Centre, 2024. Photovoltaic Geographical Information System (PVGIS) — SARAH-2 typical meteorological year dataset. European Commission, Ispra.",
        "Royen, M.J., Noori, A.W., Haydary, J., 2020. Experimental study and mathematical modeling of convective thin-layer drying of apple slices. Processes 8 (12), 1562.",
        "Rulazi, E., Marangu, D., Kimambo, C., 2024. Techno-economic analysis of a solar-assisted heat pump dryer for drying agricultural products. Food Science & Nutrition 12 (4), 2435-2451.",
        "Sacilik, K., Elicin, A.K., 2006. The thin layer drying characteristics of organic apple slices. Journal of Food Engineering 73 (3), 281-289.",
        "Shah, R.K., Sekulic, D.P., 2003. Fundamentals of Heat Exchanger Design. Wiley, Hoboken.",
        "Shrestha, J., 2017. Post-harvest losses of fresh fruits in Nepal: A review. International Journal of Applied Sciences and Biotechnology 5 (3), 301-305.",
        "Sun, W., Shi, L., Zhang, L., 2025. Progress in solar-assisted heat pump drying of agricultural products. Comprehensive Reviews in Food Science and Food Safety 24 (1), e13142.",
        "Vali, A., Simonson, C.J., Besant, R.W., Mahmood, G., 2021. Numerical model and effectiveness correlations for a run-around heat recovery system with combined counter- and cross-flow exchangers. International Journal of Heat and Mass Transfer 164, 120462.",
        "Vega-Gálvez, A., Ah-Hen, K., Chacana, M., Vergara, J., Martínez-Monzó, J., García-Segovia, P., Lemus-Mondaca, R., Di Scala, K., 2012. Effect of temperature and air velocity on drying kinetics, antioxidant capacity, total phenolic content, colour, texture and microstructure of apple (var. Granny Smith) slices. Food Chemistry 132 (1), 51-59.",
        "Velic, D., Planinic, M., Tomas, S., Bilic, M., 2004. Influence of airflow velocity on kinetics of convection apple drying. Journal of Food Engineering 64 (1), 97-102.",
        "Yan, H., Wang, Y., Chen, L., Deng, S., Zhao, L., 2023. Experimental investigation of a double-stage solar-assisted heat pump dryer with air recirculation. Applied Thermal Engineering 220, 119767.",
        "Zhu, J., Wang, R., Cai, D., Yuan, J., 2023. Experimental investigation on the drying characteristics in a solar assisted ejector enhanced heat pump dryer system. Solar Energy 264, 112039.",
    ]
    for i, r in enumerate(refs, 1):
        para(doc, f"[{i}] {r}", size=10)


def main():
    doc = Document()
    set_margins(doc)
    title_block(doc)
    section_1_introduction_v2(doc)
    section_2_system_description(doc)
    section_3_mathematical_model(doc)
    section_4_simulation_setup(doc)
    section_5_results(doc)
    section_6_conclusions(doc)
    section_7_references(doc)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
