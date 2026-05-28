"""Export the SAHPD component-sizing table to Word (SAHPD_Sizing.docx).

Reference case: Config A, Kathmandu, r = 0 (worst-case ambient drives peak HP duty).
HRX sizing read from Config D1, Kathmandu. Data pulled live from outputs/.
"""
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL

ROOT = Path(__file__).resolve().parents[1]
A_CSV = ROOT / "outputs" / "config_A" / "kathmandu" / "baseline_r0.0.csv"
D1_CSV = ROOT / "outputs" / "config_D1" / "kathmandu" / "hrx_eps0.70.csv"
OUT = ROOT / "SAHPD_Sizing.docx"


def fmt(x, d=2):
    try:
        return f"{float(x):.{d}f}"
    except Exception:
        return "-"


def main():
    a = pd.read_csv(A_CSV)
    d = pd.read_csv(D1_CSV)

    Q_cond_max = a["Q_cond_kW"].max()
    Q_cond_mean = a["Q_cond_kW"].mean()
    W_comp_max = a["W_comp_kW"].max()
    W_comp_mean = a["W_comp_kW"].mean()
    Q_evap_max = a["Q_evap_kW"].max()
    T_cond = a["T_cond_C"].iloc[0]
    T_exh_max = a["T_exhaust_C"].max()
    omega_chamber = a["omega_to_chamber"].mean()

    Q_HRX_max = d["Q_HRX_kW"].max()
    Q_HRX_mean = d["Q_HRX_kW"].mean()
    dT_hot = d["T_exhaust_C"] - d["T_exh_cooled_C"]

    doc = Document()
    for sec in doc.sections:
        sec.left_margin = Cm(2); sec.right_margin = Cm(2)

    title = doc.add_heading("SAHPD Component Sizing", level=1)
    doc.add_paragraph(
        "Reference: Config A (HP-only), Kathmandu, r = 0 — the coldest ambient of the three simulated sites, "
        "which drives peak HP duty. Every other configuration (B, C, D, E) has lower electrical demand, "
        "so this sizing is the worst-case envelope for purchasing decisions. "
        "HRX duty is taken from Config D1 (Kathmandu, eps_HRX = 0.70)."
    )
    doc.add_paragraph(
        "Data source: outputs/config_A/kathmandu/baseline_r0.0.csv and "
        "outputs/config_D1/kathmandu/hrx_eps0.70.csv (post 2026-04-08 iterative-evap bugfix, "
        "regenerated 2026-04-21)."
    )

    # ---- Main sizing table ----
    doc.add_heading("1. Component sizing table", level=2)
    cols = ["#", "Component", "Quantity", "Value from simulation", "Design pick (with margin)"]
    rows = [
        ("1", "Heat pump / reversible split AC", "Peak Q_cond (heating duty)",
         f"{fmt(Q_cond_max,2)} kW", "≥ 4.5 kW heating capacity (≈ 1.5 ton / 18,000 BTU/h)"),
        ("", "", "Mean Q_cond", f"{fmt(Q_cond_mean,2)} kW", ""),
        ("", "", "Peak W_comp (electric)", f"{fmt(W_comp_max,2)} kW",
         "230 V, 10 A single-phase circuit"),
        ("", "", "Mean W_comp", f"{fmt(W_comp_mean,3)} kW", ""),
        ("", "", "Peak Q_evap (cooling)", f"{fmt(Q_evap_max,2)} kW", "≈ 1 ton cooling equivalent"),
        ("", "", "Refrigerant T_cond (fixed)", f"{fmt(T_cond,0)} °C",
         "Within standard split AC envelope"),
        ("", "", "Refrigerant T_evap target", "5 °C", ""),
        ("", "", "Mean COP (simulation)", "~3.6", "Inverter, 5-star preferred"),
        ("2", "Supply air temperature", "T_to_chamber", "45 °C (constant)",
         "Set-point at controller"),
        ("3", "Airflow (blower / inline fan)", "Mass flow m_da", "0.098 kg/s", ""),
        ("", "", "Volumetric flow (ρ = 0.943)", "~374 m³/h",
         "≥ 450 m³/h axial or centrifugal inline fan"),
        ("", "", "Total static pressure", "210 Pa", "Fan curve ≥ 250 Pa at design CFM"),
        ("", "", "Fan electric power", "~25 W continuous", "Fan efficiency η ≥ 0.55"),
        ("4", "Drying chamber", "Inner cross-section", "0.79 m × 0.12 m (air gap)",
         "10 trays stacked, 6 cm pitch"),
        ("", "", "Tray area (each)", "0.625 m² (≈ 0.79 × 0.79 m)",
         "Perforated food-grade mesh"),
        ("", "", "Chamber volume (trays + ducts)", "~0.8 m³",
         "Insulate walls (50 mm polyurethane)"),
        ("5", "Exhaust duct", "Peak T_exhaust", f"{fmt(T_exh_max,1)} °C",
         "Ø 150 mm flexible duct, short run"),
        ("", "", "Supply humidity ratio ω", f"{fmt(omega_chamber,4)}",
         "Dry supply, within expected envelope"),
        ("6", "Heat-recovery exchanger (HRX)", "Peak Q_HRX", f"{fmt(Q_HRX_max,2)} kW",
         "Counter-flow air-to-air plate, ε ≥ 0.70"),
        ("", "", "Mean Q_HRX", f"{fmt(Q_HRX_mean,2)} kW",
         "Rated ≥ 3 kW @ 400 m³/h each side"),
        ("", "", "ΔT hot side", f"{fmt(dT_hot.min(),0)}–{fmt(dT_hot.max(),0)} K (mean {fmt(dT_hot.mean(),0)} K)",
         "Equal supply + exhaust ducting"),
        ("", "", "Face volumetric flow per side", "~400 m³/h", ""),
    ]

    t = doc.add_table(rows=1 + len(rows), cols=len(cols))
    t.style = "Light Grid Accent 1"
    t.autofit = True
    for i, c in enumerate(cols):
        cell = t.rows[0].cells[i]
        cell.text = c
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for i, row in enumerate(rows, 1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = val
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)

    # ---- Purchasing checklist ----
    doc.add_heading("2. Purchasing checklist (AC unit)", level=2)
    for item in [
        "Nominal cooling 18,000 BTU/h AND heating capacity ≥ 4.5 kW at 7 °C outdoor.",
        "Reversible heat-pump (4-way valve); NOT a cooling-only unit.",
        "DC inverter compressor (capacity modulation).",
        "Refrigerant R32 preferred over R410A (lower GWP, single-component, matches modelling).",
        "Ask whether evaporator and condenser can be mounted together inside the dryer loop "
        "with a custom fan (critical for research instrumentation).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # ---- Nepal market options ----
    doc.add_heading("3. Nepal-market options (1.5-ton reversible inverter split)", level=2)
    doc.add_paragraph(
        "All widely sold in Kathmandu. Confirm current nameplate heating capacity and availability "
        "with the retailer before purchase."
    )
    models = [
        ("Midea MSAGC-18HRFN1 1.5-ton 5-star",
         "18,000 BTU/h cool, 20,000 BTU/h heat, full-DC inverter — closest match", "Trade Nepal, Electrosewa"),
        ("Stellar SA-15IS1HPA/HE 1.5-ton", "Built for Nepali winters, labelled cooling + heating",
         "Himelectronics (Stellar)"),
        ("LG 1.5-ton Dual Inverter (wall mount)",
         "Common 1.5-ton inverter; confirm SKU supports heating mode", "AC Ghar, Daraz"),
        ("Gree 1.5-ton wall mount", "Regularly available; verify reversible model", "AC Ghar"),
        ("CG 1.5-ton DC Inverter wall-mount", "Local brand, serviceable parts", "AC Ghar"),
        ("Samsung AR18TSHZRWKNRC", "5-in-1 convertible, DC inverter, reversible", "Trade Nepal"),
    ]
    mt = doc.add_table(rows=1 + len(models), cols=3)
    mt.style = "Light Grid Accent 1"
    for i, h in enumerate(["Model", "Notes", "Where to check"]):
        mt.rows[0].cells[i].text = h
        for p in mt.rows[0].cells[i].paragraphs:
            for r in p.runs: r.bold = True
    for i, row in enumerate(models, 1):
        for j, v in enumerate(row):
            mt.rows[i].cells[j].text = v

    # ---- HRX note ----
    doc.add_heading("4. Heat-recovery exchanger (HRX) sourcing", level=2)
    doc.add_paragraph(
        "Counter-flow air-to-air plate HRX units are not commonly stocked in the Nepali market. "
        "Two realistic sourcing paths:"
    )
    doc.add_paragraph(
        "Import a residential HRV core (sensible-only, NOT ERV) sized 300–500 m³/h at ≥ 70 % "
        "effectiveness. Suppliers: ERI Corporation (Italy), Heatex (Sweden), Holtop (China), "
        "Airwoods (China), Zeco Aircon (India). Zeco is the closest option logistically.",
        style="List Number")
    doc.add_paragraph(
        "Fabricate locally: a workshop in Balaju can build a 0.5 × 0.5 × 0.3 m counter-flow box "
        "from brazed aluminium or corrugated polypropylene plates. Target ≥ 2 m² plate area to "
        "hit ε = 0.70 at 400 m³/h. Validate effectiveness with a thermocouple traverse before use.",
        style="List Number")
    doc.add_paragraph(
        "Avoid ERVs with enthalpy wheels: moisture transfer across the plate would break the "
        "sensible-only assumption in the Config D/E simulations and confound the comparison."
    )

    # ---- Instrumentation ----
    doc.add_heading("5. Instrumentation to budget for", level=2)
    for item in [
        "PID thermostat with supply-air sensor (override the AC's stock room sensor).",
        "Four T/RH sensors minimum (SHT35 or equivalent): ambient, chamber inlet, exhaust, evaporator inlet.",
        "Two clamp-on kWh meters: one for compressor, one for fan, to validate SEC.",
        "Pressure taps + manometer across evaporator and condenser coils (verify the 80 Pa each assumed by the fan-power model).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
