"""Generate thesis/Section1_Config_A.docx — standalone §1 on Config A (HP-only).

Story beats: system description, open-loop baseline across three sites,
recirculation sweep (r=0..1), seasonal variation, evaporator strategy note.
Concise, table-heavy, objective language. No D3. No biased framing.
Data pulled live from outputs/master_summary.csv and Config A CSVs.
"""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL

ROOT = Path(__file__).resolve().parents[1]
SUM = ROOT / "outputs" / "master_summary.csv"
OUT = ROOT / "thesis" / "Section1_Config_A.docx"


def fmt(x, d=3):
    try:
        if x != x:
            return "-"
        return f"{float(x):.{d}f}"
    except Exception:
        return "-"


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = True
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            c = t.rows[i].cells[j]
            c.text = str(v)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
    return t


def main():
    df = pd.read_csv(SUM)
    a = df[df["config"] == "A"].copy()

    doc = Document()
    for s in doc.sections:
        s.left_margin = Cm(2); s.right_margin = Cm(2)

    doc.add_heading("Section 1  Config A: Heat-Pump-Only Baseline", level=1)
    doc.add_paragraph(
        "Config A is the reference case. A single-stage vapour-compression heat pump (R134a, "
        "eta_is = 0.75) heats ambient air to the chamber set-point (T_set = 45 deg C) and removes "
        "moisture from 3.0 kg of apple slices (X0 = 6.5, X_f = 0.10 db) distributed across ten trays. "
        "Every other configuration in this thesis adds one or more free-energy components on top of "
        "this baseline, so Config A sets the ceiling against which their benefits are measured."
    )

    # 1.1 System description
    doc.add_heading("1.1  Air paths", level=2)
    doc.add_paragraph(
        "Two operating modes are possible and differ in whether exhaust air is reused."
    )
    doc.add_paragraph("Open-loop (r = 0).", style="List Bullet")
    doc.add_paragraph(
        "Ambient air enters the condenser, is heated to T_set, passes through the chamber, "
        "and is expelled. The evaporator draws heat from a separate ambient stream; its "
        "saturation temperature tracks ambient (T_evap ~ T_amb - 10 K).",
        style="List Bullet")
    doc.add_paragraph("Closed-loop (r > 0).", style="List Bullet")
    doc.add_paragraph(
        "Fraction r of the exhaust is mixed with fresh ambient, dehumidified at the evaporator "
        "(T_evap fixed at 5 deg C), then reheated at the condenser. The first law is enforced: "
        "Q_cond = Q_evap + W_comp.",
        style="List Bullet")

    # 1.2 Open-loop baseline
    doc.add_heading("1.2  Open-loop baseline (r = 0)", level=2)
    ann0 = a[(a["season"] == "annual") & (a["r_recirc"] == 0.0) & (a["vpd_bypass"] == False)]
    rows = []
    for _, r in ann0.iterrows():
        rows.append([
            r["location"].title(),
            fmt(r["time_h"], 2),
            fmt(r["SEC_kWh_per_kg"], 3),
            fmt(r["SMER_kg_per_kWh"], 3),
            fmt(r["W_comp_kWh"], 2),
            fmt(r["COP_mean"], 2),
            fmt(r["m_water_kg"], 2),
        ])
    doc.add_paragraph("Table 1.1  Config A annual-TMY performance, r = 0.")
    add_table(doc, ["Site", "Time (h)", "SEC (kWh/kg)", "SMER (kg/kWh)",
                    "W_comp (kWh)", "COP_mean", "m_water (kg)"], rows)

    doc.add_paragraph(
        "Annual SEC spans 0.543 kWh/kg at Biratnagar to 0.717 kWh/kg at Kathmandu, "
        "a 32 % range. Two factors drive the gap. First, Biratnagar's warmer annual mean "
        "ambient (T_amb ~ 25 deg C) gives a smaller condenser-to-evaporator lift "
        "(35 K vs. 45 K at Kathmandu), lifting COP from 3.63 to 4.43. Second, the lower "
        "altitude at Biratnagar (72 m, ~100.5 kPa) gives denser air and higher mass flow "
        "at the same face velocity, which reduces the time-integrated compressor load. "
        "Drying time sits in a narrow band (13.9-14.5 h) because T_to_chamber is fixed "
        "at 45 deg C at all sites, so the kinetic driving force is identical."
    )

    # 1.2b seasonal
    seas = a[(a["r_recirc"] == 0.0) & (a["vpd_bypass"] == False) &
             (a["season"].isin(["autumn_oct_nov", "winter_dec_jan", "spring_mar_apr"]))]
    pivot = seas.pivot_table(index="location", columns="season",
                             values="SEC_kWh_per_kg").round(3)
    pivot = pivot[["winter_dec_jan", "spring_mar_apr", "autumn_oct_nov"]]
    rows = []
    for loc, row in pivot.iterrows():
        rows.append([loc.title(), fmt(row["winter_dec_jan"], 3),
                     fmt(row["spring_mar_apr"], 3), fmt(row["autumn_oct_nov"], 3)])
    doc.add_paragraph("Table 1.2  Config A seasonal SEC (kWh/kg), r = 0.")
    add_table(doc, ["Site", "Winter (Dec-Jan)", "Spring (Mar-Apr)",
                    "Autumn (Oct-Nov)"], rows)
    doc.add_paragraph(
        "Seasonal SEC is 15-30 % below the annual-TMY value at all three sites. "
        "The annual TMY includes pre-monsoon and monsoon months when ambient RH exceeds "
        "80 %, raising the humidity-suppression term alpha_RH in the Midilli kinetics and "
        "extending drying. The three drying seasons used here (autumn, winter, spring) "
        "correspond to the typical post-harvest window for apple in Nepal."
    )

    # 1.3 Recirculation sweep
    doc.add_heading("1.3  Recirculation sweep", level=2)
    rs = a[(a["season"] == "annual") & (a["vpd_bypass"] == False) &
           (a["r_recirc"].isin([0.0, 0.3, 0.5, 0.7, 0.9, 1.0]))]
    rs = rs.sort_values(["location", "r_recirc"])
    pivot = rs.pivot_table(index="location", columns="r_recirc",
                           values="SEC_kWh_per_kg").round(3)
    rows = []
    for loc, row in pivot.iterrows():
        rows.append([loc.title()] + [fmt(row[c], 3) for c in pivot.columns])
    doc.add_paragraph("Table 1.3  Config A annual-TMY SEC (kWh/kg) vs. recirculation fraction r.")
    add_table(doc, ["Site"] + [f"r = {c:.1f}" for c in pivot.columns], rows)

    doc.add_paragraph(
        "The response to recirculation is site-dependent. At Kathmandu, SEC falls from "
        "0.717 at r = 0 to a minimum of 0.669 at r = 0.9 (6.7 % reduction). At Biratnagar "
        "the trend inverts: SEC rises from 0.543 to 0.753 at r = 0.9 (39 % increase). "
        "Taplejung sits between the two, with a shallow 14 % increase at r = 0.9."
    )
    doc.add_paragraph(
        "The crossover is set by the evaporator operating point. In open-loop, T_evap "
        "adapts to ambient (T_evap ~ T_amb - 10 K), giving 15 deg C at Biratnagar and "
        "5 deg C at Kathmandu. In closed-loop T_evap is fixed at 5 deg C regardless of "
        "site. At warm sites this forces the heat pump to pull the saturation temperature "
        "lower than ambient conditions require, increasing the pressure ratio and the "
        "compressor work per kg of water removed. At cool sites the fixed T_evap is "
        "already close to the open-loop value, and the dehumidification benefit of mixing "
        "slightly cooler recirculated air into the condenser inlet produces a net SEC gain."
    )
    doc.add_paragraph(
        "Between r = 0.9 and r = 1.0 the change is below 1 % at every site. Full closure "
        "removes no fresh air to dilute the humidity build-up, so the steady-state humidity "
        "and SEC are set by the evaporator's dehumidification rate rather than by r itself."
    )

    # 1.4 evap strategy
    doc.add_heading("1.4  Evaporator temperature strategy", level=2)
    doc.add_paragraph(
        "All closed-loop runs reported above use a fixed T_evap = 5 deg C. An onset-tracking "
        "strategy, in which T_evap is set a few K below the instantaneous mixed-air dew point, "
        "was investigated to reduce the late-drying compressor load once the air is already dry. "
        "The detail is covered under Section 5 (control strategies); the conclusion that applies "
        "here is that under the humidity-suppressed kinetics used in this work (alpha_RH = 1.75) "
        "the fixed strategy remains competitive at Kathmandu and Taplejung, with onset-tracking "
        "giving a 6-10 % gain only at Biratnagar."
    )

    # Takeaway
    doc.add_heading("1.5  Takeaways", level=2)
    doc.add_paragraph(
        "Three properties of Config A propagate through the rest of the thesis.",
    )
    doc.add_paragraph(
        "Drying time is fixed by T_to_chamber and ambient humidity, not by configuration. "
        "Every configuration that reaches 45 deg C chamber inlet will dry in 14-15 h under "
        "annual TMY. Differences between configurations appear in SEC and SMER, not in time.",
        style="List Bullet")
    doc.add_paragraph(
        "SEC is site-driven. Biratnagar (warm, low altitude) reaches 0.543 kWh/kg open-loop, "
        "Kathmandu reaches 0.717, Taplejung 0.566. Any configuration's reported SEC must be "
        "read against its site to be meaningful.",
        style="List Bullet")
    doc.add_paragraph(
        "The open-loop path is the strongest baseline at warm sites. Closed-loop recirculation "
        "only reduces SEC when ambient is cold enough that the fixed T_evap = 5 deg C sits at or "
        "below the open-loop evaporator setpoint. This constraint recurs when Config B's solar "
        "collector is combined with recirculation in Section 2.",
        style="List Bullet")

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
