"""Generate Results & Discussion Section 2: VPD Bypass Optimisation.

Creates a Word document covering:
  2.1  The Late-Drying Inefficiency Problem
  2.2  VPD Bypass Strategy — Mechanism and Implementation
  2.3  Config A (r=0.9) — Condenser-Direct Bypass Optimisation
  2.4  Config D1/D2 — Exhaust Bypass Optimisation
  2.5  Config E1/E2 — Exhaust Bypass Optimisation
  2.6  Optimal Threshold Selection: Balancing SEC and Drying Time
  2.7  Summary of VPD Bypass Impact

Usage:
    python scripts/write_section2_vpd_bypass.py
"""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = Path(__file__).resolve().parents[1]
FIGURES_DIR = PROJECT_ROOT / "figures" / "thesis"


# ── helpers ────────────────────────────────────────────────────────────


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(size)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    return table


def add_figure_placeholder(doc, label, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[DIAGRAM: {label}]")
    run.bold = True
    run.font.size = Pt(12)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.italic = True
    return cap


def fmt(val, decimals=3):
    return f"{val:.{decimals}f}"


def pct(base, new):
    """Percent improvement (positive = new is lower)."""
    if base == 0:
        return "-"
    return f"{(base - new) / base * 100:.1f}%"


# ── document sections ──────────────────────────────────────────────────


def write_section_2_1(doc):
    """2.1 The late-drying inefficiency problem."""
    add_heading(doc, "2.1  The Late-Drying Inefficiency Problem", level=2)

    add_para(doc, (
        "Section 1.7 demonstrated that the evaporator\u2019s dehumidification "
        "effectiveness declines sharply as drying progresses. In late-stage "
        "drying, the product has low moisture content, the exhaust air is "
        "only marginally more humid than the inlet, and the evaporator "
        "continues cooling the air without meaningful moisture removal. "
        "This wastes compressor energy on cooling that must then be undone "
        "by the condenser."
    ))

    add_para(doc, (
        "The same problem affects the open-loop D and E configurations. "
        "In late drying, the chamber exhaust carries very little additional "
        "moisture. The heat recovery exchanger (HRX) pre-warms the inlet "
        "air using this exhaust, but the benefit diminishes when there is "
        "little temperature or humidity difference between streams. "
        "Meanwhile, the heat pump continues to heat ambient air from the "
        "HRX outlet to T_set, consuming energy that scales with drying time "
        "rather than moisture removal."
    ))

    add_para(doc, (
        "Two questions arise: (1) Can we detect when the air is no longer "
        "being productively used? (2) Can we reduce energy consumption "
        "during these periods without unacceptably extending drying time?"
    ))


def write_section_2_2(doc):
    """2.2 VPD bypass strategy — mechanism and implementation."""
    add_heading(doc, "2.2  VPD Bypass Strategy \u2014 Mechanism", level=2)

    add_para(doc, (
        "Two related bypass strategies were developed, each using a "
        "Vapour Pressure Deficit (VPD) criterion to detect when the drying "
        "air is under-utilised."
    ))

    add_para(doc, "Condenser-direct bypass (Config A, r > 0)", bold=True)

    add_figure_placeholder(doc,
        "Config A condenser-direct bypass schematic",
        "Figure 2.1: Config A condenser-direct bypass. Left: normal evaporator "
        "path. Right: bypass mode \u2014 warm exhaust routes directly to condenser, "
        "evaporator runs on ambient as heat source only.")

    add_para(doc, (
        "The condenser penalty fraction estimates how much the evaporator "
        "improves the air\u2019s drying potential compared to simply recirculating "
        "the exhaust directly:"
    ))

    add_para(doc, (
        "cond_penalty = (VPD_post_evap \u2212 VPD_exhaust) / VPD_post_evap"
    ), italic=True)

    add_para(doc, (
        "where VPD is evaluated at T_set for both streams. When this "
        "fraction falls below a threshold (i.e. the evaporator barely "
        "improves drying potential), the system switches to condenser-direct "
        "mode: warm exhaust (~43 \u00b0C) is routed directly to the condenser, "
        "which needs only ~2 K of reheating (Q_cond \u2248 0.2 kW vs ~2 kW in "
        "normal mode). The evaporator switches to drawing heat from ambient "
        "air (sensible heat source only, no dehumidification)."
    ))

    add_para(doc, (
        "Humidity accumulates in the closed loop during bypass. When the "
        "condenser penalty rises above 3\u00d7 the threshold, the system switches "
        "back to normal evaporator mode for dehumidification. A minimum "
        "dwell time of 300 s prevents rapid oscillation."
    ))

    add_para(doc, "Exhaust bypass (Configs D/E, r = 0)", bold=True)

    add_figure_placeholder(doc,
        "Config D/E exhaust bypass schematic",
        "Figure 2.2: Config D exhaust bypass. Left: normal HRX path. "
        "Right: bypass mode \u2014 warm exhaust routes directly to condenser, "
        "bypassing the HRX.")

    add_para(doc, (
        "For open-loop D and E configurations, the VPD utilisation metric "
        "measures the fraction of the inlet air\u2019s drying potential that is "
        "consumed by the chamber:"
    ))

    add_para(doc, (
        "utilisation = 1 \u2212 VPD_exhaust / VPD_inlet"
    ), italic=True)

    add_para(doc, (
        "High utilisation indicates the air is picking up significant "
        "moisture (early drying). Low utilisation means the air passes "
        "through nearly unchanged (late drying). When utilisation falls "
        "below the threshold, the warm exhaust (~40 \u00b0C) is routed directly "
        "to the condenser, bypassing the HRX. The condenser provides a "
        "small temperature lift (~5 K), dramatically reducing compressor "
        "work. When utilisation rises above 3\u00d7 the threshold, the system "
        "reverts to normal HRX operation. Dwell time is 600 s."
    ))


def write_section_2_3(doc):
    """2.3 Config A condenser-direct bypass optimisation."""
    add_heading(doc, "2.3  Config A (r = 0.9): Condenser-Direct Bypass", level=2)

    add_para(doc, (
        "A sweep of the condenser penalty threshold was conducted for "
        "Config A at r = 0.9 across both locations. Table 2.1 presents "
        "SEC and drying time as a function of threshold."
    ))

    headers = ["Threshold", "KTM SEC", "KTM time (h)", "KTM \u0394SEC",
               "BTN SEC", "BTN time (h)", "BTN \u0394SEC"]
    rows = [
        ["OFF",  "0.669", "14.9", "\u2014",     "0.753", "14.6", "\u2014"],
        ["0.01", "0.669", "14.9", "0%",          "0.753", "14.6", "0%"],
        ["0.02", "0.650", "14.9", "\u22122.8%",  "0.753", "14.6", "0%"],
        ["0.05", "0.536", "15.6", "\u221219.7%", "0.617", "15.3", "\u221218.1%"],
        ["0.10", "0.365", "19.1", "\u221245.4%", "0.395", "19.8", "\u221247.5%"],
        ["0.15", "0.331", "30.9", "\u221250.5%", "0.336", "28.8", "\u221255.4%"],
        ["0.20", "0.349", "48.5", "\u221247.9%", "0.303", "30.9", "\u221259.7%"],
        ["0.30", "0.383", "72.0", "\u221242.7%", "0.364", "72.0", "\u221251.6%"],
    ]
    add_table(doc, headers, rows)
    add_para(doc,
        "Table 2.1: Config A (r = 0.9) condenser-direct bypass sweep. "
        "\u0394SEC is improvement relative to no bypass.",
        italic=True, size=9)

    add_para(doc, (
        "The results reveal a dramatic trade-off between SEC and drying time. "
        "At threshold = 0.10, SEC is reduced by 45\u201348% with a moderate drying "
        "time increase from ~15 h to ~19 h (+30%). At threshold = 0.15, SEC "
        "improves by 50\u201355% but drying time doubles to 29\u201331 h. Above 0.20, "
        "the system spends too long in bypass mode, humidity accumulates, and "
        "the simulation approaches the 72 h limit."
    ))

    add_para(doc, (
        "The mechanism is intuitive: during bypass, the compressor work "
        "drops from ~2 kW (heating air from 12 \u00b0C to 45 \u00b0C) to ~0.2 kW "
        "(heating from 43 \u00b0C to 45 \u00b0C). However, no dehumidification occurs, "
        "so chamber RH rises and the drying rate slows. The system oscillates "
        "between bypass (low energy, slow drying) and normal (high energy, "
        "fast drying) modes, with the threshold determining the duty cycle."
    ))

    add_para(doc, (
        "For a practical system where drying time is constrained to ~16 h "
        "(one operational day), threshold = 0.05 provides the best balance: "
        "19\u201320% SEC reduction with only ~1 h of additional drying time."
    ))


def write_section_2_4(doc):
    """2.4 Config D1/D2 exhaust bypass optimisation."""
    add_heading(doc, "2.4  Configs D1 and D2: Exhaust Bypass Optimisation", level=2)

    add_para(doc, (
        "Tables 2.2 and 2.3 present the VPD utilisation threshold sweep "
        "for Configs D1 and D2 respectively."
    ))

    # D1 table
    headers = ["Threshold", "KTM SEC", "KTM time (h)", "KTM \u0394SEC",
               "BTN SEC", "BTN time (h)", "BTN \u0394SEC"]
    rows_d1 = [
        ["OFF",  "0.365", "13.8", "\u2014",     "0.293", "14.4", "\u2014"],
        ["0.01", "0.331", "14.0", "\u22129.3%",  "0.264", "14.6", "\u221210.2%"],
        ["0.02", "0.314", "14.2", "\u221214.1%", "0.253", "14.8", "\u221213.7%"],
        ["0.05", "0.287", "15.1", "\u221221.4%", "0.238", "15.8", "\u221219.1%"],
        ["0.10", "0.268", "17.2", "\u221226.7%", "0.225", "18.0", "\u221223.2%"],
        ["0.15", "0.257", "20.5", "\u221229.6%", "0.221", "22.3", "\u221224.6%"],
        ["0.20", "0.258", "25.3", "\u221229.3%", "0.229", "30.3", "\u221222.0%"],
        ["0.30", "0.337", "72.0", "\u22127.8%",  "0.298", "72.0", "\u2212*"],
    ]
    add_table(doc, headers, rows_d1)
    add_para(doc,
        "Table 2.2: Config D1 (HRX, ambient evaporator) VPD exhaust bypass sweep.",
        italic=True, size=9)

    # D2 table
    rows_d2 = [
        ["OFF",  "0.354", "13.8", "\u2014",     "0.282", "14.4", "\u2014"],
        ["0.01", "0.327", "14.0", "\u22127.7%",  "0.258", "14.6", "\u22128.5%"],
        ["0.02", "0.311", "14.2", "\u221212.0%", "0.250", "14.8", "\u221211.5%"],
        ["0.05", "0.289", "15.1", "\u221218.4%", "0.239", "15.8", "\u221215.4%"],
        ["0.10", "0.273", "17.2", "\u221222.9%", "0.231", "18.0", "\u221217.9%"],
        ["0.15", "0.267", "20.5", "\u221224.5%", "0.232", "22.3", "\u221217.6%"],
        ["0.20", "0.273", "25.3", "\u221222.9%", "0.247", "30.3", "\u221212.4%"],
        ["0.30", "0.389", "72.0", "+9.8%",        "0.348", "72.0", "+23.3%"],
    ]
    add_table(doc, headers, rows_d2)
    add_para(doc,
        "Table 2.3: Config D2 (HRX, dynamic ambient compensation evaporator) "
        "VPD exhaust bypass sweep.",
        italic=True, size=9)

    add_para(doc, (
        "Both D1 and D2 show substantial SEC improvement from exhaust bypass, "
        "with diminishing returns above threshold = 0.10. The pattern is "
        "consistent across locations: at threshold = 0.05, SEC improves by "
        "15\u201321% with only 1\u20132 h of additional drying time (from ~14 h to "
        "~15\u201316 h). At threshold = 0.10, improvement reaches 18\u201327% but "
        "drying time extends to 17\u201318 h."
    ))

    add_para(doc, (
        "D1 and D2 respond similarly to the bypass, with D2 showing slightly "
        "smaller relative gains. This is expected: D2\u2019s dynamic ambient "
        "compensation at the evaporator already improves the heat pump\u2019s "
        "COP (by using cooled exhaust rather than raw ambient at the "
        "evaporator), so the marginal benefit of bypassing is smaller."
    ))

    add_para(doc, (
        "At threshold \u2265 0.30, the simulation exceeds 72 h for both configs, "
        "indicating that the system spends nearly all its time in bypass mode "
        "and drying rate is insufficient to reach the target moisture content."
    ))


def write_section_2_5(doc):
    """2.5 Config E1/E2 exhaust bypass optimisation."""
    add_heading(doc, "2.5  Configs E1 and E2: Exhaust Bypass Optimisation", level=2)

    add_para(doc, (
        "Tables 2.4 and 2.5 present the VPD bypass sweep for the solar-assisted "
        "HRX configurations E1 and E2 (A_c = 10 m\u00b2)."
    ))

    headers = ["Threshold", "KTM SEC", "KTM time (h)", "KTM \u0394SEC",
               "BTN SEC", "BTN time (h)", "BTN \u0394SEC"]

    # E1 table
    rows_e1 = [
        ["OFF",  "0.220", "13.8", "\u2014",     "0.141", "14.4", "\u2014"],
        ["0.01", "0.186", "14.0", "\u221215.5%", "0.111", "14.6", "\u221221.2%"],
        ["0.02", "0.169", "14.2", "\u221223.2%", "0.103", "14.8", "\u221226.9%"],
        ["0.05", "0.152", "15.1", "\u221230.9%", "0.101", "15.8", "\u221228.5%"],
        ["0.10", "0.155", "17.2", "\u221229.7%", "0.106", "18.0", "\u221225.0%"],
        ["0.15", "0.151", "20.5", "\u221231.4%", "0.110", "22.3", "\u221221.7%"],
        ["0.20", "0.150", "25.3", "\u221231.6%", "0.114", "30.3", "\u221219.4%"],
        ["0.30", "0.235", "72.0", "+6.5%",       "0.184", "72.0", "+30.8%"],
    ]
    add_table(doc, headers, rows_e1)
    add_para(doc,
        "Table 2.4: Config E1 (HRX + Solar, ambient evaporator, A_c = 10 m\u00b2) "
        "VPD exhaust bypass sweep.",
        italic=True, size=9)

    # E2 table
    rows_e2 = [
        ["OFF",  "0.197", "13.8", "\u2014",     "0.129", "14.4", "\u2014"],
        ["0.01", "0.170", "14.0", "\u221213.7%", "0.105", "14.6", "\u221218.6%"],
        ["0.02", "0.157", "14.2", "\u221220.3%", "0.099", "14.8", "\u221223.6%"],
        ["0.05", "0.144", "15.1", "\u221226.9%", "0.097", "15.8", "\u221224.7%"],
        ["0.10", "0.147", "17.2", "\u221225.2%", "0.102", "18.0", "\u221220.7%"],
        ["0.15", "0.146", "20.5", "\u221225.9%", "0.108", "22.3", "\u221216.5%"],
        ["0.20", "0.147", "25.3", "\u221225.6%", "0.112", "30.3", "\u221213.6%"],
        ["0.30", "0.232", "72.0", "+17.8%",      "0.182", "72.0", "+41.4%"],
    ]
    add_table(doc, headers, rows_e2)
    add_para(doc,
        "Table 2.5: Config E2 (HRX + Solar, dynamic evaporator, A_c = 10 m\u00b2) "
        "VPD exhaust bypass sweep.",
        italic=True, size=9)

    add_para(doc, (
        "The E configurations show the sharpest knee in the SEC\u2013time curve "
        "at threshold = 0.05. At this threshold:"
    ))

    add_bullet(doc,
        "E1: SEC reduces by 29\u201331% (KTM: 0.220 \u2192 0.152; BTN: 0.141 \u2192 0.101 kWh/kg) "
        "with only 1\u20132 h extra drying time."
    )
    add_bullet(doc,
        "E2: SEC reduces by 25\u201327% (KTM: 0.197 \u2192 0.144; BTN: 0.129 \u2192 0.097 kWh/kg) "
        "with only 1\u20132 h extra drying time."
    )

    add_para(doc, (
        "Above threshold = 0.05, the E configs show diminishing returns with "
        "rapidly increasing drying time. This is because the solar collector "
        "already provides substantial free energy; during bypass, the system "
        "loses access to the HRX\u2019s heat recovery while the solar benefit "
        "alone is insufficient to maintain the drying rate. The optimal "
        "threshold for E configs (0.05) is lower than for D configs (0.10\u20130.15) "
        "because the solar energy makes the system more sensitive to "
        "operational disruptions."
    ))

    add_para(doc, (
        "E2 at Biratnagar with threshold = 0.05 achieves SEC = 0.097 kWh/kg, "
        "which represents the lowest specific energy consumption across all "
        "configurations and conditions tested in this study. This is 86% lower "
        "than the Config A baseline (0.717 kWh/kg at r = 0) and 82% lower than "
        "Config A with recirculation (0.543 kWh/kg at r = 0)."
    ))


def write_section_2_6(doc):
    """2.6 Optimal threshold selection."""
    add_heading(doc, "2.6  Optimal Threshold Selection: Balancing SEC and Drying Time", level=2)

    add_para(doc, (
        "Table 2.6 summarises the recommended thresholds, selected to achieve "
        "substantial SEC reduction while keeping drying time within practical "
        "limits (no more than ~2 h increase over the baseline ~14 h cycle)."
    ))

    headers = ["Config", "Bypass type", "Recommended\nthreshold",
               "KTM SEC\n(vs baseline)", "BTN SEC\n(vs baseline)",
               "Time increase"]
    rows = [
        ["A (r=0.9)", "Cond-direct", "0.05",
         "0.536 (\u221220%)", "0.617 (\u221218%)", "+1 h"],
        ["D1", "Exhaust", "0.05",
         "0.287 (\u221221%)", "0.238 (\u221219%)", "+1\u20132 h"],
        ["D2", "Exhaust", "0.05",
         "0.289 (\u221218%)", "0.239 (\u221215%)", "+1\u20132 h"],
        ["E1 (10 m\u00b2)", "Exhaust", "0.05",
         "0.152 (\u221231%)", "0.101 (\u221229%)", "+1\u20132 h"],
        ["E2 (10 m\u00b2)", "Exhaust", "0.05",
         "0.144 (\u221227%)", "0.097 (\u221225%)", "+1\u20132 h"],
    ]
    add_table(doc, headers, rows)
    add_para(doc,
        "Table 2.6: Recommended VPD bypass thresholds balancing SEC and drying time.",
        italic=True, size=9)

    add_para(doc, (
        "A uniform threshold of 0.05 is recommended across all configurations. "
        "While D configs show marginally better SEC at 0.10\u20130.15, the drying "
        "time penalty (3\u20136 h additional) makes these impractical for daily batch "
        "operations where the dryer must complete a cycle within one working period."
    ))

    add_para(doc, (
        "For Config A, the threshold operates on a different metric "
        "(condenser penalty rather than VPD utilisation), but the 0.05 value "
        "provides a comparable balance. Higher thresholds (0.10\u20130.15) are "
        "available for applications where energy cost is critical and a "
        "longer drying period is acceptable."
    ))


def write_section_2_7(doc):
    """2.7 Summary of VPD bypass impact."""
    add_heading(doc, "2.7  Summary: VPD Bypass Impact on System Performance", level=2)

    add_para(doc, (
        "Table 2.7 provides a comprehensive comparison of all configurations "
        "with and without VPD bypass at the recommended threshold of 0.05, "
        "alongside the Config A open-loop baseline."
    ))

    headers = ["Config", "VPD", "KTM SEC", "BTN SEC", "KTM time (h)", "BTN time (h)"]
    rows = [
        ["A (r=0)", "OFF", "0.717", "0.543", "13.8", "14.4"],
        ["A (r=0.9)", "OFF", "0.669", "0.753", "14.9", "14.6"],
        ["A (r=0.9)", "0.05", "0.536", "0.617", "15.6", "15.3"],
        ["D1", "OFF", "0.365", "0.293", "13.8", "14.4"],
        ["D1", "0.05", "0.287", "0.238", "15.1", "15.8"],
        ["D2", "OFF", "0.354", "0.282", "13.8", "14.4"],
        ["D2", "0.05", "0.289", "0.239", "15.1", "15.8"],
        ["E1 (10 m\u00b2)", "OFF", "0.220", "0.141", "13.8", "14.4"],
        ["E1 (10 m\u00b2)", "0.05", "0.152", "0.101", "15.1", "15.8"],
        ["E2 (10 m\u00b2)", "OFF", "0.197", "0.129", "13.8", "14.4"],
        ["E2 (10 m\u00b2)", "0.05", "0.144", "0.097", "15.1", "15.8"],
    ]
    add_table(doc, headers, rows)
    add_para(doc,
        "Table 2.7: Summary of all configurations with and without VPD bypass "
        "(threshold = 0.05). SEC in kWh/kg.",
        italic=True, size=9)

    add_para(doc, (
        "Key findings from the VPD bypass optimisation:"
    ))

    add_bullet(doc,
        "VPD bypass is universally beneficial: every configuration tested "
        "shows SEC improvement at threshold = 0.05, with 15\u201331% reduction "
        "depending on configuration and location."
    )
    add_bullet(doc,
        "The benefit is greatest for E configs (25\u201331%) because the solar "
        "collector reduces the HP\u2019s share of total energy, making the "
        "bypass\u2019s compressor savings a larger fraction of the total."
    )
    add_bullet(doc,
        "Drying time increases by only 1\u20132 h at threshold = 0.05, from "
        "~14 h to ~15\u201316 h \u2014 well within a single daily batch cycle."
    )
    add_bullet(doc,
        "The SEC\u2013time trade-off is non-linear: doubling the threshold from "
        "0.05 to 0.10 typically gains only 5\u20138% additional SEC improvement "
        "but adds 2\u20133 h of drying time. Tripling to 0.15 adds minimal SEC "
        "gain but extends drying by 5\u20137 h."
    )
    add_bullet(doc,
        "At aggressive thresholds (\u22650.30), all configurations fail to "
        "complete drying within 72 h, confirming that the bypass must be "
        "used judiciously."
    )
    add_bullet(doc,
        "The best absolute performance is E2 at Biratnagar with VPD bypass: "
        "SEC = 0.097 kWh/kg. This represents a 13.5\u00d7 reduction in specific "
        "energy from the worst-case scenario (Config A, r = 0.9, BTN, "
        "SEC = 0.753 kWh/kg with no bypass) and a 5.6\u00d7 reduction from the "
        "simple open-loop baseline (Config A, r = 0, BTN, SEC = 0.543 kWh/kg)."
    )

    add_para(doc, (
        "The VPD bypass strategy is a controls-level optimisation that "
        "requires no additional hardware \u2014 only a humidity sensor at the "
        "chamber exhaust and control logic to actuate a damper. Combined "
        "with the hardware-level improvements of heat recovery (D configs) "
        "and solar integration (E configs), it enables SEC values below "
        "0.10 kWh/kg at favourable locations, approaching the theoretical "
        "minimum for heat-pump-assisted convective drying."
    ))


# ── main ───────────────────────────────────────────────────────────────


def main():
    doc = Document()

    add_heading(doc, "Results and Discussion", level=0)
    add_heading(doc, "Section 2: VPD Bypass Optimisation", level=1)
    add_para(doc, (
        "This section investigates a controls-level strategy to reduce "
        "specific energy consumption (SEC) by bypassing the heat pump\u2019s "
        "evaporator or the heat recovery exchanger during periods of "
        "low drying potential. The strategy is applied to Config A "
        "(closed-loop), D1/D2 (HRX + HP), and E1/E2 (HRX + Solar + HP). "
        "A systematic threshold sweep identifies the optimal operating "
        "point that balances SEC reduction against drying time increase."
    ))

    write_section_2_1(doc)
    write_section_2_2(doc)
    write_section_2_3(doc)
    write_section_2_4(doc)
    write_section_2_5(doc)
    write_section_2_6(doc)
    write_section_2_7(doc)

    # Save
    out_dir = PROJECT_ROOT / "thesis"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / "Section2_VPD_Bypass.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
