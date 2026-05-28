"""Generate unified Results & Discussion chapter.

Single document covering all configurations analysed so far:
  1. Config A (HP-Only Baseline)
  2. Config B (Solar + HP Series)
  3. Config C (Solar Cascade)
  4. Evaporator Temperature and Bypass Optimisation

Usage:
    python scripts/write_results_discussion.py
"""

import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_ROOT = PROJECT_ROOT / "outputs"
FIGURES_DIR = PROJECT_ROOT / "figures" / "thesis"


# ── helpers ────────────────────────────────────────────────────────────

TABLE_NUM = [0]  # mutable counter
FIG_NUM = [0]


def next_table():
    TABLE_NUM[0] += 1
    return TABLE_NUM[0]


def next_fig():
    FIG_NUM[0] += 1
    return FIG_NUM[0]


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(size)
    return p


def add_table(doc, headers, rows, caption=None):
    tnum = next_table()
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Table {tnum}: {caption}")
        run.font.size = Pt(9)
        run.italic = True
    return table


def add_figure_placeholder(doc, label, caption=None):
    fnum = next_fig()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[DIAGRAM: {label}]")
    run.bold = True
    run.font.size = Pt(12)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Figure {fnum}: {caption}")
        run.font.size = Pt(9)
        run.italic = True


def fmt(val, decimals=3):
    if pd.isna(val):
        return "\u2014"
    return f"{val:.{decimals}f}"


def pct_diff(base, new):
    if pd.isna(base) or pd.isna(new) or base == 0:
        return "\u2014"
    return f"{(base - new) / base * 100:.1f}%"


def load_data():
    df = pd.read_csv(OUTPUT_ROOT / "master_summary.csv")
    df = df[~df["filename"].str.contains("_s4", na=False)]
    return df


def get_row(df, config, loc, season, r, vpd=False, area=0.0):
    mask = (
        (df["config"] == config)
        & (df["location"] == loc)
        & (df["season"] == season)
        & (df["r_recirc"] == r)
        & (df["vpd_bypass"] == vpd)
    )
    if area > 0:
        mask &= df["solar_area_m2"] == area
    rows = df[mask]
    return rows.iloc[0] if not rows.empty else None


LOC_SHORT = {"kathmandu": "KTM", "biratnagar": "BTN", "taplejung": "TPJ"}
LOCS = ["kathmandu", "biratnagar", "taplejung"]
SEASONS = ["annual", "autumn_oct_nov", "winter_dec_jan", "spring_mar_apr"]
SEASON_NAMES = {
    "annual": "Annual TMY",
    "autumn_oct_nov": "Autumn (Oct\u2013Nov)",
    "winter_dec_jan": "Winter (Dec\u2013Jan)",
    "spring_mar_apr": "Spring (Mar\u2013Apr)",
}


# ======================================================================
#  CHAPTER 1: CONFIG A
# ======================================================================

def write_config_A(doc, df):
    add_heading(doc, "1  Config A: Heat-Pump-Only Baseline", level=1)

    # ── 1.1 System Description ──
    add_heading(doc, "1.1  System Description and Air Paths", level=2)

    add_para(doc, (
        "Config A is the baseline heat-pump convective dryer (HPCD) with no "
        "solar assistance or heat recovery. It operates in one of two modes "
        "depending on the recirculation ratio r."
    ))

    add_para(doc, "Open-loop mode (r = 0)", bold=True)

    add_figure_placeholder(doc,
        "Config A open-loop air path",
        "Config A open-loop (r = 0) schematic. Process air and evaporator "
        "air are separate streams.")

    add_bullet(doc,
        "Process air path: Ambient air enters the condenser, is heated from "
        "T_amb to T_set (45 \u00b0C), passes through the drying chamber where it "
        "picks up moisture, and is exhausted. The air is not dehumidified "
        "\u2014 it enters the chamber at ambient humidity ratio."
    )
    add_bullet(doc,
        "Evaporator heat source: A separate ambient air stream (or the "
        "outdoor environment) provides the low-temperature heat source. "
        "The evaporator saturation temperature is T_evap = T_amb \u2212 10 K, "
        "adapting to local conditions every timestep."
    )
    add_bullet(doc,
        "Temperature lift: The compressor lifts from T_evap = T_amb \u2212 10 \u00b0C "
        "to T_cond = T_set + 10 = 55 \u00b0C. The lift is 65 \u2212 T_amb [\u00b0C], so "
        "warmer ambient air means a smaller lift and higher COP."
    )

    add_para(doc, "Closed-loop mode (r > 0)", bold=True)

    add_figure_placeholder(doc,
        "Config A closed-loop air path",
        "Config A closed-loop (r > 0) schematic. Recirculated exhaust is "
        "mixed with fresh air, dehumidified at the evaporator, then "
        "reheated at the condenser.")

    add_para(doc, (
        "In closed-loop mode, a fraction r of the chamber exhaust is "
        "recirculated and mixed with fresh ambient air before entering "
        "the evaporator:"
    ))

    add_bullet(doc,
        "Mixing: \u03c9_mix = r\u00b7\u03c9_exh + (1\u2212r)\u00b7\u03c9_amb; "
        "h_mix = r\u00b7h_exh + (1\u2212r)\u00b7h_amb."
    )
    add_bullet(doc,
        "Evaporator (dehumidification): The mixed air passes over the coil "
        "at T_coil = T_evap + 3 K, where T_evap = 5 \u00b0C (fixed). Effectiveness "
        "\u03b5_evap = 0.85 determines the outlet temperature. If the outlet "
        "falls below the mixed air\u2019s dew point, moisture condenses and the "
        "air exits saturated."
    )
    add_bullet(doc,
        "Condenser (reheating): The cold, dehumidified air is heated to "
        "T_set = 45 \u00b0C. The first law is enforced: Q_cond = Q_evap + W_comp."
    )

    add_para(doc, (
        "A key difference: in open-loop, T_evap adapts to ambient (T_amb \u2212 10 K). "
        "In closed-loop, T_evap is fixed at 5 \u00b0C. This distinction drives "
        "the performance crossover discussed in Section 1.3."
    ))

    # ── 1.2 Open-Loop Baseline ──
    add_heading(doc, "1.2  Open-Loop Baseline (r = 0)", level=2)

    t_amb_means = {"kathmandu": 9.8, "biratnagar": 18.8, "taplejung": 12.0}
    headers = ["Location", "Elev. (m)", "T_amb mean (\u00b0C)", "SEC (kWh/kg)",
               "Time (h)", "COP mean", "W_comp (kWh)"]
    rows = []
    for loc in LOCS:
        row = get_row(df, "A", loc, "annual", 0.0)
        if row is not None:
            elev = {"kathmandu": "1 350", "biratnagar": "72", "taplejung": "1 820"}[loc]
            rows.append([
                LOC_SHORT[loc], elev, f"{t_amb_means[loc]:.1f}",
                fmt(row["SEC_kWh_per_kg"]),
                fmt(row["time_h"], 1),
                fmt(row["COP_mean"], 2),
                fmt(row["W_comp_kWh"], 1),
            ])
    add_table(doc, headers, rows,
              "Config A open-loop (r = 0) annual-TMY performance.")

    a_ktm = get_row(df, "A", "kathmandu", "annual", 0.0)
    a_btn = get_row(df, "A", "biratnagar", "annual", 0.0)

    add_para(doc, (
        f"SEC ranges from {fmt(a_btn['SEC_kWh_per_kg'])} kWh/kg at Biratnagar "
        f"to {fmt(a_ktm['SEC_kWh_per_kg'])} kWh/kg at Kathmandu, a difference "
        f"of {pct_diff(a_ktm['SEC_kWh_per_kg'], a_btn['SEC_kWh_per_kg'])}. "
        f"Two factors explain the gap:"
    ))

    add_bullet(doc,
        f"Temperature lift: At Biratnagar (T_amb \u2248 19 \u00b0C), the lift is "
        f"\u224846 K, yielding COP = {fmt(a_btn['COP_mean'], 2)}. At Kathmandu "
        f"(T_amb \u2248 10 \u00b0C), the lift is \u224855 K and COP = "
        f"{fmt(a_ktm['COP_mean'], 2)}."
    )
    add_bullet(doc,
        "Atmospheric pressure: Biratnagar at 72 m operates at \u2248100.5 kPa; "
        "Kathmandu at 1 350 m at \u224886.1 kPa. Lower pressure reduces air "
        "density and affects psychrometric state."
    )

    add_para(doc, (
        f"Drying time is nearly identical ({fmt(a_ktm['time_h'], 1)}"
        f"\u2013{fmt(a_btn['time_h'], 1)} h) because the chamber inlet is "
        f"T_set = 45 \u00b0C at all locations. The difference between locations "
        f"is purely in electricity consumption."
    ))

    # Seasonal table
    headers_s = ["Location", "Annual", "Autumn", "Winter", "Spring"]
    rows_s = []
    for loc in LOCS:
        row_data = [LOC_SHORT[loc]]
        for season in SEASONS:
            d = get_row(df, "A", loc, season, 0.0)
            row_data.append(fmt(d["SEC_kWh_per_kg"]) if d is not None else "\u2014")
        rows_s.append(row_data)
    add_table(doc, headers_s, rows_s,
              "Config A (r = 0) SEC (kWh/kg) by season.")

    a_btn_spr = get_row(df, "A", "biratnagar", "spring_mar_apr", 0.0)
    a_btn_win = get_row(df, "A", "biratnagar", "winter_dec_jan", 0.0)

    if a_btn_spr is not None and a_btn_win is not None:
        add_para(doc, (
            f"Seasonal variation is substantial. At Biratnagar, spring SEC "
            f"({fmt(a_btn_spr['SEC_kWh_per_kg'])}) is "
            f"{pct_diff(a_btn_win['SEC_kWh_per_kg'], a_btn_spr['SEC_kWh_per_kg'])} "
            f"lower than winter ({fmt(a_btn_win['SEC_kWh_per_kg'])}), driven "
            f"by the higher spring ambient temperature."
        ))

    # ── 1.3 Recirculation Effect ──
    add_heading(doc, "1.3  Effect of Recirculation", level=2)

    headers = ["Location", "r", "SEC (kWh/kg)", "Time (h)", "COP mean",
               "T_evap mean (\u00b0C)"]
    t_evap_r0 = {"kathmandu": -0.2, "biratnagar": 8.8, "taplejung": 2.0}
    rows = []
    for loc in LOCS:
        for r in [0.0, 0.9, 1.0]:
            row = get_row(df, "A", loc, "annual", r)
            if row is not None:
                t_evap = t_evap_r0[loc] if r == 0 else 5.0
                rows.append([
                    LOC_SHORT[loc], f"{r:.1f}",
                    fmt(row["SEC_kWh_per_kg"]),
                    fmt(row["time_h"], 1),
                    fmt(row["COP_mean"], 2),
                    f"{t_evap:.1f}",
                ])
    add_table(doc, headers, rows,
              "Config A annual-TMY performance at r = 0, 0.9, and 1.0.")

    add_para(doc, "The recirculation crossover", bold=True)

    add_para(doc, (
        "Recirculation improves SEC at Kathmandu but worsens it at "
        "Biratnagar. The mechanism is rooted in the evaporator temperature:"
    ))

    add_bullet(doc,
        "Open-loop (r = 0): T_evap = T_amb \u2212 10 K. At Biratnagar "
        "(T_amb \u2248 19 \u00b0C), T_evap \u2248 9 \u00b0C. At Kathmandu "
        "(T_amb \u2248 10 \u00b0C), T_evap \u2248 0 \u00b0C."
    )
    add_bullet(doc,
        "Closed-loop (r > 0): T_evap is fixed at 5 \u00b0C regardless of "
        "ambient conditions. COP \u2248 4.04."
    )

    add_para(doc, (
        "The crossover occurs at T_amb \u2248 15 \u00b0C, where the open-loop "
        "T_evap (T_amb \u2212 10 = 5 \u00b0C) equals the fixed closed-loop value. "
        "Below this threshold, recirculation raises T_evap (e.g. from 0 \u00b0C "
        "to 5 \u00b0C at Kathmandu) and improves COP. Above it, recirculation "
        "lowers T_evap (e.g. from 9 \u00b0C to 5 \u00b0C at Biratnagar) and "
        "degrades COP."
    ))

    add_para(doc, (
        "A second reinforcing effect is the condenser inlet temperature. "
        "At r = 0, air enters the condenser at T_amb. At r = 0.9, the "
        "evaporator cools mixed air to T_after_evap \u2248 12 \u00b0C. At "
        "Biratnagar, this means the condenser receives 12 \u00b0C instead "
        "of 19 \u00b0C \u2014 7 K of additional reheating. At Kathmandu, the "
        "condenser receives 12 \u00b0C instead of 10 \u00b0C \u2014 slightly warmer."
    ))

    a_ktm_r0 = get_row(df, "A", "kathmandu", "annual", 0.0)
    a_ktm_r09 = get_row(df, "A", "kathmandu", "annual", 0.9)
    a_btn_r0 = get_row(df, "A", "biratnagar", "annual", 0.0)
    a_btn_r09 = get_row(df, "A", "biratnagar", "annual", 0.9)

    if all(x is not None for x in [a_ktm_r0, a_ktm_r09, a_btn_r0, a_btn_r09]):
        add_para(doc, (
            f"At Kathmandu, r = 0.9 reduces annual SEC by "
            f"{pct_diff(a_ktm_r0['SEC_kWh_per_kg'], a_ktm_r09['SEC_kWh_per_kg'])} "
            f"(from {fmt(a_ktm_r0['SEC_kWh_per_kg'])} to "
            f"{fmt(a_ktm_r09['SEC_kWh_per_kg'])} kWh/kg). At Biratnagar, "
            f"r = 0.9 increases SEC from {fmt(a_btn_r0['SEC_kWh_per_kg'])} to "
            f"{fmt(a_btn_r09['SEC_kWh_per_kg'])}."
        ))

    add_para(doc, (
        "The difference between r = 0.9 and r = 1.0 is negligible (\u22641%), "
        "indicating that the system\u2019s energy balance is dominated by the "
        "internal heat pump cycle rather than the small fresh-air fraction."
    ))


# ======================================================================
#  CHAPTER 2: CONFIG B
# ======================================================================

def write_config_B(doc, df):
    add_heading(doc, "2  Config B: Solar-Assisted HP (Series)", level=1)

    # ── 2.1 System Description ──
    add_heading(doc, "2.1  System Description", level=2)

    add_figure_placeholder(doc,
        "Config B air path",
        "Config B series air path. Ambient air passes through the solar "
        "collector, then the condenser, before entering the chamber.")

    add_para(doc, (
        "Config B adds a flat-plate solar collector (A_c = 10 m\u00b2) in series "
        "with the heat pump. The collector is modelled using the "
        "Hottel\u2013Whillier\u2013Bliss equation with an optical efficiency "
        "\u03c4\u03b1 = 0.75 and an overall heat loss coefficient "
        "U_L = 5.0 W/m\u00b2K. These values represent a typical single-glazed "
        "flat-plate air collector (Duffie & Beckman, 2013; \u03c4\u03b1 = 0.75\u20130.80, "
        "U_L = 4\u20136 W/m\u00b2K for single glass). The heat removal factor F_R "
        "is computed dynamically at each timestep using the "
        "NTU\u2013effectiveness method, which accounts for the actual air mass "
        "flow rate and collector area."
    ))

    add_para(doc, (
        "Open-loop (r = 0): Ambient \u2192 Solar \u2192 Condenser \u2192 Chamber. "
        "The solar collector pre-heats ambient air; the condenser provides "
        "the remaining lift to T_set. At night, the system degrades "
        "gracefully to Config A behaviour."
    ))

    add_para(doc, (
        "Closed-loop (r > 0): Mix \u2192 Evaporator \u2192 Solar \u2192 Condenser "
        "\u2192 Chamber. The evaporator dehumidifies the recirculated air, "
        "then the solar collector reheats it before the condenser provides "
        "the final lift."
    ))

    # ── 2.2 Open-Loop Results ──
    add_heading(doc, "2.2  Open-Loop Results (r = 0, A_c = 10 m\u00b2)", level=2)

    headers = ["Location", "SEC (kWh/kg)", "Time (h)", "COP",
               "Solar frac.", "W_comp (kWh)", "\u0394SEC vs A"]
    rows = []
    for loc in LOCS:
        b = get_row(df, "B", loc, "annual", 0.0, area=10.0)
        a = get_row(df, "A", loc, "annual", 0.0)
        if b is not None and a is not None:
            rows.append([
                LOC_SHORT[loc],
                fmt(b["SEC_kWh_per_kg"]),
                fmt(b["time_h"], 1),
                fmt(b["COP_mean"], 2),
                fmt(b["solar_fraction"], 2),
                fmt(b["W_comp_kWh"], 1),
                pct_diff(a["SEC_kWh_per_kg"], b["SEC_kWh_per_kg"]),
            ])
    add_table(doc, headers, rows,
              "Config B (r = 0, A_c = 10 m\u00b2) annual-TMY performance.")

    b_ktm = get_row(df, "B", "kathmandu", "annual", 0.0, area=10.0)
    b_btn = get_row(df, "B", "biratnagar", "annual", 0.0, area=10.0)
    a_ktm = get_row(df, "A", "kathmandu", "annual", 0.0)
    a_btn = get_row(df, "A", "biratnagar", "annual", 0.0)

    if all(x is not None for x in [b_ktm, b_btn, a_ktm, a_btn]):
        add_para(doc, (
            f"Solar integration reduces SEC by "
            f"{pct_diff(a_ktm['SEC_kWh_per_kg'], b_ktm['SEC_kWh_per_kg'])} "
            f"at Kathmandu and "
            f"{pct_diff(a_btn['SEC_kWh_per_kg'], b_btn['SEC_kWh_per_kg'])} "
            f"at Biratnagar. Drying time is unchanged (\u224814 h) because the "
            f"chamber inlet remains at T_set = 45 \u00b0C. The benefit is purely "
            f"in reduced compressor work."
        ))

    # Seasonal table
    headers_s = ["Location", "Annual", "Autumn", "Winter", "Spring"]
    rows_s = []
    for loc in LOCS:
        row_data = [LOC_SHORT[loc]]
        for season in SEASONS:
            d = get_row(df, "B", loc, season, 0.0, area=10.0)
            row_data.append(fmt(d["SEC_kWh_per_kg"]) if d is not None else "\u2014")
        rows_s.append(row_data)
    add_table(doc, headers_s, rows_s,
              "Config B (r = 0, A_c = 10 m\u00b2) SEC by season.")

    # ── 2.3 Recirculation ──
    add_heading(doc, "2.3  Effect of Recirculation on Config B", level=2)

    headers = ["Location", "r", "SEC (kWh/kg)", "Time (h)", "COP", "Solar frac."]
    rows = []
    for loc in LOCS:
        for r in [0.0, 0.9, 1.0]:
            b = get_row(df, "B", loc, "annual", r, area=10.0)
            if b is not None:
                rows.append([
                    LOC_SHORT[loc], f"{r:.1f}",
                    fmt(b["SEC_kWh_per_kg"]),
                    fmt(b["time_h"], 1),
                    fmt(b["COP_mean"], 2),
                    fmt(b["solar_fraction"], 2),
                ])
    add_table(doc, headers, rows,
              "Config B (A_c = 10 m\u00b2) at r = 0, 0.9, and 1.0.")

    b_btn_r0 = get_row(df, "B", "biratnagar", "annual", 0.0, area=10.0)
    b_btn_r09 = get_row(df, "B", "biratnagar", "annual", 0.9, area=10.0)

    if b_btn_r0 is not None and b_btn_r09 is not None:
        add_para(doc, (
            f"Recirculation degrades Config B at all locations. At Biratnagar, "
            f"r = 0.9 increases SEC from {fmt(b_btn_r0['SEC_kWh_per_kg'])} to "
            f"{fmt(b_btn_r09['SEC_kWh_per_kg'])} kWh/kg. The penalty is more "
            f"severe than in Config A because of the solar\u2013evaporator "
            f"interaction:"
        ))

    add_bullet(doc,
        "The evaporator cools the mixed air from ~35 \u00b0C to ~12 \u00b0C. The "
        "solar collector then receives this cold air and must reheat it, "
        "largely compensating for the cooling the evaporator just performed."
    )
    add_bullet(doc,
        "The compressor does extra work to cool air that the solar "
        "collector then undoes. The solar fraction decreases because more "
        "compressor energy is consumed."
    )

    add_para(doc, (
        "For solar-assisted HP dryers in series, open-loop operation is "
        "strongly preferred. This finding raises the question: can the "
        "solar collector be placed differently to avoid this interaction?"
    ))


# ======================================================================
#  CHAPTER 3: CONFIG C
# ======================================================================

def write_config_C(doc, df):
    add_heading(doc, "3  Config C: Solar Cascade Configurations", level=1)

    add_para(doc, (
        "Config B demonstrated that a solar collector in series with the "
        "condenser reduces SEC by 32\u201347%. A natural question is whether "
        "rearranging the solar collector\u2019s position can improve performance "
        "further. Config C explores two cascade arrangements where the "
        "solar collector interacts with the evaporator side of the cycle."
    ))

    # ── 3.1 Variants ──
    add_heading(doc, "3.1  C1 and C2: Two Cascade Arrangements", level=2)

    add_para(doc, "Config C1: Mix before solar", bold=True)

    add_figure_placeholder(doc,
        "Config C1 air path",
        "Config C1 air path. Recirculated exhaust mixes with ambient, "
        "then passes through the solar collector, evaporator, and condenser.")

    add_para(doc, (
        "r > 0: [r\u00d7Exhaust + (1\u2212r)\u00d7Ambient] \u2192 Solar \u2192 "
        "Evaporator \u2192 Condenser \u2192 Chamber"
    ), italic=True)
    add_para(doc, (
        "r = 0: Ambient \u2192 Solar \u2192 Evaporator \u2192 Condenser \u2192 Chamber"
    ), italic=True)

    add_para(doc, (
        "The solar collector heats the mixed air, which then enters the "
        "evaporator. The evaporator must cool this warm, solar-heated air "
        "back down to the coil surface temperature \u2014 undoing the solar gain."
    ))

    add_para(doc, "Config C2: Mix after solar", bold=True)

    add_figure_placeholder(doc,
        "Config C2 air path",
        "Config C2 air path. Fresh ambient air passes through the solar "
        "collector, then mixes with exhaust before the evaporator.")

    add_para(doc, (
        "r > 0: Ambient \u2192 Solar \u2192 [Mix + r\u00d7Exhaust] \u2192 "
        "Evaporator \u2192 Condenser \u2192 Chamber"
    ), italic=True)
    add_para(doc, (
        "r = 0: Ambient \u2192 Solar \u2192 Evaporator \u2192 Condenser \u2192 Chamber"
    ), italic=True)

    add_para(doc, (
        "The solar collector always sees fresh ambient air (cold inlet "
        "\u2192 high collector efficiency). At r = 0, both C1 and C2 reduce "
        "to the same path. In both cases, the evaporator sits between the "
        "solar collector and condenser \u2014 a placement that has profound "
        "consequences."
    ))

    # ── 3.2 C1 Results ──
    add_heading(doc, "3.2  Config C1 Results: The Cost of Wrong Ordering", level=2)

    headers = ["Config", "Location", "A_c (m\u00b2)", "SEC (kWh/kg)",
               "Time (h)", "COP", "Solar frac."]
    rows = []
    for loc in LOCS:
        a = get_row(df, "A", loc, "annual", 0.0)
        if a is not None:
            rows.append(["A (r=0)", LOC_SHORT[loc], "\u2014",
                         fmt(a["SEC_kWh_per_kg"]), fmt(a["time_h"], 1),
                         fmt(a["COP_mean"], 2), "\u2014"])
        b = get_row(df, "B", loc, "annual", 0.0, area=10.0)
        if b is not None:
            rows.append(["B (r=0)", LOC_SHORT[loc], "10",
                         fmt(b["SEC_kWh_per_kg"]), fmt(b["time_h"], 1),
                         fmt(b["COP_mean"], 2), fmt(b["solar_fraction"], 2)])
        c1 = get_row(df, "C1", loc, "annual", 0.0, area=10.0)
        if c1 is not None:
            rows.append(["C1 (r=0)", LOC_SHORT[loc], "10",
                         fmt(c1["SEC_kWh_per_kg"]), fmt(c1["time_h"], 1),
                         fmt(c1["COP_mean"], 2), fmt(c1["solar_fraction"], 2)])
    add_table(doc, headers, rows,
              "Config C1 (r = 0, A_c = 10 m\u00b2) vs Config A and B.")

    c1_ktm = get_row(df, "C1", "kathmandu", "annual", 0.0, area=10.0)
    c1_btn = get_row(df, "C1", "biratnagar", "annual", 0.0, area=10.0)
    a_ktm = get_row(df, "A", "kathmandu", "annual", 0.0)

    if c1_ktm is not None and a_ktm is not None:
        add_para(doc, (
            f"Config C1 produces alarming results. At Kathmandu with 10 m\u00b2, "
            f"SEC = {fmt(c1_ktm['SEC_kWh_per_kg'])} \u2014 nearly identical to "
            f"Config A ({fmt(a_ktm['SEC_kWh_per_kg'])}) despite a solar fraction "
            f"of {fmt(c1_ktm['solar_fraction'], 2)}. Worse, drying time explodes "
            f"to {fmt(c1_ktm['time_h'], 1)} h, more than double the baseline."
        ))

    add_para(doc, (
        "The explanation is straightforward: the solar collector heats air "
        "that the evaporator immediately cools back down. At night, with "
        "no solar gain, the evaporator still cools ambient air unnecessarily "
        "before the condenser reheats it. The solar energy is largely wasted, "
        "and the evaporator\u2019s cooling step extends drying time by forcing "
        "the condenser to work harder."
    ))

    # ── 3.3 C2 Results ──
    add_heading(doc, "3.3  Config C2 Results: Solar-Boosted COP", level=2)

    headers = ["Config", "Location", "SEC (kWh/kg)", "Time (h)", "COP",
               "Solar frac.", "\u0394SEC vs A"]
    rows = []
    for loc in LOCS:
        a = get_row(df, "A", loc, "annual", 0.0)
        b = get_row(df, "B", loc, "annual", 0.0, area=10.0)
        c2 = get_row(df, "C2", loc, "annual", 0.0, area=10.0)
        if c2 is not None and a is not None:
            rows.append([
                "C2 (r=0)", LOC_SHORT[loc],
                fmt(c2["SEC_kWh_per_kg"]), fmt(c2["time_h"], 1),
                fmt(c2["COP_mean"], 2), fmt(c2["solar_fraction"], 2),
                pct_diff(a["SEC_kWh_per_kg"], c2["SEC_kWh_per_kg"]),
            ])
        if b is not None and a is not None:
            rows.append([
                "B (r=0)", LOC_SHORT[loc],
                fmt(b["SEC_kWh_per_kg"]), fmt(b["time_h"], 1),
                fmt(b["COP_mean"], 2), fmt(b["solar_fraction"], 2),
                pct_diff(a["SEC_kWh_per_kg"], b["SEC_kWh_per_kg"]),
            ])
    add_table(doc, headers, rows,
              "Config C2 vs B (r = 0, A_c = 10 m\u00b2, annual TMY).")

    c2_ktm = get_row(df, "C2", "kathmandu", "annual", 0.0, area=10.0)
    c2_btn = get_row(df, "C2", "biratnagar", "annual", 0.0, area=10.0)
    b_ktm = get_row(df, "B", "kathmandu", "annual", 0.0, area=10.0)
    b_btn = get_row(df, "B", "biratnagar", "annual", 0.0, area=10.0)

    if all(x is not None for x in [c2_ktm, c2_btn, b_ktm, b_btn]):
        add_para(doc, (
            f"C2 performs far better than C1, with normal drying times "
            f"({fmt(c2_ktm['time_h'], 1)} h) and notably high COP "
            f"({fmt(c2_ktm['COP_mean'], 2)} at KTM, "
            f"{fmt(c2_btn['COP_mean'], 2)} at BTN). The solar collector "
            f"pre-heats the evaporator source air, raising T_evap and "
            f"reducing the temperature lift."
        ))

        add_para(doc, (
            f"However, C2 still does not match Config B. At Kathmandu, "
            f"B achieves {fmt(b_ktm['SEC_kWh_per_kg'])} versus C2\u2019s "
            f"{fmt(c2_ktm['SEC_kWh_per_kg'])}. The reason is fundamental: "
            f"C2 improves COP but does not reduce Q_cond. The condenser "
            f"still heats air from the evaporator outlet to 45 \u00b0C. In "
            f"Config B, the solar collector directly reduces Q_cond by "
            f"pre-heating the condenser inlet. Since W_comp depends on "
            f"both Q_cond and COP, reducing Q_cond (Config B) is more "
            f"effective than only increasing COP (Config C2)."
        ))

    # ── 3.4 Design Lessons ──
    add_heading(doc, "3.4  Lessons: Where to Put the Solar Collector", level=2)

    add_para(doc, (
        "The ranking is consistent across all locations: B > C2 > C1. "
        "Three design principles emerge:"
    ))

    add_bullet(doc,
        "Solar energy is most effective on the condenser side. Config B\u2019s "
        "series arrangement outperforms the cascade arrangements by "
        "13\u201325% in SEC."
    )
    add_bullet(doc,
        "Component ordering matters critically. Placing the evaporator "
        "after the solar collector (C1/C2) wastes solar energy on air "
        "that must be cooled for dehumidification."
    )
    add_bullet(doc,
        "COP improvement alone is insufficient. C2 achieves COP values "
        "2\u20133\u00d7 higher than B, yet has worse SEC because Q_cond is unchanged."
    )

    add_para(doc, (
        "These principles motivate the Config D family, which introduces "
        "heat recovery to reduce the condenser load \u2014 the same mechanism "
        "that makes Config B effective, but using waste heat from the "
        "exhaust rather than solar energy. Heat recovery works at night "
        "and in all weather, addressing Config B\u2019s dependence on irradiance."
    ))


# ======================================================================
#  CHAPTER 4: EVAPORATOR TEMPERATURE AND BYPASS OPTIMISATION
# ======================================================================

def write_optimisation(doc, df):
    add_heading(doc, "4  Evaporator Temperature and Bypass Optimisation", level=1)

    add_para(doc, (
        "Before introducing the more advanced D and E configurations, "
        "this section addresses two optimisation questions arising from "
        "the Config A analysis: (1) Can the fixed evaporator temperature "
        "be improved? (2) Can the late-drying inefficiency be mitigated "
        "by bypassing the evaporator?"
    ))

    # ── 4.1 Late-drying inefficiency ──
    add_heading(doc, "4.1  The Late-Drying Inefficiency", level=2)

    add_para(doc, (
        "Analysis of the simulation time-series for Config A (r = 0.9) "
        "reveals that dehumidification effectiveness declines sharply "
        "as drying progresses:"
    ))

    headers = ["Drying phase", "\u03c9_mix (g/kg)", "T_dp (\u00b0C)",
               "T_after_evap (\u00b0C)", "\u0394\u03c9 removed (g/kg)"]
    rows = [
        ["Early (1 h)",  "15.6", "20.8", "10.4", "7.6"],
        ["Mid (7 h)",    "11.1", "15.5", "12.4", "2.1"],
        ["Late (13 h)",  "9.6",  "13.3", "13.1", "0.15"],
    ]
    add_table(doc, headers, rows,
              "Evaporator dehumidification over the drying cycle "
              "(Config A, r = 0.9, Biratnagar).")

    add_para(doc, (
        "By hour 13, the mixed air\u2019s dew point (13.3 \u00b0C) nearly equals "
        "the evaporator outlet (13.1 \u00b0C). The evaporator is still cooling "
        "air from ~35 \u00b0C to ~13 \u00b0C \u2014 consuming compressor energy \u2014 but "
        "removing almost no moisture. The condenser must then reheat from "
        "13 \u00b0C to 45 \u00b0C: 22 K of wasted cooling and reheating."
    ))

    # ── 4.2 Adaptive T_evap investigation ──
    add_heading(doc, "4.2  Adaptive Evaporator Temperature: A Negative Result", level=2)

    add_para(doc, (
        "An adaptive strategy was investigated that sets T_evap relative "
        "to the condensation onset temperature \u2014 the coil temperature at "
        "which the evaporator outlet just reaches the dew point:"
    ))

    add_para(doc, (
        "T_evap,coil,onset = T_mix \u2212 (T_mix \u2212 T_dp,mix) / \u03b5_evap"
    ), italic=True)

    add_para(doc, (
        "The evaporator is set \u0394 K below onset: T_evap = T_onset \u2212 \u0394. "
        "A small \u0394 provides minimal subcooling; a large \u0394 approaches "
        "the aggressive dehumidification of the fixed strategy. A sweep "
        "from \u0394 = 2 to 20 K was conducted for Config A at r = 0.9."
    ))

    headers_d = ["\u0394 (K)", "KTM SEC", "KTM time (h)", "BTN SEC", "BTN time (h)"]
    rows_d = [
        ["Fixed 5\u00b0C", "0.669", "14.9", "0.753", "14.6"],
        ["2",  "0.680", "22.4", "0.678", "20.3"],
        ["4",  "0.700", "16.9", "0.747", "16.1"],
        ["6",  "0.731", "14.7", "0.799", "14.2"],
        ["10", "0.806", "13.7", "0.891", "13.3"],
        ["20", "0.892", "13.2", "0.996", "12.8"],
    ]
    add_table(doc, headers_d, rows_d,
              "Onset-tracking adaptive T_evap sweep (Config A, r = 0.9). "
              "Fixed baseline included for comparison.")

    add_para(doc, (
        "At Kathmandu, the fixed baseline wins outright. At Biratnagar, "
        "\u0394 = 2 K provides a 10% SEC improvement but extends drying time "
        "from 14.6 h to 20.3 h \u2014 a 39% increase. For practical daily-batch "
        "operations, this trade-off is unacceptable."
    ))

    add_para(doc, "Why adaptive T_evap fails", bold=True)

    add_para(doc, (
        "The drying kinetics model includes an RH-dependent suppression "
        "term with exponent \u03b1_RH = 1.75. Higher chamber RH directly "
        "slows the drying rate. When the evaporator is raised above 5 \u00b0C "
        "(less dehumidification), chamber RH increases, the drying rate "
        "decreases, and extended fan and compressor hours offset any COP "
        "improvement."
    ))

    add_para(doc, (
        "This is specific to systems with RH-dependent kinetics. In a "
        "temperature-limited system, reducing subcooling would improve COP "
        "without affecting drying time. For apple slices, with their strong "
        "humidity sensitivity, aggressive dehumidification at T_evap = 5 \u00b0C "
        "is near-optimal."
    ))

    # ── 4.3 Condenser-direct bypass ──
    add_heading(doc, "4.3  Condenser-Direct Bypass: A Better Approach", level=2)

    add_para(doc, (
        "Rather than weakening dehumidification (raising T_evap), the "
        "bypass strategy eliminates the cooling\u2013reheating cycle entirely "
        "when dehumidification is no longer needed."
    ))

    add_figure_placeholder(doc,
        "Config A condenser-direct bypass",
        "Condenser-direct bypass. Left: normal evaporator path. Right: "
        "bypass \u2014 warm exhaust routes directly to condenser; evaporator "
        "runs on ambient as heat source only.")

    add_para(doc, (
        "The condenser penalty fraction estimates how much the evaporator "
        "improves drying potential:"
    ))

    add_para(doc, (
        "cond_penalty = (VPD_post_evap \u2212 VPD_exhaust) / VPD_post_evap"
    ), italic=True)

    add_para(doc, (
        "When this drops below a threshold, the system switches to bypass: "
        "warm exhaust (~43 \u00b0C) goes directly to the condenser (Q_cond "
        "\u2248 0.2 kW vs ~2 kW in normal mode). Humidity accumulates in the "
        "loop; when the penalty exceeds 3\u00d7 the threshold, the system "
        "reverts to normal dehumidification. A physics-based dwell time "
        "prevents chattering: the minimum time between mode switches is "
        "computed from the humidity accumulation rate "
        "(\u03c4 = \u0394penalty / |d(penalty)/dt|), clamped between 300 s "
        "(compressor protection) and 7 200 s (prevent lock-in). This "
        "ensures the system waits long enough for humidity to meaningfully "
        "change before switching again."
    ))

    headers_v = ["Threshold", "KTM SEC", "KTM time (h)",
                 "BTN SEC", "BTN time (h)"]
    rows_v = [
        ["OFF",  "0.669", "14.9", "0.753", "14.6"],
        ["0.02", "0.650", "14.9", "0.753", "14.6"],
        ["0.05", "0.536", "15.6", "0.617", "15.3"],
        ["0.10", "0.365", "19.1", "0.395", "19.8"],
        ["0.15", "0.331", "30.9", "0.336", "28.8"],
        ["0.20", "0.349", "48.5", "0.303", "30.9"],
    ]
    add_table(doc, headers_v, rows_v,
              "Config A (r = 0.9) condenser-direct bypass threshold sweep.")

    add_para(doc, (
        "The trade-off is clear. At threshold = 0.05, SEC drops by 18\u201320% "
        "with only ~1 h of additional drying time. At 0.10, the reduction "
        "reaches 45\u201348% but drying extends to ~19 h. Beyond 0.15, drying "
        "time doubles and SEC gains plateau."
    ))

    add_para(doc, (
        "The mechanism is intuitive: during bypass, compressor work drops "
        "from ~2 kW to ~0.2 kW (heating from 43 \u00b0C to 45 \u00b0C instead of "
        "12 \u00b0C to 45 \u00b0C). However, no dehumidification occurs, so chamber "
        "RH rises and drying slows. The system oscillates between bypass "
        "(low energy, slow drying) and normal (high energy, fast drying), "
        "with the threshold controlling the duty cycle."
    ))

    add_para(doc, (
        "For daily-batch operations (target \u226416 h), threshold = 0.05 "
        "provides the best balance. This condenser-direct bypass concept "
        "is extended to D and E configurations as an exhaust bypass "
        "strategy in subsequent sections."
    ))


# ======================================================================
#  CHAPTER 5: CONFIG D — HEAT RECOVERY
# ======================================================================

def write_config_D(doc, df):
    add_heading(doc, "5  Config D: Heat Recovery Exchanger", level=1)

    add_para(doc, (
        "The Config C analysis established that reducing the condenser load "
        "is the most effective use of supplementary energy. Config B achieves "
        "this with solar pre-heating, but its performance depends on "
        "irradiance and degrades at night. Config D takes a different "
        "approach: recovering waste heat from the chamber exhaust using a "
        "counter-flow heat recovery exchanger (HRX). This works continuously "
        "\u2014 day and night, in all weather \u2014 because it exploits the "
        "temperature difference between the warm exhaust and cool ambient "
        "air."
    ))

    # ── 5.1 System Description ──
    add_heading(doc, "5.1  System Description and Variants", level=2)

    add_para(doc, (
        "All D configurations use a counter-flow plate HRX with "
        "effectiveness \u03b5_HRX = 0.70. The HRX transfers sensible heat "
        "from the warm exhaust stream (hot side) to the incoming ambient "
        "air (cold side). The outlet temperatures are:"
    ))

    add_para(doc, (
        "T_amb,heated = T_amb + \u03b5 \u00b7 (T_exhaust \u2212 T_amb)"
    ), italic=True)
    add_para(doc, (
        "T_exh,cooled = T_exhaust \u2212 \u03b5 \u00b7 (T_exhaust \u2212 T_amb)"
    ), italic=True)

    add_para(doc, (
        "If the cooled exhaust falls below its dew point, condensation "
        "occurs on the exhaust side, recovering latent heat as well. "
        "The ambient side experiences only sensible heating (no moisture "
        "addition). All D configs operate in open-loop (r = 0)."
    ))

    add_para(doc, "Config D1: Ambient preheating, separate evaporator", bold=True)

    add_figure_placeholder(doc,
        "Config D1 air path",
        "Config D1 air path. Ambient air is preheated by the HRX, then "
        "heated to T_set by the condenser. Exhaust is cooled by the HRX "
        "and expelled. Evaporator draws from a separate ambient stream.")

    add_para(doc, (
        "Condenser stream: Ambient \u2192 HRX (cold side) \u2192 Condenser "
        "\u2192 Chamber"
    ), italic=True)
    add_para(doc, (
        "Exhaust stream: Chamber \u2192 HRX (hot side) \u2192 Expelled"
    ), italic=True)
    add_para(doc, (
        "Evaporator: Separate ambient air (T_evap = T_amb \u2212 10 K)"
    ), italic=True)

    add_para(doc, (
        "The HRX pre-heats the ambient air before the condenser. If the "
        "exhaust is at 40 \u00b0C and ambient at 10 \u00b0C, the HRX delivers "
        "air at 10 + 0.70 \u00d7 30 = 31 \u00b0C to the condenser. The condenser "
        "only needs to provide the remaining 14 K lift to 45 \u00b0C, compared "
        "to 35 K without heat recovery."
    ))

    add_para(doc, "Config D2: Dynamic ambient compensation at evaporator", bold=True)

    add_figure_placeholder(doc,
        "Config D2 air path",
        "Config D2 air path. Same condenser path as D1. The evaporator "
        "draws from cooled exhaust (post-HRX), supplemented with ambient "
        "air when the exhaust alone cannot supply Q_evap.")

    add_para(doc, (
        "D2 uses the same condenser path as D1, but routes the cooled "
        "exhaust (post-HRX) to the evaporator as a heat source. Since "
        "the cooled exhaust is still warmer than raw ambient "
        "(T_exh,cooled \u2248 19 \u00b0C vs T_amb \u2248 10 \u00b0C at Kathmandu), "
        "the evaporator operates at a higher T_evap and better COP. "
        "When the exhaust stream alone cannot supply Q_evap, ambient "
        "air is dynamically mixed in to make up the deficit. This "
        "creates a circular dependency between the air mix and the HP "
        "operating point, resolved by a fixed-point iteration "
        "described in Section 6.1 under Config E2."
    ))

    add_para(doc, "Config D3: Swapped routing (humidity risk)", bold=True)

    add_para(doc, (
        "D3 reverses the stream assignment: exhaust goes through the HRX "
        "and then to the condenser; ambient goes through the HRX and then "
        "to the evaporator. This gives the evaporator a warm source "
        "(HRX-heated ambient) for high COP, but sends partially "
        "dehumidified exhaust air to the chamber \u2014 introducing a "
        "humidity penalty."
    ))

    # ── 5.2 D1 and D2 Results ──
    add_heading(doc, "5.2  D1 and D2 Results", level=2)

    headers = ["Config", "Location", "SEC (kWh/kg)", "Time (h)", "COP",
               "\u0394SEC vs A(r=0)"]
    rows = []
    for loc in LOCS:
        a = get_row(df, "A", loc, "annual", 0.0)
        for cfg_name in ["D1", "D2"]:
            d = get_row(df, cfg_name, loc, "annual", 0.0)
            if d is not None and a is not None:
                rows.append([
                    cfg_name, LOC_SHORT[loc],
                    fmt(d["SEC_kWh_per_kg"]),
                    fmt(d["time_h"], 1),
                    fmt(d["COP_mean"], 2),
                    pct_diff(a["SEC_kWh_per_kg"], d["SEC_kWh_per_kg"]),
                ])
    add_table(doc, headers, rows,
              "Config D1 and D2 annual-TMY performance vs Config A (r = 0).")

    d1_ktm = get_row(df, "D1", "kathmandu", "annual", 0.0)
    d2_ktm = get_row(df, "D2", "kathmandu", "annual", 0.0)
    d1_btn = get_row(df, "D1", "biratnagar", "annual", 0.0)
    d2_btn = get_row(df, "D2", "biratnagar", "annual", 0.0)
    a_ktm = get_row(df, "A", "kathmandu", "annual", 0.0)
    a_btn = get_row(df, "A", "biratnagar", "annual", 0.0)

    if all(x is not None for x in [d1_ktm, d2_ktm, d1_btn, d2_btn, a_ktm, a_btn]):
        add_para(doc, (
            f"Heat recovery delivers dramatic SEC reductions. D1 achieves "
            f"{fmt(d1_ktm['SEC_kWh_per_kg'])} kWh/kg at Kathmandu "
            f"({pct_diff(a_ktm['SEC_kWh_per_kg'], d1_ktm['SEC_kWh_per_kg'])} "
            f"below Config A) and {fmt(d1_btn['SEC_kWh_per_kg'])} at Biratnagar "
            f"({pct_diff(a_btn['SEC_kWh_per_kg'], d1_btn['SEC_kWh_per_kg'])} "
            f"below Config A). Drying time is unchanged (\u224814 h) because "
            f"the chamber still receives air at T_set = 45 \u00b0C."
        ))

        add_para(doc, (
            f"D2 provides a further improvement over D1: "
            f"{fmt(d2_ktm['SEC_kWh_per_kg'])} at Kathmandu (COP = "
            f"{fmt(d2_ktm['COP_mean'], 2)}) vs D1\u2019s "
            f"{fmt(d1_ktm['SEC_kWh_per_kg'])} (COP = "
            f"{fmt(d1_ktm['COP_mean'], 2)}). The advantage comes from the "
            f"warmer evaporator source: at Kathmandu, D2\u2019s evaporator sees "
            f"cooled exhaust at \u224819 \u00b0C (T_evap \u2248 9 \u00b0C) instead of D1\u2019s "
            f"raw ambient at \u224810 \u00b0C (T_evap \u2248 0 \u00b0C). The reduced "
            f"temperature lift improves COP by {fmt(d2_ktm['COP_mean'] / d1_ktm['COP_mean'] * 100 - 100, 0)}%."
        ))

    add_para(doc, (
        "Compared to Config B (solar series, 10 m\u00b2), the D configs "
        "offer a meaningful trade-off. Config B achieves lower SEC at "
        "high-irradiance locations (BTN) but depends on weather. The D "
        "configs provide consistent performance regardless of solar "
        "conditions, making them attractive for operations that span "
        "night-time or monsoon seasons."
    ))

    # ── 5.3 D3: Humidity Risk ──
    add_heading(doc, "5.3  Config D3: The Humidity Penalty", level=2)

    headers = ["Config", "Location", "SEC (kWh/kg)", "Time (h)", "COP"]
    rows = []
    for loc in LOCS:
        d3 = get_row(df, "D3", loc, "annual", 0.0)
        if d3 is not None:
            rows.append([
                "D3", LOC_SHORT[loc],
                fmt(d3["SEC_kWh_per_kg"]),
                fmt(d3["time_h"], 1),
                fmt(d3["COP_mean"], 2),
            ])
    add_table(doc, headers, rows,
              "Config D3 annual-TMY performance.")

    d3_ktm = get_row(df, "D3", "kathmandu", "annual", 0.0)
    d3_btn = get_row(df, "D3", "biratnagar", "annual", 0.0)

    if d3_ktm is not None and d3_btn is not None:
        add_para(doc, (
            f"D3 achieves the highest COP among D variants "
            f"({fmt(d3_ktm['COP_mean'], 2)} at KTM, "
            f"{fmt(d3_btn['COP_mean'], 2)} at BTN) because the evaporator "
            f"sees warm HRX-heated ambient air. However, SEC is "
            f"substantially worse ({fmt(d3_ktm['SEC_kWh_per_kg'])} and "
            f"{fmt(d3_btn['SEC_kWh_per_kg'])} kWh/kg) and drying time "
            f"extends to {fmt(d3_ktm['time_h'], 1)}\u2013"
            f"{fmt(d3_btn['time_h'], 1)} h."
        ))

    add_para(doc, (
        "The root cause is that D3 sends the HRX-cooled exhaust \u2014 "
        "which is still humid \u2014 to the condenser and then the chamber. "
        "The chamber receives air at 45 \u00b0C but with higher humidity "
        "ratio than in D1/D2 (which send dry ambient air). The RH-dependent "
        "drying kinetics (\u03b1_RH = 1.75) directly penalise this: higher "
        "chamber RH slows drying, extending the batch time and increasing "
        "total energy consumption despite the better COP."
    ))

    add_para(doc, (
        "This reinforces the lesson from the adaptive T_evap investigation: "
        "for products with strong humidity sensitivity, the humidity of the "
        "chamber inlet air matters as much as \u2014 or more than \u2014 the "
        "thermodynamic efficiency of the heat pump cycle. D3 is not "
        "recommended for humidity-sensitive products."
    ))

    # ── 5.4 VPD Bypass for D ──
    add_heading(doc, "5.4  VPD Exhaust Bypass for D Configs", level=2)

    add_para(doc, (
        "The condenser-direct bypass developed in Section 4.3 for Config A "
        "is adapted here as an exhaust bypass for the open-loop D configs. "
        "The trigger metric is VPD utilisation \u2014 the fraction of the "
        "inlet air\u2019s drying potential consumed by the chamber:"
    ))

    add_para(doc, (
        "utilisation = 1 \u2212 VPD_exhaust / VPD_inlet"
    ), italic=True)

    add_figure_placeholder(doc,
        "Config D exhaust bypass",
        "Exhaust bypass for D configs. When VPD utilisation drops below "
        "the threshold, warm exhaust routes directly to condenser, "
        "bypassing the HRX.")

    add_para(doc, (
        "When utilisation falls below the threshold (air passes through "
        "the chamber nearly unchanged in late drying), warm exhaust "
        "(\u224840 \u00b0C) bypasses the HRX and routes directly to the "
        "condenser. The condenser provides only \u22485 K of reheating "
        "instead of the full lift from HRX-heated ambient. When "
        "utilisation rises above 3\u00d7 the threshold, the system reverts "
        "to normal HRX operation. A fixed 600 s minimum dwell time "
        "prevents oscillation."
    ))

    headers = ["Config", "Location", "SEC (no VPD)", "SEC (VPD 0.05)",
               "\u0394SEC", "Time (no VPD)", "Time (VPD)"]
    rows = []
    for loc in LOCS:
        for cfg_name in ["D1", "D2"]:
            d_off = get_row(df, cfg_name, loc, "annual", 0.0, vpd=False)
            d_on = get_row(df, cfg_name, loc, "annual", 0.0, vpd=True)
            if d_off is not None and d_on is not None:
                rows.append([
                    cfg_name, LOC_SHORT[loc],
                    fmt(d_off["SEC_kWh_per_kg"]),
                    fmt(d_on["SEC_kWh_per_kg"]),
                    pct_diff(d_off["SEC_kWh_per_kg"], d_on["SEC_kWh_per_kg"]),
                    fmt(d_off["time_h"], 1),
                    fmt(d_on["time_h"], 1),
                ])
    add_table(doc, headers, rows,
              "D1 and D2 with VPD exhaust bypass (threshold = 0.05).")

    add_para(doc, (
        "At the recommended threshold of 0.05, D1 achieves "
        "19\u201321% SEC reduction and D2 achieves 15\u201318%, both with only "
        "1\u20132 h of additional drying time. D2\u2019s smaller percentage gain "
        "reflects the fact that its baseline SEC is already lower (better "
        "COP from the warmer evaporator source), leaving less room for "
        "bypass savings. The combined effect of heat recovery plus VPD "
        "bypass brings D1 to 0.238\u20130.287 kWh/kg and D2 to "
        "0.238\u20130.288 kWh/kg \u2014 comparable to Config B\u2019s solar-assisted "
        "performance but without any dependence on solar irradiance."
    ))


# ======================================================================
#  CHAPTER 6: CONFIG E — COMBINED HRX + SOLAR
# ======================================================================

def write_config_E(doc, df):
    add_heading(doc, "6  Config E: Combined HRX + Solar", level=1)

    add_para(doc, (
        "Config E combines the two most effective strategies identified "
        "so far: heat recovery (Config D) and solar pre-heating on the "
        "condenser side (Config B). By stacking both, the condenser load "
        "is reduced from two independent sources \u2014 waste exhaust heat "
        "and solar irradiance \u2014 achieving the lowest SEC values in this "
        "study."
    ))

    # ── 6.1 System Description ──
    add_heading(doc, "6.1  System Description", level=2)

    add_para(doc, "Config E1: HRX + Solar, separate evaporator", bold=True)

    add_figure_placeholder(doc,
        "Config E1 air path",
        "Config E1 air path. Ambient air is preheated by the HRX, then "
        "by the solar collector, then by the condenser. Evaporator draws "
        "from separate ambient.")

    add_para(doc, (
        "Condenser stream: Ambient \u2192 HRX \u2192 Solar \u2192 Condenser "
        "\u2192 Chamber"
    ), italic=True)
    add_para(doc, (
        "Evaporator: Separate ambient air (as in D1)"
    ), italic=True)

    add_para(doc, (
        "E1 extends D1 by inserting the solar collector between the HRX "
        "and the condenser. The condenser receives air that has been "
        "preheated by both the HRX and the sun. On a clear day when "
        "the HRX delivers air at 31 \u00b0C and the solar collector adds "
        "another 10 K, the condenser only needs to lift from 41 \u00b0C to "
        "45 \u00b0C \u2014 a 4 K lift requiring minimal compressor work."
    ))

    add_para(doc, "Config E2: HRX + Solar, dynamic evaporator", bold=True)

    add_figure_placeholder(doc,
        "Config E2 air path",
        "Config E2 air path. Same condenser path as E1. Evaporator draws "
        "from cooled exhaust with dynamic ambient supplement (as in D2).")

    add_para(doc, (
        "E2 uses the same condenser path as E1 but routes the cooled "
        "exhaust (post-HRX) to the evaporator with dynamic ambient "
        "supplement, as in D2. The warmer evaporator source improves "
        "COP further. When exhaust alone cannot supply Q_evap, ambient "
        "air is mixed in to make up the deficit. However, changing the "
        "air mix changes the evaporator source temperature, which in "
        "turn changes the HP operating point and therefore Q_evap "
        "itself \u2014 creating a circular dependency that must be resolved "
        "iteratively."
    ))

    add_para(doc,
        "Iterative evaporator sizing algorithm", bold=True
    )

    add_para(doc, (
        "The evaporator coil temperature is set at T_evap = "
        "T_evap_source \u2212 10 K (the approach temperature), so the "
        "available heat extraction is:"
    ))

    add_para(doc, (
        "Q_avail = (m\u0307_da + m\u0307_amb,extra) \u00b7 c_p \u00b7 \u03b5_evap \u00b7 10 K"
    ), italic=True)

    add_para(doc, (
        "where m\u0307_da is the main dry-air mass flow and m\u0307_amb,extra "
        "is the supplementary ambient air. When the exhaust stream "
        "alone (m\u0307_amb,extra = 0) provides less heat than Q_evap "
        "required by the HP cycle, the algorithm proceeds as follows:"
    ))

    add_bullet(doc,
        "Step 1: Size the HP at the current T_evap_source to obtain "
        "Q_evap,needed."
    )
    add_bullet(doc,
        "Step 2: Compute Q_avail from the current air mix. If "
        "Q_avail \u2265 Q_evap,needed, converged \u2014 stop."
    )
    add_bullet(doc,
        "Step 3: Calculate the total air mass needed: "
        "m\u0307_total = Q_evap,needed / (c_p \u00b7 \u03b5_evap \u00b7 10). "
        "Set m\u0307_amb,extra = m\u0307_total \u2212 m\u0307_da."
    )
    add_bullet(doc,
        "Step 4: Update the mixed source temperature: "
        "T_evap_source = (m\u0307_da \u00b7 T_exh,cooled + m\u0307_amb,extra \u00b7 "
        "T_amb) / (m\u0307_da + m\u0307_amb,extra)."
    )
    add_bullet(doc,
        "Step 5: If |T_new \u2212 T_old| < 0.05 K, converged. Otherwise "
        "return to Step 1 (max 5 iterations)."
    )

    add_para(doc, (
        "In practice, convergence is reached within 2\u20133 iterations. "
        "The algorithm ensures that the evaporator always has a "
        "physically realisable heat source: exhaust alone when it "
        "suffices (higher COP), or a blend with ambient when the "
        "exhaust is too cold or the flow too small. This same "
        "algorithm is used in Config D2 (Section 5)."
    ))

    add_para(doc, "Config E3: Solar-priority control", bold=True)

    add_figure_placeholder(doc,
        "Config E3 air path",
        "Config E3 air path. Ambient \u2192 HRX \u2192 Condenser (variable) "
        "\u2192 Solar \u2192 Chamber. HP provides partial lift; solar finishes.")

    add_para(doc, (
        "E3 rearranges the solar collector to after the condenser: "
        "Ambient \u2192 HRX \u2192 Condenser \u2192 Solar \u2192 Chamber. The HP "
        "provides a partial temperature lift (variable T_cond), and the "
        "solar collector supplies the remaining gain to T_set. When "
        "solar irradiance alone (from the HRX output temperature) can "
        "reach T_set, the HP is switched off entirely. This solar-priority "
        "control maximises solar utilisation but means the HP operates "
        "at a variable and sometimes higher T_cond when solar contribution "
        "is low."
    ))

    # ── 6.2 E1 and E2 Results ──
    add_heading(doc, "6.2  E1 and E2 Results (A_c = 10 m\u00b2)", level=2)

    headers = ["Config", "Location", "SEC (kWh/kg)", "Time (h)", "COP",
               "Solar frac.", "\u0394SEC vs A(r=0)"]
    rows = []
    for loc in LOCS:
        a = get_row(df, "A", loc, "annual", 0.0)
        for cfg_name in ["E1", "E2"]:
            e = get_row(df, cfg_name, loc, "annual", 0.0, area=10.0)
            if e is not None and a is not None:
                rows.append([
                    cfg_name, LOC_SHORT[loc],
                    fmt(e["SEC_kWh_per_kg"]),
                    fmt(e["time_h"], 1),
                    fmt(e["COP_mean"], 2),
                    fmt(e["solar_fraction"], 2),
                    pct_diff(a["SEC_kWh_per_kg"], e["SEC_kWh_per_kg"]),
                ])
    add_table(doc, headers, rows,
              "Config E1 and E2 (A_c = 10 m\u00b2) annual-TMY performance.")

    e1_btn = get_row(df, "E1", "biratnagar", "annual", 0.0, area=10.0)
    e2_btn = get_row(df, "E2", "biratnagar", "annual", 0.0, area=10.0)
    e2_ktm = get_row(df, "E2", "kathmandu", "annual", 0.0, area=10.0)
    a_btn = get_row(df, "A", "biratnagar", "annual", 0.0)

    if all(x is not None for x in [e1_btn, e2_btn, e2_ktm, a_btn]):
        add_para(doc, (
            f"The combined HRX + solar approach produces the lowest SEC "
            f"values seen so far. E2 at Biratnagar achieves "
            f"{fmt(e2_btn['SEC_kWh_per_kg'])} kWh/kg \u2014 "
            f"{pct_diff(a_btn['SEC_kWh_per_kg'], e2_btn['SEC_kWh_per_kg'])} "
            f"below the Config A baseline \u2014 with a solar fraction of "
            f"{fmt(e2_btn['solar_fraction'], 2)} and COP of "
            f"{fmt(e2_btn['COP_mean'], 2)}. At Kathmandu, E2 achieves "
            f"{fmt(e2_ktm['SEC_kWh_per_kg'])} kWh/kg."
        ))

        add_para(doc, (
            f"E2 consistently outperforms E1 by "
            f"{pct_diff(e1_btn['SEC_kWh_per_kg'], e2_btn['SEC_kWh_per_kg'])} "
            f"at BTN, following the same D2 > D1 pattern: the warmer "
            f"evaporator source (cooled exhaust) improves COP."
        ))

    # ── 6.3 E3 Results ──
    add_heading(doc, "6.3  Config E3: Solar-Priority Control", level=2)

    headers = ["Config", "Location", "SEC (kWh/kg)", "Time (h)", "COP",
               "Solar frac."]
    rows = []
    for loc in LOCS:
        e3 = get_row(df, "E3", loc, "annual", 0.0, area=10.0)
        if e3 is not None:
            rows.append([
                "E3", LOC_SHORT[loc],
                fmt(e3["SEC_kWh_per_kg"]),
                fmt(e3["time_h"], 1),
                fmt(e3["COP_mean"], 2),
                fmt(e3["solar_fraction"], 2),
            ])
    add_table(doc, headers, rows,
              "Config E3 (A_c = 10 m\u00b2) annual-TMY performance.")

    e3_btn = get_row(df, "E3", "biratnagar", "annual", 0.0, area=10.0)
    e3_ktm = get_row(df, "E3", "kathmandu", "annual", 0.0, area=10.0)

    if e3_btn is not None and e3_ktm is not None and e2_btn is not None:
        add_para(doc, (
            f"E3 achieves SEC = {fmt(e3_btn['SEC_kWh_per_kg'])} at BTN "
            f"and {fmt(e3_ktm['SEC_kWh_per_kg'])} at KTM, with higher COP "
            f"({fmt(e3_btn['COP_mean'], 2)} and {fmt(e3_ktm['COP_mean'], 2)}) "
            f"than E1 or E2. The higher COP reflects the solar-priority "
            f"control: when the HP runs, it targets a lower T_cond (partial "
            f"lift), and when solar is abundant the HP is off entirely."
        ))

        add_para(doc, (
            f"However, E3 does not beat E2 in SEC at Biratnagar "
            f"({fmt(e3_btn['SEC_kWh_per_kg'])} vs {fmt(e2_btn['SEC_kWh_per_kg'])}). "
            f"The reason is solar clipping: when the solar collector is placed "
            f"after the condenser, it receives already-warm air (\u224840 \u00b0C), "
            f"and the collector\u2019s efficiency drops because thermal losses "
            f"scale with (T_in \u2212 T_amb). In E1/E2, the solar collector "
            f"receives HRX-heated air at \u224830 \u00b0C \u2014 a colder inlet that "
            f"preserves collector efficiency."
        ))

    # ── 6.4 Solar Area Sweep ──
    add_heading(doc, "6.4  Solar Collector Area Sweep (E2)", level=2)

    add_para(doc, (
        "To determine the practical collector size, a sweep from 2 to "
        "20 m\u00b2 was conducted for E2."
    ))

    headers = ["A_c (m\u00b2)", "KTM SEC", "KTM solar frac.",
               "BTN SEC", "BTN solar frac."]
    rows = []
    for area in [2.0, 4.0, 6.0, 8.0, 10.0, 15.0, 20.0]:
        e2k = get_row(df, "E2", "kathmandu", "annual", 0.0, area=area)
        e2b = get_row(df, "E2", "biratnagar", "annual", 0.0, area=area)
        if e2k is not None and e2b is not None:
            rows.append([
                f"{area:.0f}",
                fmt(e2k["SEC_kWh_per_kg"]),
                fmt(e2k["solar_fraction"], 2),
                fmt(e2b["SEC_kWh_per_kg"]),
                fmt(e2b["solar_fraction"], 2),
            ])
    add_table(doc, headers, rows,
              "E2 SEC and solar fraction vs collector area (annual TMY).")

    add_para(doc, (
        "Diminishing returns are evident. At Biratnagar, increasing A_c "
        "from 6 to 10 m\u00b2 reduces SEC by only 0.018 kWh/kg (from 0.147 "
        "to 0.129), while the first 6 m\u00b2 reduces SEC by 0.067 (from "
        "0.214 to 0.147). At Kathmandu, the marginal benefit of each "
        "additional m\u00b2 is larger due to lower baseline irradiance but "
        "still flattens beyond 10 m\u00b2."
    ))

    add_para(doc, (
        "For practical installations, 6\u201310 m\u00b2 provides the best "
        "cost\u2013performance balance. Larger collectors yield diminishing "
        "SEC reductions while increasing capital cost and roof area "
        "requirements."
    ))

    # ── 6.5 VPD Bypass for E ──
    add_heading(doc, "6.5  VPD Exhaust Bypass for E Configs", level=2)

    add_para(doc, (
        "The same VPD exhaust bypass strategy from Section 5.4 is applied "
        "to E1 and E2 (threshold = 0.05). E3 is excluded because its "
        "solar-priority control already manages HP operation dynamically."
    ))

    headers = ["Config", "Location", "SEC (no VPD)", "SEC (VPD 0.05)",
               "\u0394SEC", "Time (VPD)"]
    rows = []
    for loc in LOCS:
        for cfg_name in ["E1", "E2"]:
            e_off = get_row(df, cfg_name, loc, "annual", 0.0, vpd=False, area=10.0)
            e_on = get_row(df, cfg_name, loc, "annual", 0.0, vpd=True, area=10.0)
            if e_off is not None and e_on is not None:
                rows.append([
                    cfg_name, LOC_SHORT[loc],
                    fmt(e_off["SEC_kWh_per_kg"]),
                    fmt(e_on["SEC_kWh_per_kg"]),
                    pct_diff(e_off["SEC_kWh_per_kg"], e_on["SEC_kWh_per_kg"]),
                    fmt(e_on["time_h"], 1),
                ])
    add_table(doc, headers, rows,
              "E1 and E2 (10 m\u00b2) with VPD exhaust bypass (threshold = 0.05).")

    e2_vpd_btn = get_row(df, "E2", "biratnagar", "annual", 0.0, vpd=True, area=10.0)
    e2_vpd_ktm = get_row(df, "E2", "kathmandu", "annual", 0.0, vpd=True, area=10.0)

    if e2_vpd_btn is not None and e2_vpd_ktm is not None:
        add_para(doc, (
            f"VPD bypass delivers 25\u201331% SEC improvement for E1 and "
            f"21\u201327% for E2, with only 1\u20132 h extra drying time. "
            f"The best result is E2 + VPD at Biratnagar: "
            f"SEC = {fmt(e2_vpd_btn['SEC_kWh_per_kg'])} kWh/kg "
            f"({fmt(e2_vpd_btn['time_h'], 1)} h drying time). At Kathmandu, "
            f"E2 + VPD achieves {fmt(e2_vpd_ktm['SEC_kWh_per_kg'])} kWh/kg."
        ))

    add_para(doc, (
        "The optimal VPD threshold for E configs is 0.05 \u2014 lower than "
        "the 0.10\u20130.15 that minimises SEC for D configs (Section 5.4 "
        "sweep). This is because the solar collector already reduces "
        "the HP\u2019s energy share substantially; aggressive bypass pushes "
        "drying time up without proportional SEC gains. At 0.05, the "
        "system briefly enters bypass during late drying, captures "
        "the easy compressor savings, and returns to normal before "
        "the drying rate deteriorates significantly."
    ))

    add_para(doc, (
        "These results represent the culmination of the configuration "
        "hierarchy: starting from Config A\u2019s 0.543\u20130.717 kWh/kg, "
        "progressive addition of solar (B), heat recovery (D), their "
        "combination (E), and controls optimisation (VPD bypass) "
        "achieves 0.097\u20130.144 kWh/kg \u2014 an 80\u201387% reduction in "
        "specific energy consumption."
    ))


# ======================================================================
#  CHAPTER 7: GRAND SYNTHESIS AND RECOMMENDATIONS
# ======================================================================

def write_synthesis(doc, df):
    add_heading(doc, "7  Grand Synthesis and Recommendations", level=1)

    add_para(doc, (
        "Sections 1\u20136 analysed each configuration on its own terms. "
        "This section places all of them on a single axis so the "
        "trade-offs become directly visible, then converts the findings "
        "into practical recommendations."
    ))

    # ── 7.1 Grand comparison table ──
    add_heading(doc, "7.1  Complete Performance Comparison", level=2)

    add_para(doc, (
        "Table below ranks every configuration analysed in this study "
        "by specific energy consumption at the two production-scale "
        "locations (Kathmandu and Biratnagar). Each row uses the best "
        "operational setup explored for that configuration: A at r = 0 "
        "(open-loop), B/C/E at 10 m\u00b2 collector, and VPD variants at "
        "threshold = 0.05."
    ))

    # Build grand table
    configs_spec = [
        ("A (r = 0)",          "A",  0.0, 0.0,  False),
        ("A (r = 0.9)",        "A",  0.9, 0.0,  False),
        ("B (10 m\u00b2)",     "B",  0.0, 10.0, False),
        ("C1 (10 m\u00b2)",    "C1", 0.0, 10.0, False),
        ("C2 (10 m\u00b2)",    "C2", 0.0, 10.0, False),
        ("D1",                 "D1", 0.0, 0.0,  False),
        ("D2",                 "D2", 0.0, 0.0,  False),
        ("D3",                 "D3", 0.0, 0.0,  False),
        ("D1 + VPD",           "D1", 0.0, 0.0,  True),
        ("D2 + VPD",           "D2", 0.0, 0.0,  True),
        ("E1 (10 m\u00b2)",    "E1", 0.0, 10.0, False),
        ("E2 (10 m\u00b2)",    "E2", 0.0, 10.0, False),
        ("E3 (10 m\u00b2)",    "E3", 0.0, 10.0, False),
        ("E1 + VPD (10 m\u00b2)", "E1", 0.0, 10.0, True),
        ("E2 + VPD (10 m\u00b2)", "E2", 0.0, 10.0, True),
    ]

    a_ktm = get_row(df, "A", "kathmandu", "annual", 0.0)
    a_btn = get_row(df, "A", "biratnagar", "annual", 0.0)

    table_rows = []
    for label, cfg, r, area, vpd in configs_spec:
        k = get_row(df, cfg, "kathmandu", "annual", r, vpd=vpd, area=area)
        b = get_row(df, cfg, "biratnagar", "annual", r, vpd=vpd, area=area)
        if k is None or b is None:
            continue
        table_rows.append({
            "label": label,
            "sec_ktm": k["SEC_kWh_per_kg"],
            "sec_btn": b["SEC_kWh_per_kg"],
            "time_ktm": k["time_h"],
            "time_btn": b["time_h"],
            "cop_ktm": k["COP_mean"],
            "cop_btn": b["COP_mean"],
            "sf_ktm": k["solar_fraction"],
            "sf_btn": b["solar_fraction"],
        })

    # Sort by average SEC (best to worst)
    table_rows.sort(key=lambda x: (x["sec_ktm"] + x["sec_btn"]) / 2)

    headers = ["Rank", "Configuration", "KTM SEC", "BTN SEC",
               "Avg SEC", "KTM time", "BTN time",
               "\u0394 vs A (KTM)", "\u0394 vs A (BTN)"]
    rows = []
    for i, r in enumerate(table_rows, start=1):
        avg = (r["sec_ktm"] + r["sec_btn"]) / 2
        rows.append([
            str(i),
            r["label"],
            fmt(r["sec_ktm"]),
            fmt(r["sec_btn"]),
            fmt(avg),
            fmt(r["time_ktm"], 1),
            fmt(r["time_btn"], 1),
            pct_diff(a_ktm["SEC_kWh_per_kg"], r["sec_ktm"]) if a_ktm is not None else "\u2014",
            pct_diff(a_btn["SEC_kWh_per_kg"], r["sec_btn"]) if a_btn is not None else "\u2014",
        ])

    add_table(doc, headers, rows,
              "Grand comparison of all configurations, ranked from "
              "best to worst by average SEC (annual TMY, KTM + BTN).")

    # ── 7.2 Narrative ranking ──
    add_heading(doc, "7.2  Ranking Narrative", level=2)

    best = table_rows[0]
    worst = table_rows[-1]
    a_row = next((r for r in table_rows if r["label"] == "A (r = 0)"), None)

    if a_row is not None:
        add_para(doc, (
            f"The spread between worst and best is striking. The "
            f"baseline Config A (r = 0) consumes "
            f"{fmt(a_row['sec_ktm'])}\u2013{fmt(a_row['sec_btn'])} kWh/kg "
            f"depending on climate, while the best configuration \u2014 "
            f"{best['label']} \u2014 achieves "
            f"{fmt(best['sec_ktm'])}\u2013{fmt(best['sec_btn'])} kWh/kg, "
            f"a {pct_diff(a_row['sec_ktm'], best['sec_ktm'])}\u2013"
            f"{pct_diff(a_row['sec_btn'], best['sec_btn'])} reduction."
        ))

    add_para(doc, "Tier 1 \u2014 Best performers (SEC < 0.20 avg):", bold=True)
    add_para(doc, (
        "The top tier is dominated by Config E variants. E2 + VPD "
        "leads the ranking because it combines all three gains: "
        "(i) heat recovery reduces condenser load; (ii) solar "
        "preheating supplies a large fraction of the remaining load; "
        "(iii) VPD exhaust bypass cuts compressor work during late "
        "drying when the chamber air is no longer carrying moisture "
        "efficiently. E1 and E3 follow close behind."
    ))

    add_para(doc, "Tier 2 \u2014 Solid performers (SEC 0.20\u20130.35 avg):", bold=True)
    add_para(doc, (
        "Config D variants (with and without VPD) and some E variants "
        "at lower solar-fraction sites sit here. D1/D2 + VPD is "
        "notable: it matches solar-assisted Config B at Kathmandu "
        "without any solar hardware \u2014 heat recovery alone delivers "
        "roughly the same magnitude of savings as a 10 m\u00b2 collector "
        "at a moderate-irradiance site."
    ))

    add_para(doc, "Tier 3 \u2014 Mediocre performers (SEC 0.35\u20130.55 avg):", bold=True)
    add_para(doc, (
        "Config B (solar + HP series), D3 (swapped HRX), and C2 sit "
        "here. B is climate-dependent \u2014 excellent at BTN, average "
        "at KTM. D3 loses its evaporator-side COP advantage to the "
        "humidity penalty at the chamber inlet. C2 works but is "
        "strictly inferior to B."
    ))

    add_para(doc, "Tier 4 \u2014 Worst performers:", bold=True)
    add_para(doc, (
        "Config A and C1 are at the bottom. A is the minimum-capital "
        "reference; C1 fails entirely at r = 0 (drying stalls at low "
        "T_evap because the condenser cannot reach T_set from cool "
        "ambient alone). C1 is a cautionary tale rather than a "
        "candidate system."
    ))

    # ── 7.3 Practical recommendations ──
    add_heading(doc, "7.3  Practical Recommendations", level=2)

    add_para(doc, (
        "Configuration choice depends on site conditions, operational "
        "schedule, and capital budget. The following recommendations "
        "distil the study findings into concrete guidance."
    ))

    add_para(doc, "For high-irradiance sites (Biratnagar-like):", bold=True)
    add_bullet(doc,
        "E2 + VPD (10 m\u00b2): lowest SEC (0.097 kWh/kg), strong "
        "solar fraction (0.91), compact footprint. Preferred when "
        "roof area \u2265 10 m\u00b2 is available."
    )
    add_bullet(doc,
        "E2 (10 m\u00b2) without VPD: if control complexity is a "
        "concern; 0.129 kWh/kg, simpler operation."
    )
    add_bullet(doc,
        "B (10 m\u00b2): if HRX retrofit is not feasible; 0.287 "
        "kWh/kg, simplest solar-assisted layout."
    )

    add_para(doc, "For moderate-irradiance sites (Kathmandu-like):", bold=True)
    add_bullet(doc,
        "E2 + VPD (10 m\u00b2): still the best choice; 0.144 "
        "kWh/kg, SF \u2248 0.78."
    )
    add_bullet(doc,
        "D2 + VPD: if solar is unavailable or unreliable; 0.288 "
        "kWh/kg, fully weather-independent."
    )
    add_bullet(doc,
        "D1 + VPD: marginally higher SEC than D2 but simpler "
        "evaporator plumbing (no exhaust routing to evap)."
    )

    add_para(doc, "For continuous / night operation:", bold=True)
    add_bullet(doc,
        "D1 + VPD or D2 + VPD: solar contribution drops to zero at "
        "night; HRX provides constant, climate-independent energy "
        "recovery. Operating at night with E-series configurations "
        "collapses performance to D-series levels."
    )

    add_para(doc, "For minimum-capital retrofits:", bold=True)
    add_bullet(doc,
        "VPD bypass on Config A (r = 0.9): 7% SEC reduction at zero "
        "hardware cost \u2014 purely a controls change."
    )
    add_bullet(doc,
        "HRX alone (D1): the biggest single energy saving from one "
        "piece of hardware (\u224850% SEC reduction vs A at KTM). If "
        "only one retrofit is possible, this is it."
    )

    # ── 7.4 Design takeaways ──
    add_heading(doc, "7.4  Key Design Takeaways", level=2)

    add_para(doc, "1. Heat recovery outperforms solar in capital-efficiency.", bold=True)
    add_para(doc, (
        "At Kathmandu, a 10 m\u00b2 solar collector (Config B) brings SEC "
        "from 0.717 to 0.488. An HRX (Config D1) brings the same "
        "baseline down to 0.365 \u2014 a larger improvement from a "
        "simpler, cheaper device. Where both are feasible they are "
        "additive; where only one is, the HRX is the better first buy."
    ))

    add_para(doc, "2. Solar placement matters as much as solar area.", bold=True)
    add_para(doc, (
        "C1, C2, B, and E all use the same 10 m\u00b2 collector but "
        "achieve SEC from 0.894 down to 0.097 depending on where the "
        "collector sits in the air loop. Placing the collector on the "
        "cold-side (condenser inlet) preserves collector efficiency "
        "and maximises useful gain. Placing it on already-warm air "
        "(E3 after condenser, C1 before evaporator) is wasteful or "
        "counter-productive."
    ))

    add_para(doc, "3. Humidity at the chamber inlet dominates kinetics.", bold=True)
    add_para(doc, (
        "The Midilli \u03b1_RH = 1.75 humidity exponent means that any "
        "design choice raising chamber inlet RH has a direct, "
        "nonlinear drying-time penalty. This is why D3 (high COP but "
        "humid inlet) loses to D1/D2, and why the adaptive-T_evap "
        "investigation (Section 4) produced a negative result. For "
        "humidity-sensitive products, dehumidification is not a "
        "luxury \u2014 it is the primary driver of drying rate."
    ))

    add_para(doc, "4. VPD bypass is nearly-free optimisation.", bold=True)
    add_para(doc, (
        "Exhaust bypass at threshold = 0.05 delivers 15\u201331% SEC "
        "reduction across all compatible configurations with "
        "<10% additional drying time, and requires only a single "
        "damper and a humidity sensor. It is the cheapest "
        "efficiency intervention in the study."
    ))

    add_para(doc, "5. COP alone is an incomplete metric.", bold=True)
    add_para(doc, (
        "C2 achieves COP = 7.76 at BTN (highest in the study) but "
        "SEC = 0.381 kWh/kg \u2014 worse than E1\u2019s 0.141 kWh/kg with "
        "COP = 4.43. High COP wasted on redundant heating (because "
        "the solar collector has already raised the air to T_set) "
        "does not translate to low SEC. System-level SEC, not "
        "component-level COP, is the correct optimisation target."
    ))


# ======================================================================
#  CHAPTER 8: LIMITATIONS AND ASSUMPTIONS
# ======================================================================

def write_limitations(doc, df):
    add_heading(doc, "8  Limitations and Assumptions", level=1)

    add_para(doc, (
        "The results presented in this study are subject to several "
        "simplifying assumptions. Understanding their scope is "
        "essential for interpreting the SEC numbers and for framing "
        "follow-up experimental validation work."
    ))

    # ── 8.1 Weather ──
    add_heading(doc, "8.1  Weather and Climate Data", level=2)

    add_bullet(doc,
        "Single TMY (typical meteorological year) dataset per location. "
        "Year-to-year variability in irradiance, temperature, and "
        "humidity is not captured. Real-world SEC for a specific "
        "batch may differ depending on that year\u2019s weather."
    )
    add_bullet(doc,
        "Hourly resolution interpolated to the simulation timestep. "
        "Sub-hourly cloud transients (minutes-scale intermittency) "
        "are smoothed. The Hottel\u2013Whillier\u2013Bliss quasi-steady "
        "model assumes the collector reaches steady state within "
        "each timestep; in reality, thermal inertia of the absorber "
        "and header damps short fluctuations."
    )
    add_bullet(doc,
        "Only three Nepal locations analysed (KTM, BTN, TPJ). "
        "Generalisation to sites with substantially different "
        "climates (e.g., coastal humid tropics, arid continental) "
        "requires re-running with local TMY data."
    )

    # ── 8.2 Thermal and fluid mechanics ──
    add_heading(doc, "8.2  Thermal and Flow Simplifications", level=2)

    add_bullet(doc,
        "No duct or chamber wall losses. All heat delivered by the "
        "condenser is assumed to reach the product. In practice, "
        "insulation quality, duct runs, and chamber wall thermal "
        "mass would absorb 5\u201315% of delivered heat."
    )
    add_bullet(doc,
        "Constant fan power. The fan is modelled as a fixed "
        "consumption term; real fans have load-dependent efficiency "
        "and flow-rate coupling to duct pressure drop, which scales "
        "with air path complexity (higher in D/E with HRX)."
    )
    add_bullet(doc,
        "No startup transients or thermal mass of the dryer "
        "structure. Each timestep is treated as steady-state. "
        "For batch drying of several hours, startup transients "
        "are a minor (<5%) contribution but not zero."
    )
    add_bullet(doc,
        "HRX modelled as a fixed-effectiveness (\u03b5 = 0.70) counter-flow "
        "exchanger. Fouling, condensation-induced performance "
        "degradation, and frost on the cold side are not modelled."
    )
    add_bullet(doc,
        "No condensation-to-liquid mass tracking on the HRX cold "
        "stream. When exhaust air cools below its dew point in the "
        "HRX, latent heat release is accounted for energetically, "
        "but the dropout water is not tracked as a separate output."
    )

    # ── 8.3 Product and kinetics ──
    add_heading(doc, "8.3  Product and Kinetics", level=2)

    add_bullet(doc,
        "Fixed apple slice geometry (6 mm thickness). Thinner slices "
        "would dry faster with different K/Ea values; thicker slices "
        "would be diffusion-limited and require a shrinkage model. "
        "Cross-product generalisation (e.g., to cardamom, chilli, "
        "ginger) requires re-fitting the Midilli kinetic parameters."
    )
    add_bullet(doc,
        "Uniform tray loading assumed. No spatial gradients within "
        "trays or between trays are modelled; all ten trays are "
        "treated as receiving identical air conditions."
    )
    add_bullet(doc,
        "Midilli kinetics (K_ref = 1.63 \u00d7 10\u207b\u2074 /s, Ea/R = 2711 K, "
        "\u03b1_RH = 1.75) fitted from phase-2 experimental data with "
        "R\u00b2 = 0.90. The 10% unexplained variance propagates into "
        "the drying-time predictions."
    )
    add_bullet(doc,
        "No product quality metrics (colour, vitamin retention, "
        "rehydration capacity) are simulated. SEC optimisation is "
        "purely an energy-efficiency exercise; thermal damage from "
        "prolonged drying or humid chambers is not penalised "
        "beyond the kinetic slowdown."
    )

    # ── 8.4 Heat pump ──
    add_heading(doc, "8.4  Heat Pump Cycle", level=2)

    add_bullet(doc,
        "Fixed compressor isentropic efficiency (\u03b7_is = 0.75). "
        "Real scroll/rotary compressors have efficiency curves that "
        "vary with pressure ratio; at high T_cond or low T_evap, "
        "\u03b7_is typically drops to 0.65\u20130.70."
    )
    add_bullet(doc,
        "R134a refrigerant assumed throughout. Other refrigerants "
        "(R290, R744, R1234yf) would shift optimal operating "
        "envelopes and might change configuration rankings at the "
        "extremes of the operating window."
    )
    add_bullet(doc,
        "Variable-speed compressor with infinite turndown assumed. "
        "Real systems have minimum stable speeds and may cycle on/off, "
        "incurring cycling losses that are not captured here."
    )
    add_bullet(doc,
        "Fixed evaporator and condenser effectiveness "
        "(\u03b5_evap, \u03b5_cond) with a constant approach temperature. "
        "Real coil behaviour depends on air-side fouling, refrigerant "
        "charge, and partial condensation at the evaporator."
    )

    # ── 8.5 Economics ──
    add_heading(doc, "8.5  What the Model Does Not Address", level=2)

    add_bullet(doc,
        "Capital cost is not quantified. The rankings above assume "
        "hardware exists; practical decisions must weigh the capital "
        "cost of a 10 m\u00b2 solar collector or an HRX against the SEC "
        "reduction achieved and the electricity tariff."
    )
    add_bullet(doc,
        "Maintenance burden is not modelled. HRX and solar collectors "
        "require periodic cleaning; compressor service intervals "
        "scale with operating hours."
    )
    add_bullet(doc,
        "Control system realism. The VPD bypass, solar-priority "
        "control (E3), and dynamic evaporator mixing (D2/E2) are "
        "modelled as idealised state-machines; real deployments "
        "would face sensor noise, actuation delays, and mode-"
        "chattering that the 600 s hysteresis dwell partially "
        "but not fully mitigates."
    )

    add_para(doc, (
        "These limitations do not change the qualitative conclusions "
        "of the study \u2014 the performance hierarchy E > D > B > C > A "
        "is robust to all of them. They do, however, mean the "
        "absolute SEC numbers should be interpreted as "
        "best-case engineering estimates. Validation against a "
        "physical prototype is the logical next step and would "
        "primarily refine: (a) HRX effectiveness under realistic "
        "condensation, (b) duct losses for the larger E-series air "
        "paths, and (c) compressor efficiency curves across the "
        "operating envelope."
    ))


# ======================================================================
#  MAIN
# ======================================================================

def main():
    df = load_data()
    doc = Document()

    # ── Chapter title ──
    add_heading(doc, "Results and Discussion", level=0)

    add_para(doc, (
        "This chapter presents the simulation results for ten SAHPD "
        "configurations of increasing complexity. Config A establishes "
        "the heat-pump-only baseline. Config B adds a flat-plate solar "
        "collector in series. Config C explores alternative solar "
        "placements that reveal fundamental design principles. An "
        "evaporator temperature investigation and VPD-based bypass "
        "strategy optimise closed-loop performance. Config D introduces "
        "heat recovery, and Config E combines heat recovery with solar "
        "integration to achieve the lowest energy consumption. Each "
        "section builds on the insights of the previous, tracing a "
        "path from the simplest dryer to the most efficient system. "
        "Section 7 synthesises all findings into a unified ranking and "
        "practical recommendations, and Section 8 documents the "
        "assumptions and limitations that bound the conclusions."
    ))

    write_config_A(doc, df)
    write_config_B(doc, df)
    write_config_C(doc, df)
    write_optimisation(doc, df)
    write_config_D(doc, df)
    write_config_E(doc, df)
    write_synthesis(doc, df)
    write_limitations(doc, df)

    # Save
    out_dir = PROJECT_ROOT / "thesis"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / "Results_Discussion.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")

    for cfg in ["A", "B", "C1", "C2", "D1", "D2", "D3", "E1", "E2", "E3"]:
        n = len(df[df["config"] == cfg])
        print(f"  {cfg}: {n} rows")


if __name__ == "__main__":
    main()
