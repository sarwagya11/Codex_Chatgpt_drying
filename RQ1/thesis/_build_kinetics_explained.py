"""One-shot writer for Kinetics_Explained.docx in the thesis folder."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def P(doc, text, bold=False, mono=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold: r.bold = True
    if italic: r.italic = True
    if mono:
        r.font.name = "Consolas"
        r.font.size = Pt(10)
    return p


def CODE(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    return p


def TABLE(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]
        c.text = h
        for r in c.paragraphs[0].runs:
            r.bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.rows[i].cells[j].text = str(v)
    return t


doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ---------------------------------------------------------------- TITLE
title = doc.add_heading("Kinetics Explained", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub.add_run("M1 drying-kinetics model: data, fit, and runtime evaluation in the SAHPD RQ1 simulator")
sub_run.italic = True

# ---------------------------------------------------------------- 1 OVERVIEW
H(doc, "1. Overview", 1)
P(doc,
  "The kinetics module has two distinct phases. A one-time fit happens once per Python "
  "process, taking 13 thin-layer experimental drying curves and condensing them into "
  "five physical parameters. A per-step evaluation then plugs the live air state at "
  "each tray into those cached parameters to compute the rate constant K, which drives "
  "the first-order moisture update. The fit runs once. The evaluation runs ~10^5 times "
  "per simulation.")

# ---------------------------------------------------------------- 2 RAW DATA
H(doc, "2. The 13 experimental curves", 1)
P(doc,
  "Each row of outputs/phase2/phase2_targets.csv lists one thin-layer drying run at a "
  "fixed operating point (T, v, thickness, RH). Together they span four physical axes:")

TABLE(doc,
      ["dataset", "T (°C)", "v (m/s)", "d (mm)", "RH_mid (%)", "axis"],
      [
          ("T_40_v1p1",              40, 1.1,  6, 42.5, "T sweep"),
          ("T_45_v1p1",              45, 1.1,  6, 42.5, "T sweep"),
          ("T_50_v1p1_tempsweep",    50, 1.1,  6, 42.5, "T sweep"),
          ("T50_v0p60",              50, 0.6,  6, 42.5, "v sweep"),
          ("T50_v0p85",              50, 0.85, 6, 42.5, "v sweep"),
          ("T50_v1p1_velsweep",      50, 1.1,  6, 42.5, "v sweep"),
          ("T50_v1p1_t_4mm",         50, 1.1,  4, 42.5, "d sweep"),
          ("T50_v1p1_t_6mm",         50, 1.1,  6, 42.5, "d sweep"),
          ("T50_v1p1_t_8mm",         50, 1.1,  8, 42.5, "d sweep"),
          ("T50_v1p1_t_10mm",        50, 1.1, 10, 42.5, "d sweep"),
          ("T50V1p1_RH_25-28%",      50, 1.1,  6, 26.5, "RH sweep"),
          ("T50V1p1_RH_35-38%",      50, 1.1,  6, 36.5, "RH sweep"),
          ("T50V1p1_RH_40-45%",      50, 1.1,  6, 42.5, "RH sweep"),
      ])
P(doc,
  "Reference state (T_ref = 50 °C, v_ref = 1.1 m/s, d_ref = 6 mm) is the intersection "
  "point that appears in every sub-sweep. Total observations across the 13 curves: 1,573 "
  "rows of (time_min, X_db).")
P(doc, "A raw row looks like:")
CODE(doc,
     "time_min, X_db\n"
     "1.92,     6.497\n"
     "5.20,     6.407\n"
     "9.77,     6.326\n"
     "13.70,    6.236")

# ---------------------------------------------------------------- 3 PAVA
H(doc, "3. PAVA isotonic smoothing", 1)
P(doc,
  "For each curve the simulator first computes MR_raw(t) = X_db(t) / X_db(0) and clips "
  "it to [0, 1.1]. Sensor noise (dust, weighing wobble) can produce small up-ticks where "
  "MR locally rises. The first-order model requires MR to be monotone decreasing, so the "
  "Pool-Adjacent-Violators Algorithm (PAVA) pools any non-monotonic blip with its "
  "neighbours and replaces them with their weighted average until the sequence is "
  "monotone. The averages are preserved; the noise is removed.")
P(doc, "Implementation: kinetics.py:_pava_monotone (lines 119–136).")

# ---------------------------------------------------------------- 4 MODEL
H(doc, "4. The M1 model", 1)
P(doc, "M1 is a five-parameter, first-order Arrhenius form with explicit RH, velocity, and thickness corrections:")
CODE(doc,
     "K(T, RH, v, d) = K_ref · exp((Ea/R) · (1/T_ref − 1/T))\n"
     "                       · exp(−α_RH · RH/100)\n"
     "                       · (v / v_ref)^γ_v\n"
     "                       · (d_ref / d)^δ_d\n"
     "\n"
     "MR(t) = exp(−K · t_seconds)")
P(doc,
  "Five unknowns: K_ref, Ea/R, α_RH, γ_v, δ_d. Everything else (T, v, d, RH, t) is "
  "measured. The reference constants T_ref = 50 °C, v_ref = 1.1 m/s, d_ref = 6 mm are "
  "hard-coded in kinetics.py:111–113.")

P(doc, "Concrete row example. Take T=45 °C, v=1.1, d=6, RH=42.5% at t = 5.20 min:")
CODE(doc,
     "v/v_ref = 1 → velocity term vanishes\n"
     "d_ref/d = 1 → thickness term vanishes\n"
     "\n"
     "K_45 = K_ref · exp((Ea/R) · (1/323.15 − 1/318.15)) · exp(−α_RH · 0.425)\n"
     "     = K_ref · exp(−4.86e-5 · (Ea/R))            · exp(−0.425·α_RH)\n"
     "\n"
     "MR_model = exp(−K_45 · 312)        # 5.20 min = 312 s\n"
     "MR_obs   = 6.407 / 6.497 = 0.986\n"
     "\n"
     "residual = MR_model − MR_obs       (one number per row)")
P(doc,
  "Five unknowns, one residual per measurement, 1,573 measurements across all 13 curves.")

# ---------------------------------------------------------------- 5 NLS
H(doc, "5. The curve fit (nonlinear least squares)", 1)
P(doc,
  "scipy.optimize.least_squares with method 'trf' (trust-region reflective). The five "
  "parameters are stacked into a vector p = (ln K_ref, Ea/R, α_RH, γ_v, δ_d). The "
  "initial guess is _M1_NOM = (ln 1.9e-4, 2711, 1.75, 0.44, 0.66) and the parameters "
  "live inside the box [_M1_LO, _M1_HI] defined in kinetics.py:114–116.")
P(doc, "Each iteration:")
P(doc,
  "1. Pick current p.\n"
  "2. For every observation across all 13 curves, compute r_i(p) = MR_model_i(p) − MR_obs_i.\n"
  "3. Concatenate all residuals into one vector R(p) of length 1,573.\n"
  "4. Compute the Jacobian J = ∂R/∂p (numerical by default).\n"
  "5. Solve the Gauss-Newton step δp = (JᵀJ)⁻¹ JᵀR, project onto the box, scale by a trust-region radius.\n"
  "6. Apply, recompute R(p), grow or shrink the trust region depending on whether the step helped.\n"
  "7. Stop when xtol and ftol (both 1e-12) say further moves do not reduce ‖R‖² meaningfully.")
P(doc,
  "Critical detail: each curve does not get its own K. The same five-parameter formula "
  "must explain every curve simultaneously, which is what forces the parameters to take "
  "physical meaning. Fitting one curve at a time would give 13 independent K values; "
  "this fit gives one function K(T, RH, v, d) that reproduces all 13.")

# ---------------------------------------------------------------- 6 RESULT
H(doc, "6. What the fit discovered", 1)
TABLE(doc,
      ["Parameter", "Value", "Meaning"],
      [
          ("K_ref",    "2.097 × 10⁻⁴ s⁻¹", "Rate at ref state (50 °C, 1.1 m/s, 6 mm, RH=0)"),
          ("Ea/R",     "3,738 K",          "Ea = 31.08 kJ/mol; T sensitivity"),
          ("α_RH",     "1.965",            "Drying ×0.14 when RH=1.0 vs RH=0"),
          ("γ_v",      "0.401",            "Drying ×(v/1.1)^0.4 with velocity"),
          ("δ_d",      "0.589",            "Drying ×(6/d)^0.59 (thinner = faster)"),
          ("RMSE_MR",  "0.04685",          "Mean prediction error across 1,573 points"),
      ])
P(doc, "Physical reading of each parameter:", bold=True)
P(doc,
  "• Characteristic time at the reference state: τ = 1/K_ref ≈ 79 min for a thin layer.\n"
  "• +10 °C (40 → 50): K speeds up by exp(3738 · (1/313.15 − 1/323.15)) = exp(0.369) ≈ 1.45×.\n"
  "• RH 20% → 80%: K drops by exp(−1.965 · 0.6) = 0.308 → 3.2× slower.\n"
  "• Doubling velocity: 2^0.401 ≈ 1.32× faster (modest; above 0.5 m/s the boundary layer is already thin).\n"
  "• Halving thickness (6 → 3 mm): 2^0.589 ≈ 1.5× faster (shorter internal diffusion path).\n"
  "• RMSE_MR = 0.047 over a 0–1 scale = 5% mean error; acceptable for five parameters explaining 1,573 measurements across four physical axes.")

# ---------------------------------------------------------------- 7 CONFIDENCE
H(doc, "7. Confidence in the fit", 1)
P(doc,
  "The fit also reports σ² = SS_res / (n_obs − n_par) = 2.20 × 10⁻³. From the profile "
  "likelihood:")
P(doc,
  "• 95% CI for Ea: [28.56, 33.60] kJ/mol — overlaps the published apple-drying range "
  "(19.96 – 30.93 kJ/mol from Aktaş, Mishra, et al.). The values are physically plausible, "
  "not curve-fitting artefacts.\n"
  "• Leave-one-curve-out cross-validation: RMSE_MR = 0.0528 on held-out curves vs 0.0469 "
  "on the training set. ~13 % degradation — no overfitting collapse.")

# ---------------------------------------------------------------- 8 RUNTIME
H(doc, "8. How the simulator uses K, every step", 1)
P(doc,
  "At each tray, every 60 s of sim time, with the current inlet air state (T_air, RH_air) "
  "and the cached parameters:")
CODE(doc,
     "ln K = ln(2.097e-4)\n"
     "     + 3738 · (1/323.15 − 1/T_K)        # T term\n"
     "     − 1.965 · RH_frac                  # RH term\n"
     "     + 0.401 · ln(1.1 / 1.1)            # = 0 (v matches v_ref)\n"
     "     + 0.589 · ln(6 / 6)                # = 0 (d matches d_ref)\n"
     "K  = max(K_min, exp(ln K))\n"
     "\n"
     "ΔX = K · (X − X_eq) · 60                # X_eq = 0 since GAB is off\n"
     "X_new = max(X − ΔX, X_eq)\n"
     "dm_w_kinetic = m_p_dry · ΔX             # kg water removed by kinetics")

P(doc, "Worked examples at the two extremes seen in the SAHPD chamber:", bold=True)

P(doc, "Tray 0 inlet, T=45 °C, RH=0.30:", bold=True)
CODE(doc,
     "K = 2.097e-4 · exp(3738·(1/323.15 − 1/318.15)) · exp(−1.965·0.30)\n"
     "  = 2.097e-4 · 0.834 · 0.555\n"
     "  = 9.7 × 10⁻⁵ s⁻¹\n"
     "\n"
     "At X = 6.5, dX/dt ≈ 6.3 × 10⁻⁴ kg/(kg·s)\n"
     "Per 60-s step: ΔX ≈ 0.038")

P(doc, "Tray 9 inlet, T=29 °C, RH=1.0 (the saturated front in the 5 cm geometry):", bold=True)
CODE(doc,
     "K = 2.097e-4 · exp(3738·(1/323.15 − 1/302.15)) · exp(−1.965·1.0)\n"
     "  = 2.097e-4 · 0.405 · 0.140\n"
     "  = 1.19 × 10⁻⁵ s⁻¹    (8× weaker than tray 0)")
P(doc,
  "Kinetic ΔX would be 0.0046 per step, but the air-capacity gate (compute_dm_w_air_capacity) "
  "returns 0 because the air is already at RH=1.0 and cannot accept any more moisture. The "
  "tray freezes until upstream trays dry enough that fresher air reaches it. That is the "
  "drying-front behaviour seen in the 5 cm / 0.9 m/s smoke run.")

# ---------------------------------------------------------------- 9 GATE
H(doc, "9. The min(kinetic, air-capacity) gate", 1)
P(doc,
  "Per-step water removed is min(kinetic, air-capacity). The air-capacity gate "
  "(kinetics.py:compute_dm_w_air_capacity) bisects on Δω until the resulting outlet "
  "RH equals RH_out_max_frac (default ≈ 0.85). It returns m_da · Δω_max · dt. When the "
  "incoming air is already saturated, Δω_max = 0 and the kinetic K becomes irrelevant. "
  "This is the physical reason the propagating drying front exists in deep beds.")

# ---------------------------------------------------------------- 10 STORAGE
H(doc, "10. Variable map", 1)
TABLE(doc,
      ["Variable", "Lives in", "Set when", "Used by"],
      [
          ("_PARAMETRIC_PARAMS",                       "kinetics.py module global", "First sim call", "Every later keff_from_state"),
          ("_M1_T_REF_C=50, _M1_V_REF=1.1, _M1_D_REF=6", "kinetics.py:111–113",       "Source constant", "Both fit and evaluation"),
          ("cfg.v_ms, cfg.thickness_mm",               "KineticsConfig",            "At sim startup",  "keff_from_state for v/d terms"),
          ("cfg.RH_out_max_frac",                      "KineticsConfig",            "At sim startup",  "Air-capacity gate (≈0.85)"),
          ("X_eq_local per tray",                      "runtime loop",              "Per step",        "First-order update (=0 since GAB removed)"),
      ])

# ---------------------------------------------------------------- 11 SUMMARY
H(doc, "11. End-to-end summary", 1)
P(doc,
  "13 raw curves  →  PAVA cleaning  →  five-parameter NLS over all 1,573 points  →  "
  "cached parameters in module global  →  per-step K(T, RH) lookup at each tray  →  "
  "first-order moisture update  →  air-capacity gate  →  tray ΔX. "
  "Fit runs once per Python process. Evaluation runs ~10⁵ times per simulation but is "
  "just exp/log arithmetic on the cached parameters, so it is essentially free.")

# ----- save
out = Path(__file__).parent / "Kinetics_Explained.docx"
doc.save(out)
print(f"Saved: {out}")
