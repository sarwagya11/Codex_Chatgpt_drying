"""Build VPD Bypass Mechanism explanation docx."""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path("VPD_Bypass_Mechanism.docx")
IMG = Path("plots/_audit/step4_vpd_sweep.png")

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("VPD Bypass Mechanism")
run.bold = True
run.font.size = Pt(18)

doc.add_paragraph()

# ------- Paragraph 1 -------
h = doc.add_paragraph()
r = h.add_run("1. What VPD is and its equation")
r.bold = True
r.font.size = Pt(13)

doc.add_paragraph(
    "Vapour pressure deficit (VPD) is a measure of how \"thirsty\" air is, i.e. how much more water "
    "vapour it could hold before becoming saturated. For moist air at temperature T, the saturation "
    "vapour pressure P_sat(T) sets the maximum partial pressure of water vapour that the air can "
    "sustain; the actual partial pressure of water vapour in the air is p_v. The deficit between "
    "these two is the VPD:"
)
eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq.add_run("VPD(T) = P_sat(T) − p_v        [Pa]").italic = True

doc.add_paragraph(
    "A high VPD means the air can absorb a lot of moisture and dry product efficiently; a VPD of zero "
    "means the air is saturated and cannot dry anything. In our simulator we evaluate the partial "
    "vapour pressure from the humidity ratio ω using"
)
eq2 = doc.add_paragraph()
eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq2.add_run("p_v = ω · P_atm / (0.622 + ω)").italic = True

doc.add_paragraph(
    "so a single humidity-ratio value (which is what the chamber model tracks) is enough to compute "
    "the VPD at any chosen reference temperature. For apple drying at our setpoint of 45 °C, "
    "P_sat(45 °C) ≈ 9.59 kPa, and that is the reference saturation pressure we use throughout the "
    "bypass logic."
)

# ------- Paragraph 2 -------
h = doc.add_paragraph()
r = h.add_run("2. Why we toggle the VPD bypass, and what \"meaningful\" means here")
r.bold = True
r.font.size = Pt(13)

doc.add_paragraph(
    "We use the VPD signal to decide, in real time, whether the air leaving the drying chamber still "
    "has enough remaining thirst to justify the full air-conditioning chain (heat-recovery exchanger, "
    "evaporator, condenser). To make this decision in a fair way, we compare two VPDs at the same "
    "reference temperature (the setpoint, 45 °C) so that the comparison isolates the humidity loading "
    "of the air and removes the trivial effect of temperature on saturation pressure:"
)
for line in [
    "VPD_in  = P_sat(T_set) − p_v(ω_ambient)     (how thirsty fresh ambient air would be if heated to 45 °C)",
    "VPD_exh = P_sat(T_set) − p_v(ω_exhaust)     (how thirsty the chamber exhaust is at 45 °C)",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(line).italic = True

doc.add_paragraph(
    "From these we form a single dimensionless number, the VPD utilization:"
)
eq3 = doc.add_paragraph()
eq3.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq3.add_run("utilization = (VPD_in − VPD_exh) / VPD_in  =  1 − VPD_exh / VPD_in        [0, 1]").italic = True

doc.add_paragraph(
    "Interpretation: utilization ≈ 1 means the exhaust is nearly saturated, i.e. the chamber pulled "
    "almost all of the available drying potential out of the air in one pass. This corresponds to the "
    "early, constant-rate drying phase where the apple surface is wet and the bottleneck is on the "
    "air side. Utilization ≈ 0 means the exhaust looks almost identical to the fresh inlet, i.e. the "
    "air went through the chamber barely touched; this is the late, falling-rate phase where moisture "
    "is locked inside the apple and the bottleneck is internal diffusion, not air thirst."
)
doc.add_paragraph(
    "Important nuance: utilization is a single-pass figure, not a system-wide \"useful work\" number. "
    "Even during the constant-rate phase utilization is only about 0.15 in our geometry, because the "
    "chamber is not a saturator (the air spends only a few seconds in contact with the apples, the "
    "skin offers diffusion resistance even when wet, and the tray-stack residence time is short). "
    "What matters is the relative drop: 0.15 in the early phase versus 0.05 in the late phase is the "
    "signal that the bottleneck has moved from the air side to inside the apple. In the falling-rate "
    "regime, paying full compressor and fan energy to heat and dehumidify fresh air is wasteful "
    "because the chamber cannot use that thirst quickly enough. The bypass logic is the controller's "
    "way of recognising that shift and reducing the energy spend accordingly."
)

# ------- Paragraph 3 -------
h = doc.add_paragraph()
r = h.add_run("3. How the air path changes, with a real example from Config E2 at Biratnagar")
r.bold = True
r.font.size = Pt(13)

doc.add_paragraph(
    "The air follows two different paths depending on whether the VPD bypass is OFF or ON. I drew "
    "both below using only the component names so the topology change is easy to see."
)

doc.add_paragraph().add_run("Bypass OFF (normal operation):").bold = True
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.25)
p.add_run("Ambient → HRX → Solar → Condenser → Chamber → Exhaust → HRX → out").font.name = "Consolas"

doc.add_paragraph().add_run("Bypass ON (for E2):").bold = True
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.25)
p.add_run("Exhaust → Solar → Condenser → Chamber → Exhaust (recirculates)").font.name = "Consolas"

doc.add_paragraph(
    "For E3 the bypass path is Exhaust → Condenser → Solar → Chamber because in E3 the collector "
    "sits downstream of the condenser. For D1 and D2 there is no solar collector at all, so the "
    "bypass simplifies to Exhaust → Condenser → Chamber. In every case the HRX is taken out of the "
    "loop and no fresh outside air is drawn in. The condenser only has to lift the air by a small "
    "amount because the exhaust is already near 40 °C and the solar collector pushes it further, so "
    "the compressor does very little work in this mode."
)
doc.add_paragraph(
    "To make this concrete I picked one full simulation run from Biratnagar (Config E2, 10 m² solar "
    "collector, threshold 0.05) and pulled the per-timestep values straight from the CSV. The vapour "
    "pressure deficits below are both evaluated at 45 °C so the inlet and the exhaust are directly "
    "comparable."
)

# Table 1
hdr = ["Time", "Drying phase", "MR (avg)", "ω_inlet [kg/kg]", "ω_exhaust [kg/kg]",
       "Δω [g/kg]", "VPD_in [Pa]", "VPD_exh [Pa]", "Utilization", "Bypass"]
rows = [
    ("1.00 h",  "Early, surface-wet apples",          "0.82", "0.00846", "0.01632", "+7.9", "8237", "7012", "0.148", "OFF"),
    ("5.00 h",  "Transition",                         "0.29", "0.00835", "0.01293", "+4.6", "8237", "7522", "0.087", "OFF"),
    ("7.05 h",  "Just before flip",                   "0.14", "0.00828", "0.01092", "+2.6", "8262", "7849", "0.050", "OFF"),
    ("7.07 h",  "Bypass turns ON",                    "0.14", "0.00829", "0.01131", "+3.0", "8261", "7787", "0.050", "ON"),
    ("7.33 h",  "Bypass turns OFF (humidity recov.)", "0.13", "0.00835", "0.01551", "+7.2", "8237", "7156", "0.153", "OFF"),
    ("13.33 h", "Deep falling-rate",                  "0.01", "0.00906", "0.01113", "+2.1", "8140", "7817", "0.046", "OFF"),
]
tbl = doc.add_table(rows=1+len(rows), cols=len(hdr))
tbl.style = "Light Grid Accent 1"
for j, h_ in enumerate(hdr):
    cell = tbl.rows[0].cells[j]
    cell.text = h_
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)
for i, row in enumerate(rows, start=1):
    for j, v in enumerate(row):
        cell = tbl.rows[i].cells[j]
        cell.text = v
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)

doc.add_paragraph(
    "Reading the table from top to bottom: in the first hour the apples are still surface-wet, so air "
    "entering the chamber at ω = 0.00846 kg/kg leaves at ω = 0.01632 kg/kg, picking up nearly 8 grams "
    "of extra water for every kilogram of dry air; the VPD drops from 8237 Pa to 7012 Pa, which gives "
    "utilization 0.148. Five hours in the chamber only loads the air with +4.6 g/kg, and by hour 7 "
    "only +2.6 g/kg, at which point utilization drops to 0.050. That is exactly the threshold and the "
    "600-second dwell timer has expired, so the controller switches bypass on. With no fresh-air "
    "dilution the exhaust accumulates moisture quickly: 15 minutes later the air is picking up +7.2 "
    "g/kg again and utilization is back up to 0.153, above the 3× hysteresis line, so the controller "
    "flips bypass off and runs a short burst of fresh-air drying to dump the trapped humidity out of "
    "the system. From hour 7 to the end of the run the system oscillates between these two states; "
    "about 48 % of the remaining timesteps are spent in bypass mode, which is the pattern we want: "
    "short fresh-air bursts to expel moisture, separated by longer low-energy recirculating intervals "
    "while the apples diffuse their internal moisture to the surface."
)

# ------- Paragraph 4 -------
h = doc.add_paragraph()
r = h.add_run("4. What the simulator does when bypass is ON, and why we picked the 0.05 threshold")
r.bold = True
r.font.size = Pt(13)

doc.add_paragraph(
    "When the bypass flag is ON, the simulator changes three things at the start of the timestep "
    "before any heat pump or chamber calculation happens. First, the air entering the condenser is "
    "set to the previous exhaust state (T_exhaust_prev, ω_exhaust_prev) instead of the HRX-heated "
    "ambient, so there is no fresh-air dilution and the recirculation loop is closed. Second, the "
    "heat-recovery exchanger is removed from the energy accounting (Q_HRX_cum_kWh only accrues when "
    "bypass is inactive), which prevents us from double-counting heat that is not actually being "
    "exchanged. Third, the heat-pump condenser is sized to lift this warm exhaust (about 39–40 °C in "
    "our run) up to the 45 °C setpoint instead of lifting cool ambient air the whole way, so the "
    "compressor work per timestep collapses to a small fraction of normal. The solar collector still "
    "runs (E2 and E1 keep solar in the bypass path), so any irradiance available continues to be "
    "used. The chamber model is then called exactly as in normal operation and the bypass decision "
    "for the next timestep is re-evaluated from the freshly updated exhaust state."
)
doc.add_paragraph(
    "The threshold itself is the only number that controls how aggressive the bypass is. A small "
    "threshold means the controller waits until the air is very under-utilised before flipping, so "
    "the system spends most of the run in normal fresh-air mode; a large threshold means the "
    "controller switches to bypass almost immediately, so the system spends most of the run "
    "recirculating. We ran a sweep on the headline case (Config E2, Biratnagar, 10 m² solar "
    "collector) at five threshold values and pulled the results from the per-run CSVs."
)

# Table 2
hdr2 = ["VPD threshold", "Drying time", "W_compressor [kWh]", "W_fan [kWh]", "Water removed [kg]", "SEC [kWh/kg]", "Time in bypass"]
rows2 = [
    ("0.02", "15.1 h", "1.521", "0.376", "19.32", "0.0982", "29.8 %"),
    ("0.05", "16.2 h", "1.450", "0.403", "19.31", "0.0959", "48.4 %"),
    ("0.10", "20.6 h", "1.486", "0.511", "19.26", "0.1037", "70.8 %"),
    ("0.15", "23.4 h", "1.497", "0.581", "19.27", "0.1078", "83.8 %"),
    ("0.20", "48.0 h", "1.402", "1.192", "19.15", "0.1354", "95.0 %"),
]
tbl2 = doc.add_table(rows=1+len(rows2), cols=len(hdr2))
tbl2.style = "Light Grid Accent 1"
for j, h_ in enumerate(hdr2):
    cell = tbl2.rows[0].cells[j]
    cell.text = h_
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)
for i, row in enumerate(rows2, start=1):
    for j, v in enumerate(row):
        cell = tbl2.rows[i].cells[j]
        cell.text = v
        # bold the recommended row
        if row[0] == "0.05":
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
        else:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

doc.add_paragraph(
    "Reading down the columns reveals the tradeoff. At a very strict threshold (0.02) the bypass "
    "almost never engages, so the dryer runs at full compressor power even in the late falling-rate "
    "phase; it finishes the batch quickly (15.1 hours) but spends 1.521 kWh of compressor work to do "
    "it, which gives SEC 0.0982 kWh/kg. At our chosen threshold (0.05) the controller accepts bypass "
    "for nearly half the run; compressor work drops to 1.450 kWh, drying time only grows by about "
    "an hour, and SEC reaches its minimum at 0.0959 kWh/kg. From there on the curve goes the wrong "
    "way. At 0.10 the bypass is on 71 % of the time; the compressor saving is already nearly "
    "exhausted but the fans now run for 20.6 hours instead of 16, so fan energy climbs from 0.403 "
    "to 0.511 kWh and SEC rises to 0.1037. At 0.20 the controller is in bypass almost all the time, "
    "the chamber barely sees fresh dry air, drying time blows out to 48 hours, fan energy nearly "
    "triples (1.192 kWh), and SEC is 41 % worse than the optimum."
)
doc.add_paragraph(
    "The shape of the tradeoff is the reason the optimum sits at 0.05 and not at 0 or 0.20. Pushing "
    "the threshold down past 0.05 forces the heat pump to work when the chamber cannot use the dry "
    "air it produces, which wastes compressor energy. Pushing it up past 0.05 makes the heat pump "
    "look cheap per timestep, but the dryer takes so much longer to finish that the fans (which "
    "run continuously) consume more total energy than the compressor savings, and the integrated "
    "SEC gets worse. The 3× hysteresis band (0.05 → 0.15) and the 600-second dwell timer are tuned "
    "to this same window: they stop the controller from chattering across the threshold at every "
    "timestep, while still letting it react within ten minutes when the chamber state genuinely "
    "changes. Across all twenty site × season cases we audited, 0.05 was best or within sweep noise "
    "of the best, and 0.20 was worst at every single one, which is why we use 0.05 as the headline "
    "operating point for the paper."
)

# ------- Paragraph 5 -------
h = doc.add_paragraph()
r = h.add_run("5. Generalising the sweep across sites and seasons")
r.bold = True
r.font.size = Pt(13)

doc.add_paragraph(
    "The sweep numbers in the table above are for one site and one season; the same trend holds "
    "across the rest of the runs. The figure below shows the same threshold sweep for two "
    "representative cases (Kathmandu over the full year and Biratnagar in autumn), with three "
    "coupled views: SEC, drying time, and the fraction of the run spent in bypass."
)

if IMG.exists():
    doc.add_picture(str(IMG), width=Inches(6.5))
    cap = doc.paragraphs[-1]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

fig_cap = doc.add_paragraph()
fig_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fig_cap.add_run(
    "Figure 1. E2 VPD threshold sweep at A_solar = 10 m². Left: specific energy consumption (SEC); "
    "middle: total drying time; right: fraction of the run spent in bypass mode. The dashed vertical "
    "line marks the recommended threshold (0.05)."
)
r.italic = True
r.font.size = Pt(10)

doc.add_paragraph(
    "Reading the three panels together makes the tradeoff explicit. The bypass-duty curve on the "
    "right is monotonic: a stricter threshold means the controller almost never engages bypass, a "
    "looser threshold means it spends most of the run recirculating. The drying-time panel in the "
    "middle is also monotonic but in the opposite direction; every step up in threshold adds hours "
    "to the batch because more time is spent in a low-power recirculation loop where the chamber "
    "dehumidifies slowly. The SEC panel on the left is the combination of those two effects, and it "
    "is the only one that has a clear minimum. To the left of 0.05 the compressor runs at full "
    "power even when the chamber cannot use the dry air, which pushes SEC up; to the right of 0.05 "
    "the dryer takes long enough that fan energy starts to dominate and SEC also climbs. The "
    "intersection of those two opposing trends sits at 0.05 for both cases shown, and the same "
    "minimum is seen at all four sites and four seasons across the full audit (twenty site × "
    "season combinations, with 0.05 best or within sweep noise of best at every one). That is why "
    "0.05 is the value we adopt as the headline operating point for the paper."
)

doc.save(OUT)
print(f"wrote {OUT}")
