"""In-place edits to VPD_Bypass_Mechanism.docx:
   1. Replace 'a few seconds in contact...' with the actual 7-second residence time.
   2. Insert a new paragraph explaining why both VPDs use the same reference temperature.
   Preserves all other user edits.
"""
from pathlib import Path
from copy import deepcopy
from docx import Document

PATH = Path("VPD_Bypass_Mechanism.docx")
doc = Document(PATH)

# --- Edit 1: replace residence-time wording in paragraph 15 -----------------
old = ("the chamber is not a saturator (the air spends only a few seconds in "
       "contact with the apples, and the tray-stack residence time is short)")
new = ("the chamber is not a saturator: the air spends only about 7 seconds in "
       "contact with the apples per pass (10 trays of 0.79 m at 1.1 m/s)")

target = doc.paragraphs[15]
found = False
for run in target.runs:
    if old in run.text:
        run.text = run.text.replace(old, new)
        found = True
        break
if not found:
    full_text = target.text
    if old in full_text:
        for run in target.runs[1:]:
            run.text = ""
        target.runs[0].text = full_text.replace(old, new)
        found = True
assert found, "could not find residence-time string in paragraph 15"

# --- Edit 2: insert new paragraph after utilization equation (para 13) -------
#     Place it before the Interpretation paragraph (para 14).
anchor = doc.paragraphs[14]   # Interpretation paragraph
src    = doc.paragraphs[15]   # use as style template (regular body paragraph)

new_p = deepcopy(src._p)
for r in list(new_p):
    if r.tag.endswith('}r'):
        new_p.remove(r)

anchor._p.addprevious(new_p)

from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
inserted = Paragraph(new_p, anchor._parent)

inserted.add_run(
    "A note on why both VPDs are evaluated at the same 45 \u00B0C. Saturation "
    "pressure P_sat depends on temperature, so if we compared inlet and outlet "
    "at their own temperatures the VPD difference would mix two effects: the "
    "moisture the air picked up in the chamber, and the fact that the exhaust "
    "happens to be a few degrees cooler than the inlet by the time it is "
    "sampled. By forcing both states to the chamber setpoint (the temperature "
    "the air is supposed to operate at), the only thing left to change between "
    "VPD_in and VPD_exh is the partial pressure of water vapour. Utilization "
    "therefore reflects humidity loading only, which is the quantity the "
    "controller actually needs."
)

doc.save(PATH)
print(f"updated {PATH}")
