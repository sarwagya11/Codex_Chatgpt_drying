"""In-place edits round 2:
   1. Add a baseline (VPD OFF) row to the threshold sweep table.
   2. Insert a new paragraph after the table explaining WHEN each threshold
      triggers (first-flip timing + MR) so the SEC numbers make sense.
   Preserves all prior edits.
"""
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Pt

PATH = Path("VPD_Bypass_Mechanism.docx")
doc = Document(PATH)

# ---- Edit 1: prepend a baseline row to the threshold sweep table (table 1)
tbl = doc.tables[1]
hdr_row = tbl.rows[0]
first_data_row = tbl.rows[1]

new_tr = deepcopy(first_data_row._tr)
hdr_row._tr.addnext(new_tr)

baseline_values = [
    "off (baseline)",
    "14.7 h",
    "2.149",
    "0.366",
    "19.32",
    "0.1301",
    "0.0 %",
]
new_row = tbl.rows[1]
for j, val in enumerate(baseline_values):
    cell = new_row.cells[j]
    for p in cell.paragraphs:
        p._p.getparent().remove(p._p)
    p = cell.add_paragraph()
    run = p.add_run(val)
    run.bold = True
    run.font.size = Pt(9)

# ---- Edit 2: insert new paragraph after the threshold sweep table
# The "At a very strict threshold..." paragraph in the current doc is para 39.
# Insert the new explanation BEFORE that paragraph.
target_text_start = "At a very strict threshold"
anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith(target_text_start):
        anchor = p
        break
assert anchor is not None, "could not find anchor paragraph"

template = anchor
new_p = deepcopy(template._p)
for child in list(new_p):
    if child.tag.endswith('}r'):
        new_p.remove(child)
anchor._p.addprevious(new_p)
inserted = Paragraph(new_p, anchor._parent)

inserted.add_run(
    "Adding the baseline row (no VPD bypass at all) makes the rest of the table "
    "easier to read. With bypass disabled the dryer runs the full chain (HRX, "
    "solar, condenser, chamber, HRX-out) for every minute of the batch, finishes "
    "in 14.7 hours, but spends 2.149 kWh of compressor work doing it because the "
    "compressor is still trying to deliver fresh dry air long after the apples "
    "stop being able to use it; SEC lands at 0.1301 kWh/kg. Every row below that "
    "is the same run with the bypass controller switched on at a different "
    "threshold, and the key question is when in the batch the controller first "
    "decides the air is wasted. From the per-timestep CSVs the first flip "
    "happens at very different moments: at threshold = 0.02 the controller waits "
    "until t = 9.55 h (MR = 0.047, deep in the falling-rate phase) before it "
    "ever switches on; at threshold = 0.05 the first flip is at t = 7.07 h "
    "(MR = 0.137, early falling-rate); at threshold = 0.10 it triggers as early "
    "as t = 4.27 h (MR = 0.365, while the apples are still in the constant-rate "
    "phase and the air is genuinely thirsty). After that first flip the "
    "controller oscillates: every entry into bypass closes the recirculation "
    "loop, the exhaust humidifies until utilization rises above the 3\u00d7 "
    "hysteresis line, the controller flips back to fresh-air mode for a short "
    "burst to dump the moisture, and the cycle repeats. So \u201cwhen does the "
    "system go to bypass\u201d is not a single answer, it is a duty cycle, and "
    "the threshold controls two things at once: how early the first switch "
    "happens and how long the bypass intervals are between dumps."
)

doc.save(PATH)
print(f"updated {PATH}")
