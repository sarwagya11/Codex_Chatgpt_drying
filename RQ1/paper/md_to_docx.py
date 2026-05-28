"""Convert DRAFT.md to a professionally formatted DRAFT.docx.

Features:
- Times New Roman 11 pt body, justified, 1.15 line spacing
- Hierarchical headings (Title / H2 / H3 / H4) with proper sizing
- Italic group sub-headers (e.g. *Refrigerant cycle.*)
- Equations: Cambria Math, centered, subscripts/superscripts rendered from
  `_x` and `^x` patterns, equation number right-aligned at margin
- Captions for **Table N.** ... blocks
- Tables: light grid, header row bold + shaded, 9.5 pt
- Blockquotes in italic indented
"""
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(r"D:\Masters\RQ5\Codex_Chatgpt_drying\RQ1\paper\DRAFT.md")
DST = Path(r"D:\Masters\RQ5\Codex_Chatgpt_drying\RQ1\paper\DRAFT.docx")

BODY_FONT = "Times New Roman"
MATH_FONT = "Cambria Math"

EQ_TAIL = re.compile(r"\((\d+)\)\.?\s*$")
CAPTION_RE = re.compile(r"^\*\*(Table|Fig\.?|Figure)\s+\d+[.\)]\*\*\s*")
IMAGE_RE = re.compile(r"^!\[(?P<caption>[^\]]*)\]\((?P<path>[^)]+)\)\s*$")


# ---------------------------------------------------------------------------
# Low-level XML helpers
# ---------------------------------------------------------------------------

def set_cell_border(cell, color="808080", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_run_font(run, name=BODY_FONT, size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    # Force east-asian/complex-script fonts
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def paragraph_format(p, before=0, after=6, line=1.15, align=None, indent_left=0,
                     indent_first=0, keep_with_next=False):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        p.alignment = align
    if indent_left:
        fmt.left_indent = Cm(indent_left)
    if indent_first:
        fmt.first_line_indent = Cm(indent_first)
    if keep_with_next:
        fmt.keep_with_next = True


# ---------------------------------------------------------------------------
# Inline parsing (bold / italic / code)
# ---------------------------------------------------------------------------

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*\s][^*]*\*|`[^`]+`)")


def add_inline_runs(p, text, base_size=11, base_bold=False, base_italic=False,
                    font=BODY_FONT, color=None):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            set_run_font(r, font, base_size, base_bold, base_italic, color)
        tok = m.group(0)
        if tok.startswith("**"):
            r = p.add_run(tok[2:-2])
            set_run_font(r, font, base_size, True, base_italic, color)
        elif tok.startswith("`"):
            r = p.add_run(tok[1:-1])
            set_run_font(r, "Consolas", base_size - 1, base_bold, base_italic, color)
        else:
            r = p.add_run(tok[1:-1])
            set_run_font(r, font, base_size, base_bold, True, color)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        set_run_font(r, font, base_size, base_bold, base_italic, color)


# ---------------------------------------------------------------------------
# Equation rendering with _subscript and ^superscript
# ---------------------------------------------------------------------------

def _is_id_char(c):
    return c.isalnum() or ord(c) > 127  # includes Greek


def _scan_sub_super(text, start):
    """From index start, return end-exclusive index of identifier-like run."""
    n = len(text)
    j = start
    while j < n:
        c = text[j]
        if _is_id_char(c):
            j += 1
        elif c == "," and j + 1 < n and _is_id_char(text[j + 1]):
            j += 1
        elif c == "_" and j > start:  # nested subscript inside super, etc.
            j += 1
        else:
            break
    return j


def render_equation_text(p, text, size=11):
    """Walk text, emit runs in Cambria Math with sub/super-script handling."""
    n = len(text)
    i = 0
    while i < n:
        c = text[i]
        if c == "_" and i + 1 < n and (_is_id_char(text[i + 1]) or text[i + 1] == "("):
            if text[i + 1] == "(":
                depth = 1
                k = i + 2
                while k < n and depth > 0:
                    if text[k] == "(":
                        depth += 1
                    elif text[k] == ")":
                        depth -= 1
                    k += 1
                sub = text[i + 2:k - 1]
                end = k
            else:
                end = _scan_sub_super(text, i + 1)
                sub = text[i + 1:end]
            if sub:
                r = p.add_run(sub)
                set_run_font(r, MATH_FONT, size)
                r.font.subscript = True
                i = end
                continue
        if c == "^" and i + 1 < n and (_is_id_char(text[i + 1]) or text[i + 1] == "("):
            if text[i + 1] == "(":
                depth = 1
                k = i + 2
                while k < n and depth > 0:
                    if text[k] == "(":
                        depth += 1
                    elif text[k] == ")":
                        depth -= 1
                    k += 1
                sup = text[i + 2:k - 1]
                end = k
            else:
                end = _scan_sub_super(text, i + 1)
                sup = text[i + 1:end]
            if sup:
                r = p.add_run(sup)
                set_run_font(r, MATH_FONT, size)
                r.font.superscript = True
                i = end
                continue
        # Plain run until next special marker
        j = i
        while j < n and text[j] not in "_^":
            j += 1
        r = p.add_run(text[i:j])
        set_run_font(r, MATH_FONT, size)
        i = j


def add_equation_paragraph(doc, eq_text, page_width_cm=17.0):
    p = doc.add_paragraph()
    paragraph_format(p, before=4, after=4, line=1.15, indent_left=0.8)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(page_width_cm), WD_ALIGN_PARAGRAPH.RIGHT)
    m = EQ_TAIL.search(eq_text)
    label = ""
    body = eq_text
    if m:
        label = f"({m.group(1)})"
        body = eq_text[:m.start()].rstrip(" ,;")
    render_equation_text(p, body, size=11)
    if label:
        r = p.add_run("\t" + label)
        set_run_font(r, BODY_FONT, 11)


def is_equation_line(stripped):
    if not stripped or not EQ_TAIL.search(stripped):
        return False
    return ("=" in stripped) or ("≤" in stripped) or ("≥" in stripped)


# ---------------------------------------------------------------------------
# Heading styling
# ---------------------------------------------------------------------------

HEADING_SPECS = {
    1: dict(size=18, bold=True, before=0, after=10, align=WD_ALIGN_PARAGRAPH.CENTER),
    2: dict(size=14, bold=True, before=14, after=6, align=WD_ALIGN_PARAGRAPH.LEFT),
    3: dict(size=12, bold=True, before=10, after=4, align=WD_ALIGN_PARAGRAPH.LEFT),
    4: dict(size=11, bold=True, italic=True, before=8, after=3, align=WD_ALIGN_PARAGRAPH.LEFT),
    5: dict(size=11, bold=True, italic=True, before=6, after=2, align=WD_ALIGN_PARAGRAPH.LEFT),
}


def add_heading(doc, level, text):
    spec = HEADING_SPECS.get(level, HEADING_SPECS[4])
    p = doc.add_paragraph()
    paragraph_format(p, before=spec["before"], after=spec["after"], line=1.15,
                     align=spec["align"], keep_with_next=True)
    r = p.add_run(text)
    set_run_font(r, BODY_FONT, spec["size"], bold=spec.get("bold", False),
                 italic=spec.get("italic", False))


# ---------------------------------------------------------------------------
# Table rendering
# ---------------------------------------------------------------------------

def parse_table_block(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        rows.append(lines[i].strip())
        i += 1
    cells_rows = []
    for r in rows:
        cells = [c.strip() for c in r.strip("|").split("|")]
        if all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
            continue
        cells_rows.append(cells)
    return cells_rows, i


def add_table(doc, cells_rows):
    if not cells_rows:
        return
    ncols = max(len(r) for r in cells_rows)
    table = doc.add_table(rows=len(cells_rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for ri, row in enumerate(cells_rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.text = ""
            text = row[ci] if ci < len(row) else ""
            p = cell.paragraphs[0]
            paragraph_format(p, before=0, after=0, line=1.1)
            add_inline_runs(p, text, base_size=9, base_bold=(ri == 0))
            for run in p.runs:
                run.font.name = BODY_FONT
            if ri == 0:
                shade_cell(cell, "E7E6E6")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell, color="808080")
    p = doc.add_paragraph()
    paragraph_format(p, before=0, after=4)


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

def configure_document(doc):
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), BODY_FONT)
    rFonts.set(qn("w:hAnsi"), BODY_FONT)
    rFonts.set(qn("w:cs"), BODY_FONT)

    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)


def add_body_paragraph(doc, text, *, justify=True):
    if not text:
        return
    p = doc.add_paragraph()
    paragraph_format(
        p, before=0, after=6, line=1.15,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT,
    )
    add_inline_runs(p, text, base_size=11)


def add_caption(doc, text):
    p = doc.add_paragraph()
    paragraph_format(p, before=8, after=2, line=1.15, align=WD_ALIGN_PARAGRAPH.LEFT,
                     keep_with_next=True)
    add_inline_runs(p, text, base_size=10)


def add_image(doc, caption_text, image_path_str):
    img_path = (SRC.parent / image_path_str).resolve()
    if not img_path.exists():
        # fall back to caption-only paragraph
        add_caption(doc, f"[image not found: {img_path}] {caption_text}")
        return
    p = doc.add_paragraph()
    paragraph_format(p, before=8, after=2, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER,
                     keep_with_next=True)
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(15.5))
    if caption_text:
        cap = doc.add_paragraph()
        paragraph_format(cap, before=2, after=6, line=1.15,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
        # bold the "Fig. N." prefix if present
        m = re.match(r"^(Fig\.?\s+\d+\.|Figure\s+\d+\.)\s*(.*)$", caption_text)
        if m:
            r1 = cap.add_run(m.group(1) + " ")
            set_run_font(r1, BODY_FONT, 10, bold=True)
            r2 = cap.add_run(m.group(2))
            set_run_font(r2, BODY_FONT, 10)
        else:
            add_inline_runs(cap, caption_text, base_size=10)


def add_blockquote(doc, text):
    p = doc.add_paragraph()
    paragraph_format(p, before=2, after=2, line=1.15, indent_left=0.6,
                     align=WD_ALIGN_PARAGRAPH.LEFT)
    add_inline_runs(p, text, base_size=10, base_italic=True, color="555555")


def add_italic_subheader(doc, text):
    p = doc.add_paragraph()
    paragraph_format(p, before=6, after=2, line=1.15, keep_with_next=True)
    r = p.add_run(text)
    set_run_font(r, BODY_FONT, 11, bold=True, italic=True)


def add_bullet(doc, text):
    # No dot/dash marker: the source already carries roman numerals (i), (ii)...
    # Render as a hanging-indent paragraph so the numeral sticks to the left
    # margin and wrap text aligns past it.
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(2)
    fmt.line_spacing = 1.15
    fmt.left_indent = Cm(0.8)
    fmt.first_line_indent = Cm(-0.8)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline_runs(p, text, base_size=11)


def is_italic_subheader(line):
    if not (line.startswith("*") and line.endswith("*")):
        return False
    if line.startswith("**"):
        return False
    # body should be plain (no nested markdown)
    return len(line) > 2 and "*" not in line[1:-1]


def is_caption(line):
    return bool(CAPTION_RE.match(line))


def process():
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    configure_document(doc)

    para_buf = []

    def flush():
        nonlocal para_buf
        if not para_buf:
            return
        joined = " ".join(s.strip() for s in para_buf).strip()
        para_buf = []
        if not joined:
            return
        if is_caption(joined):
            add_caption(doc, joined)
        else:
            add_body_paragraph(doc, joined)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            flush()
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            flush()
            level = len(m.group(1))
            add_heading(doc, level, m.group(2).strip())
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            flush()
            cells_rows, end = parse_table_block(lines, i)
            add_table(doc, cells_rows)
            i = end
            continue

        img_m = IMAGE_RE.match(stripped)
        if img_m:
            flush()
            add_image(doc, img_m.group("caption").strip(), img_m.group("path").strip())
            i += 1
            continue

        if stripped.startswith(">"):
            flush()
            add_blockquote(doc, stripped.lstrip("> ").strip())
            i += 1
            continue

        if stripped.startswith("- "):
            flush()
            body = stripped[2:].strip()
            collected = [body]
            j = i + 1
            while j < len(lines):
                nxt = lines[j]
                if not nxt.strip():
                    break
                if nxt.lstrip().startswith(("- ", "#", "|", ">")):
                    break
                if nxt.startswith("  "):
                    collected.append(nxt.strip())
                    j += 1
                else:
                    break
            for k, piece in enumerate(collected):
                if is_equation_line(piece):
                    add_equation_paragraph(doc, piece)
                else:
                    if k == 0:
                        add_bullet(doc, piece)
                    else:
                        # continuation text inside the bullet block
                        p = doc.add_paragraph()
                        paragraph_format(p, before=0, after=2, line=1.15,
                                         indent_left=0.8)
                        add_inline_runs(p, piece, base_size=11)
            i = j
            continue

        if line.startswith("  ") and is_equation_line(stripped):
            flush()
            add_equation_paragraph(doc, stripped)
            i += 1
            continue

        if is_italic_subheader(stripped):
            flush()
            add_italic_subheader(doc, stripped.strip("*"))
            i += 1
            continue

        if not stripped:
            flush()
            i += 1
            continue

        para_buf.append(stripped)
        i += 1

    flush()
    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DST))
    print(f"Wrote: {DST}")


if __name__ == "__main__":
    process()
