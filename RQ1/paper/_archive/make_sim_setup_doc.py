"""Generate Simulation_Setup.docx from 04_simulation_setup.md."""
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

SRC = Path(__file__).parent / "04_simulation_setup.md"
DST = Path(__file__).parent / "Simulation_Setup.docx"


def add_inline(p, text):
    tokens = re.split(r"(\*\*[^*]+\*\*|_[^_]+_)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("_") and tok.endswith("_"):
            r = p.add_run(tok[1:-1]); r.italic = True
        else:
            p.add_run(tok)


def render_table(doc, rows):
    n_rows = len(rows)
    n_cols = len(rows[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(val)
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
            if i == 0:
                run.bold = True


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        line = raw
        if not line.strip():
            doc.add_paragraph(); i += 1; continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip()); run.bold = True; run.font.size = Pt(16)
            i += 1; continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip()); run.bold = True; run.font.size = Pt(13)
            i += 1; continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip()); run.bold = True; run.font.size = Pt(11)
            i += 1; continue
        if line.startswith("[Figure") and line.endswith("]"):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line); run.italic = True; run.font.size = Pt(10)
            i += 1; continue
        if line.startswith("|") and "|" in line[1:]:
            # Collect table block
            tbl = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if all(set(c) <= set("-: ") for c in cells):
                    i += 1; continue  # skip separator
                tbl.append(cells); i += 1
            if tbl:
                render_table(doc, tbl)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
            i += 1; continue
        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s", "", line))
            i += 1; continue

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(6)
        add_inline(p, line)
        i += 1

    doc.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
