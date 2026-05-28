"""Consolidate §5.1, §5.2, §5.3, §5.4 into one Results.docx with all figures."""
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

HERE = Path(__file__).parent
SRCS = [HERE / "05_results.md"]
DST = HERE / "Results.docx"

FIG_MAP = {
    "Figure 5.1":  HERE / "fig_5_1_sec_heatmap.png",
    "Figure 5.2b": HERE / "fig_5_2b_deep.png",
    "Figure 5.3a": HERE / "fig_5_3a_sec_bar.png",
    "Figure 5.3b": HERE / "fig_5_3b_B_timeseries.png",
    "Figure 5.3c": HERE / "fig_5_3c_C1_TPJ_diag.png",
    "Figure 5.3d": HERE / "fig_5_3d_C2_evap.png",
    "Figure 5.4a": HERE / "fig_5_4a_sec_bar.png",
    "Figure 5.4b": HERE / "fig_5_4b_D2_evap.png",
    "Figure 5.4c": HERE / "fig_5_4c_D3_humidity.png",
}


def add_inline(p, text):
    for tok in re.split(r"(\*\*[^*]+\*\*|_[^_]+_)", text):
        if not tok: continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("_") and tok.endswith("_"):
            r = p.add_run(tok[1:-1]); r.italic = True
        else:
            p.add_run(tok)


def render_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            para = cell.paragraphs[0]
            txt = val; bold = False
            if txt.startswith("**") and txt.endswith("**"):
                txt = txt[2:-2]; bold = True
            run = para.add_run(txt)
            run.font.size = Pt(9.5); run.font.name = "Times New Roman"
            if i == 0 or bold:
                run.bold = True


def render_file(doc, src):
    lines = src.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            doc.add_paragraph(); i += 1; continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:].strip()); r.bold = True; r.font.size = Pt(16)
            i += 1; continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line[3:].strip()); r.bold = True; r.font.size = Pt(13)
            i += 1; continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:].strip()); r.bold = True; r.font.size = Pt(11)
            i += 1; continue
        if line.startswith("[Figure"):
            inserted = False
            # match longest key first
            for key in sorted(FIG_MAP.keys(), key=len, reverse=True):
                if key in line and FIG_MAP[key].exists():
                    doc.add_picture(str(FIG_MAP[key]), width=Inches(6.3))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cr = cap.add_run(line.strip("[]"))
                    cr.italic = True; cr.font.size = Pt(9.5)
                    inserted = True; break
            if not inserted:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(line); r.italic = True; r.font.size = Pt(10)
            i += 1; continue
        if line.startswith("|") and "|" in line[1:]:
            tbl = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if all(set(c) <= set("-: ") for c in cells):
                    i += 1; continue
                tbl.append(cells); i += 1
            if tbl: render_table(doc, tbl)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet"); add_inline(p, line[2:])
            i += 1; continue
        if re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number"); add_inline(p, re.sub(r"^\d+\.\s", "", line))
            i += 1; continue
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(6)
        add_inline(p, line)
        i += 1


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"; style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
    for src in SRCS:
        render_file(doc, src)
    doc.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
