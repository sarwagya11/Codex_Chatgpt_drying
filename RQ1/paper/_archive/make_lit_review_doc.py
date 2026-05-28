"""Generate Literature_Review.docx from 01_literature_review.md."""
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

TABLE_ROWS = [
    ["Ref.", "Topology / variant", "Crop", "Site / climate", "COP", "SMER (kg/kWh)", "Headline result"],
    ["Prasertsan & Saen-saby (1998)", "HPD baseline (no solar)", "Banana", "Thailand, tropical", "n/r", "0.572", "Early agro-food HPD demonstrator"],
    ["Hawlader et al. (2003)", "Air collector upstream of condenser", "n/r (review/sim)", "Singapore, tropical", "n/r", "n/r", "Reduced condenser-side lift via solar preheat"],
    ["Hawlader & Jahangeer (2006); Hawlader, Rahman & Jahangeer (2008)", "Solar evaporator-collector (DX)", "Green beans", "Singapore, tropical", "7.0 sim / 5.0 exp", "0.65", "Collector eta 0.80–0.86; tropical upper bound"],
    ["Mortezapour et al. (2012)", "PV/T + heat pump hybrid", "Saffron", "Iran, semi-arid", "n/r", "1.16", "33% energy reduction vs baseline"],
    ["Şevik (2013)", "PV-supplied air collector + HP", "Carrot at 50 °C", "Türkiye, continental", "n/r", "n/r", "Air-collector eta 60–78%"],
    ["Rahman et al. (2013)", "Evap-collector + air-collector", "n/r (model)", "Tropical (FORTRAN)", "n/r", "n/r", "Payback ~4 yr"],
    ["Amin & Hawlader (2013)", "R134a evaporator-collector", "Multiple", "Singapore, tropical", "up to 8.0", "n/r", "Tropical-climate upper bound"],
    ["Qiu et al. (2016)", "SAHPD + HR + TES", "Radish, pepper, mushroom", "China", "3.21–3.49", "n/r", "−40.5% input energy; payback 2–6 yr"],
    ["Yahya et al. (2016)", "SAHPD", "Cassava", "Indonesia, tropical", "n/r", "0.38–0.47", "Field trial"],
    ["Kuan et al. (2019)", "SAHPD", "Mint", "Continental-cold", "n/r", "n/r", "Lower solar fraction at cold sites"],
    ["Singh et al. (2020)", "HPD with R1234yf", "Banana chips", "India", "n/r", "n/r", "Drying rate 0.205 kg/(kg·min)"],
    ["Ismaeel & Yumrutaş (2020)", "SAHPD + TES (+HRU)", "Wheat", "Türkiye, periodic model", "5.55", "9.25 (yr-5)", "+HRU adds 21.4% annual saving; eta_HRU up to 41.7%"],
    ["Li et al. (2021)", "SAHPD", "Multiple", "High-altitude China", "n/r", "n/r", "Slower payback at altitude"],
    ["Rulazi et al. (2023)", "SAHPD", "Fruit/vegetable", "East Africa", "n/r", "n/r", "Within typical 0.4–2.7 SMER cluster"],
    ["Yahya et al. (2023)", "HPD", "Paddy", "Indonesia, tropical", "n/r", "0.44", "SEC 4.69 kWh/kg"],
    ["Aacharya et al. (2024)", "Solar-only, low-e double-pass + counter-flow plate HRX", "Apple", "Dhulikhel, Nepal (mid-hill)", "n/r", "n/r", "Coll eta 89%; drying rate 107 g/(h·m²); payback 1.61 yr"],
    ["Yu et al. (2024)", "Ejector + solar evaporator HPD", "n/r", "China", "n/r", "1.40", "+28.7% MER; +54.3% exergy efficiency"],
    ["Abdullah et al. (2025)", "Dual-condenser SAHPD with hot-water solar", "Pandan herbs", "Malaysia, tropical", "6.53", "2.71", "Solar preheats one of two condensers"],
    ["Adhikari et al. (2025)", "Solar-only, CFD chamber redesign", "Apple", "Dhulikhel, Nepal (mid-hill)", "n/r", "n/r", "Uniformity index 0.569 → 0.728"],
    ["Tang et al. (2025)", "Multistage SAHPD", "Tomato", "China", "n/r", "40.7", "85.1% solar contribution (spring/autumn)"],
    ["Kang et al. (2024)", "Solar + HP + PCM TES with fuzzy multi-energy adaptive control", "Kelp", "China", "n/r", "n/r", "T setpoint 40 ± 2 °C; RH setpoint 30 ± 5%; T deviation ±2 °C (vs ±5 °C original); energy −43.0% (S-HP), −16.4% (HS-HP)"],
    ["Loemba et al. (2023)", "Review of HPD types incl. closed/open/bypass loops", "Various agro-food", "Tanzania (review)", "1.94–5.34", "0.156–9.25", "Closed-type SMER maximum at BAR ≈ 0.4; up to 80% energy reduction reported across reviewed designs"],
    ["Wang et al. (2026)", "High-T HPD + bypass-airflow control", "Bamboo", "China", "n/r", "n/r", "−36.5% input; +60.4% SMER; −57.5% SEC vs no-bypass"],
    ["This work", "Ten-topology sweep (A, B, C1, C2, D1, D2, D3, E1, E2, E3) with VPD bypass", "Apple slices", "Biratnagar 72 m, Kathmandu 1350 m, Taplejung 1820 m (Nepal)", "to be reported", "to be reported", "Unified first-law model; VPD-triggered exhaust bypass; four-season"],
]

TABLE_CAPTION = ("Table 1. Verified performance metrics for solar-assisted heat-pump and "
                 "related dryer studies cited in §1.4–§1.5. Entries marked 'n/r' (not reported) "
                 "indicate the metric is absent from the cited work; values are reproduced "
                 "directly from the source as recorded in LIT_REVIEW_LEDGER.md.")


def add_table(doc):
    p = doc.add_paragraph()
    cap = p.add_run(TABLE_CAPTION)
    cap.italic = True
    cap.font.size = Pt(10)

    n_rows = len(TABLE_ROWS)
    n_cols = len(TABLE_ROWS[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    for i, row in enumerate(TABLE_ROWS):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(val)
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            if i == 0:
                run.bold = True

SRC = Path(__file__).parent / "01_literature_review.md"
DST = Path(__file__).parent / "Literature_Review.docx"

def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(16)
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(13)
            continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.size = Pt(11)
            continue
        if line.startswith("_") and line.endswith("_") and len(line) > 2:
            p = doc.add_paragraph()
            run = p.add_run(line.strip("_"))
            run.italic = True
            run.font.size = Pt(10)
            continue

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(6)
        # parse inline italics _text_ and bold **text**
        tokens = re.split(r"(\*\*[^*]+\*\*|_[^_]+_)", line)
        for tok in tokens:
            if not tok:
                continue
            if tok.startswith("**") and tok.endswith("**"):
                r = p.add_run(tok[2:-2])
                r.bold = True
            elif tok.startswith("_") and tok.endswith("_"):
                r = p.add_run(tok[1:-1])
                r.italic = True
            else:
                p.add_run(tok)

    doc.add_page_break()
    h = doc.add_paragraph()
    hr = h.add_run("Table 1")
    hr.bold = True
    hr.font.size = Pt(13)
    add_table(doc)

    doc.save(DST)
    print(f"Wrote {DST}")

if __name__ == "__main__":
    main()
