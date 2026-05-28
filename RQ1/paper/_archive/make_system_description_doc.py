"""Generate System_Description.docx from 02_system_description.md."""
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path(__file__).parent / "02_system_description.md"
DST = Path(__file__).parent / "System_Description.docx"


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line[2:].strip())
            run.bold = True
            run.font.size = Pt(16)
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip())
            run.bold = True
            run.font.size = Pt(13)
            continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.size = Pt(11)
            continue
        if line.startswith("[Figure") and line.endswith("]"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.italic = True
            run.font.size = Pt(10)
            continue
        if line.startswith("_") and line.endswith("_") and len(line) > 2:
            p = doc.add_paragraph()
            run = p.add_run(line.strip("_"))
            run.italic = True
            run.font.size = Pt(10)
            continue

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(6)
        tokens = re.split(r"(\*\*[^*]+\*\*|_[^_]+_)", line)
        for tok in tokens:
            if not tok:
                continue
            if tok.startswith("**") and tok.endswith("**"):
                r = p.add_run(tok[2:-2])
                r.bold = True
            elif tok.startswith("_") and tok.endswith("_"):
                r = p.add_run(tok[1:-1])
                r.italic = True
            else:
                p.add_run(tok)

    doc.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
