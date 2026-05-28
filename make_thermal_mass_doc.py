from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# --- Page margins ---
section = doc.sections[0]
section.page_width    = Cm(21)
section.page_height   = Cm(29.7)
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# --- Default font ---
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)


def add_heading(text, size=13, space_before=8, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'


def add_body(text, size=11, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing  = Pt(13)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'


def add_ref(number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before      = Pt(0)
    p.paragraph_format.space_after       = Pt(2)
    p.paragraph_format.left_indent       = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.line_spacing      = Pt(11)
    run = p.add_run(f'[{number}] {text}')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'


# ============================================================
# TITLE
# ============================================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(0)
t.paragraph_format.space_after  = Pt(10)
r = t.add_run('Thermal Mass as a Building Envelope Strategy in Cold Climates')
r.bold = True
r.font.size = Pt(13)
r.font.name = 'Times New Roman'

# ============================================================
# PARAGRAPH 1 — Principles and motivation
# ============================================================
add_body(
    'Adding thermal mass to the building envelope through the use of dense, high-heat-capacity '
    'materials in walls, floors, and ceilings is an established passive strategy for reducing '
    'temperature fluctuations in the indoor environment. In cold-climate countries such as Sweden, '
    'Norway, and Finland, where heating seasons are long and outdoor temperatures can fall sharply '
    'overnight, thermal mass acts primarily as a thermal buffer: it slows the rate at which interior '
    'temperatures drop during heating interruptions or night-time set-back periods, and moderates '
    'short-duration overheating that can occur when south-facing glazing admits solar gains on '
    'winter days [1, 2]. This buffering role is distinct from, though complementary to, that of '
    'insulation, which governs the total rate of envelope heat loss. Both thermal comfort and energy '
    'performance are cited as motivations in recent construction practice and research; however, '
    'evidence consistently shows that the comfort benefit is more direct and immediate, particularly '
    'in well-insulated Nordic envelopes where the marginal energy impact of added mass is smaller [3].'
)

# ============================================================
# PARAGRAPH 2 — Concrete and CLT
# ============================================================
add_body(
    'Concrete remains the most widely used high-mass building material in Nordic residential and '
    'commercial construction, with a volumetric heat capacity of approximately 2,000 kJ/m\u00b3K. '
    'Exposed concrete soffits, heavyweight internal slabs, and prefabricated concrete sandwich '
    'facade panels are routinely specified in Swedish multi-family housing. Comparative simulation '
    'studies across cold-climate cities, including Helsinki, have shown that concrete structural '
    'systems reduce peak indoor operative temperatures by 2\u20134\u00b0C compared to lightweight '
    'timber-frame construction during the warmest periods and sustain habitable indoor conditions '
    'substantially longer during heating system failures [4, 5]. Cross-laminated timber (CLT), '
    'increasingly common in Scandinavian construction since the early 2010s, provides a more '
    'moderate but meaningful thermal mass; research finds that concrete frames consume roughly '
    '2\u20137% less annual heating and cooling energy than equivalent CLT frames, while both '
    'substantially outperform lightweight timber construction [6].'
)

# ============================================================
# PARAGRAPH 3 — PCM
# ============================================================
add_body(
    'Phase change materials (PCMs) have attracted significant recent research interest as a means '
    'of achieving high effective heat storage capacity in relatively thin assemblies. PCMs absorb '
    'large quantities of latent heat as they transition between solid and liquid phases at a defined '
    'temperature, and can be integrated into plasterboard, concrete, or wall insulation layers. '
    'Norwegian research institution SINTEF has evaluated PCM-enhanced wall assemblies under Nordic '
    'conditions [7], and studies report heating demand reductions of approximately 20% when PCM is '
    'incorporated into interior wall layers in cold European climates [8]. An optimised configuration '
    'featuring 75 mm of PCM with a melting point of 21\u00b0C in interior walls has been shown to '
    'yield annual energy cost savings of up to 18% under spot-price electricity tariffs [9]. '
    'An additional practical benefit in cold climates is that PCM integration into concrete '
    'reduces freeze-thaw degradation of the host material, providing a structural durability '
    'advantage alongside the thermal performance benefit [8].'
)

# ============================================================
# PARAGRAPH 4 — TABS and resilience
# ============================================================
add_body(
    'A growing application of concrete thermal mass in Swedish and Nordic construction is '
    'thermally activated building systems (TABS), in which water pipes embedded within heavyweight '
    'concrete floor and ceiling slabs allow controlled pre-loading of the mass using off-peak or '
    'renewable electricity. This strategy enables demand shifting away from peak hours, supporting '
    'the flexibility requirements of electricity grids with high proportions of variable wind '
    'power [10]. Research published in 2025 further frames thermal mass as a resilience asset: '
    'thermal mass capacity of the building envelope was identified as a critical parameter '
    'determining how many hours a building sustains habitable indoor temperatures during a '
    'cold-climate power outage, with high-mass buildings remaining within safe temperature bounds '
    'significantly longer than lightweight equivalents [5]. Taken together, the evidence positions '
    'thermal mass in cold-climate building envelopes as a material strategy serving comfort '
    'stability, energy system flexibility, and structural resilience, with quantified benefits '
    'that support its continued and expanding specification in Nordic construction and renovation.'
)

# ============================================================
# REFERENCES
# ============================================================
add_heading('References', size=11, space_before=6, space_after=3)

refs = [
    ('1',  'Australian Government Department of Industry, Science, Energy and Resources. Thermal Mass. '
           'Your Home Technical Manual. Available: yourhome.gov.au/passive-design/thermal-mass'),
    ('2',  'Ruud S., Lundin L. Experiences from Nine Passive Houses in Sweden: Indoor Thermal Environment '
           'and Energy Use. SP Technical Research Institute of Sweden, 2013. '
           'ResearchGate: doi:10.13140/RG.2.1.3848.8484'),
    ('3',  'Building America Solution Center, PNNL. High-Thermal-Mass Construction. '
           'Available: basc.pnnl.gov/resource-guides/high-thermal-mass-construction'),
    ('4',  'Libralato M., De Angelis A., Saro O. Improving Thermal Response of Lightweight Timber '
           'Building Envelopes During Cooling Season in Three European Locations. '
           'Journal of Cleaner Production, 2018. doi:10.1016/j.jclepro.2017.08.198'),
    ('5',  'Hamdy M., Hasan A. Enhancing Building Resilience: Maintaining Energy Efficiency and Thermal '
           'Comfort During Power Outages in Cold Climates. Building Simulation, Springer, 2025. '
           'doi:10.1007/s12273-025-1371-2'),
    ('6',  'Oak Ridge National Laboratory. Impact of Mass Wood Walls on Building Energy Use, Peak Load, '
           'and Carbon. ORNL Report, 2022. '
           'Available: info.ornl.gov/sites/publications/Files/Pub174562.pdf'),
    ('7',  'SINTEF Blog. What are Phase Change Materials? (Will They Be the Next Big Thing in Norway?). '
           'Available: blog.sintef.com/energy/phase-change-materials-pcm/'),
    ('8',  'Jayalath A. et al. Phase Change Material Integration in Concrete for Thermal Energy Storage: '
           'Techniques and Applications in Sustainable Building. Sustainable Energy Research, Springer, '
           '2024. doi:10.1186/s40807-024-00138-8'),
    ('9',  'Kishore R.A. et al. PCM-Incorporated Building Envelope for Improving Cost Savings in '
           'Residential Buildings Under Cold Climates. Lecture Notes in Civil Engineering, Springer, '
           '2024. doi:10.1007/978-981-97-8305-2_53'),
    ('10', 'Lund H. et al. Energy System and Cost Impacts of Heat Supply to Low-Energy Buildings '
           'in Sweden. Energy, Elsevier, 2023. doi:10.1016/j.energy.2023.126608'),
]

for num, text in refs:
    add_ref(num, text)

# ============================================================
# SAVE
# ============================================================
out_path = 'D:/Masters/RQ5/Codex_Chatgpt_drying/Thermal_Mass_Background.docx'
doc.save(out_path)
print('Saved:', out_path)
